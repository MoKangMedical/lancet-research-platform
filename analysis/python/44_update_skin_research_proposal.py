#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


PROJECT_DIR = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万 2"
)
OUTPUT_DOCX = PROJECT_DIR / "0308-研究方案-全球45岁及以上成年人皮肤及皮下疾病流行病学、疾病负担及趋势-更新版.docx"
OUTPUT_MD = PROJECT_DIR / "0308-研究方案-全球45岁及以上成年人皮肤及皮下疾病流行病学、疾病负担及趋势-更新版.md"

TITLE = "全球45岁及以上成年人皮肤及皮下疾病流行病学、疾病负担及趋势研究"

SECTIONS: list[tuple[str, list[str]]] = [
    (
        "一、研究背景与意义",
        [
            "皮肤及皮下疾病是全球最常见的健康问题之一，虽然致死性通常低于心脑血管疾病和肿瘤，但其长期症状、反复发作、慢性瘙痒、感染风险以及功能受限可显著影响生活质量，并带来持续的卫生资源消耗。随着全球人口老龄化进程加快，中老年人群皮肤屏障功能减退、免疫反应改变、慢病共存增多，皮肤病的发病、持续和结局均呈现更复杂的年龄相关特征。",
            "现有全球疾病负担研究多聚焦全人群或单一病种，针对45岁及以上成年人皮肤及皮下疾病总体负担、年龄和性别差异、分亚型构成以及与全球老龄化背景之间关系的系统研究仍相对不足。因此，有必要基于最新的 GBD 2023 数据，对1990至2023年全球、地区和国家层面的疾病负担进行系统梳理，并在国家层面结合标准化老龄化指标开展生态关联分析。",
            "结合当前已完成稿件和可重复获取的数据基础，本研究对原始方案做出适度调整：由原先“多种常见皮肤病分别分析并广泛连接 global aging data”的设计，优化为“以 GBD 2023 中皮肤及皮下疾病总体负担为主线，辅以分亚型分析，并采用世界银行全球老龄化指标开展国家层面生态分析”的版本，以确保研究问题清晰、数据来源稳定、方法路径可复现，并与当前论文初稿保持一致。",
        ],
    ),
    (
        "二、研究目标",
        [
            "1. 描述1990至2023年全球、地区和国家层面45岁及以上成年人皮肤及皮下疾病的发病、患病、死亡和 DALYs 负担特征。",
            "2. 采用年龄标准化率和估计年度变化百分比（EAPC）评价主要疾病负担指标的长期变化趋势。",
            "3. 分析不同年龄组、性别、地区和社会人口学发展水平（SDI）下皮肤及皮下疾病负担的异质性。",
            "4. 对皮肤及皮下疾病的主要亚型进行分层描述，比较不同亚型在发病、患病、死亡和 DALYs 中的构成差异及其时间变化。",
            "5. 引入全球老龄化指标，在国家层面探索人口老龄化背景与皮肤及皮下疾病死亡负担之间的生态学相关性。",
        ],
    ),
    (
        "三、研究内容（按当前稿件版本形成）",
        [
            "1. 全球、地区和国家负担描述：基于 GBD 2023，系统呈现45岁及以上成年人皮肤及皮下疾病在1990至2023年的发病率、患病率、死亡率和 DALYs，并报告相应的年龄标准化率。",
            "2. 趋势分析：围绕年龄标准化发病率、患病率、死亡率和 DALY 率计算 EAPC，评价全球及不同地区疾病负担变化方向与幅度。",
            "3. SDI 相关模式分析：结合 GBD 提供的 SDI 指标，比较不同发展水平地区疾病负担差异，识别高负担地区与发展梯度特征。",
            "4. 年龄和性别差异分析：在45岁及以上成年人中，分析不同年龄层与不同性别人群的流行病学差异，明确非致死性负担与死亡负担在不同亚群中的分布特点。",
            "5. 疾病亚型分析：对痤疮、斑秃、细菌性皮肤病、压疮、皮炎、真菌性皮肤病、银屑病、疥疮、病毒性皮肤病、荨麻疹、瘙痒症等主要亚型进行时间分布和构成分析。",
            "6. 全球老龄化生态分析：在国家层面整合老年人口占比、预期寿命和老年抚养比等指标，分析其与皮肤及皮下疾病年龄标准化死亡率之间的相关性，并从全球人口老龄化背景下解释疾病负担变化。",
        ],
    ),
    (
        "四、数据来源与研究对象",
        [
            "1. 疾病负担数据来源：全球疾病负担研究（Global Burden of Disease Study, GBD 2023）公开数据库，由 Institute for Health Metrics and Evaluation（IHME）维护。提取1990至2023年皮肤及皮下疾病相关指标，覆盖204个国家和地区。",
            "2. 全球老龄化数据来源：世界银行（World Bank）官方数据库，提取1990至2023年可跨国稳定比较的老龄化指标，包括65岁及以上人口占总人口比例（population ages 65 and above, % of total population）、出生时预期寿命（life expectancy at birth）和老年抚养比（old-age dependency ratio）。",
            "3. 研究对象：GBD 2023 中45岁及以上成年人，按全球、21个GBD地区、国家/地区、年龄组、性别及 SDI 分层进行分析。国家层面老龄化生态分析以2023年能够与世界银行指标匹配的国家/地区为分析单元。",
            "4. 方案调整说明：原方案中的“global aging data”现统一替换为世界银行标准化老龄化指标体系，以提高跨国可比性、数据完整性和后续论文写作的一致性。",
        ],
    ),
    (
        "五、主要变量与指标定义",
        [
            "1. 结局指标：发病数与年龄标准化发病率、患病数与年龄标准化患病率、死亡数与年龄标准化死亡率、DALYs 数与年龄标准化 DALY 率。",
            "2. 趋势指标：EAPC，用于反映1990至2023年年龄标准化率的平均年度变化幅度。",
            "3. 分层变量：年龄组、性别、地区、国家/地区、SDI 分组、皮肤及皮下疾病亚型。",
            "4. 老龄化指标：65岁及以上人口占比、出生时预期寿命、老年抚养比。上述指标主要用于国家层面生态相关性分析和全球人口背景描述。",
            "5. 不确定性表达：所有主要估计值尽可能同时报告95%不确定区间或95%置信区间，并在结果解释中区分绝对数与年龄标准化率。",
        ],
    ),
    (
        "六、统计分析方法",
        [
            "1. 描述性分析：汇总不同年份、地区、国家、年龄组和性别人群的疾病负担指标，采用表格、趋势图和构成图展示结果。",
            "2. 趋势分析：对1990至2023年的年龄标准化率进行对数线性拟合，计算 EAPC 及其95%置信区间，用于判断负担变化趋势。",
            "3. 相关性分析：在国家层面使用 Spearman 相关分析评估老龄化指标与年龄标准化死亡率之间的生态学相关性；相关分析主要用于描述关联，不作因果推断。",
            "4. 分层分析：按 SDI、年龄、性别和疾病亚型进行比较，识别高负担人群与疾病谱差异。",
            "5. 质量控制：统一国家名称与编码体系，严格区分计数指标与年龄标准化率，保留 GBD 与世界银行原始定义，确保提取、匹配和统计过程可复现。",
            "6. 软件工具：数据整理和补充分析主要使用 Python，统计分析和绘图可结合 R 完成，文稿结果部分与现有英文稿件保持一致。",
            "7. 方法调整说明：原始方案中的多元回归、Logistic 回归、ARIMA 时间序列和 Meta 分析等内容不再作为本研究当前主方案的核心方法，避免研究问题与数据结构不匹配。",
        ],
    ),
    (
        "七、预期结果与论文框架",
        [
            "1. 形成一篇以45岁及以上成年人皮肤及皮下疾病全球负担为主题的英文论文，核心框架包括：全球/地区/国家负担、SDI 模式、年龄与性别差异、亚型分析以及全球老龄化生态分析。",
            "2. 预期产出图表包括：全球和地区负担趋势图、SDI 相关图、年龄/性别分层趋势图、亚型构成图以及国家层面老龄化指标与死亡率的生态关联结果表。",
            "3. 预期结果将说明：在全球人口老龄化背景下，45岁及以上成年人皮肤及皮下疾病负担总体呈持续存在并伴随区域差异的特征，非致死性负担与死亡负担在不同年龄和亚型之间呈现明显分化。",
        ],
    ),
    (
        "八、创新点与可行性",
        [
            "1. 以45岁及以上成年人为重点研究对象，较现有全人群研究更贴近人口老龄化背景下的皮肤健康需求。",
            "2. 将 GBD 2023 疾病负担结果与世界银行老龄化指标结合，在全球尺度上补充老龄化生态解释框架。",
            "3. 研究路径与当前英文稿件完全对齐，现有数据提取、国家匹配和分析脚本均已形成，可直接支撑后续结果补充、图表优化与论文定稿。",
            "4. 研究采用公开数据库，数据来源权威、覆盖范围广、可重复性强，具备较高的实施可行性。",
        ],
    ),
    (
        "九、目前版本对应的研究内容摘要",
        [
            "当前版本论文题目可统一为：《全球45岁及以上成年人皮肤及皮下疾病流行病学、疾病负担及趋势研究：基于 GBD 2023 和全球老龄化指标的系统分析》。",
            "当前版本方法学包括五个核心模块：GBD 2023 疾病负担提取、年龄标准化率与 EAPC 趋势分析、SDI 相关模式分析、年龄与性别分层分析、疾病亚型分析，以及世界银行老龄化指标支持的国家层面生态相关性分析。",
            "当前版本结果部分应围绕五个主结果展开：全球及地区负担现状、SDI 分层差异、年龄与性别差异、亚型构成与变化、全球老龄化背景与国家层面生态关联。",
            "当前版本讨论部分重点解释人口老龄化、地区不平等、非致死性疾病负担主导、感染性和压疮相关死亡负担以及生态分析结果的公共卫生意义，同时明确 GBD 属于模型估计数据、生态分析不代表因果推断等限制。",
        ],
    ),
]


def set_run_font(run, east_asia: str, western: str, size_pt: int, bold: bool = False) -> None:
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size_pt)
    run.bold = bold


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)
    section.left_margin = Pt(90)
    section.right_margin = Pt(90)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def add_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="黑体", western="Times New Roman", size_pt=16, bold=True)


def add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="黑体", western="Times New Roman", size_pt=14, bold=True)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    run = paragraph.add_run(text)
    set_run_font(run, east_asia="宋体", western="Times New Roman", size_pt=12, bold=False)


def build_docx(output_path: Path) -> None:
    doc = Document()
    configure_document(doc)
    add_title(doc, TITLE)
    for heading, paragraphs in SECTIONS:
        add_section_heading(doc, heading)
        for paragraph in paragraphs:
            add_body_paragraph(doc, paragraph)
    doc.save(output_path)


def build_markdown(output_path: Path) -> None:
    lines = [f"# {TITLE}", ""]
    for heading, paragraphs in SECTIONS:
        lines.append(f"## {heading}")
        lines.append("")
        for paragraph in paragraphs:
            lines.append(paragraph)
            lines.append("")
    output_path.write_text("\n".join(lines), encoding="utf-8")


def main() -> None:
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    build_docx(OUTPUT_DOCX)
    build_markdown(OUTPUT_MD)
    print(f"Updated proposal written to: {OUTPUT_DOCX}")
    print(f"Markdown summary written to: {OUTPUT_MD}")


if __name__ == "__main__":
    main()
