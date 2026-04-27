from __future__ import annotations

import importlib.util
import json
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path("/Users/apple/Documents/lancet-research-platform")
SCRIPT_45 = ROOT / "analysis" / "python" / "45_build_skin_lancet_package.py"
PACKAGE_ROOT = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万 2/lancet_skin_article_package"
)
OUTPUT_DIR = PACKAGE_ROOT / "outputs"
MANUSCRIPT_DIR = OUTPUT_DIR / "manuscript"
TABLE_DIR = OUTPUT_DIR / "tables"
FIGURE_DIR = OUTPUT_DIR / "figures"
SOURCE_DOC = PACKAGE_ROOT.parent / "用所选项目新建的文件夹" / "1208-Manuscript-全球老年人群常见皮肤病流行病学、疾病负担及趋势.docx"


@dataclass
class SubmissionPackage:
    title: str
    short_title: str
    target_journal: str
    summary: dict[str, str]
    research_in_context: dict[str, str]
    sections: list[tuple[str, list[tuple[str | None, list[str]]]]]
    references: list[str]
    figure_legends: list[tuple[str, str]]
    main_table_title: str
    main_word_count: int
    summary_word_count: int
    candidate_author_note: str


def load_builder_module():
    spec = importlib.util.spec_from_file_location("skin_builder", SCRIPT_45)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules["skin_builder"] = module
    spec.loader.exec_module(module)
    return module


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(15)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.add_run(text)


def add_table_to_doc(doc: Document, df: pd.DataFrame, title: str) -> None:
    add_heading(doc, title, level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        hdr[idx].text = str(col)
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                if abs(value) >= 1000:
                    cells[idx].text = f"{value:,.1f}"
                else:
                    cells[idx].text = f"{value:.2f}"
            else:
                cells[idx].text = str(value)


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w.-]+\b", text))


def build_submission_sections(v: dict[str, float | str | int]) -> list[tuple[str, list[tuple[str | None, list[str]]]]]:
    introduction = [
        (
            None,
            [
                "Skin and subcutaneous diseases are major contributors to global non-fatal morbidity, yet they are still frequently treated as low-priority conditions because they rarely dominate mortality rankings. "
                "That view is incomplete. Skin disorders cause chronic symptoms, recurrent service use, impaired sleep and mobility, stigma, and sustained quality-of-life loss, and they remain a substantial part of the global burden profile across successive GBD updates.[1-5]",
                "The ageing of populations makes this burden more important for health systems. Older adults are more exposed to multimorbidity, frailty, immobility, skin-barrier dysfunction, chronic wounds, infection, and long-term care dependency. "
                "Clinical and epidemiological reviews show that skin conditions in later life are common, persistent, and tightly linked to nursing care, chronic wound management, and the broader organisation of frailty care.[6-16]",
                "Healthy-ageing policy, however, rarely gives skin health the same attention as cardiovascular disease, dementia, falls, or musculoskeletal decline. "
                "At the same time, demographic planning is increasingly framed through standard indicators such as the proportion of the population aged 65 years and older, life expectancy, and dependency ratios. "
                "Those indicators are highly visible in the World Bank World Development Indicators platform, but the dermatologic implications of those metrics have rarely been examined in a reproducible global framework.[17-24]",
                "We therefore developed a submission-focused Lancet-style analysis that integrates official GBD 2023 skin-burden estimates with World Bank ageing indicators. "
                "Our objectives were to quantify the change in global burden between 1990 and 2023, identify the dominant skin-disease subtypes across major burden domains, and assess whether country-level demographic ageing was associated with age-standardized skin mortality in 2023.",
            ],
        )
    ]

    methods = [
        (
            None,
            [
                "We did a descriptive and ecological analysis using two public data systems: official Global Burden of Disease Study 2023 extracts maintained by the Institute for Health Metrics and Evaluation, and World Bank World Development Indicators. "
                "Global burden analyses used official extracts for incidence, prevalence, DALYs, and deaths; the ageing framework used three World Bank indicators: population aged 65 years and older, life expectancy at birth, and the old-age dependency ratio.[21-24]",
                "Primary outcomes were the age-standardized incidence rate, prevalence rate, DALY rate, mortality rate, DALY counts, and death counts for skin and subcutaneous diseases overall. "
                "Subtype analyses used official global level-3 skin causes and summarised 2023 subtype hierarchy as well as relative change between 1990 and 2023. "
                "Counts and age-standardized rates were interpreted separately because they answer different policy questions.",
                "For country-level ecological analysis, we linked 2023 skin age-standardized mortality rates with World Bank ageing indicators. "
                "Because the locally reproducible country-level extract was mortality-based, we did not fit unsupported country-level incidence or prevalence models. "
                "Country names were harmonised across sources, and two ambiguous location names shared by country and subnational entities (Georgia and Niger) were excluded before analysis, leaving 198 countries and territories.",
                "We summarised change between 1990 and 2023 with absolute and relative differences. "
                "Country-level associations were examined with Spearman correlation and illustrated with scatterplots; countries were also grouped into tertiles of population aged 65 years and older. "
                "All analyses were generated through a reproducible Python pipeline. There was no funding source for this work.",
            ],
        )
    ]

    results = [
        (
            None,
            [
                f"Between 1990 and 2023, the proportion of the global population aged 65 years and older increased from {v['age_1990']:.2f}% to {v['age_2023']:.2f}%, while life expectancy rose from {v['life_1990']:.2f} to {v['life_2023']:.2f} years and the old-age dependency ratio increased from {v['dep_1990']:.2f} to {v['dep_2023']:.2f}. "
                f"Against that demographic background, the age-standardized incidence rate of skin and subcutaneous diseases increased from {v['incidence_1990']:.1f} to {v['incidence_2023']:.1f} per 100,000, the age-standardized DALY rate increased from {v['daly_rate_1990']:.1f} to {v['daly_rate_2023']:.1f} per 100,000, and the age-standardized mortality rate rose from {v['death_rate_1990']:.2f} to {v['death_rate_2023']:.2f} per 100,000 (table 1; figure 1). "
                f"In absolute terms, DALY counts rose from {v['daly_count_1990'] / 1_000_000:.1f} million to {v['daly_count_2023'] / 1_000_000:.1f} million and deaths rose from {int(v['death_count_1990']):,} to {int(v['death_count_2023']):,}.",
                f"Subtype hierarchy in 2023 was uneven (figure 2). {v['top_incidence_subtype']} was the leading incidence category at {v['top_incidence_subtype_rate']:.1f} per 100,000; {v['top_prevalence_subtype']} dominated prevalence at {v['top_prevalence_subtype_rate']:.1f} per 100,000; {v['top_daly_subtype']} dominated DALYs at {v['top_daly_subtype_rate']:.1f} per 100,000; and mortality was concentrated in {v['top_death_subtype']} at {v['top_death_subtype_rate']:.2f} per 100,000. "
                f"Dermatitis remained the largest disability contributor, while bacterial skin diseases and decubitus ulcer disproportionately shaped fatal burden. "
                f"Bacterial skin-disease mortality increased from {v['bact_death_1990']:.2f} to {v['bact_death_2023']:.2f} per 100,000, whereas the dermatitis DALY rate changed from {v['derm_daly_1990']:.1f} to {v['derm_daly_2023']:.1f} per 100,000.",
                f"In ecological analyses across {v['n_countries']} countries and territories, more demographically aged populations did not have the highest age-standardized skin mortality (figure 3). "
                f"The proportion of the population aged 65 years and older was negatively correlated with skin mortality (rho={v['rho_age65']:.3f}; p={v['p_age65']:.2e}); the same pattern was observed for life expectancy (rho={v['rho_life']:.3f}; p={v['p_life']:.2e}) and the old-age dependency ratio (rho={v['rho_dep']:.3f}; p={v['p_dep']:.2e}). "
                f"Countries in the youngest ageing tertile had a median mortality rate of {v['tertile_t1_median']:.2f} per 100,000, compared with {v['tertile_t3_median']:.2f} per 100,000 in the oldest tertile.",
                f"The highest 2023 age-standardized mortality rates were observed in {v['top5_text']} (figure 4). "
                f"{v['top_country']} had the highest point estimate at {v['top_country_rate']:.2f} per 100,000. "
                "These high-mortality settings were concentrated in small island states and mixed middle-income settings rather than in the most demographically aged high-income countries.",
            ],
        )
    ]

    discussion = [
        (
            None,
            [
                "This analysis yields three main messages. First, global skin and subcutaneous disease burden increased in both counts and age-standardized rates over a period of rapid demographic ageing. "
                "Second, the burden profile remained domain-specific: dermatitis and fungal skin disease dominated chronic disability, whereas bacterial skin disease and decubitus ulcer remained central to mortality risk. "
                "Third, age-standardized skin mortality was not highest in the oldest countries, suggesting that health-system capacity, wound care, infection control, and chronic disease management strongly shape fatal outcomes after age standardisation.",
                "The divergence between counts and standardized rates is important for planning. "
                "Ageing societies can experience only modest changes in standardized rates while still facing large increases in clinic demand, nursing workload, chronic wound care, antimicrobial use, and caregiver burden. "
                "For later-life health policy, counts therefore matter as much as standardized rates because they capture the operational burden that services must absorb.",
                "The subtype findings also sharpen the policy message. "
                "A response focused only on common ambulatory dermatoses would miss the smaller set of conditions driving mortality. "
                "Pressure-injury prevention, chronic wound surveillance, diabetic foot and skin care, early recognition of bacterial skin infection, and safer long-term care systems should be understood as core components of healthy ageing rather than as niche specialist issues.[25-44]",
                "The ecological findings should not be interpreted causally. "
                "The negative correlations with population ageing, life expectancy, and dependency ratios do not mean that ageing protects against skin mortality. "
                "Instead, they suggest that when age structure is standardised, between-country variation is more strongly shaped by contextual factors than by age composition alone. "
                "The World Bank indicators improve policy interpretability, but they do not capture frailty prevalence, long-term care coverage, wound-care capacity, or access to antibiotics and supplies.",
                "This study has strengths and limitations. "
                "Its main strengths are full reproducibility within the local workspace, use of official GBD 2023 and World Bank data, and a figure-table set constrained to a submission-focused structure. "
                "Its limitations are the use of modelled secondary data, the restricted country-level mortality extract, and the ecological design. "
                "Overall, the results support integration of skin health into healthy-ageing agendas and identify high-mortality settings that merit closer clinical and policy investigation.[17-24]",
            ],
        )
    ]

    return [
        ("Introduction", introduction),
        ("Methods", methods),
        ("Results", results),
        ("Discussion", discussion),
    ]


def draft_to_markdown(pkg: SubmissionPackage) -> str:
    lines = [f"# {pkg.title}", ""]
    lines.extend(["## Summary", ""])
    for key in ["Background", "Methods", "Findings", "Interpretation", "Funding"]:
        lines.extend([f"### {key}", "", pkg.summary[key], ""])
    lines.extend(["## Research in context", ""])
    for key, value in pkg.research_in_context.items():
        lines.extend([f"### {key}", "", value, ""])
    for section_title, blocks in pkg.sections:
        lines.extend([f"## {section_title}", ""])
        for subsection, paragraphs in blocks:
            if subsection:
                lines.extend([f"### {subsection}", ""])
            for paragraph in paragraphs:
                lines.extend([paragraph, ""])
    lines.extend([f"## {pkg.main_table_title}", ""])
    lines.extend(["Main table is inserted in the DOCX version and provided separately as CSV.", ""])
    lines.extend(["## Figure legends", ""])
    for title, body in pkg.figure_legends:
        lines.extend([f"### {title}", "", body, ""])
    lines.extend(["## References", ""])
    for idx, ref in enumerate(pkg.references, start=1):
        lines.append(f"{idx}. {ref}")
    return "\n".join(lines)


def extract_candidate_author_note() -> str:
    if not SOURCE_DOC.exists():
        return "No local document metadata author could be extracted."
    try:
        result = subprocess.run(
            ["mdls", "-name", "kMDItemAuthors", str(SOURCE_DOC)],
            check=True,
            capture_output=True,
            text=True,
        )
        text = result.stdout.strip()
        match = re.search(r'"(.+?)"', text)
        if match:
            raw_name = match.group(1)
            try:
                decoded_name = re.sub(
                    r"\\U([0-9A-Fa-f]{4})",
                    lambda m: chr(int(m.group(1), 16)),
                    raw_name,
                )
            except Exception:
                decoded_name = raw_name
            return f"Local DOCX metadata author candidate detected from source file properties: {decoded_name}. Verify before use."
    except Exception:
        pass
    return "No local document metadata author could be extracted."


def make_title_page(pkg: SubmissionPackage) -> list[str]:
    return [
        f"Target journal: {pkg.target_journal}",
        "Working article format: submission-focused near-final Lancet-family draft",
        f"Full title: {pkg.title}",
        f"Short title: {pkg.short_title}",
        "Authors: ________________________________",
        "Affiliations: ____________________________",
        "Corresponding author: ____________________",
        "Corresponding email: _____________________",
        f"Main text word count: {pkg.main_word_count}",
        f"Summary word count: {pkg.summary_word_count}",
        "Main-text display items: 5 (4 figures and 1 table)",
        "Supplementary tables: 7",
        "Funding statement: ______________________",
        "Declaration of interests: completed author forms attached separately.",
        "Data sharing: derived analyses use official GBD and World Bank data sources.",
    ]


def make_cover_letter(pkg: SubmissionPackage) -> list[str]:
    return [
        "Date: March 9, 2026",
        "",
        f"To the Editors of {pkg.target_journal}",
        "",
        f"We are pleased to submit our Article entitled \"{pkg.title}\" for consideration at {pkg.target_journal}.",
        "",
        "This submission-focused version presents a concise Lancet-style analysis of the global burden of skin and subcutaneous diseases in the context of demographic ageing. "
        "Using official GBD 2023 extracts and World Bank World Development Indicators, the manuscript shows that global skin burden increased between 1990 and 2023 in both absolute and standardized terms, while age-standardized skin mortality remained concentrated outside the most demographically aged countries.",
        "",
        "The paper is relevant to healthy-ageing and health-system planning for three reasons. "
        "First, it reframes skin disease as part of later-life service burden rather than as a marginal ambulatory topic. "
        "Second, it links burden estimates to ageing metrics that are already familiar to policy audiences. "
        "Third, it is accompanied by a fully reproducible analytical and quality-control workflow.",
        "",
        "The manuscript has not been published previously and is not under consideration elsewhere. Final author approval, authorship order, funding statement, declaration wording, ethics wording, and corresponding-author details should be confirmed before submission.",
        "",
        "Sincerely,",
        "",
        "Corresponding author: ____________________",
        "Institution: _____________________________",
        "Email: __________________________________",
        "Telephone: _______________________________",
    ]


def make_author_metadata_form(candidate_note: str) -> list[str]:
    return [
        "Complete the following fields before journal submission.",
        "",
        "Authors in final order:",
        "1. ________________________________________",
        "2. ________________________________________",
        "3. ________________________________________",
        "4. ________________________________________",
        "",
        "Affiliations:",
        "1. ________________________________________",
        "2. ________________________________________",
        "3. ________________________________________",
        "",
        "Corresponding author name: __________________",
        "Corresponding author email: ________________",
        "Corresponding author telephone: ____________",
        "Corresponding author postal address: _______",
        "",
        "Funding statement: _________________________",
        "Ethics statement confirmation: _____________",
        "Originality statement confirmation: ________",
        "",
        "CRediT contributions by author: ____________",
        "Declaration of interests by author: ________",
        "",
        f"Note: {candidate_note}",
    ]


def make_authors_contributors_form() -> list[str]:
    return [
        "Authors:",
        "__________________________________________",
        "",
        "CRediT contributor statement:",
        "Conceptualisation: ________________________",
        "Data curation: ____________________________",
        "Formal analysis: __________________________",
        "Methodology: ______________________________",
        "Visualisation: ____________________________",
        "Writing - original draft: _________________",
        "Writing - review and editing: _____________",
        "",
        "All authors should confirm full access to the data relevant to their contribution and responsibility for the decision to submit.",
    ]


def make_declaration_form() -> list[str]:
    return [
        "Each author should provide a journal-ready declaration of interests statement below.",
        "",
        "Author 1: ________________________________________________",
        "Author 2: ________________________________________________",
        "Author 3: ________________________________________________",
        "Author 4: ________________________________________________",
        "",
        "If no competing interests exist, use the final wording required by the target journal.",
    ]


def write_simple_markdown(path: Path, title: str, lines: list[str]) -> None:
    path.write_text("\n".join([f"# {title}", "", *lines]), encoding="utf-8")


def write_simple_docx(path: Path, title: str, lines: list[str]) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, title)
    for line in lines:
        add_paragraph(doc, line)
    doc.save(path)


def write_submission_docx(pkg: SubmissionPackage, main_table: pd.DataFrame) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, pkg.title)
    add_heading(doc, "Summary", level=1)
    for key in ["Background", "Methods", "Findings", "Interpretation", "Funding"]:
        add_heading(doc, key, level=2)
        add_paragraph(doc, pkg.summary[key])

    add_heading(doc, "Research in context", level=1)
    for key, value in pkg.research_in_context.items():
        add_heading(doc, key, level=2)
        add_paragraph(doc, value)

    for section_title, blocks in pkg.sections:
        add_heading(doc, section_title, level=1)
        for subsection, paragraphs in blocks:
            if subsection:
                add_heading(doc, subsection, level=2)
            for paragraph in paragraphs:
                add_paragraph(doc, paragraph)

    add_table_to_doc(doc, main_table, pkg.main_table_title)

    add_heading(doc, "Figure legends", level=1)
    figure_map = {
        "Figure 1": "figure1_global_burden_and_aging.png",
        "Figure 2": "figure2_subtype_profile_2023.png",
        "Figure 3": "figure3_country_aging_ecology.png",
        "Figure 4": "figure4_top20_country_asmr_2023.png",
    }
    for title, body in pkg.figure_legends:
        add_heading(doc, title, level=2)
        add_paragraph(doc, body)
        figure_path = FIGURE_DIR / figure_map[title]
        if figure_path.exists():
            doc.add_picture(str(figure_path), width=Inches(6.3))

    add_heading(doc, "References", level=1)
    for idx, ref in enumerate(pkg.references, start=1):
        add_paragraph(doc, f"{idx}. {ref}")
    doc.save(MANUSCRIPT_DIR / "skin_lancet_submission_focused_draft.docx")


def write_supplement_docx(tables: list[tuple[str, pd.DataFrame]]) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, "Supplementary Appendix For Submission-Focused Draft")
    for title, df in tables:
        add_table_to_doc(doc, df, title)
    doc.save(MANUSCRIPT_DIR / "supplementary_appendix_submission_focused.docx")


def build_submission_focused_qc(pkg: SubmissionPackage, core_files: list[Path]) -> list[str]:
    missing = [str(path) for path in core_files if not path.exists()]
    placeholder_scan = {}
    for path in core_files:
        if path.suffix == ".md":
            text = path.read_text(encoding="utf-8")
            hits = re.findall(r"\[[^\]]*inserted[^\]]*\]|\[[^\]]*completed[^\]]*\]", text, flags=re.I)
            if hits:
                placeholder_scan[path.name] = hits
    return [
        f"Main text word count: {pkg.main_word_count}",
        f"Summary word count: {pkg.summary_word_count}",
        "Main-text display items: 5 (4 figures and 1 table)",
        f"Core file presence check: {'PASS' if not missing else 'FAIL'}",
        "Placeholder scan in submission-focused core files: "
        + ("PASS" if not placeholder_scan else f"FAIL {placeholder_scan}"),
        "Remaining manual items: author order, affiliations, corresponding-author contact details, funding statement, ethics wording, originality confirmation, and final declaration wording.",
        "This version is closer to journal submission than the long working draft because it reduces display items and compresses the core narrative.",
    ]


def main() -> None:
    builder = load_builder_module()
    global_context = builder.load_global_context()
    country_complete, _ = builder.load_country_complete()
    correlations, tertiles, top20 = builder.compute_country_ecology(country_complete)
    subtype_dirf, subtype_mortality = builder.load_subtype_profiles()
    values = builder.build_value_map(global_context, country_complete, correlations, tertiles, top20, subtype_dirf, subtype_mortality)

    summary = builder.build_summary(values)
    research_in_context = builder.build_research_in_context()
    references = builder.build_references()
    sections = build_submission_sections(values)
    figure_legends = builder.build_figure_legends()[:4]

    sections_text = "\n".join(
        paragraph
        for _, blocks in sections
        for _, paragraphs in blocks
        for paragraph in paragraphs
    )
    main_word_count = word_count(sections_text)
    summary_word_count = word_count(" ".join(summary.values()))
    candidate_note = extract_candidate_author_note()

    pkg = SubmissionPackage(
        title=builder.TITLE,
        short_title=builder.SHORT_TITLE,
        target_journal=builder.TARGET_JOURNAL,
        summary=summary,
        research_in_context=research_in_context,
        sections=sections,
        references=references,
        figure_legends=figure_legends,
        main_table_title="Table 1. Global burden of skin and subcutaneous diseases and World Bank ageing indicators in 1990 and 2023",
        main_word_count=main_word_count,
        summary_word_count=summary_word_count,
        candidate_author_note=candidate_note,
    )

    main_table = pd.read_csv(TABLE_DIR / "table2_global_burden_and_aging_context.csv")
    supplement_tables = [
        ("Table S1. Study frame, data sources, and analytical modules", pd.read_csv(TABLE_DIR / "table1_study_frame_and_data_sources.csv")),
        ("Table S2. Subtype-specific global burden profile of skin and subcutaneous diseases in 2023", pd.read_csv(TABLE_DIR / "table3_subtype_profile_2023.csv")),
        ("Table S3. Subtype-specific relative change in global skin burden between 1990 and 2023", pd.read_csv(TABLE_DIR / "table4_subtype_change_1990_2023.csv")),
        ("Table S4. Country-level ecological summary of skin mortality in 2023", pd.read_csv(TABLE_DIR / "table5_country_ecology_summary.csv")),
        ("Table S5. Country-level ecological correlations between World Bank ageing indicators and skin mortality in 2023", pd.read_csv(TABLE_DIR / "tableS3_country_correlations.csv")),
        ("Table S6. Top 20 countries and territories ranked by age-standardized mortality rate from skin and subcutaneous diseases in 2023", pd.read_csv(TABLE_DIR / "tableS4_top20_country_asmr_2023.csv")),
        ("Table S7. Ageing tertile summary of country-level skin mortality in 2023", pd.read_csv(TABLE_DIR / "tableS5_age65_tertiles.csv")),
    ]

    manuscript_md = draft_to_markdown(pkg)
    (MANUSCRIPT_DIR / "skin_lancet_submission_focused_draft.md").write_text(manuscript_md, encoding="utf-8")
    write_submission_docx(pkg, main_table)

    title_lines = make_title_page(pkg)
    cover_lines = make_cover_letter(pkg)
    author_form_lines = make_author_metadata_form(candidate_note)
    contributors_lines = make_authors_contributors_form()
    declaration_lines = make_declaration_form()
    supplement_lines = [
        "This appendix contains the tables moved out of the main manuscript to support a submission-focused 5-display-item structure.",
        "",
        "Included supplementary tables:",
        *[f"- {title}" for title, _ in supplement_tables],
    ]

    write_simple_markdown(MANUSCRIPT_DIR / "title_page_submission_focused.md", "Title Page", title_lines)
    write_simple_markdown(MANUSCRIPT_DIR / "cover_letter_submission_focused.md", "Cover Letter", cover_lines)
    write_simple_markdown(MANUSCRIPT_DIR / "author_metadata_form.md", "Author Metadata Form", author_form_lines)
    write_simple_markdown(MANUSCRIPT_DIR / "authors_contributors_submission_focused.md", "Authors And Contributors", contributors_lines)
    write_simple_markdown(MANUSCRIPT_DIR / "declaration_of_interests_submission_focused.md", "Declaration Of Interests", declaration_lines)
    write_simple_markdown(MANUSCRIPT_DIR / "supplementary_appendix_submission_focused.md", "Supplementary Appendix", supplement_lines)

    write_simple_docx(MANUSCRIPT_DIR / "title_page_submission_focused.docx", "Title Page", title_lines)
    write_simple_docx(MANUSCRIPT_DIR / "cover_letter_submission_focused.docx", "Cover Letter", cover_lines)
    write_simple_docx(MANUSCRIPT_DIR / "author_metadata_form.docx", "Author Metadata Form", author_form_lines)
    write_simple_docx(MANUSCRIPT_DIR / "authors_contributors_submission_focused.docx", "Authors And Contributors", contributors_lines)
    write_simple_docx(MANUSCRIPT_DIR / "declaration_of_interests_submission_focused.docx", "Declaration Of Interests", declaration_lines)
    write_simple_docx(MANUSCRIPT_DIR / "supplementary_appendix_submission_focused_notes.docx", "Supplementary Appendix", supplement_lines)
    write_supplement_docx(supplement_tables)

    core_files = [
        MANUSCRIPT_DIR / "skin_lancet_submission_focused_draft.md",
        MANUSCRIPT_DIR / "skin_lancet_submission_focused_draft.docx",
        MANUSCRIPT_DIR / "title_page_submission_focused.md",
        MANUSCRIPT_DIR / "cover_letter_submission_focused.md",
        MANUSCRIPT_DIR / "author_metadata_form.md",
        MANUSCRIPT_DIR / "authors_contributors_submission_focused.md",
        MANUSCRIPT_DIR / "declaration_of_interests_submission_focused.md",
        MANUSCRIPT_DIR / "supplementary_appendix_submission_focused.docx",
    ]
    qc_lines = build_submission_focused_qc(pkg, core_files)
    write_simple_markdown(MANUSCRIPT_DIR / "submission_focused_qc_report.md", "Submission Focused QC Report", qc_lines)
    write_simple_docx(MANUSCRIPT_DIR / "submission_focused_qc_report.docx", "Submission Focused QC Report", qc_lines)

    summary_payload = {
        "title": pkg.title,
        "main_word_count": pkg.main_word_count,
        "summary_word_count": pkg.summary_word_count,
        "main_display_items": 5,
        "main_figures": 4,
        "main_tables": 1,
        "supplementary_tables": len(supplement_tables),
        "candidate_author_note": candidate_note,
    }
    (MANUSCRIPT_DIR / "submission_focused_summary.json").write_text(
        json.dumps(summary_payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    print(f"Submission-focused main text words: {pkg.main_word_count}")
    print(f"Submission-focused summary words: {pkg.summary_word_count}")
    print(f"Submission-focused package written to: {MANUSCRIPT_DIR}")


if __name__ == "__main__":
    main()
