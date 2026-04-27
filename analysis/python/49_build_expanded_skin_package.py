#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches


ROOT = Path("/Users/apple/Documents/lancet-research-platform")
SCRIPT_45 = ROOT / "analysis" / "python" / "45_build_skin_lancet_package.py"
PACKAGE_ROOT = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万 2/lancet_skin_article_package"
)
OUTPUT_DIR = PACKAGE_ROOT / "outputs"
MANUSCRIPT_DIR = OUTPUT_DIR / "manuscript"
FIGURE_DIR = OUTPUT_DIR / "figures"


def load_builder_module():
    spec = importlib.util.spec_from_file_location("skin_builder_long", SCRIPT_45)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules["skin_builder_long"] = module
    spec.loader.exec_module(module)
    return module


def write_expanded_docx(builder, package, main_tables: list[tuple[str, object]], output_path: Path) -> None:
    manuscript_doc = Document()
    builder.configure_doc(manuscript_doc)
    builder.add_title(manuscript_doc, builder.TITLE)

    builder.add_doc_heading(manuscript_doc, "Summary", level=1)
    for key in ["Background", "Methods", "Findings", "Interpretation", "Funding"]:
        builder.add_doc_heading(manuscript_doc, key, level=2)
        builder.add_doc_paragraph(manuscript_doc, package.summary[key])

    builder.add_doc_heading(manuscript_doc, "Research in context", level=1)
    for key, value in package.research_in_context.items():
        builder.add_doc_heading(manuscript_doc, key, level=2)
        builder.add_doc_paragraph(manuscript_doc, value)

    for section_title, blocks in package.sections:
        builder.add_doc_heading(manuscript_doc, section_title, level=1)
        for subsection, paragraphs in blocks:
            if subsection:
                builder.add_doc_heading(manuscript_doc, subsection, level=2)
            for paragraph in paragraphs:
                builder.add_doc_paragraph(manuscript_doc, paragraph)

    for table_title, table_df in main_tables:
        builder.add_table_to_doc(manuscript_doc, table_df, table_title)

    builder.add_doc_heading(manuscript_doc, "Figure legends", level=1)
    figure_map = {
        "Figure 1": "figure1_global_burden_and_aging.png",
        "Figure 2": "figure2_subtype_profile_2023.png",
        "Figure 3": "figure3_country_aging_ecology.png",
        "Figure 4": "figure4_top20_country_asmr_2023.png",
        "Figure 5": "figure5_subtype_trends_1990_2023.png",
    }
    for title, body in package.figure_legends:
        builder.add_doc_heading(manuscript_doc, title, level=2)
        builder.add_doc_paragraph(manuscript_doc, body)
        figure_path = FIGURE_DIR / figure_map[title]
        if figure_path.exists():
            manuscript_doc.add_picture(str(figure_path), width=Inches(6.3))

    builder.add_doc_heading(manuscript_doc, "References", level=1)
    for idx, ref in enumerate(package.selected_references, start=1):
        builder.add_doc_paragraph(manuscript_doc, f"{idx}. {ref}")

    manuscript_doc.save(output_path)


def build_title_page(builder, package) -> list[str]:
    return [
        f"Target journal: {builder.TARGET_JOURNAL}",
        "Working article format: expanded long-form Lancet-family draft",
        f"Full title: {builder.TITLE}",
        f"Short title: {builder.SHORT_TITLE}",
        "Authors: ________________________________",
        "Affiliations: ____________________________",
        "Corresponding author: ____________________",
        "Corresponding email: _____________________",
        f"Main text word count: {package.main_word_count}",
        f"Summary word count: {package.summary_word_count}",
        "Main-text display items: 10 (5 figures and 5 tables)",
        "Supplementary tables: 5",
        "Funding statement: ______________________",
        "Declaration of interests: completed author forms attached separately.",
        "Data sharing: derived analyses use official GBD 2023 and World Bank WDI data sources.",
        "Version note: this is the expanded approximately 5000-word main-text version requested for full narrative development.",
    ]


def build_cover_letter(builder, values: dict[str, float | str | int]) -> list[str]:
    return [
        "Date: March 9, 2026",
        "",
        f"To the Editors of {builder.TARGET_JOURNAL}",
        "",
        f"We are pleased to submit our Article entitled \"{builder.TITLE}\" for consideration at {builder.TARGET_JOURNAL}.",
        "",
        "This expanded long-form version retains a fully developed narrative structure so that the epidemiologic findings, the World Bank ageing framework, and the clinical-policy implications can be presented in sufficient depth for editorial and author review. "
        "The manuscript integrates official GBD 2023 skin-burden estimates with World Bank World Development Indicators and frames skin disease as part of the service burden of healthy ageing rather than as an isolated dermatology topic.",
        "",
        f"In the current analysis, the global proportion of people aged 65 years and older increased from {values['age_1990']:.2f}% in 1990 to {values['age_2023']:.2f}% in 2023. "
        f"Over the same period, the age-standardized incidence rate increased from {values['incidence_1990']:.1f} to {values['incidence_2023']:.1f} per 100,000, DALY counts increased from {values['daly_count_1990'] / 1_000_000:.1f} million to {values['daly_count_2023'] / 1_000_000:.1f} million, and deaths increased from {int(values['death_count_1990']):,} to {int(values['death_count_2023']):,}. "
        f"Higher population ageing remained inversely associated with age-standardized skin mortality at country level, and the highest mortality was observed in {values['top3_text']}.",
        "",
        "We believe the paper will interest readers because it links dermatologic burden to population ageing using a policy-familiar demographic source, distinguishes disability-dominant and mortality-dominant skin subtypes, and identifies high-mortality settings that may warrant stronger capacity in wound care, infection control, and long-term care prevention.",
        "",
        "The manuscript has not been published previously and is not under consideration elsewhere. Final author approval, authorship order, funding statement, ethics wording, originality confirmation, and declaration wording should be completed before submission.",
        "",
        "Sincerely,",
        "",
        "Corresponding author: ____________________",
        "Institution: _____________________________",
        "Email: __________________________________",
        "Telephone: _______________________________",
    ]


def build_qc_lines(package, core_files: list[Path], render_summary: dict[str, object] | None) -> list[str]:
    missing = [str(path) for path in core_files if not path.exists()]
    placeholder_scan: dict[str, list[str]] = {}
    placeholder_pattern = re.compile(
        r"\[To be inserted\]|\[Corresponding author name\]|\[Degrees\]|\[Department and institution\]|\[Postal address\]|\[Email\]|\[Telephone\]",
        flags=re.I,
    )
    for path in core_files:
        if path.suffix == ".md" and path.exists():
            text = path.read_text(encoding="utf-8")
            hits = placeholder_pattern.findall(text)
            if hits:
                placeholder_scan[path.name] = hits

    render_lines = []
    if isinstance(render_summary, dict) and render_summary:
        if render_summary.get("available"):
            render_lines.append(
                f"- Render pipeline available: yes; documents rendered={render_summary.get('rendered_count', 'unknown')}/{render_summary.get('document_count', 'unknown')}; total pages={render_summary.get('page_count', 'unknown')}"
            )
            for payload in render_summary.get("documents", []):
                if isinstance(payload, dict):
                    doc_name = Path(str(payload.get("docx", "unknown"))).name
                    status = "PASS" if payload.get("ok") else "FAIL"
                    page_count = payload.get("page_count", "unknown")
                    render_lines.append(f"- {doc_name}: {status}, pages={page_count}")
        else:
            render_lines.append("- Render pipeline unavailable in the local environment.")

    target_ok = 4800 <= package.main_word_count <= 5400
    lines = [
        f"Main text word count: {package.main_word_count}",
        f"Summary word count: {package.summary_word_count}",
        "Main-text display items: 10 (5 figures and 5 tables)",
        f"Target range around 5000 words: {'PASS' if target_ok else 'FAIL'}",
        f"Core file presence check: {'PASS' if not missing else 'FAIL'}",
        "Placeholder scan in expanded-version core files: "
        + ("PASS" if not placeholder_scan else f"FAIL {placeholder_scan}"),
        "Rendered page QA:",
    ]
    if render_lines:
        lines.extend(render_lines)
    else:
        lines.append("- Render summary unavailable.")
    lines.extend(
        [
            "Remaining manual items: author order, affiliations, corresponding-author contact details, funding statement, ethics wording, originality confirmation, and final declaration wording.",
            "This version restores the full long-form narrative while preserving the cleaned author-side metadata files generated in the submission-focused package.",
        ]
    )
    return lines


def main() -> None:
    builder = load_builder_module()
    builder.ensure_dirs()
    builder.configure_matplotlib()

    global_context = builder.load_global_context()
    country_complete, ambiguous_names = builder.load_country_complete()
    correlations, tertiles, top20 = builder.compute_country_ecology(country_complete)
    subtype_dirf, subtype_mortality = builder.load_subtype_profiles()

    table1 = builder.build_table1_study_frame()
    table2 = builder.build_main_table(global_context)
    result_tables = builder.build_main_result_tables(subtype_dirf, subtype_mortality, correlations, tertiles, top20)
    builder.make_figure1(global_context)
    builder.make_figure2(subtype_dirf, subtype_mortality)
    builder.make_figure3(country_complete, correlations)
    builder.make_figure4(top20)
    builder.make_figure5(subtype_dirf, subtype_mortality)

    values = builder.build_value_map(
        global_context,
        country_complete,
        correlations,
        tertiles,
        top20,
        subtype_dirf,
        subtype_mortality,
    )
    package = builder.build_draft_package(values)

    main_tables = [
        ("Table 1. Study frame, data sources, and analytical modules", table1),
        ("Table 2. Global burden of skin and subcutaneous diseases and World Bank ageing indicators in 1990 and 2023", table2),
        ("Table 3. Subtype-specific global burden profile of skin and subcutaneous diseases in 2023", result_tables["subtype_2023"]),
        ("Table 4. Subtype-specific relative change in global skin burden between 1990 and 2023", result_tables["subtype_change"]),
        ("Table 5. Country-level ecological summary of skin mortality in 2023", result_tables["ecology_table"]),
    ]

    manuscript_md_path = MANUSCRIPT_DIR / "skin_lancet_expanded_5000w_draft.md"
    manuscript_docx_path = MANUSCRIPT_DIR / "skin_lancet_expanded_5000w_draft.docx"
    title_md_path = MANUSCRIPT_DIR / "title_page_expanded_5000w.md"
    title_docx_path = MANUSCRIPT_DIR / "title_page_expanded_5000w.docx"
    cover_md_path = MANUSCRIPT_DIR / "cover_letter_expanded_5000w.md"
    cover_docx_path = MANUSCRIPT_DIR / "cover_letter_expanded_5000w.docx"
    qc_md_path = MANUSCRIPT_DIR / "expanded_5000w_qc_report.md"
    qc_docx_path = MANUSCRIPT_DIR / "expanded_5000w_qc_report.docx"
    summary_json_path = MANUSCRIPT_DIR / "expanded_5000w_summary.json"

    manuscript_md_path.write_text(builder.draft_to_markdown(package), encoding="utf-8")
    write_expanded_docx(builder, package, main_tables, manuscript_docx_path)

    title_lines = build_title_page(builder, package)
    cover_lines = build_cover_letter(builder, values)
    builder.write_simple_markdown(title_md_path, "Title Page", title_lines)
    builder.write_simple_markdown(cover_md_path, "Cover Letter", cover_lines)
    builder.write_simple_docx(title_docx_path, "Title Page", title_lines)
    builder.write_simple_docx(cover_docx_path, "Cover Letter", cover_lines)

    render_summary = builder.render_docx_collection(
        [manuscript_docx_path, title_docx_path, cover_docx_path],
        MANUSCRIPT_DIR / "rendered_pages_expanded_5000w",
    )

    core_files = [
        manuscript_md_path,
        manuscript_docx_path,
        title_md_path,
        title_docx_path,
        cover_md_path,
        cover_docx_path,
        MANUSCRIPT_DIR / "author_metadata_form.docx",
        MANUSCRIPT_DIR / "authors_contributors_submission_focused.docx",
        MANUSCRIPT_DIR / "declaration_of_interests_submission_focused.docx",
    ]
    qc_lines = build_qc_lines(package, core_files, render_summary)
    builder.write_simple_markdown(qc_md_path, "Expanded 5000w QC Report", qc_lines)
    builder.write_simple_docx(qc_docx_path, "Expanded 5000w QC Report", qc_lines)

    summary_payload = {
        "title": builder.TITLE,
        "main_word_count": package.main_word_count,
        "summary_word_count": package.summary_word_count,
        "main_figures": 5,
        "main_tables": 5,
        "supplementary_tables": 5,
        "countries_in_ecology": int(values["n_countries"]),
        "ambiguous_country_names_excluded": ambiguous_names,
        "render_summary": render_summary,
    }
    summary_json_path.write_text(
        json.dumps(summary_payload, indent=2, ensure_ascii=False),
        encoding="utf-8",
    )

    print(f"Expanded main text words: {package.main_word_count}")
    print(f"Expanded summary words: {package.summary_word_count}")
    print(f"Expanded package written to: {MANUSCRIPT_DIR}")


if __name__ == "__main__":
    main()
