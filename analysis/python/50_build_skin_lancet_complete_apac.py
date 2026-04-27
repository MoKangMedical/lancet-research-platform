#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import json
import re
import sys
from dataclasses import dataclass
from pathlib import Path

import geopandas as gpd
import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import requests
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from shapely.ops import transform as shapely_transform


ROOT = Path("/Users/apple/Desktop/lancet-research-platform")
ANALYSIS_PY = ROOT / "analysis" / "python"
if str(ANALYSIS_PY) not in sys.path:
    sys.path.insert(0, str(ANALYSIS_PY))

from lib.rendering import render_docx_collection


SCRIPT_45 = ANALYSIS_PY / "45_build_skin_lancet_package.py"
PROJECT_DIR = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万"
)
PACKAGE_ROOT = PROJECT_DIR / "lancet_skin_article_package"
OUTPUT_DIR = PACKAGE_ROOT / "outputs"
FIGURE_DIR = OUTPUT_DIR / "figures"
TABLE_DIR = OUTPUT_DIR / "tables"
MANUSCRIPT_DIR = OUTPUT_DIR / "manuscript"
AGING_DIR = PACKAGE_ROOT / "aging_analysis_outputs"
APAC_RESULTS_DIR = PACKAGE_ROOT / "apac_results_tool_outputs"
OFFICIAL_APAC_ASR_PATH = APAC_RESULTS_DIR / "skin_apac_official_asr_2023.csv"
NATURAL_EARTH_PATH = Path(
    "/Users/apple/Documents/.venvs/data-analytics/lib/python3.14/site-packages/pyogrio/tests/fixtures/naturalearth_lowres/naturalearth_lowres.shp"
)

TITLE = (
    "Global burden of skin and subcutaneous diseases in the context of population ageing, "
    "1990-2023: a systematic analysis of GBD 2023 and World Bank ageing indicators"
)
SHORT_TITLE = "Skin burden and population ageing"
TARGET_JOURNAL = "The Lancet Healthy Longevity"

COLOR_NAVY = "#003A70"
COLOR_STEEL = "#7F9DB9"
COLOR_PALE_BLUE = "#DCE6F2"
COLOR_SLATE = "#5A6B7A"
COLOR_RUST = "#B55D4C"
COLOR_TEAL = "#2E6F62"
COLOR_GOLD = "#C08B30"
COLOR_GREY = "#ECEFF3"
COLOR_ROW = "#F7F9FB"
COLOR_GRID = "#C8D0D9"
COLOR_TEXT = "#1F2933"

WORLD_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "lancet_blue", ["#F6F8FB", "#B4C8DB", COLOR_NAVY]
)
AGE_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "lancet_green", ["#F6FAF8", "#A8C7BB", COLOR_TEAL]
)
LIFE_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "lancet_rust", ["#FCF7F2", "#D9A98B", COLOR_RUST]
)
PREVALENCE_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "lancet_prev", ["#FFF8F3", "#E1B99D", COLOR_RUST]
)
INCIDENCE_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "lancet_inc", ["#F6FAF8", "#A8C7BB", COLOR_TEAL]
)


@dataclass
class PackageArtifacts:
    main_docx: Path
    main_md: Path
    title_docx: Path
    title_md: Path
    cover_docx: Path
    cover_md: Path
    supp_docx: Path
    supp_md: Path
    qc_docx: Path
    qc_md: Path
    summary_json: Path


def load_builder_module():
    spec = importlib.util.spec_from_file_location("skin_builder_complete", SCRIPT_45)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules["skin_builder_complete"] = module
    spec.loader.exec_module(module)
    module.ROOT = ROOT
    module.DIRF_PATH = ROOT / "data" / "silver" / "gbd" / "gbd2023_dirf_global_core_tidy.csv"
    module.MORTALITY_PATH = ROOT / "data" / "silver" / "gbd" / "gbd2023_mortality_s7_both_sex_long.csv"
    module.COLOR_MAIN = COLOR_NAVY
    module.COLOR_ACCENT = COLOR_RUST
    module.COLOR_GREEN = COLOR_TEAL
    module.COLOR_GOLD = COLOR_GOLD
    module.COLOR_PURPLE = "#6C728A"
    module.COLOR_GREY = COLOR_SLATE
    module.PANEL_COLORS = [
        module.COLOR_MAIN,
        module.COLOR_ACCENT,
        module.COLOR_GREEN,
        module.COLOR_GOLD,
        module.COLOR_PURPLE,
        module.COLOR_GREY,
    ]
    module.PROJECT_DIR = PROJECT_DIR
    module.OUTPUT_DIR = PACKAGE_ROOT
    module.FIGURE_DIR = FIGURE_DIR
    module.TABLE_DIR = TABLE_DIR
    module.MANUSCRIPT_DIR = MANUSCRIPT_DIR
    module.GLOBAL_CONTEXT_CANDIDATES = [
        PROJECT_DIR / "aging_analysis_outputs" / "skin_aging_global_context_1990_2023.csv",
        PACKAGE_ROOT / "aging_analysis_outputs" / "skin_aging_global_context_1990_2023.csv",
        PACKAGE_ROOT / "aging_analysis_outputs_test" / "skin_aging_global_context_1990_2023.csv",
    ]
    module.COUNTRY_COMPLETE_CANDIDATES = [
        PROJECT_DIR / "aging_analysis_outputs" / "skin_aging_2023_country_complete.csv",
        PACKAGE_ROOT / "aging_analysis_outputs" / "skin_aging_2023_country_complete.csv",
        PACKAGE_ROOT / "aging_analysis_outputs_test" / "skin_aging_2023_country_complete.csv",
    ]
    return module


def fetch_apac_reference() -> pd.DataFrame:
    rows: list[dict[str, object]] = []
    for region_code, region_label in [("EAS", "East Asia & Pacific"), ("SAS", "South Asia")]:
        url = f"https://api.worldbank.org/v2/region/{region_code}/country?format=json&per_page=500"
        payload = requests.get(url, timeout=30).json()[1]
        for item in payload:
            rows.append(
                {
                    "wb_iso3": item["id"],
                    "apac_region_code": region_code,
                    "apac_region_label": region_label,
                    "apac_name": item["name"],
                    "longitude": pd.to_numeric(item.get("longitude"), errors="coerce"),
                    "latitude": pd.to_numeric(item.get("latitude"), errors="coerce"),
                    "capital_city": item.get("capitalCity"),
                }
            )
    return pd.DataFrame(rows).drop_duplicates(subset=["wb_iso3"]).reset_index(drop=True)


def load_world_shapes() -> gpd.GeoDataFrame:
    return gpd.read_file(NATURAL_EARTH_PATH)[["name", "iso_a3", "geometry"]].copy()


def shift_negative_longitude(geometry):
    def _shift(x, y, z=None):
        if np.isscalar(x):
            x_new = x + 360 if x < 0 else x
            return (x_new, y) if z is None else (x_new, y, z)
        x_arr = np.asarray(x)
        x_new = np.where(x_arr < 0, x_arr + 360, x_arr)
        return (x_new, y) if z is None else (x_new, y, z)

    return shapely_transform(_shift, geometry)


def build_apac_datasets(country_complete: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    apac_ref = fetch_apac_reference()
    apac = country_complete.merge(apac_ref, on="wb_iso3", how="inner").copy()
    official = load_official_apac_rates()
    apac = apac.merge(official, on=["location_id", "wb_iso3"], how="left", validate="one_to_one")
    if apac["deaths_asr_2023"].isna().any():
        missing = apac.loc[apac["deaths_asr_2023"].isna(), "gbd_name"].tolist()
        raise RuntimeError(
            "Official APAC GBD Results Tool export is incomplete. Missing deaths ASR for: "
            + ", ".join(missing)
        )
    apac.loc[:, "asmr_2023"] = apac["deaths_asr_2023"]
    apac.loc[:, "asmr_2023_lower"] = apac["deaths_asr_2023_lower"]
    apac.loc[:, "asmr_2023_upper"] = apac["deaths_asr_2023_upper"]
    world = load_world_shapes().copy()
    polygon_iso = set(world["iso_a3"])
    apac.loc[:, "plot_method"] = np.where(apac["wb_iso3"].isin(polygon_iso), "polygon", "point")
    apac = apac.sort_values("asmr_2023", ascending=False).reset_index(drop=True)
    return apac, apac_ref


def load_official_apac_rates() -> pd.DataFrame:
    if not OFFICIAL_APAC_ASR_PATH.exists():
        raise RuntimeError(
            f"Missing official APAC results file: {OFFICIAL_APAC_ASR_PATH}. "
            "Run 52_fetch_skin_apac_results_tool.py first."
        )
    df = pd.read_csv(OFFICIAL_APAC_ASR_PATH)
    wide = (
        df.pivot_table(
            index=["location_id", "wb_iso3"],
            columns="measure_name",
            values=["val", "lower", "upper"],
            aggfunc="first",
        )
        .sort_index(axis=1)
        .reset_index()
    )
    renamed_columns = []
    for col in wide.columns:
        if col == ("location_id", ""):
            renamed_columns.append("location_id")
        elif col == ("wb_iso3", ""):
            renamed_columns.append("wb_iso3")
        else:
            stat, measure = col
            renamed_columns.append(
                f"{measure.lower()}_asr_2023" if stat == "val" else f"{measure.lower()}_asr_2023_{stat}"
            )
    wide.columns = renamed_columns
    return wide


def top_rate_text(apac: pd.DataFrame, value_col: str, n: int = 3) -> str:
    subset = apac.sort_values(value_col, ascending=False).head(n)
    return "; ".join(f"{row.gbd_name} ({getattr(row, value_col):.2f})" for row in subset.itertuples())


def add_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_style(cell, *, bold: bool = False, color: str = COLOR_TEXT, size: float = 9.0) -> None:
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.bold = bold
            run.font.size = Pt(size)
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            run.font.color.rgb = RGBColor.from_string(color.replace("#", ""))


def style_lancet_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                add_shading(cell, COLOR_PALE_BLUE.replace("#", ""))
                set_cell_text_style(cell, bold=True, color=COLOR_NAVY, size=8.8)
            else:
                if row_idx % 2 == 0:
                    add_shading(cell, COLOR_ROW.replace("#", ""))
                set_cell_text_style(cell, size=8.5)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    style.font.size = Pt(10.5)


def add_title(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(8)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.color.rgb = RGBColor.from_string(COLOR_NAVY.replace("#", ""))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.color.rgb = RGBColor.from_string(COLOR_NAVY.replace("#", ""))


def add_paragraph(doc: Document, text: str, *, size: float = 10.5) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.22
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(COLOR_TEXT.replace("#", ""))


def add_table_caption(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(COLOR_NAVY.replace("#", ""))


def add_figure_caption(doc: Document, title: str, legend: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(COLOR_NAVY.replace("#", ""))
    add_paragraph(doc, legend, size=9.8)


def add_dataframe_table(doc: Document, title: str, df: pd.DataFrame) -> None:
    add_table_caption(doc, title)
    table = doc.add_table(rows=1, cols=len(df.columns))
    headers = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        headers[idx].text = str(col)
    for row in df.itertuples(index=False):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            if isinstance(value, float):
                if abs(value) >= 1000:
                    cells[idx].text = f"{value:,.1f}"
                else:
                    cells[idx].text = f"{value:.2f}"
            else:
                cells[idx].text = str(value)
    style_lancet_table(table)


def save_figure(fig: plt.Figure, stem: str) -> tuple[Path, Path]:
    png_path = FIGURE_DIR / f"{stem}.png"
    pdf_path = FIGURE_DIR / f"{stem}.pdf"
    fig.tight_layout()
    fig.savefig(png_path, dpi=300, bbox_inches="tight")
    fig.savefig(pdf_path, dpi=300, bbox_inches="tight")
    plt.close(fig)
    return png_path, pdf_path


def configure_matplotlib() -> None:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.labelcolor": COLOR_TEXT,
            "xtick.color": COLOR_TEXT,
            "ytick.color": COLOR_TEXT,
            "axes.edgecolor": COLOR_GRID,
            "axes.titleweight": "bold",
            "figure.dpi": 160,
            "savefig.dpi": 300,
        }
    )


def _draw_apac_map_panel(
    ax,
    apac: pd.DataFrame,
    value_col: str,
    cmap,
    title: str,
    label_col: str | None = None,
):
    world = load_world_shapes()
    polygons = apac[apac["plot_method"] == "polygon"][["wb_iso3", value_col]].copy()
    points = apac[apac["plot_method"] == "point"].copy()
    points.loc[:, "plot_lon"] = np.where(points["longitude"] < 0, points["longitude"] + 360, points["longitude"])

    poly_gdf = world.merge(polygons, left_on="iso_a3", right_on="wb_iso3", how="inner")
    poly_gdf.loc[:, "geometry"] = poly_gdf["geometry"].apply(shift_negative_longitude)
    norm = mpl.colors.Normalize(vmin=float(apac[value_col].min()), vmax=float(apac[value_col].max()))
    poly_gdf.plot(ax=ax, color="#F5F6F8", edgecolor="#D0D6DD", linewidth=0.5)
    poly_gdf.plot(ax=ax, column=value_col, cmap=cmap, norm=norm, edgecolor="white", linewidth=0.35)
    ax.scatter(
        points["plot_lon"],
        points["latitude"],
        s=38 + points[value_col].rank(pct=True) * 55,
        c=points[value_col],
        cmap=cmap,
        norm=norm,
        edgecolors="white",
        linewidths=0.45,
        zorder=5,
    )
    if label_col:
        label_rows = apac.sort_values(value_col, ascending=False).head(6).copy()
        label_rows.loc[:, "plot_lon"] = np.where(
            label_rows["longitude"] < 0, label_rows["longitude"] + 360, label_rows["longitude"]
        )
        for row in label_rows.itertuples(index=False):
            ax.text(
                row.plot_lon + 1.3,
                row.latitude + 0.4,
                getattr(row, label_col),
                fontsize=6.8,
                color=COLOR_TEXT,
                zorder=6,
            )
    ax.set_xlim(60, 205)
    ax.set_ylim(-28, 38)
    ax.set_xticks([])
    ax.set_yticks([])
    ax.set_title(title, loc="left", color=COLOR_NAVY, fontsize=10.5)
    for spine in ax.spines.values():
        spine.set_visible(True)
        spine.set_color(COLOR_GRID)
        spine.set_linewidth(0.8)
    return mpl.cm.ScalarMappable(norm=norm, cmap=cmap)


def make_world_apac_mortality_map(country_complete: pd.DataFrame, apac: pd.DataFrame) -> None:
    world = load_world_shapes().copy()
    merged_world = world.merge(
        country_complete[["wb_iso3", "asmr_2023"]],
        left_on="iso_a3",
        right_on="wb_iso3",
        how="left",
    )
    apac_polygons = merged_world[merged_world["iso_a3"].isin(apac[apac["plot_method"] == "polygon"]["wb_iso3"])].copy()
    apac_polygons.loc[:, "geometry"] = apac_polygons["geometry"].apply(shift_negative_longitude)
    point_df = apac[apac["plot_method"] == "point"].copy()
    point_df.loc[:, "plot_lon"] = np.where(
        point_df["longitude"] < 0, point_df["longitude"] + 360, point_df["longitude"]
    )

    vmax = max(15, float(country_complete["asmr_2023"].quantile(0.98)))
    norm = mpl.colors.Normalize(vmin=0, vmax=vmax)

    fig = plt.figure(figsize=(14.2, 10.4))
    gs = fig.add_gridspec(2, 2, width_ratios=[1.12, 1.0], height_ratios=[1.0, 1.0], wspace=0.12, hspace=0.18)
    ax_world = fig.add_subplot(gs[0, 0])
    ax_death = fig.add_subplot(gs[0, 1])
    ax_prev = fig.add_subplot(gs[1, 0])
    ax_inc = fig.add_subplot(gs[1, 1])

    for ax in [ax_world, ax_death, ax_prev, ax_inc]:
        ax.set_facecolor("white")

    merged_world.plot(ax=ax_world, color="#F3F5F7", edgecolor="#C8D0D9", linewidth=0.4)
    merged_world.dropna(subset=["asmr_2023"]).plot(
        ax=ax_world, column="asmr_2023", cmap=WORLD_CMAP, norm=norm, edgecolor="white", linewidth=0.25
    )
    ax_world.set_title("A  Global mortality ASR, 2023", loc="left", color=COLOR_NAVY, fontsize=10.5)
    ax_world.set_axis_off()

    death_sm = _draw_apac_map_panel(
        ax_death,
        apac,
        "asmr_2023",
        WORLD_CMAP,
        "B  Asia-Pacific mortality ASR, 2023",
        "gbd_short_name",
    )
    prev_sm = _draw_apac_map_panel(
        ax_prev,
        apac,
        "prevalence_asr_2023",
        PREVALENCE_CMAP,
        "C  Asia-Pacific prevalence ASR, 2023",
    )
    inc_sm = _draw_apac_map_panel(
        ax_inc,
        apac,
        "incidence_asr_2023",
        INCIDENCE_CMAP,
        "D  Asia-Pacific incidence ASR, 2023",
    )

    top_labels = apac.sort_values("asmr_2023", ascending=False).head(8).copy()
    top_labels.loc[:, "plot_lon"] = np.where(
        top_labels["longitude"] < 0, top_labels["longitude"] + 360, top_labels["longitude"]
    )
    projected_centroids = gpd.GeoSeries(
        apac_polygons.to_crs(3857).geometry.centroid,
        crs=3857,
    ).to_crs(apac_polygons.crs)
    poly_centroids = pd.DataFrame(
        {
            "wb_iso3": apac_polygons["iso_a3"].to_numpy(),
            "cx": projected_centroids.x.to_numpy(),
            "cy": projected_centroids.y.to_numpy(),
        }
    )
    top_labels = top_labels.merge(poly_centroids, on="wb_iso3", how="left")
    top_labels.loc[:, "label_x"] = top_labels["cx"].fillna(top_labels["plot_lon"])
    top_labels.loc[:, "label_y"] = top_labels["cy"].fillna(top_labels["latitude"])
    label_offsets = {
        "ASM": (2.2, 0.4),
        "KIR": (2.0, 1.0),
        "MHL": (2.2, 0.1),
        "FSM": (1.6, 0.6),
        "MYS": (1.2, 0.3),
        "MNP": (1.0, 0.2),
        "BGD": (1.1, 0.6),
    }
    for row in top_labels.itertuples(index=False):
        dx, dy = label_offsets.get(row.wb_iso3, (1.2, 0.4))
        ax_death.text(
            row.label_x + dx,
            row.label_y + dy,
            row.gbd_short_name,
            fontsize=7.2,
            color=COLOR_TEXT,
            zorder=6,
        )

    world_sm = mpl.cm.ScalarMappable(norm=norm, cmap=WORLD_CMAP)
    for sm, ax, label in [
        (world_sm, ax_world, "Mortality ASR per 100,000"),
        (death_sm, ax_death, "Mortality ASR per 100,000"),
        (prev_sm, ax_prev, "Prevalence ASR per 100,000"),
        (inc_sm, ax_inc, "Incidence ASR per 100,000"),
    ]:
        cbar = fig.colorbar(sm, ax=ax, fraction=0.03, pad=0.015)
        cbar.set_label(label, color=COLOR_TEXT, fontsize=8.5)
        cbar.outline.set_edgecolor(COLOR_GRID)
        cbar.ax.tick_params(labelsize=7.5)
    fig.subplots_adjust(left=0.03, right=0.96, top=0.92, bottom=0.05)

    fig.suptitle(
        "Figure 4. Global mortality and Asia-Pacific distribution of deaths, prevalence, and incidence in 2023",
        y=0.985,
        fontsize=13.2,
        color=COLOR_NAVY,
        fontweight="bold",
    )
    save_figure(fig, "figure4_global_apac_mortality_map")


def make_top20_bar(top20: pd.DataFrame) -> None:
    data = top20.sort_values("asmr_2023", ascending=True).copy()
    fig, ax = plt.subplots(figsize=(9.5, 7.8))
    lower = data["asmr_2023"] - data["asmr_2023_lower"]
    upper = data["asmr_2023_upper"] - data["asmr_2023"]
    ax.barh(data["gbd_short_name"], data["asmr_2023"], color=COLOR_RUST, alpha=0.9)
    ax.errorbar(
        data["asmr_2023"],
        data["gbd_short_name"],
        xerr=[lower, upper],
        fmt="none",
        ecolor=COLOR_SLATE,
        elinewidth=1,
        capsize=2,
    )
    ax.set_xlabel("Age-standardized mortality rate per 100,000", color=COLOR_TEXT)
    ax.set_title("Figure S1. Top 20 countries and territories by skin mortality in 2023", color=COLOR_NAVY)
    save_figure(fig, "figureS1_top20_country_asmr_2023")


def make_apac_top15_bar(apac: pd.DataFrame) -> None:
    make_apac_top15_rate_bar(
        apac,
        "asmr_2023",
        "Figure S2. Highest mortality settings in Asia-Pacific, 2023",
        "Age-standardized mortality rate per 100,000",
        "figureS2_apac_top15_asmr_2023",
        COLOR_NAVY,
        lower_col="asmr_2023_lower",
        upper_col="asmr_2023_upper",
    )


def make_apac_top15_rate_bar(
    apac: pd.DataFrame,
    value_col: str,
    title: str,
    x_label: str,
    stem: str,
    color: str,
    lower_col: str | None = None,
    upper_col: str | None = None,
) -> None:
    data = apac.sort_values(value_col, ascending=True).tail(15).copy()
    fig, ax = plt.subplots(figsize=(9.2, 6.8))
    ax.barh(data["gbd_short_name"], data[value_col], color=color, alpha=0.9)
    if lower_col and upper_col:
        lower = data[value_col] - data[lower_col]
        upper = data[upper_col] - data[value_col]
        ax.errorbar(
            data[value_col],
            data["gbd_short_name"],
            xerr=[lower, upper],
            fmt="none",
            ecolor=COLOR_SLATE,
            elinewidth=1,
            capsize=2,
        )
    ax.set_xlabel(x_label, color=COLOR_TEXT)
    ax.set_title(title, color=COLOR_NAVY)
    save_figure(fig, stem)


def make_apac_scatter(apac: pd.DataFrame) -> None:
    fig, ax = plt.subplots(figsize=(7.4, 5.4))
    region_palette = {"East Asia & Pacific": COLOR_NAVY, "South Asia": COLOR_GOLD}
    for region, subset in apac.groupby("apac_region_label"):
        ax.scatter(
            subset["age65_pct"],
            subset["asmr_2023"],
            s=50,
            alpha=0.8,
            color=region_palette.get(region, COLOR_SLATE),
            edgecolor="white",
            linewidth=0.4,
            label=region,
        )
    top_labels = apac.sort_values("asmr_2023", ascending=False).head(8)
    for row in top_labels.itertuples(index=False):
        ax.text(row.age65_pct + 0.3, row.asmr_2023 + 0.15, row.gbd_short_name, fontsize=7.5, color=COLOR_TEXT)
    ax.set_xlabel("Population aged 65 years and older, %")
    ax.set_ylabel("Age-standardized mortality rate per 100,000")
    ax.set_title("Figure S5. Asia-Pacific mortality and ageing structure", color=COLOR_NAVY)
    ax.legend(frameon=False, fontsize=8)
    save_figure(fig, "figureS5_apac_age65_scatter")


def _plot_apac_map(
    apac: pd.DataFrame,
    value_col: str,
    title: str,
    cmap,
    stem: str,
    legend_label: str,
) -> None:
    fig, ax = plt.subplots(figsize=(10.4, 5.8))
    sm = _draw_apac_map_panel(ax, apac, value_col, cmap, title)
    cbar = fig.colorbar(sm, ax=ax, fraction=0.03, pad=0.015)
    cbar.set_label(legend_label, color=COLOR_TEXT)
    cbar.outline.set_edgecolor(COLOR_GRID)
    cbar.ax.tick_params(labelsize=8)
    save_figure(fig, stem)


def main_doc_markdown(
    summary: dict[str, str],
    research_in_context: dict[str, str],
    sections: list[tuple[str, list[tuple[str | None, list[str]]]]],
    references: list[str],
    table_titles: list[str],
    figure_entries: list[tuple[str, str]],
) -> str:
    lines = [f"# {TITLE}", "", "## Summary", ""]
    for key in ["Background", "Methods", "Findings", "Interpretation", "Funding"]:
        lines.extend([f"### {key}", "", summary[key], ""])
    lines.extend(["## Research in context", ""])
    for heading, text in research_in_context.items():
        lines.extend([f"### {heading}", "", text, ""])
    for section_title, blocks in sections:
        lines.extend([f"## {section_title}", ""])
        for subsection, paragraphs in blocks:
            if subsection:
                lines.extend([f"### {subsection}", ""])
            for paragraph in paragraphs:
                lines.extend([paragraph, ""])
    lines.extend(["## References", ""])
    for idx, ref in enumerate(references, start=1):
        lines.append(f"{idx}. {ref}")
    lines.extend(["", "## Tables", ""])
    for title in table_titles:
        lines.append(f"- {title}")
    lines.extend(["", "## Figure Legends And Figures", ""])
    for title, legend in figure_entries:
        lines.extend([f"### {title}", "", legend, ""])
    return "\n".join(lines)


def build_title_page(main_word_count: int, summary_word_count: int) -> list[str]:
    return [
        f"Target journal: {TARGET_JOURNAL}",
        "Working article format: enhanced long-form Lancet-family draft",
        f"Full title: {TITLE}",
        f"Short title: {SHORT_TITLE}",
        "Authors: ________________________________",
        "Affiliations: ____________________________",
        "Corresponding author: ____________________",
        "Corresponding email: _____________________",
        f"Main text word count: {main_word_count}",
        f"Summary word count: {summary_word_count}",
        "Main-text display items: 10 (5 figures and 5 tables)",
        "Supplementary figures: 7",
        "Supplementary tables: 11",
        "Funding statement: ______________________",
        "Declaration of interests: completed author forms attached separately.",
        "Data sharing: derived analyses use official GBD 2023 and World Bank WDI data sources.",
        "Formatting note: all tables and figures are positioned after the reference list in the review-copy manuscript.",
    ]


def build_cover_letter(values: dict[str, float | str | int], apac: pd.DataFrame) -> list[str]:
    top_death_text = top_rate_text(apac, "asmr_2023", 3)
    top_prev_text = top_rate_text(apac, "prevalence_asr_2023", 3)
    top_inc_text = top_rate_text(apac, "incidence_asr_2023", 3)
    return [
        "Date: March 9, 2026",
        "",
        f"To the Editors of {TARGET_JOURNAL}",
        "",
        f"We are pleased to submit our Article entitled \"{TITLE}\" for consideration at {TARGET_JOURNAL}.",
        "",
        "This enhanced long-form version presents a structured Lancet-style analysis of the global burden of skin and subcutaneous diseases in the context of demographic ageing, with additional geographic emphasis on Asia-Pacific settings. "
        "The manuscript integrates official GBD 2023 burden estimates with World Bank World Development Indicators and positions skin disease as part of the service burden of healthy ageing rather than as an isolated dermatology topic.",
        "",
        f"In the current analysis, the global proportion of people aged 65 years and older increased from {values['age_1990']:.2f}% in 1990 to {values['age_2023']:.2f}% in 2023. "
        f"Over the same interval, the age-standardized incidence rate increased from {values['incidence_1990']:.1f} to {values['incidence_2023']:.1f} per 100,000, DALY counts increased from {values['daly_count_1990'] / 1_000_000:.1f} million to {values['daly_count_2023'] / 1_000_000:.1f} million, and deaths increased from {int(values['death_count_1990']):,} to {int(values['death_count_2023']):,}. "
        f"Within Asia-Pacific settings defined by World Bank East Asia & Pacific and South Asia regions, the highest age-standardized mortality rates were observed in {top_death_text}; the highest prevalence rates in {top_prev_text}; and the highest incidence rates in {top_inc_text}.",
        "",
        "We believe the manuscript will interest readers because it links dermatologic burden to demographic transition using a policy-familiar demographic source, distinguishes disability-dominant and mortality-dominant skin subtypes, and adds a geographically explicit Asia-Pacific extension based on official authenticated GBD Results Tool exports for deaths, prevalence, and incidence while remaining fully reproducible.",
        "",
        "The manuscript has not been published previously and is not under consideration elsewhere. Final author approval, authorship order, funding statement, ethics wording, originality confirmation, and declaration wording should be completed before submission.",
        "",
        "Sincerely,",
        "",
        "Corresponding author: ____________________",
        "Institution: _____________________________",
        "Email: __________________________________",
        "Telephone: _______________________________",
    ]


def build_apac_tables(apac: pd.DataFrame) -> dict[str, pd.DataFrame]:
    apac_table = apac[
        [
            "gbd_name",
            "wb_iso3",
            "apac_region_label",
            "asmr_2023",
            "prevalence_asr_2023",
            "incidence_asr_2023",
            "age65_pct",
            "life_expectancy",
            "old_age_dependency",
            "plot_method",
        ]
    ].rename(
        columns={
            "gbd_name": "Location",
            "wb_iso3": "ISO3",
            "apac_region_label": "APAC region",
            "asmr_2023": "Mortality ASR 2023",
            "prevalence_asr_2023": "Prevalence ASR 2023",
            "incidence_asr_2023": "Incidence ASR 2023",
            "age65_pct": "Population aged 65+, %",
            "life_expectancy": "Life expectancy, years",
            "old_age_dependency": "Old-age dependency ratio",
            "plot_method": "Map plotting method",
        }
    )
    top15 = apac.nlargest(15, "asmr_2023")[
        ["gbd_name", "wb_iso3", "apac_region_label", "asmr_2023", "age65_pct", "life_expectancy"]
    ].rename(
        columns={
            "gbd_name": "Location",
            "wb_iso3": "ISO3",
            "apac_region_label": "APAC region",
            "asmr_2023": "Mortality ASR 2023",
            "age65_pct": "Population aged 65+, %",
            "life_expectancy": "Life expectancy, years",
        }
    )
    top15_prev = apac.nlargest(15, "prevalence_asr_2023")[
        ["gbd_name", "wb_iso3", "apac_region_label", "prevalence_asr_2023", "age65_pct", "life_expectancy"]
    ].rename(
        columns={
            "gbd_name": "Location",
            "wb_iso3": "ISO3",
            "apac_region_label": "APAC region",
            "prevalence_asr_2023": "Prevalence ASR 2023",
            "age65_pct": "Population aged 65+, %",
            "life_expectancy": "Life expectancy, years",
        }
    )
    top15_inc = apac.nlargest(15, "incidence_asr_2023")[
        ["gbd_name", "wb_iso3", "apac_region_label", "incidence_asr_2023", "age65_pct", "life_expectancy"]
    ].rename(
        columns={
            "gbd_name": "Location",
            "wb_iso3": "ISO3",
            "apac_region_label": "APAC region",
            "incidence_asr_2023": "Incidence ASR 2023",
            "age65_pct": "Population aged 65+, %",
            "life_expectancy": "Life expectancy, years",
        }
    )
    map_coverage = apac[
        ["gbd_name", "wb_iso3", "apac_region_label", "plot_method", "longitude", "latitude"]
    ].rename(
        columns={
            "gbd_name": "Location",
            "wb_iso3": "ISO3",
            "apac_region_label": "APAC region",
            "plot_method": "Plot method",
            "longitude": "Longitude",
            "latitude": "Latitude",
        }
    )
    region_summary = (
        apac.groupby("apac_region_label", observed=False)
        .agg(
            countries=("wb_iso3", "count"),
            median_mortality_asr=("asmr_2023", "median"),
            mean_mortality_asr=("asmr_2023", "mean"),
            median_prevalence_asr=("prevalence_asr_2023", "median"),
            mean_prevalence_asr=("prevalence_asr_2023", "mean"),
            median_incidence_asr=("incidence_asr_2023", "median"),
            mean_incidence_asr=("incidence_asr_2023", "mean"),
            median_age65_pct=("age65_pct", "median"),
            median_life_expectancy=("life_expectancy", "median"),
        )
        .reset_index()
        .rename(
            columns={
                "apac_region_label": "APAC region",
                "countries": "Locations",
                "median_mortality_asr": "Median mortality ASR 2023",
                "mean_mortality_asr": "Mean mortality ASR 2023",
                "median_prevalence_asr": "Median prevalence ASR 2023",
                "mean_prevalence_asr": "Mean prevalence ASR 2023",
                "median_incidence_asr": "Median incidence ASR 2023",
                "mean_incidence_asr": "Mean incidence ASR 2023",
                "median_age65_pct": "Median population aged 65+, %",
                "median_life_expectancy": "Median life expectancy, years",
            }
        )
    )
    return {
        "tableS6": apac_table,
        "tableS7": map_coverage,
        "tableS8": top15,
        "tableS9": region_summary,
        "tableS10": top15_prev,
        "tableS11": top15_inc,
    }


def augment_sections(
    sections: list[tuple[str, list[tuple[str | None, list[str]]]]],
    apac: pd.DataFrame,
) -> list[tuple[str, list[tuple[str | None, list[str]]]]]:
    top_death_text = top_rate_text(apac, "asmr_2023", 5)
    top_prev_text = top_rate_text(apac, "prevalence_asr_2023", 3)
    top_inc_text = top_rate_text(apac, "incidence_asr_2023", 3)
    polygon_n = int((apac["plot_method"] == "polygon").sum())
    point_n = int((apac["plot_method"] == "point").sum())
    apac_n = len(apac)
    median_death = float(apac["asmr_2023"].median())
    max_death = float(apac["asmr_2023"].max())
    median_prev = float(apac["prevalence_asr_2023"].median())
    max_prev = float(apac["prevalence_asr_2023"].max())
    median_inc = float(apac["incidence_asr_2023"].median())
    max_inc = float(apac["incidence_asr_2023"].max())

    new_sections: list[tuple[str, list[tuple[str | None, list[str]]]]] = []
    for section_title, blocks in sections:
        mutable_blocks = [(subsection, list(paragraphs)) for subsection, paragraphs in blocks]
        if section_title == "Methods":
            insert_at = len(mutable_blocks) - 2 if len(mutable_blocks) >= 2 else len(mutable_blocks)
            mutable_blocks.insert(
                insert_at,
                (
                    "Asia-Pacific geographic extension",
                    [
                        f"To provide a geographic extension relevant to rapidly ageing Asian and Pacific health systems, we defined Asia-Pacific operationally as countries and territories classified by the World Bank in East Asia & Pacific and South Asia. "
                        f"The regional country list was obtained from the World Bank country API on March 9, 2026 and merged to the harmonized mortality-ageing dataset using ISO3 codes. "
                        f"This yielded {apac_n} locations with complete mortality and ageing data and stable GBD location identifiers.",
                        "To extend the atlas beyond mortality, we additionally extracted official country-level age-standardized rates for deaths, prevalence, and incidence from the authenticated GBD Results Tool on March 9, 2026. "
                        "The query used GBD 2023 version 8352, cause id 653 (skin and subcutaneous diseases), both sexes, age-standardized rates, and 2023 country-level outputs. "
                        "These official Results Tool exports were used only for the Asia-Pacific geographic supplement, whereas the global descriptive and ecological analyses retained the previously locked local reproducible DIRF and mortality extracts.",
                        f"Country polygons were plotted with a Natural Earth low-resolution world layer, and small island territories or special administrative regions missing from that polygon layer were added as point markers using World Bank country longitude and latitude fields. "
                        f"In the final Asia-Pacific map set, {polygon_n} locations were rendered as polygons and {point_n} as point markers. "
                        "This approach preserves geographic completeness while making explicit the difference between polygon and point-based rendering. "
                        "All three Asia-Pacific rate maps therefore share a common country frame and differ only in measure definition.",
                    ],
                ),
            )
        if section_title == "Results":
            insert_at = max(len(mutable_blocks) - 1, 0)
            mutable_blocks.insert(
                insert_at,
                (
                    "Asia-Pacific rate geography",
                    [
                        f"The Asia-Pacific geographic extension showed pronounced within-region heterogeneity in 2023 across deaths, prevalence, and incidence (figure 4; supplementary figures S2-S7 and tables S6-S11). "
                        f"Across {apac_n} locations in World Bank East Asia & Pacific and South Asia, the median age-standardized mortality rate was {median_death:.2f} per 100,000 and the maximum was {max_death:.2f} per 100,000. "
                        f"For prevalence, the median was {median_prev:.1f} per 100,000 and the maximum was {max_prev:.1f} per 100,000; for incidence, the median was {median_inc:.1f} per 100,000 and the maximum was {max_inc:.1f} per 100,000. "
                        f"The highest mortality estimates were observed in {top_death_text}; the highest prevalence estimates in {top_prev_text}; and the highest incidence estimates in {top_inc_text}.",
                        "The four-panel atlas makes two patterns immediately visible. "
                        "First, the heaviest mortality burden clustered in Pacific island and maritime settings rather than in the oldest continental economies of the region. "
                        "Second, the non-fatal burden was also elevated in many Pacific and maritime settings, but the ordering was not identical across measures, suggesting that settings with high prevalence or incidence do not necessarily have the highest standardized mortality.",
                        "The Asia-Pacific supplement also clarified why map design matters for this topic. "
                        "Many of the highest mortality settings are small islands or territories that can disappear in coarse polygon maps. "
                        "Representing unmatched small territories as point markers preserved those observations in the figure rather than excluding them. "
                        "That choice is especially important here because several of the highest mortality, prevalence, and incidence estimates in the region came from jurisdictions that are geographically small but epidemiologically prominent.",
                    ],
                ),
            )
        if section_title == "Discussion":
            if mutable_blocks and mutable_blocks[0][0] is None:
                mutable_blocks[0][1].append(
                    "The Asia-Pacific geographic supplement reinforces the main interpretation of the paper. "
                    "It shows that high standardized mortality can cluster in small islands, territories, and mixed middle-income settings where wound care, infection management, referral pathways, and continuity of chronic care may be more fragile than simple age structure would suggest. "
                    "At the same time, the newly integrated official prevalence and incidence outputs indicate that the non-fatal dermatologic burden is also concentrated in several Pacific and Southeast Asian settings, expanding the policy relevance from end-stage mortality prevention to broader outpatient, chronic-care, and prevention capacity."
                )
                mutable_blocks[0][1].append(
                    "For readers concerned with healthy ageing in the Asia-Pacific region, the implication is therefore twofold: population ageing will increase dermatologic demand, and some settings already combine high non-fatal burden with disproportionately high mortality. "
                    "That combination supports an integrated response spanning prevention, chronic disease management, infection control, wound care, referral access, and long-term service planning rather than a narrow mortality-only framing."
                )
        new_sections.append((section_title, [(subsection, paragraphs) for subsection, paragraphs in mutable_blocks]))
    return new_sections


def write_main_docx(
    path: Path,
    summary: dict[str, str],
    research_in_context: dict[str, str],
    sections: list[tuple[str, list[tuple[str | None, list[str]]]]],
    references: list[str],
    main_tables: list[tuple[str, pd.DataFrame]],
    main_figures: list[tuple[str, str, Path]],
) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, TITLE)
    add_heading(doc, "Summary", level=1)
    for key in ["Background", "Methods", "Findings", "Interpretation", "Funding"]:
        add_heading(doc, key, level=2)
        add_paragraph(doc, summary[key])

    add_heading(doc, "Research in context", level=1)
    for key, value in research_in_context.items():
        add_heading(doc, key, level=2)
        add_paragraph(doc, value)

    for section_title, blocks in sections:
        add_heading(doc, section_title, level=1)
        for subsection, paragraphs in blocks:
            if subsection:
                add_heading(doc, subsection, level=2)
            for paragraph in paragraphs:
                add_paragraph(doc, paragraph)

    add_heading(doc, "References", level=1)
    for idx, ref in enumerate(references, start=1):
        add_paragraph(doc, f"{idx}. {ref}")

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Tables", level=1)
    for title, df in main_tables:
        add_dataframe_table(doc, title, df)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "Figure Legends And Figures", level=1)
    for idx, (title, legend, image_path) in enumerate(main_figures):
        if idx > 0:
            doc.add_page_break()
        add_figure_caption(doc, title, legend)
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6.7))

    doc.save(path)


def write_support_docx(path: Path, title: str, lines: list[str]) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, title)
    for line in lines:
        if not line:
            doc.add_paragraph()
        else:
            add_paragraph(doc, line)
    doc.save(path)


def write_supplementary_docx(
    path: Path,
    supp_figures: list[tuple[str, str, Path]],
    supp_tables: list[tuple[str, pd.DataFrame]],
) -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, "Supplementary Appendix")
    add_heading(doc, "Supplementary Figures", level=1)
    for idx, (title, legend, image_path) in enumerate(supp_figures):
        if idx > 0:
            doc.add_page_break()
        add_figure_caption(doc, title, legend)
        if image_path.exists():
            doc.add_picture(str(image_path), width=Inches(6.55))
    doc.add_page_break()
    add_heading(doc, "Supplementary Tables", level=1)
    for title, df in supp_tables:
        add_dataframe_table(doc, title, df)
    doc.save(path)


def build_qc_lines(
    main_docx: Path,
    supp_docx: Path,
    main_word_count: int,
    summary_word_count: int,
    main_tables: list[tuple[str, pd.DataFrame]],
    supp_tables: list[tuple[str, pd.DataFrame]],
    main_figures: list[tuple[str, str, Path]],
    supp_figures: list[tuple[str, str, Path]],
    apac: pd.DataFrame,
    render_summary: dict[str, object],
) -> list[str]:
    order_status = "FAIL"
    table_count_status = "FAIL"
    try:
        doc = Document(main_docx)
        headings = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        refs_idx = next(i for i, t in enumerate(headings) if t == "References")
        tables_idx = next(i for i, t in enumerate(headings) if t == "Tables")
        figs_idx = next(i for i, t in enumerate(headings) if t == "Figure Legends And Figures")
        order_status = "PASS" if refs_idx < tables_idx < figs_idx else "FAIL"
        table_count_status = "PASS" if len(doc.tables) == len(main_tables) else "FAIL"
    except Exception:
        pass

    render_lines = []
    if render_summary.get("available"):
        render_lines.append(
            f"- Render pipeline available: yes; documents rendered={render_summary.get('rendered_count')}/{render_summary.get('document_count')}; total pages={render_summary.get('page_count')}"
        )
        for item in render_summary.get("documents", []):
            doc_name = Path(str(item.get("docx", ""))).name
            render_lines.append(f"- {doc_name}: {'PASS' if item.get('ok') else 'FAIL'}, pages={item.get('page_count')}")
    else:
        render_lines.append("- Render pipeline unavailable.")

    lines = [
        f"Main text word count: {main_word_count}",
        f"Summary word count: {summary_word_count}",
        "Main-text display items: 10 (5 figures and 5 tables)",
        f"Supplementary figures: {len(supp_figures)}",
        f"Supplementary tables: {len(supp_tables)}",
        f"APAC locations with complete data: {len(apac)}",
        f"APAC polygon locations: {int((apac['plot_method'] == 'polygon').sum())}",
        f"APAC point-marker locations: {int((apac['plot_method'] == 'point').sum())}",
        f"Main-manuscript order check (References -> Tables -> Figures): {order_status}",
        f"Main-manuscript table count check: {table_count_status}",
        f"Main figure file existence check: {'PASS' if all(path.exists() for _, _, path in main_figures) else 'FAIL'}",
        f"Supplementary figure file existence check: {'PASS' if all(path.exists() for _, _, path in supp_figures) else 'FAIL'}",
        f"Official APAC Results Tool extract present: {'PASS' if OFFICIAL_APAC_ASR_PATH.exists() else 'FAIL'}",
        "Data-support note: the Asia-Pacific deaths, prevalence, and incidence maps use official authenticated GBD Results Tool country-level age-standardized rates exported on March 9, 2026; the global sections remain anchored to the locked local reproducible DIRF and mortality datasets.",
        "Rendered page QA:",
    ]
    lines.extend(render_lines)
    lines.extend(
        [
            "Remaining manual items: author order, affiliations, corresponding-author contact details, funding statement, ethics wording, originality confirmation, and final declaration wording.",
            "This enhanced package places all main tables and figures after the reference list and adds an Asia-Pacific geographic extension plus supplementary figures and tables.",
        ]
    )
    return lines


def main() -> None:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    TABLE_DIR.mkdir(parents=True, exist_ok=True)
    MANUSCRIPT_DIR.mkdir(parents=True, exist_ok=True)

    builder = load_builder_module()
    configure_matplotlib()

    global_context = pd.read_csv(AGING_DIR / "skin_aging_global_context_1990_2023.csv")
    global_context = global_context[global_context["year_id"].isin([1990, 2010, 2020, 2023])].copy()
    country_complete, ambiguous_names = builder.load_country_complete()
    correlations, tertiles, top20 = builder.compute_country_ecology(country_complete)
    subtype_dirf, subtype_mortality = builder.load_subtype_profiles()

    table1 = builder.build_table1_study_frame()
    table2 = builder.build_main_table(global_context)
    result_tables = builder.build_main_result_tables(subtype_dirf, subtype_mortality, correlations, tertiles, top20)

    builder.make_figure1(global_context)
    builder.make_figure2(subtype_dirf, subtype_mortality)
    builder.make_figure3(country_complete, correlations)
    builder.make_figure5(subtype_dirf, subtype_mortality)

    values = builder.build_value_map(
        global_context,
        country_complete,
        correlations,
        tertiles,
        top20,
        subtype_dirf,
        subtype_mortality,
    )
    summary = builder.build_summary(values)
    research_in_context = builder.build_research_in_context()
    references = builder.build_references()
    sections = augment_sections(builder.build_main_sections(values), *build_apac_datasets(country_complete)[:1])

    apac, _ = build_apac_datasets(country_complete)
    make_world_apac_mortality_map(country_complete, apac)
    make_top20_bar(top20)
    make_apac_top15_bar(apac)
    make_apac_top15_rate_bar(
        apac,
        "prevalence_asr_2023",
        "Figure S3. Highest prevalence settings in Asia-Pacific, 2023",
        "Age-standardized prevalence rate per 100,000",
        "figureS3_apac_top15_prevalence_2023",
        COLOR_RUST,
    )
    make_apac_top15_rate_bar(
        apac,
        "incidence_asr_2023",
        "Figure S4. Highest incidence settings in Asia-Pacific, 2023",
        "Age-standardized incidence rate per 100,000",
        "figureS4_apac_top15_incidence_2023",
        COLOR_TEAL,
    )
    make_apac_scatter(apac)
    _plot_apac_map(
        apac,
        "age65_pct",
        "Figure S6. Population aged 65 years and older in Asia-Pacific, 2023",
        AGE_CMAP,
        "figureS6_apac_age65_map",
        "Population aged 65+, %",
    )
    _plot_apac_map(
        apac,
        "life_expectancy",
        "Figure S7. Life expectancy at birth in Asia-Pacific, 2023",
        LIFE_CMAP,
        "figureS7_apac_life_expectancy_map",
        "Life expectancy at birth, years",
    )

    main_tables = [
        ("Table 1. Study frame, data sources, and analytical modules", table1),
        ("Table 2. Global burden of skin and subcutaneous diseases and World Bank ageing indicators in 1990 and 2023", table2),
        ("Table 3. Subtype-specific global burden profile of skin and subcutaneous diseases in 2023", result_tables["subtype_2023"]),
        ("Table 4. Subtype-specific relative change in global skin burden between 1990 and 2023", result_tables["subtype_change"]),
        ("Table 5. Country-level ecological summary of skin mortality in 2023", result_tables["ecology_table"]),
    ]

    apac_tables = build_apac_tables(apac)
    supplementary_tables = [
        ("Table S1. Long-format subtype-specific global burden profile of skin and subcutaneous diseases in 2023", result_tables["subtype_2023_long"]),
        ("Table S2. Long-format subtype-specific change in global skin burden between 1990 and 2023", result_tables["subtype_change_long"]),
        ("Table S3. Country-level ecological correlations between World Bank ageing indicators and skin mortality in 2023", result_tables["correlations"]),
        ("Table S4. Top 20 countries and territories ranked by age-standardized mortality rate from skin and subcutaneous diseases in 2023", result_tables["top20"]),
        ("Table S5. Ageing tertile summary of country-level skin mortality in 2023", result_tables["tertiles"]),
        ("Table S6. Asia-Pacific country-level deaths, prevalence, incidence, and ageing indicators in 2023", apac_tables["tableS6"]),
        ("Table S7. Asia-Pacific map coverage and plotting method in 2023", apac_tables["tableS7"]),
        ("Table S8. Highest mortality settings in Asia-Pacific in 2023", apac_tables["tableS8"]),
        ("Table S9. Asia-Pacific regional summary statistics for deaths, prevalence, and incidence in 2023", apac_tables["tableS9"]),
        ("Table S10. Highest prevalence settings in Asia-Pacific in 2023", apac_tables["tableS10"]),
        ("Table S11. Highest incidence settings in Asia-Pacific in 2023", apac_tables["tableS11"]),
    ]

    for filename, df in {
        "tableS6_apac_country_death_prevalence_incidence_aging_2023.csv": apac_tables["tableS6"],
        "tableS7_apac_map_coverage_2023.csv": apac_tables["tableS7"],
        "tableS8_apac_top15_mortality_2023.csv": apac_tables["tableS8"],
        "tableS9_apac_region_summary_2023.csv": apac_tables["tableS9"],
        "tableS10_apac_top15_prevalence_2023.csv": apac_tables["tableS10"],
        "tableS11_apac_top15_incidence_2023.csv": apac_tables["tableS11"],
    }.items():
        df.to_csv(TABLE_DIR / filename, index=False)

    main_figure_legends = [
        builder.build_figure_legends()[0],
        builder.build_figure_legends()[1],
        builder.build_figure_legends()[2],
        (
            "Figure 4",
            "Four-panel atlas showing global mortality and Asia-Pacific country-level age-standardized rates in 2023. Panel A shows global mortality; panel B shows Asia-Pacific mortality; panel C shows Asia-Pacific prevalence; and panel D shows Asia-Pacific incidence. "
            "Asia-Pacific panels use official authenticated GBD Results Tool exports for cause id 653 (skin and subcutaneous diseases), both sexes, and age-standardized rate outputs. Country polygons show mapped locations available in the Natural Earth layer, and point markers denote small islands or territories added using World Bank coordinates because polygon geometry was not available in the low-resolution base map.",
        ),
        builder.build_figure_legends()[4],
    ]
    main_figures = [
        ("Figure 1", main_figure_legends[0][1], FIGURE_DIR / "figure1_global_burden_and_aging.png"),
        ("Figure 2", main_figure_legends[1][1], FIGURE_DIR / "figure2_subtype_profile_2023.png"),
        ("Figure 3", main_figure_legends[2][1], FIGURE_DIR / "figure3_country_aging_ecology.png"),
        ("Figure 4", main_figure_legends[3][1], FIGURE_DIR / "figure4_global_apac_mortality_map.png"),
        ("Figure 5", main_figure_legends[4][1], FIGURE_DIR / "figure5_subtype_trends_1990_2023.png"),
    ]
    supp_figures = [
        (
            "Figure S1",
            "Countries and territories with the highest age-standardized mortality rates from skin and subcutaneous diseases in 2023. Bars show point estimates and whiskers show lower and upper uncertainty bounds.",
            FIGURE_DIR / "figureS1_top20_country_asmr_2023.png",
        ),
        (
            "Figure S2",
            "Highest mortality settings in Asia-Pacific in 2023. Bars show official age-standardized mortality rates from the authenticated GBD Results Tool export, with uncertainty bounds.",
            FIGURE_DIR / "figureS2_apac_top15_asmr_2023.png",
        ),
        (
            "Figure S3",
            "Highest prevalence settings in Asia-Pacific in 2023. Bars show official age-standardized prevalence rates from the authenticated GBD Results Tool export.",
            FIGURE_DIR / "figureS3_apac_top15_prevalence_2023.png",
        ),
        (
            "Figure S4",
            "Highest incidence settings in Asia-Pacific in 2023. Bars show official age-standardized incidence rates from the authenticated GBD Results Tool export.",
            FIGURE_DIR / "figureS4_apac_top15_incidence_2023.png",
        ),
        (
            "Figure S5",
            "Asia-Pacific relationship between population ageing and skin mortality in 2023. Points are colored by World Bank region grouping used for the regional supplement.",
            FIGURE_DIR / "figureS5_apac_age65_scatter.png",
        ),
        (
            "Figure S6",
            "Population aged 65 years and older across Asia-Pacific settings in 2023. Point markers indicate locations not represented as polygons in the low-resolution base map.",
            FIGURE_DIR / "figureS6_apac_age65_map.png",
        ),
        (
            "Figure S7",
            "Life expectancy at birth across Asia-Pacific settings in 2023. Point markers indicate locations not represented as polygons in the low-resolution base map.",
            FIGURE_DIR / "figureS7_apac_life_expectancy_map.png",
        ),
    ]

    section_paragraphs = [
        paragraph
        for _, blocks in sections
        for _, paragraphs in blocks
        for paragraph in paragraphs
    ]
    main_word_count = builder.word_count(" ".join(section_paragraphs))
    summary_word_count = builder.word_count(" ".join(summary.values()))

    table_titles = [title for title, _ in main_tables]
    figure_entries = [(title, legend) for title, legend, _ in main_figures]

    artifacts = PackageArtifacts(
        main_docx=MANUSCRIPT_DIR / "skin_lancet_complete_apac_5000w.docx",
        main_md=MANUSCRIPT_DIR / "skin_lancet_complete_apac_5000w.md",
        title_docx=MANUSCRIPT_DIR / "title_page_complete_apac.docx",
        title_md=MANUSCRIPT_DIR / "title_page_complete_apac.md",
        cover_docx=MANUSCRIPT_DIR / "cover_letter_complete_apac.docx",
        cover_md=MANUSCRIPT_DIR / "cover_letter_complete_apac.md",
        supp_docx=MANUSCRIPT_DIR / "supplementary_appendix_complete_apac.docx",
        supp_md=MANUSCRIPT_DIR / "supplementary_appendix_complete_apac.md",
        qc_docx=MANUSCRIPT_DIR / "complete_apac_qc_report.docx",
        qc_md=MANUSCRIPT_DIR / "complete_apac_qc_report.md",
        summary_json=MANUSCRIPT_DIR / "complete_apac_summary.json",
    )

    artifacts.main_md.write_text(
        main_doc_markdown(summary, research_in_context, sections, references, table_titles, figure_entries),
        encoding="utf-8",
    )
    write_main_docx(
        artifacts.main_docx,
        summary,
        research_in_context,
        sections,
        references,
        main_tables,
        main_figures,
    )

    title_lines = build_title_page(main_word_count, summary_word_count)
    cover_lines = build_cover_letter(values, apac)
    artifacts.title_md.write_text("# Title Page\n\n" + "\n".join(title_lines), encoding="utf-8")
    artifacts.cover_md.write_text("# Cover Letter\n\n" + "\n".join(cover_lines), encoding="utf-8")
    write_support_docx(artifacts.title_docx, "Title Page", title_lines)
    write_support_docx(artifacts.cover_docx, "Cover Letter", cover_lines)

    supp_md_lines = ["# Supplementary Appendix", "", "## Supplementary Figures", ""]
    for title, legend, _ in supp_figures:
        supp_md_lines.extend([f"### {title}", "", legend, ""])
    supp_md_lines.extend(["## Supplementary Tables", ""])
    for title, _ in supplementary_tables:
        supp_md_lines.append(f"- {title}")
    artifacts.supp_md.write_text("\n".join(supp_md_lines), encoding="utf-8")
    write_supplementary_docx(artifacts.supp_docx, supp_figures, supplementary_tables)

    render_summary = render_docx_collection(
        [artifacts.main_docx, artifacts.supp_docx, artifacts.title_docx, artifacts.cover_docx],
        MANUSCRIPT_DIR / "rendered_pages_complete_apac",
    )

    qc_lines = build_qc_lines(
        artifacts.main_docx,
        artifacts.supp_docx,
        main_word_count,
        summary_word_count,
        main_tables,
        supplementary_tables,
        main_figures,
        supp_figures,
        apac,
        render_summary,
    )
    artifacts.qc_md.write_text("# Complete APAC QC Report\n\n" + "\n".join(qc_lines), encoding="utf-8")
    write_support_docx(artifacts.qc_docx, "Complete APAC QC Report", qc_lines)

    summary_payload = {
        "title": TITLE,
        "main_word_count": main_word_count,
        "summary_word_count": summary_word_count,
        "main_figures": len(main_figures),
        "main_tables": len(main_tables),
        "supplementary_figures": len(supp_figures),
        "supplementary_tables": len(supplementary_tables),
        "countries_in_ecology": int(values["n_countries"]),
        "ambiguous_country_names_excluded": ambiguous_names,
        "apac_locations": len(apac),
        "apac_polygon_locations": int((apac["plot_method"] == "polygon").sum()),
        "apac_point_locations": int((apac["plot_method"] == "point").sum()),
        "render_summary": render_summary,
    }
    artifacts.summary_json.write_text(json.dumps(summary_payload, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"Main text words: {main_word_count}")
    print(f"Summary words: {summary_word_count}")
    print(f"APAC locations: {len(apac)}")
    print(f"Package written to: {MANUSCRIPT_DIR}")


if __name__ == "__main__":
    main()
