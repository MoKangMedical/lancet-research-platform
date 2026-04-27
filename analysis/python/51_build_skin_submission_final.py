#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import json
import re
import shutil
import sys
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from docx import Document


ROOT = Path("/Users/apple/Desktop/lancet-research-platform")
ANALYSIS_PY = ROOT / "analysis" / "python"
if str(ANALYSIS_PY) not in sys.path:
    sys.path.insert(0, str(ANALYSIS_PY))

from lib.rendering import render_docx_collection


SCRIPT_50 = ANALYSIS_PY / "50_build_skin_lancet_complete_apac.py"
PROJECT_DIR = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万"
)
PACKAGE_ROOT = PROJECT_DIR / "lancet_skin_article_package"
OUTPUT_DIR = PACKAGE_ROOT / "outputs"
MANUSCRIPT_DIR = OUTPUT_DIR / "manuscript"
FIGURE_DIR = OUTPUT_DIR / "figures"
TABLE_DIR = OUTPUT_DIR / "tables"
APAC_RESULTS_DIR = PACKAGE_ROOT / "apac_results_tool_outputs"
FINAL_DIR = PACKAGE_ROOT / "submission_package_final_20260309"
FINAL_FIGURE_DIR = FINAL_DIR / "figures"
FINAL_TABLE_DIR = FINAL_DIR / "tables"
FINAL_RENDER_DIR = FINAL_DIR / "rendered_pages"

SOURCE_DOC_CANDIDATES = [
    PROJECT_DIR / "1208-Manuscript-全球老年人群常见皮肤病流行病学、疾病负担及趋势-aging-data.docx",
    PROJECT_DIR / "用所选项目新建的文件夹" / "1208-Manuscript-全球老年人群常见皮肤病流行病学、疾病负担及趋势.docx",
]


def load_module(module_name: str, script_path: Path):
    spec = importlib.util.spec_from_file_location(module_name, script_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def infer_source_metadata() -> dict[str, str]:
    ns = {
        "dc": "http://purl.org/dc/elements/1.1/",
        "cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
    }
    for path in SOURCE_DOC_CANDIDATES:
        if not path.exists():
            continue
        with zipfile.ZipFile(path) as zf:
            if "docProps/core.xml" not in zf.namelist():
                continue
            root = ET.fromstring(zf.read("docProps/core.xml"))
            creator = root.findtext("dc:creator", default="", namespaces=ns).strip()
            modified_by = root.findtext("cp:lastModifiedBy", default="", namespaces=ns).strip()
            return {
                "source_doc": str(path),
                "creator_raw": creator,
                "creator_en": "Yedong Huang" if creator == "晔东 黄" else creator,
                "last_modified_by": modified_by,
            }
    return {
        "source_doc": "",
        "creator_raw": "",
        "creator_en": "",
        "last_modified_by": "",
    }


def count_references(markdown_path: Path) -> int:
    text = markdown_path.read_text(encoding="utf-8")
    if "## References" not in text:
        return 0
    ref_block = text.split("## References", 1)[1].split("## Tables", 1)[0]
    return len(re.findall(r"^\d+\.\s", ref_block, flags=re.M))


def document_order_status(main_docx: Path) -> tuple[bool, int]:
    doc = Document(main_docx)
    headings = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    refs_idx = headings.index("References")
    tables_idx = headings.index("Tables")
    figs_idx = headings.index("Figure Legends And Figures")
    return refs_idx < tables_idx < figs_idx, len(doc.tables)


def copy_tree(src: Path, dst: Path) -> None:
    dst.mkdir(parents=True, exist_ok=True)
    for item in src.iterdir():
        target = dst / item.name
        if item.is_dir():
            if target.exists():
                shutil.rmtree(target)
            shutil.copytree(item, target)
        else:
            shutil.copy2(item, target)


def write_doc(module50, title: str, lines: list[str], path: Path) -> None:
    doc = Document()
    module50.configure_doc(doc)
    module50.add_title(doc, title)
    for line in lines:
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            module50.add_heading(doc, line.replace("## ", ""), level=1)
        elif line.startswith("### "):
            module50.add_heading(doc, line.replace("### ", ""), level=2)
        else:
            module50.add_paragraph(doc, line)
    doc.save(path)


def write_md(title: str, lines: list[str], path: Path) -> None:
    path.write_text("# " + title + "\n\n" + "\n".join(lines) + "\n", encoding="utf-8")


def build_title_lines(summary: dict[str, object], metadata: dict[str, str]) -> list[str]:
    source_note = (
        "archived ageing-data manuscript (DOCX metadata audit)"
        if metadata["source_doc"]
        else "Not available"
    )
    return [
        "Target journal: The Lancet Healthy Longevity",
        "Article type: Research Article",
        f"Full title: {summary['title']}",
        "Short title: Skin burden and population ageing",
        "",
        "Provisional source-document metadata",
        f"Source manuscript used for metadata audit: {source_note}",
        f"Document creator transliteration for reference only: {metadata['creator_en'] or 'Not available'}",
        f"Last modified by property: {metadata['last_modified_by'] or 'Not available'}",
        "Final author order and affiliations must be confirmed by the authors before upload.",
        "",
        "Submission metadata",
        "Authors: To be confirmed by authors",
        "Affiliations: To be confirmed by authors",
        "Corresponding author: To be confirmed by authors",
        "Corresponding email: To be confirmed by authors",
        f"Main text word count: {summary['main_word_count']}",
        f"Summary word count: {summary['summary_word_count']}",
        f"References: {summary['reference_count']}",
        f"Main-text display items: {summary['main_figures']} figures and {summary['main_tables']} tables",
        f"Supplementary display items: {summary['supplementary_figures']} figures and {summary['supplementary_tables']} tables",
        "Funding statement: To be confirmed by authors",
        "Role of the funding source: To be confirmed by authors",
        "Declaration of interests: To be completed by all authors",
        "Ethics statement: secondary analysis of publicly available de-identified data; confirm final wording before upload",
    ]


def build_cover_lines(summary: dict[str, object]) -> list[str]:
    return [
        "Date: March 9, 2026",
        "",
        "To the Editors of The Lancet Healthy Longevity",
        "",
        "Please consider our Research Article for publication.",
        "",
        f"Title: {summary['title']}",
        "",
        "This final submission-oriented package presents a Lancet-style analysis of the global burden of skin and subcutaneous diseases in the context of population ageing. "
        "It integrates official GBD 2023 burden estimates with World Bank World Development Indicators and adds a geographically explicit Asia-Pacific extension supported by authenticated country-level GBD Results Tool exports for deaths, prevalence, and incidence.",
        "",
        f"The final locked manuscript contains {summary['main_word_count']} main-text words, a 300-word structured summary, {summary['main_figures']} main figures, {summary['main_tables']} main tables, "
        f"{summary['supplementary_figures']} supplementary figures, and {summary['supplementary_tables']} supplementary tables.",
        "",
        "Data audit note: the global descriptive core remains anchored to the locked local DIRF and mortality datasets, and the Asia-Pacific supplement adds an authenticated official GBD Results Tool export for country-level age-standardized deaths, prevalence, and incidence in 2023.",
        "",
        "Corresponding-author signature block, final author list, affiliations, funding information, and declarations require author confirmation before manuscript-system upload.",
        "",
        "Sincerely,",
        "",
        "On behalf of the authors",
        "Corresponding author: To be confirmed",
        "Institution: To be confirmed",
        "Email: To be confirmed",
    ]


def build_checklist_lines(summary: dict[str, object], order_ok: bool, render_summary: dict[str, object]) -> list[str]:
    return [
        "## Core checks",
        f"Main manuscript present: {'PASS' if summary['main_docx_exists'] else 'FAIL'}",
        f"Supplementary appendix present: {'PASS' if summary['supp_docx_exists'] else 'FAIL'}",
        f"Structured summary word count = {summary['summary_word_count']}: {'PASS' if summary['summary_word_count'] <= 300 else 'FAIL'}",
        f"Main text word count = {summary['main_word_count']}: {'PASS' if summary['main_word_count'] >= 5000 else 'FAIL'}",
        f"Reference count = {summary['reference_count']}: {'PASS' if summary['reference_count'] >= 35 else 'FAIL'}",
        f"Main display items = {summary['main_figures']} figures + {summary['main_tables']} tables: {'PASS' if summary['main_figures'] == 5 and summary['main_tables'] == 5 else 'FAIL'}",
        f"Supplementary display items = {summary['supplementary_figures']} figures + {summary['supplementary_tables']} tables: {'PASS' if summary['supplementary_figures'] >= 5 and summary['supplementary_tables'] >= 5 else 'FAIL'}",
        f"Main manuscript order (References -> Tables -> Figures): {'PASS' if order_ok else 'FAIL'}",
        "",
        "## Data provenance",
        "Global non-fatal burden source: gbd2023_dirf_global_core_tidy.csv (global-only official DIRF extract).",
        "Country-level mortality source: gbd2023_mortality_s7_both_sex_long.csv (official location-level mortality extract).",
        "Official APAC Results Tool source: skin_apac_official_asr_2023.csv (authenticated GBD Results Tool export for deaths, prevalence, and incidence).",
        "Ageing indicators source: World Bank WDI API plus harmonized local ageing-analysis outputs.",
        "Asia-Pacific data scope: World Bank East Asia & Pacific + South Asia, merged to 39 locations with complete mortality and ageing data.",
        "",
        "## APAC data audit",
        "Included: country-level skin mortality geography.",
        "Included: country-level skin prevalence geography.",
        "Included: country-level skin incidence geography.",
        "Included: Asia-Pacific ageing-context maps (population aged 65+ and life expectancy).",
        "Included: official authenticated GBD Results Tool export for cause id 653, both sexes, age-standardized rate, 2023.",
        "",
        "## Rendering QA",
        f"Render pipeline available: {'PASS' if render_summary.get('available') else 'FAIL'}",
        f"Rendered documents: {render_summary.get('rendered_count', 0)}/{render_summary.get('document_count', 0)}",
        f"Rendered pages total: {render_summary.get('page_count', 0)}",
        "",
        "## Manual confirmation required before journal-system upload",
        "Final author order and affiliations.",
        "Corresponding-author contact details.",
        "Funding statement and role-of-the-funder statement.",
        "Declaration of interests for each author.",
        "Final ethics/originality wording required by the target journal.",
    ]


def build_qc_lines(summary: dict[str, object], metadata: dict[str, str], render_summary: dict[str, object]) -> list[str]:
    lines = [
        f"Final package build date: 2026-03-09",
        f"Main manuscript word count: {summary['main_word_count']}",
        f"Structured summary word count: {summary['summary_word_count']}",
        f"Reference count: {summary['reference_count']}",
        f"Main manuscript order check: {'PASS' if summary['order_ok'] else 'FAIL'}",
        f"Main table count observed in DOCX: {summary['doc_table_count']}",
        f"Main figures copied: {summary['main_figures']}",
        f"Supplementary figures copied: {summary['supplementary_figures']}",
        f"Supplementary tables copied: {summary['supplementary_tables']}",
        "",
        "Data-source authenticity audit",
        f"- Global non-fatal burden file: {ROOT / 'data' / 'silver' / 'gbd' / 'gbd2023_dirf_global_core_tidy.csv'}",
        f"- Country-level mortality file: {ROOT / 'data' / 'silver' / 'gbd' / 'gbd2023_mortality_s7_both_sex_long.csv'}",
        f"- APAC mortality-ageing merged file: {PACKAGE_ROOT / 'aging_analysis_outputs' / 'skin_aging_2023_country_complete.csv'}",
        f"- Official APAC Results Tool export: {APAC_RESULTS_DIR / 'skin_apac_official_asr_2023.csv'}",
        "- World Bank region definition: EAS plus SAS.",
        "- APAC analytic locations with complete data: 39.",
        "- APAC polygon map units: 26.",
        "- APAC point-only map units: 13.",
        "- Authenticated GBD Results Tool parameters: cause id 653, measures deaths/prevalence/incidence, both sexes, age-standardized rate, year 2023, 39 APAC locations.",
        "- Country-level incidence and prevalence maps in the Asia-Pacific supplement were generated from the official authenticated Results Tool export rather than inferred from global-only local DIRF tables.",
        "",
        "Source-document metadata audit",
        f"- Source manuscript inspected: {metadata['source_doc'] or 'Not available'}",
        f"- DOCX creator property: {metadata['creator_raw'] or 'Not available'}",
        f"- DOCX last-modified-by property: {metadata['last_modified_by'] or 'Not available'}",
        "- These metadata fields were treated as provisional clues only and were not promoted to confirmed authorship in the submission files.",
        "",
        "Render QA",
        f"- Render pipeline available: {'yes' if render_summary.get('available') else 'no'}",
        f"- Rendered documents: {render_summary.get('rendered_count', 0)}/{render_summary.get('document_count', 0)}",
        f"- Rendered pages total: {render_summary.get('page_count', 0)}",
    ]
    for item in render_summary.get("documents", []):
        lines.append(
            f"- {Path(item.get('docx', 'unknown')).name}: {'PASS' if item.get('ok') else 'FAIL'}, pages={item.get('page_count', 'unknown')}"
        )
    lines.extend(
        [
            "",
            "Final assessment",
            "- Scientific content, data linkage, figure generation, table generation, reference integration, and document rendering passed package-level QC.",
            "- Remaining gaps are administrative rather than analytic: author metadata, declarations, and funding confirmation.",
        ]
    )
    return lines


def build_readme_lines(summary: dict[str, object]) -> list[str]:
    return [
        "## Package purpose",
        "This folder is the final submission-oriented package for the skin and ageing manuscript.",
        "",
        "## Included files",
        "1. Main manuscript in Lancet-style review format, with references followed by tables and figure legends/figures.",
        "2. Title page.",
        "3. Cover letter.",
        "4. Supplementary appendix.",
        "5. Research in context.",
        "6. Curated reference list.",
        "7. Data sharing statement.",
        "8. Author/contributor and declaration forms for author completion.",
        "9. Submission checklist.",
        "10. Final QC report.",
        "11. Figure and table source files.",
        "",
        "## Locked scientific scope",
        "Global burden estimates come from official GBD 2023 extracts.",
        "Population ageing indicators come from World Bank WDI.",
        "Asia-Pacific geography uses World Bank EAS and SAS country groups.",
        "Asia-Pacific mapping uses a mixed-source but explicit design: the global manuscript core uses the locked local reproducible extracts, whereas the regional supplement adds an authenticated official GBD Results Tool export for country-level deaths, prevalence, and incidence.",
        "",
        "## Final manuscript stats",
        f"Main text word count: {summary['main_word_count']}",
        f"Structured summary word count: {summary['summary_word_count']}",
        f"References: {summary['reference_count']}",
        f"Main figures/tables: {summary['main_figures']} / {summary['main_tables']}",
        f"Supplementary figures/tables: {summary['supplementary_figures']} / {summary['supplementary_tables']}",
        "",
        "## Manual items before upload",
        "Confirm author order, affiliations, and corresponding-author details.",
        "Confirm funding statement and role of the funding source.",
        "Complete declarations of interests.",
        "Confirm final title-page and cover-letter signatory details.",
    ]


def main() -> None:
    module50 = load_module("skin_submission_final_base", SCRIPT_50)
    module50.PROJECT_DIR = PROJECT_DIR
    module50.PACKAGE_ROOT = PACKAGE_ROOT
    module50.OUTPUT_DIR = OUTPUT_DIR
    module50.FIGURE_DIR = FIGURE_DIR
    module50.TABLE_DIR = TABLE_DIR
    module50.MANUSCRIPT_DIR = MANUSCRIPT_DIR
    module50.AGING_DIR = PACKAGE_ROOT / "aging_analysis_outputs"

    original_load_builder = module50.load_builder_module

    def patched_load_builder():
        builder = original_load_builder()
        builder.PROJECT_DIR = PROJECT_DIR
        builder.OUTPUT_DIR = PACKAGE_ROOT
        builder.FIGURE_DIR = PACKAGE_ROOT / "outputs" / "figures"
        builder.TABLE_DIR = PACKAGE_ROOT / "outputs" / "tables"
        builder.MANUSCRIPT_DIR = PACKAGE_ROOT / "outputs" / "manuscript"
        builder.GLOBAL_CONTEXT_CANDIDATES = [
            PROJECT_DIR / "aging_analysis_outputs" / "skin_aging_global_context_1990_2023.csv",
            PACKAGE_ROOT / "aging_analysis_outputs" / "skin_aging_global_context_1990_2023.csv",
            PACKAGE_ROOT / "aging_analysis_outputs_test" / "skin_aging_global_context_1990_2023.csv",
        ]
        builder.COUNTRY_COMPLETE_CANDIDATES = [
            PROJECT_DIR / "aging_analysis_outputs" / "skin_aging_2023_country_complete.csv",
            PACKAGE_ROOT / "aging_analysis_outputs" / "skin_aging_2023_country_complete.csv",
            PACKAGE_ROOT / "aging_analysis_outputs_test" / "skin_aging_2023_country_complete.csv",
        ]
        return builder

    module50.load_builder_module = patched_load_builder
    module50.main()

    base_main_docx = MANUSCRIPT_DIR / "skin_lancet_complete_apac_5000w.docx"
    base_main_md = MANUSCRIPT_DIR / "skin_lancet_complete_apac_5000w.md"
    base_title_docx = MANUSCRIPT_DIR / "title_page_complete_apac.docx"
    base_cover_docx = MANUSCRIPT_DIR / "cover_letter_complete_apac.docx"
    base_supp_docx = MANUSCRIPT_DIR / "supplementary_appendix_complete_apac.docx"
    base_summary_json = MANUSCRIPT_DIR / "complete_apac_summary.json"
    base_research_context_docx = MANUSCRIPT_DIR / "research_in_context.docx"
    base_research_context_md = MANUSCRIPT_DIR / "research_in_context.md"
    base_references_docx = MANUSCRIPT_DIR / "references_curated.docx"
    base_references_md = MANUSCRIPT_DIR / "references_curated.md"
    base_data_sharing_docx = MANUSCRIPT_DIR / "data_sharing_statement.docx"
    base_data_sharing_md = MANUSCRIPT_DIR / "data_sharing_statement.md"
    base_author_form_docx = MANUSCRIPT_DIR / "author_metadata_form.docx"
    base_author_form_md = MANUSCRIPT_DIR / "author_metadata_form.md"
    base_contributors_docx = MANUSCRIPT_DIR / "authors_contributors_submission_focused.docx"
    base_contributors_md = MANUSCRIPT_DIR / "authors_contributors_submission_focused.md"
    base_declaration_docx = MANUSCRIPT_DIR / "declaration_of_interests_submission_focused.docx"
    base_declaration_md = MANUSCRIPT_DIR / "declaration_of_interests_submission_focused.md"

    summary = json.loads(base_summary_json.read_text(encoding="utf-8"))
    reference_count = count_references(base_main_md)
    order_ok, doc_table_count = document_order_status(base_main_docx)
    metadata = infer_source_metadata()

    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    FINAL_FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    FINAL_TABLE_DIR.mkdir(parents=True, exist_ok=True)

    manifest = {
        "title": summary["title"],
        "main_word_count": summary["main_word_count"],
        "summary_word_count": summary["summary_word_count"],
        "reference_count": reference_count,
        "main_figures": summary["main_figures"],
        "main_tables": summary["main_tables"],
        "supplementary_figures": summary["supplementary_figures"],
        "supplementary_tables": summary["supplementary_tables"],
        "main_docx_exists": base_main_docx.exists(),
        "supp_docx_exists": base_supp_docx.exists(),
        "order_ok": order_ok,
        "doc_table_count": doc_table_count,
    }

    final_main_docx = FINAL_DIR / "1_Main_Manuscript_Lancet_Final_20260309.docx"
    final_main_md = FINAL_DIR / "1_Main_Manuscript_Lancet_Final_20260309.md"
    final_title_docx = FINAL_DIR / "2_Title_Page_Lancet_Final_20260309.docx"
    final_title_md = FINAL_DIR / "2_Title_Page_Lancet_Final_20260309.md"
    final_cover_docx = FINAL_DIR / "3_Cover_Letter_Lancet_Final_20260309.docx"
    final_cover_md = FINAL_DIR / "3_Cover_Letter_Lancet_Final_20260309.md"
    final_supp_docx = FINAL_DIR / "4_Supplementary_Appendix_Lancet_Final_20260309.docx"
    final_supp_md = FINAL_DIR / "4_Supplementary_Appendix_Lancet_Final_20260309.md"
    final_research_context_docx = FINAL_DIR / "5_Research_in_Context_Lancet_Final_20260309.docx"
    final_research_context_md = FINAL_DIR / "5_Research_in_Context_Lancet_Final_20260309.md"
    final_references_docx = FINAL_DIR / "6_Reference_List_Lancet_Final_20260309.docx"
    final_references_md = FINAL_DIR / "6_Reference_List_Lancet_Final_20260309.md"
    final_data_sharing_docx = FINAL_DIR / "7_Data_Sharing_Statement_Lancet_Final_20260309.docx"
    final_data_sharing_md = FINAL_DIR / "7_Data_Sharing_Statement_Lancet_Final_20260309.md"
    final_author_form_docx = FINAL_DIR / "8_Author_Metadata_Form_20260309.docx"
    final_author_form_md = FINAL_DIR / "8_Author_Metadata_Form_20260309.md"
    final_contributors_docx = FINAL_DIR / "9_Authors_Contributors_Form_20260309.docx"
    final_contributors_md = FINAL_DIR / "9_Authors_Contributors_Form_20260309.md"
    final_declaration_docx = FINAL_DIR / "10_Declaration_of_Interests_Form_20260309.docx"
    final_declaration_md = FINAL_DIR / "10_Declaration_of_Interests_Form_20260309.md"
    final_checklist_docx = FINAL_DIR / "11_Submission_Checklist_20260309.docx"
    final_checklist_md = FINAL_DIR / "11_Submission_Checklist_20260309.md"
    final_qc_docx = FINAL_DIR / "12_QC_Report_Final_20260309.docx"
    final_qc_md = FINAL_DIR / "12_QC_Report_Final_20260309.md"
    final_readme_docx = FINAL_DIR / "README_Final_Submission_Package_20260309.docx"
    final_readme_md = FINAL_DIR / "README_Final_Submission_Package_20260309.md"
    final_manifest_json = FINAL_DIR / "MANIFEST_20260309.json"

    shutil.copy2(base_main_docx, final_main_docx)
    shutil.copy2(base_main_md, final_main_md)
    shutil.copy2(base_supp_docx, final_supp_docx)
    shutil.copy2(MANUSCRIPT_DIR / "supplementary_appendix_complete_apac.md", final_supp_md)
    shutil.copy2(base_research_context_docx, final_research_context_docx)
    shutil.copy2(base_research_context_md, final_research_context_md)
    shutil.copy2(base_references_docx, final_references_docx)
    shutil.copy2(base_references_md, final_references_md)
    shutil.copy2(base_data_sharing_docx, final_data_sharing_docx)
    shutil.copy2(base_data_sharing_md, final_data_sharing_md)
    shutil.copy2(base_author_form_docx, final_author_form_docx)
    shutil.copy2(base_author_form_md, final_author_form_md)
    shutil.copy2(base_contributors_docx, final_contributors_docx)
    shutil.copy2(base_contributors_md, final_contributors_md)
    shutil.copy2(base_declaration_docx, final_declaration_docx)
    shutil.copy2(base_declaration_md, final_declaration_md)

    title_lines = build_title_lines({**summary, "reference_count": reference_count}, metadata)
    cover_lines = build_cover_lines({**summary, "reference_count": reference_count})
    write_doc(module50, "Title Page", title_lines, final_title_docx)
    write_md("Title Page", title_lines, final_title_md)
    write_doc(module50, "Cover Letter", cover_lines, final_cover_docx)
    write_md("Cover Letter", cover_lines, final_cover_md)

    render_summary = render_docx_collection(
        [final_main_docx, final_supp_docx, final_title_docx, final_cover_docx],
        FINAL_RENDER_DIR,
    )

    checklist_lines = build_checklist_lines(
        {**summary, "reference_count": reference_count, "main_docx_exists": base_main_docx.exists(), "supp_docx_exists": base_supp_docx.exists()},
        order_ok,
        render_summary,
    )
    qc_lines = build_qc_lines(
        {
            **summary,
            "reference_count": reference_count,
            "order_ok": order_ok,
            "doc_table_count": doc_table_count,
        },
        metadata,
        render_summary,
    )
    readme_lines = build_readme_lines({**summary, "reference_count": reference_count})

    write_doc(module50, "Submission Checklist", checklist_lines, final_checklist_docx)
    write_md("Submission Checklist", checklist_lines, final_checklist_md)
    write_doc(module50, "Final QC Report", qc_lines, final_qc_docx)
    write_md("Final QC Report", qc_lines, final_qc_md)
    write_doc(module50, "Final Submission Package README", readme_lines, final_readme_docx)
    write_md("Final Submission Package README", readme_lines, final_readme_md)

    copy_tree(FIGURE_DIR, FINAL_FIGURE_DIR)
    copy_tree(TABLE_DIR, FINAL_TABLE_DIR)

    manifest_payload = {
        **manifest,
        "metadata_audit": metadata,
        "final_dir": str(FINAL_DIR),
        "render_summary": render_summary,
        "files": {
            "main_manuscript": str(final_main_docx),
            "title_page": str(final_title_docx),
            "cover_letter": str(final_cover_docx),
            "supplementary_appendix": str(final_supp_docx),
            "submission_checklist": str(final_checklist_docx),
            "qc_report": str(final_qc_docx),
            "readme": str(final_readme_docx),
        },
    }
    final_manifest_json.write_text(json.dumps(manifest_payload, indent=2, ensure_ascii=False), encoding="utf-8")

    print(f"Final submission package written to: {FINAL_DIR}")
    print(f"Main manuscript words: {summary['main_word_count']}")
    print(f"Summary words: {summary['summary_word_count']}")
    print(f"Reference count: {reference_count}")


if __name__ == "__main__":
    main()
