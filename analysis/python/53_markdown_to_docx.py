#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert simple Markdown to a styled DOCX document.")
    parser.add_argument("--input", required=True)
    parser.add_argument("--output", required=True)
    return parser.parse_args()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.start_type = WD_SECTION.CONTINUOUS

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    normal.font.size = Pt(11)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")

    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(13)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(11.5)
    doc.styles["Heading 3"].font.bold = True


def add_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.3
    paragraph.add_run(text)


def add_list(doc: Document, text: str, style: str) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.add_run(text)


def flush_buffer(doc: Document, buffer: list[str]) -> None:
    if not buffer:
        return
    text = " ".join(part.strip() for part in buffer if part.strip())
    if text:
        add_paragraph(doc, text)
    buffer.clear()


def convert_markdown(markdown_text: str, doc: Document) -> None:
    buffer: list[str] = []
    in_code_block = False

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()

        if line.startswith("```"):
            flush_buffer(doc, buffer)
            in_code_block = not in_code_block
            continue
        if in_code_block:
            paragraph = doc.add_paragraph(style="Normal")
            run = paragraph.add_run(line)
            run.font.name = "Courier New"
            run.font.size = Pt(10)
            continue
        if not line.strip():
            flush_buffer(doc, buffer)
            continue
        if line.startswith("# "):
            flush_buffer(doc, buffer)
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            flush_buffer(doc, buffer)
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            flush_buffer(doc, buffer)
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if re.match(r"^\d+\.\s+", line):
            flush_buffer(doc, buffer)
            add_list(doc, re.sub(r"^\d+\.\s+", "", line), "List Number")
            continue
        if line.startswith("- "):
            flush_buffer(doc, buffer)
            add_list(doc, line[2:].strip(), "List Bullet")
            continue
        buffer.append(line)

    flush_buffer(doc, buffer)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()

    doc = Document()
    configure_document(doc)
    convert_markdown(input_path.read_text(encoding="utf-8"), doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    main()
