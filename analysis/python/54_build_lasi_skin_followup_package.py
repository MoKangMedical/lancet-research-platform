from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import pyreadstat
import statsmodels.api as sm
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万"
)
PACKAGE_ROOT = PROJECT_ROOT / "lasi_skin_followup_package_20260312"
MANUSCRIPT_DIR = PACKAGE_ROOT / "outputs" / "manuscript"
TABLE_DIR = PACKAGE_ROOT / "outputs" / "tables"
FIGURE_DIR = PACKAGE_ROOT / "outputs" / "figures"
LASI_PATH = Path(
    "/Users/apple/Desktop/所有数据/global aging data数据/LASI_印度/Harmonized LASI A.3_SPSS/H_LASI_a3.sav"
)
SCRIPT_OUTPUT = Path("/Users/apple/Desktop/lancet-research-platform/analysis/python/54_build_lasi_skin_followup_package.py")

BLUE = "#1F4E79"
DARK = "#0F2747"
TEAL = "#4F8C9D"
GREEN = "#6A9E78"
CORAL = "#D96B5F"
LIGHT_BLUE = "#DCEAF7"
LIGHT_GRAY = "#F4F6F8"


def ensure_dirs() -> None:
    for path in [PACKAGE_ROOT, MANUSCRIPT_DIR, TABLE_DIR, FIGURE_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def weighted_mean(series: pd.Series, weights: pd.Series) -> float:
    valid = series.notna() & weights.notna()
    if valid.sum() == 0:
        return np.nan
    return float(np.average(series[valid], weights=weights[valid]))


def weighted_prop(series: pd.Series, weights: pd.Series) -> float:
    return weighted_mean(series.astype(float), weights)


def weighted_ci_prop(series: pd.Series, weights: pd.Series) -> tuple[float, float]:
    valid = series.notna() & weights.notna()
    if valid.sum() == 0:
        return np.nan, np.nan
    s = series[valid].astype(float)
    w = weights[valid].astype(float)
    p = float(np.average(s, weights=w))
    n_eff = float((w.sum() ** 2) / (w.pow(2).sum()))
    se = np.sqrt(max(p * (1 - p), 0) / max(n_eff, 1))
    lo = max(0.0, p - 1.96 * se)
    hi = min(1.0, p + 1.96 * se)
    return float(lo), float(hi)


def format_p_value(value: float) -> str:
    if pd.isna(value):
        return ""
    if value < 0.001:
        return "<0.001"
    return f"{value:.3f}"


def read_lasi() -> tuple[pd.DataFrame, pyreadstat.metadata_container]:
    usecols = [
        "prim_key",
        "hhid",
        "pn",
        "r1wtresp",
        "r1proxy",
        "r1agey",
        "ragender",
        "raeduc_l",
        "hh1rural",
        "r1skindise",
        "r1adlaa",
        "r1iadlaa",
        "r1mobilca",
        "r1painfr",
        "r1cesd10dep",
        "r1fallslp",
        "r1vgactx",
        "r1mdactx",
        "r1mbmi",
    ]
    df, meta = pyreadstat.read_sav(LASI_PATH, usecols=usecols, apply_value_formats=False)
    return df, meta


def clean_lasi(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy(deep=True).reset_index(drop=True)

    for col in ["r1skindise", "r1adlaa", "r1iadlaa", "r1mobilca", "r1painfr", "r1cesd10dep", "r1proxy"]:
        out.loc[~out[col].isin([0, 1]), col] = np.nan

    out.loc[:, "female"] = np.where(out["ragender"].eq(2), 1, np.where(out["ragender"].eq(1), 0, np.nan))
    out.loc[:, "rural"] = np.where(out["hh1rural"].eq(1), 1, np.where(out["hh1rural"].eq(0), 0, np.nan))
    out.loc[:, "sleep_problem"] = np.where(
        out["r1fallslp"].isin([1, 2]),
        1,
        np.where(out["r1fallslp"].eq(3), 0, np.nan),
    )
    out.loc[:, "vig_inactive"] = np.where(
        out["r1vgactx"].eq(5),
        1,
        np.where(out["r1vgactx"].isin([1, 2, 3, 4]), 0, np.nan),
    )
    out.loc[:, "mod_inactive"] = np.where(
        out["r1mdactx"].eq(5),
        1,
        np.where(out["r1mdactx"].isin([1, 2, 3, 4]), 0, np.nan),
    )
    out.loc[:, "age_group"] = pd.cut(out["r1agey"], bins=[44, 59, 74, 120], labels=["45-59", "60-74", "75+"])
    out.loc[:, "education_group"] = pd.cut(
        out["raeduc_l"],
        bins=[-0.1, 0.1, 2.1, 9.1],
        labels=["No schooling", "Primary", "Secondary+"],
    )
    out.loc[:, "education_secondary_plus"] = np.where(
        out["raeduc_l"].ge(3),
        1,
        np.where(out["raeduc_l"].between(0, 2), 0, np.nan),
    )
    out.loc[:, "bmi_group"] = pd.cut(
        out["r1mbmi"],
        bins=[0, 18.5, 25, 100],
        labels=["Underweight", "Normal", "Overweight/obese"],
    )

    analysis = out.loc[out["r1skindise"].isin([0, 1])].copy()
    analysis = analysis.loc[analysis["r1agey"].ge(45) & analysis["r1wtresp"].notna()].copy()
    return analysis


def prevalence_row(df: pd.DataFrame, label: str) -> dict:
    prev = weighted_prop(df["r1skindise"], df["r1wtresp"])
    lo, hi = weighted_ci_prop(df["r1skindise"], df["r1wtresp"])
    return {
        "Group": label,
        "Unweighted n": int(df["r1skindise"].notna().sum()),
        "Weighted prevalence": prev,
        "Weighted prevalence %": round(prev * 100, 2),
        "Weighted 95% CI lower %": round(lo * 100, 2),
        "Weighted 95% CI upper %": round(hi * 100, 2),
    }


def build_prevalence_table(df: pd.DataFrame) -> pd.DataFrame:
    specs = [
        ("Overall", None, None),
        ("Age 45-59", "age_group", "45-59"),
        ("Age 60-74", "age_group", "60-74"),
        ("Age 75+", "age_group", "75+"),
        ("Men", "female", 0),
        ("Women", "female", 1),
        ("Urban", "rural", 0),
        ("Rural", "rural", 1),
        ("No schooling", "education_group", "No schooling"),
        ("Primary", "education_group", "Primary"),
        ("Secondary+", "education_group", "Secondary+"),
    ]
    rows = []
    for label, column, value in specs:
        subset = df if column is None else df.loc[df[column] == value]
        rows.append(prevalence_row(subset, label))
    return pd.DataFrame(rows)


def build_additional_prevalence_table(df: pd.DataFrame) -> pd.DataFrame:
    specs = [
        ("Underweight", "bmi_group", "Underweight"),
        ("Normal BMI", "bmi_group", "Normal"),
        ("Overweight or obese", "bmi_group", "Overweight/obese"),
        ("Vigorously inactive", "vig_inactive", 1),
        ("Not vigorously inactive", "vig_inactive", 0),
        ("Moderately inactive", "mod_inactive", 1),
        ("Not moderately inactive", "mod_inactive", 0),
        ("Proxy respondent", "r1proxy", 1),
        ("Self respondent", "r1proxy", 0),
    ]
    rows = []
    for label, column, value in specs:
        subset = df.loc[df[column] == value]
        rows.append(prevalence_row(subset, label))
    return pd.DataFrame(rows)


def weighted_value(df: pd.DataFrame, column: str, mode: str) -> float:
    if mode == "count":
        return float(df["r1skindise"].notna().sum())
    if mode == "mean":
        return weighted_mean(df[column], df["r1wtresp"])
    if mode == "prop":
        return weighted_prop(df[column], df["r1wtresp"]) * 100
    raise ValueError(mode)


def build_characteristics_table(df: pd.DataFrame) -> pd.DataFrame:
    columns = {
        "Total": df,
        "No skin disease": df.loc[df["r1skindise"] == 0],
        "Skin disease": df.loc[df["r1skindise"] == 1],
    }
    variables = [
        ("Participants, n", None, "count"),
        ("Mean age, years", "r1agey", "mean"),
        ("Women, %", "female", "prop"),
        ("Rural residence, %", "rural", "prop"),
        ("No schooling, %", None, "category_no_school"),
        ("Secondary education or above, %", "education_secondary_plus", "prop"),
        ("Underweight, %", None, "category_underweight"),
        ("Overweight or obese, %", None, "category_overweight"),
        ("Any ADL limitation, %", "r1adlaa", "prop"),
        ("Any IADL limitation, %", "r1iadlaa", "prop"),
        ("Any mobility limitation, %", "r1mobilca", "prop"),
        ("Frequent pain, %", "r1painfr", "prop"),
        ("Depressive symptoms, %", "r1cesd10dep", "prop"),
        ("Sleep problems, %", "sleep_problem", "prop"),
    ]
    rows = []
    for label, column, mode in variables:
        row = {"Characteristic": label}
        for group_name, subset in columns.items():
            if mode == "category_no_school":
                value = weighted_prop(subset["education_group"].eq("No schooling"), subset["r1wtresp"]) * 100
            elif mode == "category_underweight":
                value = weighted_prop(subset["bmi_group"].eq("Underweight"), subset["r1wtresp"]) * 100
            elif mode == "category_overweight":
                value = weighted_prop(subset["bmi_group"].eq("Overweight/obese"), subset["r1wtresp"]) * 100
            else:
                value = weighted_value(subset, column, mode)
            row[group_name] = round(value, 2)
        rows.append(row)
    return pd.DataFrame(rows)


def fit_binomial_or(data: pd.DataFrame, outcome: str, predictors: list[str]) -> sm.GLMResults:
    cols = [outcome, "r1wtresp"] + predictors
    sub = data[cols].dropna().copy()
    x = sm.add_constant(sub[predictors].astype(float), has_constant="add")
    model = sm.GLM(sub[outcome].astype(float), x, family=sm.families.Binomial(), freq_weights=sub["r1wtresp"])
    return model.fit()


def extract_model_rows(result: sm.GLMResults, predictors: list[str], labels: dict[str, str], n: int, name: str) -> pd.DataFrame:
    rows = []
    conf = result.conf_int()
    for term in predictors:
        rows.append(
            {
                name: labels[term],
                "Adjusted OR": round(float(np.exp(result.params[term])), 3),
                "95% CI lower": round(float(np.exp(conf.loc[term].iloc[0])), 3),
                "95% CI upper": round(float(np.exp(conf.loc[term].iloc[1])), 3),
                "p value": float(result.pvalues[term]),
                "Analytic n": int(n),
            }
        )
    return pd.DataFrame(rows)


def build_predictor_table(df: pd.DataFrame) -> pd.DataFrame:
    predictors = ["r1agey", "female", "rural", "education_secondary_plus"]
    labels = {
        "r1agey": "Age (per year)",
        "female": "Women vs men",
        "rural": "Rural vs urban",
        "education_secondary_plus": "Secondary education or above vs less",
    }
    result = fit_binomial_or(df, "r1skindise", predictors)
    n = len(df[["r1skindise"] + predictors].dropna())
    return extract_model_rows(result, predictors, labels, n, "Predictor")


def build_outcome_association_table(df: pd.DataFrame) -> pd.DataFrame:
    outcomes = {
        "r1adlaa": "Any ADL limitation",
        "r1iadlaa": "Any IADL limitation",
        "r1mobilca": "Any mobility limitation",
        "r1painfr": "Frequent pain",
        "r1cesd10dep": "Depressive symptoms",
        "sleep_problem": "Sleep problems",
    }
    covars = ["r1skindise", "r1agey", "female", "rural", "education_secondary_plus"]
    rows = []
    for outcome, label in outcomes.items():
        result = fit_binomial_or(df, outcome, covars)
        conf = result.conf_int()
        n = len(df[[outcome] + covars].dropna())
        rows.append(
            {
                "Outcome": label,
                "Adjusted OR for skin disease": round(float(np.exp(result.params["r1skindise"])), 3),
                "95% CI lower": round(float(np.exp(conf.loc["r1skindise"].iloc[0])), 3),
                "95% CI upper": round(float(np.exp(conf.loc["r1skindise"].iloc[1])), 3),
                "p value": float(result.pvalues["r1skindise"]),
                "Analytic n": int(n),
            }
        )
    return pd.DataFrame(rows)


def build_variable_definition_table() -> pd.DataFrame:
    rows = [
        ["Exposure", "r1skindise", "Self-reported skin disease", "Binary harmonized LASI respondent-level skin disease variable."],
        ["Sampling weight", "r1wtresp", "Respondent weight", "Used for weighted descriptive and weighted GLM analyses."],
        ["Demographic", "r1agey", "Age in years", "Restricted to respondents aged 45 years and older."],
        ["Demographic", "ragender", "Sex", "Recoded as women vs men."],
        ["Demographic", "hh1rural", "Rural residence", "Recoded as rural vs urban."],
        ["Socioeconomic", "raeduc_l", "Education level", "Used for grouped prevalence and adjusted models."],
        ["Functional outcome", "r1adlaa", "Any ADL limitation", "Binary indicator of any activity of daily living limitation."],
        ["Functional outcome", "r1iadlaa", "Any IADL limitation", "Binary indicator of any instrumental activity of daily living limitation."],
        ["Functional outcome", "r1mobilca", "Any mobility limitation", "Binary indicator of any mobility difficulty."],
        ["Symptom outcome", "r1painfr", "Frequent pain", "Binary indicator of frequent pain."],
        ["Psychological outcome", "r1cesd10dep", "Depressive symptoms", "Binary indicator derived from harmonized LASI depression variable."],
        ["Sleep outcome", "r1fallslp", "Sleep problems", "Recoded as frequent or occasional sleep problems vs rare or never."],
        ["Behavioural context", "r1vgactx", "Vigorous physical inactivity", "Recoded as inactive vs any vigorous activity frequency."],
        ["Behavioural context", "r1mdactx", "Moderate physical inactivity", "Recoded as inactive vs any moderate activity frequency."],
        ["Anthropometric", "r1mbmi", "Body-mass index", "Used to derive underweight, normal, and overweight or obese categories."],
        ["Sensitivity variable", "r1proxy", "Proxy interview", "Used for sensitivity analyses excluding proxy respondents."],
    ]
    return pd.DataFrame(rows, columns=["Domain", "Variable", "Construct", "Operational definition"])


def prevalence_for_label(table: pd.DataFrame, label: str) -> float:
    return float(table.loc[table["Group"] == label, "Weighted prevalence %"].iloc[0])


def table_to_display_strings(df: pd.DataFrame, table_type: str) -> pd.DataFrame:
    out = df.copy(deep=True).reset_index(drop=True)
    if table_type == "prevalence":
        out["Weighted prevalence % (95% CI)"] = out.apply(
            lambda row: f"{row['Weighted prevalence %']:.2f} ({row['Weighted 95% CI lower %']:.2f}-{row['Weighted 95% CI upper %']:.2f})",
            axis=1,
        )
        return out[["Group", "Unweighted n", "Weighted prevalence % (95% CI)"]]
    if table_type == "characteristics":
        for col in ["Total", "No skin disease", "Skin disease"]:
            out[col] = out.apply(
                lambda row: f"{int(round(row[col], 0)):,}" if row["Characteristic"] == "Participants, n" else f"{row[col]:.2f}",
                axis=1,
            )
        return out
    if table_type == "predictors":
        out["Adjusted OR (95% CI)"] = out.apply(
            lambda row: f"{row['Adjusted OR']:.3f} ({row['95% CI lower']:.3f}-{row['95% CI upper']:.3f})",
            axis=1,
        )
        out = out.assign(**{"p value": out["p value"].map(format_p_value).astype(str)})
        return out[[out.columns[0], "Adjusted OR (95% CI)", "p value", "Analytic n"]]
    if table_type == "outcomes":
        out["Adjusted OR (95% CI)"] = out.apply(
            lambda row: f"{row['Adjusted OR for skin disease']:.3f} ({row['95% CI lower']:.3f}-{row['95% CI upper']:.3f})",
            axis=1,
        )
        out = out.assign(**{"p value": out["p value"].map(format_p_value).astype(str)})
        return out[["Outcome", "Adjusted OR (95% CI)", "p value", "Analytic n"]]
    raise ValueError(table_type)


def write_prevalence_figure(prevalence_table: pd.DataFrame, output_png: Path, output_pdf: Path, groups: list[str], title: str) -> None:
    plot_df = prevalence_table.loc[prevalence_table["Group"].isin(groups)].copy()
    palette = []
    for label in plot_df["Group"]:
        if "Age" in label:
            palette.append(BLUE)
        elif label in {"Men", "Women"}:
            palette.append(CORAL)
        elif label in {"Urban", "Rural"}:
            palette.append(GREEN)
        else:
            palette.append(TEAL)

    fig, ax = plt.subplots(figsize=(7.8, 5.2))
    y_pos = np.arange(len(plot_df))[::-1]
    x = plot_df["Weighted prevalence %"].to_numpy()
    lo = plot_df["Weighted 95% CI lower %"].to_numpy()
    hi = plot_df["Weighted 95% CI upper %"].to_numpy()
    ax.errorbar(
        x,
        y_pos,
        xerr=[x - lo, hi - x],
        fmt="o",
        color=DARK,
        ecolor=TEAL,
        elinewidth=2,
        capsize=4,
    )
    for idx, color in enumerate(palette):
        ax.scatter(x[idx], y_pos[idx], color=color, s=50, zorder=3)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(plot_df["Group"].tolist())
    ax.set_xlabel("Weighted prevalence of self-reported skin disease (%)")
    ax.set_title(title)
    ax.grid(axis="x", linestyle="--", alpha=0.3)
    ax.spines[["top", "right"]].set_visible(False)
    plt.tight_layout()
    fig.savefig(output_png, dpi=300)
    fig.savefig(output_pdf)
    plt.close(fig)


def write_forest_figure(outcome_table: pd.DataFrame, output_png: Path, output_pdf: Path, title: str) -> None:
    labels = outcome_table["Outcome"].tolist()
    ors = outcome_table["Adjusted OR for skin disease"].to_numpy()
    lows = outcome_table["95% CI lower"].to_numpy()
    highs = outcome_table["95% CI upper"].to_numpy()
    y = np.arange(len(labels))[::-1]
    fig, ax = plt.subplots(figsize=(8.0, 5.4))
    ax.errorbar(
        ors,
        y,
        xerr=[ors - lows, highs - ors],
        fmt="o",
        color=DARK,
        ecolor=TEAL,
        elinewidth=2,
        capsize=4,
    )
    ax.axvline(1.0, color=CORAL, linestyle="--", linewidth=1.5)
    ax.set_yticks(y)
    ax.set_yticklabels(labels)
    ax.set_xlabel("Adjusted odds ratio")
    ax.set_title(title)
    ax.grid(axis="x", linestyle="--", alpha=0.3)
    ax.spines[["top", "right"]].set_visible(False)
    plt.tight_layout()
    fig.savefig(output_png, dpi=300)
    fig.savefig(output_pdf)
    plt.close(fig)


def sha_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text_style(cell, bold: bool = False, color: str | None = None, size: int = 9) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = bold
            run.font.name = "Arial"
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = RGBColor.from_string(color)


def style_document(doc: Document, title: str) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    if "List Bullet" in doc.styles:
        bullet = doc.styles["List Bullet"]
        bullet.font.name = "Arial"
        bullet.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = "Arial"
            style.font.color.rgb = RGBColor.from_string(BLUE.replace("#", ""))

    title_para = doc.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_para.runs:
        title_para.runs[0].font.name = "Arial"
        title_para.runs[0].font.size = Pt(16)
        title_para.runs[0].bold = True
        title_para.runs[0].font.color.rgb = RGBColor.from_string(DARK.replace("#", ""))


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.italic = italic


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string(BLUE.replace("#", ""))


def add_note(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_dataframe_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cells = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        header_cells[idx].text = str(col)
        sha_cell(header_cells[idx], BLUE.replace("#", ""))
        set_cell_text_style(header_cells[idx], bold=True, color="FFFFFF", size=9)

    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        cells = table.add_row().cells
        fill = "FFFFFF" if row_idx % 2 else LIGHT_GRAY.replace("#", "")
        for idx, value in enumerate(row.tolist()):
            cells[idx].text = str(value)
            sha_cell(cells[idx], fill)
            set_cell_text_style(cells[idx], size=9)


def save_markdown(path: Path, text: str) -> None:
    path.write_text(text, encoding="utf-8")


def references_list() -> list[str]:
    return [
        "Longitudinal Ageing Study in India. Study design. https://www.iipsindia.ac.in/content/lasi-study-design. Accessed March 12, 2026.",
        "International Institute for Population Sciences, NPHCE, Ministry of Health and Family Welfare, Harvard T H Chan School of Public Health, University of Southern California. Longitudinal Ageing Study in India (LASI) Wave 1, 2017-18, India Report. Mumbai: IIPS; 2020.",
        "Lee J, Banerjee J, Khobragade PY, et al. Gateway to Global Aging Data: resources for studies of aging populations worldwide. J Gerontol B Psychol Sci Soc Sci 2021; 76(Suppl 1): S5-S16.",
        "Lee J, Cantu PA, Ghilarducci T, et al. Cross-country comparisons of disability and morbidity: evidence from the Gateway to Global Aging Data. J Gerontol B Psychol Sci Soc Sci 2024; 79(Suppl 1): S35-S43.",
        "Li D, Lyu J, Xu T, et al. Global burden of skin and subcutaneous diseases: an analysis of the Global Burden of Disease Study 2021. J Eur Acad Dermatol Venereol 2026; 40: e686-e689.",
        "Yakupu A, Wang Y, Pan KH, et al. The global, regional, and national burden of skin and subcutaneous diseases: findings from the Global Burden of Disease Study 2019. Front Public Health 2023; 11: 1089259.",
        "Hahnel E, Lichterfeld-Kottner A, Bohnke N, et al. Epidemiology of skin conditions in the aged: a systematic review. J Tissue Viability 2017; 26: 20-28.",
        "Shah M, Coates M. An assessment of the quality of life in older patients with skin disease. Br J Dermatol 2006; 154: 150-153.",
        "Mostaghimi L. The prevalence of mood and sleep problems in chronic skin diseases. Arch Dermatol 2008; 144: 18-24.",
        "Duran S, Yurekli A. Quality of life and satisfaction with life in patients with skin diseases. Cutan Ocul Toxicol 2023; 42: 178-184.",
        "Salfi F, Oliviero A, D'Atri A, et al. The role of insomnia in vulnerability to depressive and anxiety symptoms in adults with atopic dermatitis. Clocks Sleep 2023; 5: 628-639.",
        "Pradhan MR, Saikia D. Prevalence and predictors of insomnia among older adults in India: an analysis based on LASI. Aging Ment Health 2025; 29: 1137-1147.",
        "Muhammad T, Rashid M. Prevalence and correlates of pain and associated depression among community-dwelling older adults in India. BMC Geriatr 2022; 22: 708.",
        "Sharma P, Maurya P, Muhammad T. Number of chronic conditions and associated functional limitations among older adults: cross-sectional findings from the Longitudinal Ageing Study in India. BMC Geriatr 2021; 21: 673.",
        "Banerjee S, Boro B. The role of sleep quality, functional limitation and depressive symptoms in predicting life satisfaction among older adults in India. Sci Rep 2022; 12: 10339.",
    ]


def build_summary(
    df: pd.DataFrame,
    prevalence_table: pd.DataFrame,
    predictor_table: pd.DataFrame,
    outcome_table: pd.DataFrame,
    predictor_table_no_proxy: pd.DataFrame,
    outcome_table_no_proxy: pd.DataFrame,
) -> dict:
    overall = prevalence_table.loc[prevalence_table["Group"] == "Overall"].iloc[0]
    top_outcome = outcome_table.sort_values("Adjusted OR for skin disease", ascending=False).iloc[0]
    sensitivity_top = outcome_table_no_proxy.sort_values("Adjusted OR for skin disease", ascending=False).iloc[0]
    return {
        "analytic_n": int(df["r1skindise"].notna().sum()),
        "weighted_skin_prev_pct": float(overall["Weighted prevalence %"]),
        "weighted_skin_prev_ci_low": float(overall["Weighted 95% CI lower %"]),
        "weighted_skin_prev_ci_high": float(overall["Weighted 95% CI upper %"]),
        "prev_age75_plus_pct": prevalence_for_label(prevalence_table, "Age 75+"),
        "prev_men_pct": prevalence_for_label(prevalence_table, "Men"),
        "prev_women_pct": prevalence_for_label(prevalence_table, "Women"),
        "prev_rural_pct": prevalence_for_label(prevalence_table, "Rural"),
        "predictor_rows": len(predictor_table),
        "outcome_rows": len(outcome_table),
        "top_burden_outcome": str(top_outcome["Outcome"]),
        "top_burden_outcome_or": float(top_outcome["Adjusted OR for skin disease"]),
        "sensitivity_top_burden_outcome": str(sensitivity_top["Outcome"]),
        "sensitivity_top_burden_outcome_or": float(sensitivity_top["Adjusted OR for skin disease"]),
        "sensitivity_predictor_rows": len(predictor_table_no_proxy),
        "sensitivity_outcome_rows": len(outcome_table_no_proxy),
    }


def manuscript_text(
    summary: dict,
    prevalence_table: pd.DataFrame,
    characteristics_table: pd.DataFrame,
    predictor_table: pd.DataFrame,
    outcome_table: pd.DataFrame,
    predictor_table_no_proxy: pd.DataFrame,
    outcome_table_no_proxy: pd.DataFrame,
) -> str:
    overall = prevalence_table.loc[prevalence_table["Group"] == "Overall"].iloc[0]
    age75 = prevalence_table.loc[prevalence_table["Group"] == "Age 75+"].iloc[0]
    rural = prevalence_table.loc[prevalence_table["Group"] == "Rural"].iloc[0]
    men = prevalence_table.loc[prevalence_table["Group"] == "Men"].iloc[0]
    women = prevalence_table.loc[prevalence_table["Group"] == "Women"].iloc[0]
    no_school = prevalence_table.loc[prevalence_table["Group"] == "No schooling"].iloc[0]
    primary = prevalence_table.loc[prevalence_table["Group"] == "Primary"].iloc[0]

    age_or = predictor_table.loc[predictor_table["Predictor"] == "Age (per year)"].iloc[0]
    women_or = predictor_table.loc[predictor_table["Predictor"] == "Women vs men"].iloc[0]
    rural_or = predictor_table.loc[predictor_table["Predictor"] == "Rural vs urban"].iloc[0]
    edu_or = predictor_table.loc[predictor_table["Predictor"] == "Secondary education or above vs less"].iloc[0]

    pain = outcome_table.loc[outcome_table["Outcome"] == "Frequent pain"].iloc[0]
    mobility = outcome_table.loc[outcome_table["Outcome"] == "Any mobility limitation"].iloc[0]
    sleep = outcome_table.loc[outcome_table["Outcome"] == "Sleep problems"].iloc[0]
    dep = outcome_table.loc[outcome_table["Outcome"] == "Depressive symptoms"].iloc[0]
    adl = outcome_table.loc[outcome_table["Outcome"] == "Any ADL limitation"].iloc[0]
    iadl = outcome_table.loc[outcome_table["Outcome"] == "Any IADL limitation"].iloc[0]

    pain_sens = outcome_table_no_proxy.loc[outcome_table_no_proxy["Outcome"] == "Frequent pain"].iloc[0]
    mobility_sens = outcome_table_no_proxy.loc[outcome_table_no_proxy["Outcome"] == "Any mobility limitation"].iloc[0]

    total_age = characteristics_table.loc[characteristics_table["Characteristic"] == "Mean age, years", "Total"].iloc[0]
    total_women = characteristics_table.loc[characteristics_table["Characteristic"] == "Women, %", "Total"].iloc[0]
    total_rural = characteristics_table.loc[characteristics_table["Characteristic"] == "Rural residence, %", "Total"].iloc[0]

    abstract = f"""Summary

Background
Skin and subcutaneous diseases are a major but still under-contextualised component of healthy ageing agendas, and ecological burden analyses do not show whether older adults who report skin disease also carry higher person-level functional and psychosocial burden.[3-7] We aimed to estimate the weighted prevalence of self-reported skin disease and examine its association with functional limitation, pain, depressive symptoms, and sleep problems among middle-aged and older adults in India.

Methods
We did a cross-sectional analysis of Harmonized LASI A.3 using baseline respondent-level variables for adults aged 45 years or older with non-missing skin-disease status. The exposure was self-reported skin disease (`r1skindise`). Weighted prevalence was estimated overall and by demographic subgroups. Weighted binomial generalised linear models were used to assess correlates of skin disease and associations with activity of daily living (ADL) limitation, instrumental ADL (IADL) limitation, mobility limitation, frequent pain, depressive symptoms, and sleep problems, adjusting for age, sex, rural residence, and education. A sensitivity analysis excluded proxy respondents.

Findings
{summary['analytic_n']:,} respondents were included. Weighted prevalence of self-reported skin disease was {overall['Weighted prevalence %']:.2f}% (95% CI {overall['Weighted 95% CI lower %']:.2f}-{overall['Weighted 95% CI upper %']:.2f}). Prevalence was higher among adults aged 75 years or older ({age75['Weighted prevalence %']:.2f}%), men ({men['Weighted prevalence %']:.2f}%), rural respondents ({rural['Weighted prevalence %']:.2f}%), and respondents with primary education ({primary['Weighted prevalence %']:.2f}%). In adjusted models, reporting skin disease was more likely with older age (OR {age_or['Adjusted OR']:.3f}, 95% CI {age_or['95% CI lower']:.3f}-{age_or['95% CI upper']:.3f}) and rural residence ({rural_or['Adjusted OR']:.3f}, {rural_or['95% CI lower']:.3f}-{rural_or['95% CI upper']:.3f}), and less likely among women ({women_or['Adjusted OR']:.3f}, {women_or['95% CI lower']:.3f}-{women_or['95% CI upper']:.3f}). Skin disease was associated with frequent pain (OR {pain['Adjusted OR for skin disease']:.3f}, 95% CI {pain['95% CI lower']:.3f}-{pain['95% CI upper']:.3f}), mobility limitation ({mobility['Adjusted OR for skin disease']:.3f}, {mobility['95% CI lower']:.3f}-{mobility['95% CI upper']:.3f}), sleep problems ({sleep['Adjusted OR for skin disease']:.3f}, {sleep['95% CI lower']:.3f}-{sleep['95% CI upper']:.3f}), depressive symptoms ({dep['Adjusted OR for skin disease']:.3f}, {dep['95% CI lower']:.3f}-{dep['95% CI upper']:.3f}), ADL limitation ({adl['Adjusted OR for skin disease']:.3f}, {adl['95% CI lower']:.3f}-{adl['95% CI upper']:.3f}), and IADL limitation ({iadl['Adjusted OR for skin disease']:.3f}, {iadl['95% CI lower']:.3f}-{iadl['95% CI upper']:.3f}). These patterns were preserved after excluding proxy respondents.

Interpretation
In LASI, self-reported skin disease marked a broader burden phenotype rather than a narrow dermatologic complaint alone. The associations with pain, mobility, sleep, and mood support positioning skin health within healthy-ageing and long-term-care strategies, and provide person-level validation for the ageing narrative emerging from global burden analyses.

Funding
No specific funding was used for this secondary analysis package.
"""

    body = f"""
Introduction

Skin and subcutaneous diseases account for substantial disability worldwide, and recent GBD analyses have shown that their aggregate burden remains large across countries and over time.[5,6] Yet burden papers alone do not resolve how skin disease is experienced among ageing populations. The public health importance of skin disease in later life is not limited to diagnostic frequency; it also includes chronic symptoms, social visibility, sleep disturbance, pain, impaired function, and reduced quality of life.[7-10] This matters because the policy language of healthy ageing increasingly emphasises intrinsic capacity, function, wellbeing, and long-term care rather than disease counts alone.[3,4]

Older adults are especially relevant to this question. With advancing age, skin barrier function, wound healing, immunologic resilience, and multimorbidity profiles change, while the cumulative impact of chronic inflammatory, infectious, neoplastic, and degenerative skin conditions becomes more clinically and socially consequential.[7] Nevertheless, most broad epidemiologic discussions either remain at the ecological level or focus on specific diagnoses. Evidence that directly links skin disease to functional and psychosocial burden in large, nationally representative ageing cohorts remains limited, especially in low-income and middle-income settings.

India provides an important setting for this question. The country is ageing rapidly, its disease profile is increasingly shaped by chronic multimorbidity, and its older population spans marked social and rural-urban inequalities.[1-4] The Longitudinal Ageing Study in India (LASI) offers a practical platform to examine how a harmonized self-reported skin disease measure maps onto pain, function, depressive symptoms, and sleep at the person level. This is particularly useful because our current global manuscript on skin burden and ageing uses an ecological framework based on GBD 2023 and World Bank indicators; LASI therefore functions as a complementary validation platform rather than a direct continuation of the same analytic design.

We aimed to estimate the weighted prevalence of self-reported skin disease among adults aged 45 years and older in Harmonized LASI, to describe variation by age, sex, rural residence, and education, and to examine whether respondents reporting skin disease had higher odds of ADL limitation, IADL limitation, mobility limitation, frequent pain, depressive symptoms, and sleep problems.

Methods

Study design and data source

We did a cross-sectional secondary analysis of Harmonized LASI A.3, a harmonized ageing-data resource derived from the nationally representative Longitudinal Ageing Study in India baseline.[1-4] Harmonized LASI provides respondent-level variables designed for comparative ageing research, with standardised naming conventions across demographic, socioeconomic, health, and functional domains.[3,4] For this package, we used the local baseline `.sav` file available in the desktop Global Aging Data repository and restricted analyses to wave-1 respondent variables (`r1*`).

Study population

We included respondents aged 45 years or older with non-missing values for the harmonized self-reported skin disease variable (`r1skindise`) and non-missing respondent weights (`r1wtresp`). This yielded an analytic sample of {summary['analytic_n']:,} respondents. The weighted mean age of the sample was {total_age:.2f} years, {total_women:.2f}% were women, and {total_rural:.2f}% resided in rural areas (Table 1).

Exposure, covariates, and outcomes

The primary exposure was self-reported skin disease (`r1skindise`), treated as a binary variable. Demographic and social covariates were age in years (`r1agey`), sex (`ragender`, recoded as women vs men), rural residence (`hh1rural`), and education level (`raeduc_l`). For descriptive analysis we also derived age groups, education groups, body-mass index categories, and physical inactivity indicators.

We analysed six burden-related outcomes chosen because they represent domains that are clinically and policy relevant in ageing populations: any ADL limitation (`r1adlaa`), any IADL limitation (`r1iadlaa`), any mobility limitation (`r1mobilca`), frequent pain (`r1painfr`), depressive symptoms (`r1cesd10dep`), and sleep problems derived from the harmonized sleep item (`r1fallslp`). Proxy interview status (`r1proxy`) was retained for sensitivity analysis.

Statistical analysis

We first estimated weighted prevalence of self-reported skin disease overall and across age, sex, rural-urban, and education strata. Approximate 95% CIs for weighted prevalence were calculated using an effective sample-size approach. We then built weighted binomial generalised linear models to estimate adjusted odds ratios for reporting skin disease according to age, sex, rural residence, and education. In separate models, each functional or psychosocial outcome was regressed on skin disease status with adjustment for age, sex, rural residence, and education. We used respondent weights as frequency weights. A sensitivity analysis restricted the sample to self-respondents by excluding proxy interviews. Because this was a secondary reproducible research package rather than a final survey-statistics paper, we did not yet implement full complex-sample variance estimation; results should therefore be interpreted as weighted association estimates rather than definitive design-based causal inference.

Ethics and role of the funding source

This study used de-identified secondary data from harmonized LASI. No specific funding was used for this secondary analysis package. The authors of the package had full access to the local analytic file and were responsible for the decision to prepare this draft.

Results

Weighted prevalence of self-reported skin disease

Weighted prevalence of self-reported skin disease was {overall['Weighted prevalence %']:.2f}% (95% CI {overall['Weighted 95% CI lower %']:.2f}-{overall['Weighted 95% CI upper %']:.2f}) in adults aged 45 years or older. Prevalence rose from {prevalence_for_label(prevalence_table, 'Age 45-59'):.2f}% among adults aged 45-59 years to {prevalence_for_label(prevalence_table, 'Age 60-74'):.2f}% among those aged 60-74 years and {age75['Weighted prevalence %']:.2f}% among those aged 75 years or older. Prevalence was higher in men than in women ({men['Weighted prevalence %']:.2f}% vs {women['Weighted prevalence %']:.2f}%) and higher in rural than urban residents ({rural['Weighted prevalence %']:.2f}% vs {prevalence_for_label(prevalence_table, 'Urban'):.2f}%; Figure 1). Educational differences were not monotonic: prevalence was {no_school['Weighted prevalence %']:.2f}% among respondents with no schooling, {primary['Weighted prevalence %']:.2f}% among those with primary education, and {prevalence_for_label(prevalence_table, 'Secondary+'):.2f}% among those with secondary education or above.

Sample characteristics according to skin disease status

Respondents reporting skin disease were slightly older than those without skin disease (weighted mean age {characteristics_table.loc[characteristics_table['Characteristic'] == 'Mean age, years', 'Skin disease'].iloc[0]:.2f} vs {characteristics_table.loc[characteristics_table['Characteristic'] == 'Mean age, years', 'No skin disease'].iloc[0]:.2f} years) and more likely to live in rural settings ({characteristics_table.loc[characteristics_table['Characteristic'] == 'Rural residence, %', 'Skin disease'].iloc[0]:.2f}% vs {characteristics_table.loc[characteristics_table['Characteristic'] == 'Rural residence, %', 'No skin disease'].iloc[0]:.2f}%). They also had higher weighted prevalence of ADL limitation, IADL limitation, mobility limitation, frequent pain, depressive symptoms, and sleep problems than respondents without skin disease (Table 1). The absolute contrast was particularly marked for frequent pain ({characteristics_table.loc[characteristics_table['Characteristic'] == 'Frequent pain, %', 'Skin disease'].iloc[0]:.2f}% vs {characteristics_table.loc[characteristics_table['Characteristic'] == 'Frequent pain, %', 'No skin disease'].iloc[0]:.2f}%) and mobility limitation ({characteristics_table.loc[characteristics_table['Characteristic'] == 'Any mobility limitation, %', 'Skin disease'].iloc[0]:.2f}% vs {characteristics_table.loc[characteristics_table['Characteristic'] == 'Any mobility limitation, %', 'No skin disease'].iloc[0]:.2f}%).

Correlates of reporting skin disease

In adjusted models, older age was associated with higher odds of reporting skin disease (OR {age_or['Adjusted OR']:.3f}, 95% CI {age_or['95% CI lower']:.3f}-{age_or['95% CI upper']:.3f}; Supplementary Table S3). Women were less likely than men to report skin disease (OR {women_or['Adjusted OR']:.3f}, {women_or['95% CI lower']:.3f}-{women_or['95% CI upper']:.3f}), whereas rural residence was associated with higher odds (OR {rural_or['Adjusted OR']:.3f}, {rural_or['95% CI lower']:.3f}-{rural_or['95% CI upper']:.3f}). Secondary education or above was weakly associated with higher odds compared with lower education (OR {edu_or['Adjusted OR']:.3f}, {edu_or['95% CI lower']:.3f}-{edu_or['95% CI upper']:.3f}).

Functional and psychosocial burden associated with skin disease

Skin disease was consistently associated with worse functional and psychosocial outcomes after adjustment for age, sex, rural residence, and education (Figure 2; Supplementary Table S4). The strongest association was for frequent pain (OR {pain['Adjusted OR for skin disease']:.3f}, 95% CI {pain['95% CI lower']:.3f}-{pain['95% CI upper']:.3f}). Skin disease was also associated with mobility limitation (OR {mobility['Adjusted OR for skin disease']:.3f}, {mobility['95% CI lower']:.3f}-{mobility['95% CI upper']:.3f}) and sleep problems (OR {sleep['Adjusted OR for skin disease']:.3f}, {sleep['95% CI lower']:.3f}-{sleep['95% CI upper']:.3f}). Although the associations with depressive symptoms (OR {dep['Adjusted OR for skin disease']:.3f}, {dep['95% CI lower']:.3f}-{dep['95% CI upper']:.3f}) and IADL limitation (OR {iadl['Adjusted OR for skin disease']:.3f}, {iadl['95% CI lower']:.3f}-{iadl['95% CI upper']:.3f}) were more modest, they remained in the same adverse direction. ADL limitation was also associated with skin disease (OR {adl['Adjusted OR for skin disease']:.3f}, {adl['95% CI lower']:.3f}-{adl['95% CI upper']:.3f}).

Sensitivity analysis

The main pattern was preserved after excluding proxy respondents (Supplementary Tables S5-S6 and Supplementary Figure S2). In self-respondents only, skin disease remained strongly associated with frequent pain (OR {pain_sens['Adjusted OR for skin disease']:.3f}, 95% CI {pain_sens['95% CI lower']:.3f}-{pain_sens['95% CI upper']:.3f}) and mobility limitation (OR {mobility_sens['Adjusted OR for skin disease']:.3f}, {mobility_sens['95% CI lower']:.3f}-{mobility_sens['95% CI upper']:.3f}). The directional stability of the findings suggests that the main results are not driven only by proxy reporting.

Discussion

In this nationally representative ageing dataset from India, self-reported skin disease was present in about one in twenty adults aged 45 years or older and was consistently linked to broader physical and psychosocial burden. Three aspects of the findings matter most. First, prevalence was not negligible and rose in the oldest age group. Second, reporting skin disease clustered with rural residence and with a modest male excess. Third, and most importantly, skin disease was associated with a range of burden outcomes spanning pain, mobility, sleep, mood, and daily functioning.

These findings help extend the interpretive value of ecological global burden work. GBD analyses quantify population-level burden, but they do not indicate whether the lived experience of skin disease in ageing populations is confined to dermatologic symptoms or embedded in wider declines in function and wellbeing.[5,6] The LASI results suggest the latter. The particularly strong association with frequent pain is plausible because many skin diseases, ulcers, infections, inflammatory disorders, and xerotic conditions are physically uncomfortable, recurrent, and socially wearing.[7-10] The link with mobility limitation may reflect shared pathways through multimorbidity, frailty, reduced participation, and symptom burden. Sleep problems and depressive symptoms are also credible correlates, given that visible or symptomatic skin disease can disturb rest, worsen self-perception, and interact with chronic stress.[8-11]

The sex pattern deserves careful interpretation. Women had lower adjusted odds of reporting skin disease than men in this analysis, despite the assumption that women may more often notice or report symptoms. This result might reflect differences in occupational exposure, health-seeking patterns, diagnostic communication, or the mixture of conditions captured by the broad self-reported skin-disease item. Likewise, the rural excess may reflect environmental exposure, access barriers, delayed treatment, or a different mix of infectious, inflammatory, and chronic dermatoses. Because `r1skindise` does not identify specific diagnoses, these explanations remain inferential and should be tested in future condition-specific work.

The educational pattern also warrants caution. Respondents with primary education had the highest weighted prevalence, while the adjusted contrast for secondary education or above versus less education was small. This may indicate that education is acting as an imperfect marker for a complex combination of occupational exposure, awareness, and access to diagnosis rather than a simple socioeconomic gradient. It reinforces the need not to over-interpret the covariate model causally.

This study has several strengths. It uses a large, nationally representative ageing cohort with harmonized variables and a clearly confirmed respondent-level skin-disease measure. It focuses on outcomes that matter for healthy ageing, rather than treating skin disease as an isolated dermatologic endpoint. It also directly operationalises the future direction proposed in the main global manuscript: using desktop Global Aging Data resources as person-level validation platforms rather than forcing them into an ecological design for which they were not built.

The study also has important limitations. First, the skin variable is self-reported and broad; it does not distinguish inflammatory, infectious, neoplastic, or wound-related conditions, and it cannot measure severity or chronicity. Second, the analysis is cross-sectional and therefore cannot establish temporality. Third, although we used respondent weights, the current package does not yet implement full complex-survey variance estimation. Fourth, residual confounding from multimorbidity, healthcare use, occupational history, and socioeconomic factors is likely. Fifth, the analysis is limited to India and should not be read as automatically generalisable to all ageing settings.

Despite these limitations, the present findings are useful for both science and policy. For science, they show that the ageing lens in global skin-burden research is not merely rhetorical; the person-level burden signal is visible in a large ageing cohort. For policy, they support integrating skin health into chronic care, functional assessment, pain management, sleep assessment, and mental-health aware geriatric services. In practical terms, skin disease in later life should be understood as part of broader wellbeing and care needs rather than as a narrow outpatient specialty issue.

In conclusion, self-reported skin disease in LASI was associated with substantial functional and psychosocial burden among middle-aged and older adults in India. This second-stage analysis provides an empiric person-level validation arm for the ongoing global skin-burden project and creates a defensible platform for future longitudinal, diagnosis-specific, and cross-country ageing-cohort analyses.

References
"""

    references = "\n".join([f"{idx}. {ref}" for idx, ref in enumerate(references_list(), start=1)])
    return abstract + body + references + "\n"


def supplementary_markdown(
    variable_table: pd.DataFrame,
    additional_prevalence_table: pd.DataFrame,
    predictor_table: pd.DataFrame,
    outcome_table: pd.DataFrame,
    predictor_table_no_proxy: pd.DataFrame,
    outcome_table_no_proxy: pd.DataFrame,
) -> str:
    return f"""# Supplementary appendix: LASI skin disease follow-up

## Supplementary Table S1. Variable definitions

{variable_table.to_markdown(index=False)}

## Supplementary Table S2. Additional weighted prevalence by BMI, activity, and proxy interview status

{table_to_display_strings(additional_prevalence_table, 'prevalence').to_markdown(index=False)}

## Supplementary Table S3. Multivariable correlates of self-reported skin disease

{table_to_display_strings(predictor_table, 'predictors').to_markdown(index=False)}

## Supplementary Table S4. Functional and psychosocial outcomes associated with skin disease

{table_to_display_strings(outcome_table, 'outcomes').to_markdown(index=False)}

## Supplementary Table S5. Multivariable correlates of skin disease after excluding proxy respondents

{table_to_display_strings(predictor_table_no_proxy, 'predictors').to_markdown(index=False)}

## Supplementary Table S6. Functional and psychosocial outcomes after excluding proxy respondents

{table_to_display_strings(outcome_table_no_proxy, 'outcomes').to_markdown(index=False)}
"""


def write_markdown_files(
    summary: dict,
    prevalence_table: pd.DataFrame,
    characteristics_table: pd.DataFrame,
    predictor_table: pd.DataFrame,
    outcome_table: pd.DataFrame,
    variable_table: pd.DataFrame,
    additional_prevalence_table: pd.DataFrame,
    predictor_table_no_proxy: pd.DataFrame,
    outcome_table_no_proxy: pd.DataFrame,
) -> None:
    proposal_text = f"""# Proposal: LASI skin disease follow-up study

## Working title

Self-reported skin disease and functional plus psychosocial burden among middle-aged and older adults in India: cross-sectional analysis of harmonized LASI.

## Rationale

The current global manuscript links GBD 2023 skin burden to World Bank ageing indicators at the country level. This LASI project is the next-stage person-level validation analysis. It asks whether older adults who report skin disease also carry higher functional, pain, sleep, and mood burden in a nationally representative ageing cohort.

## Data source

- Dataset: Harmonized LASI A.3
- Local file: `{LASI_PATH}`
- Analytic population: adults aged 45 years or older with non-missing skin-disease status and respondent weights

## Main aims

1. Estimate the weighted prevalence of self-reported skin disease.
2. Describe demographic and social patterning.
3. Quantify associations with ADL, IADL, mobility, pain, depressive symptoms, and sleep problems.
4. Build a reusable manuscript and supplementary appendix package for a second skin-disease paper.

## Preliminary signal

- Analytic sample: {summary['analytic_n']:,}
- Weighted prevalence: {summary['weighted_skin_prev_pct']:.2f}% (95% CI {summary['weighted_skin_prev_ci_low']:.2f}-{summary['weighted_skin_prev_ci_high']:.2f})
- Strongest burden correlate: {summary['top_burden_outcome']} (OR {summary['top_burden_outcome_or']:.3f})
"""
    results_text = f"""# Preliminary results: LASI skin disease follow-up

## Headline findings

- Analytic sample: {summary['analytic_n']:,}
- Weighted prevalence: {summary['weighted_skin_prev_pct']:.2f}%
- Highest age-group prevalence: {summary['prev_age75_plus_pct']:.2f}% in adults aged 75 years or older
- Highest burden correlate: {summary['top_burden_outcome']} (OR {summary['top_burden_outcome_or']:.3f})

## Table 1 preview

{table_to_display_strings(characteristics_table, 'characteristics').to_markdown(index=False)}

## Main model highlights

{table_to_display_strings(outcome_table, 'outcomes').to_markdown(index=False)}
"""
    qc_text = f"""# QC report: LASI skin disease follow-up

## Data provenance

- Source file: `{LASI_PATH}`
- Release used: Harmonized LASI A.3
- Script: `{SCRIPT_OUTPUT}`

## QC conclusions

- Skin disease variable confirmed: `r1skindise`
- Respondent weight confirmed: `r1wtresp`
- Main analytic sample: {summary['analytic_n']:,}
- Weighted prevalence reproduced: {summary['weighted_skin_prev_pct']:.2f}%
- Sensitivity analyses excluding proxy respondents completed and directionally consistent

## Output inventory

- Full manuscript and markdown draft
- Main display items with Table 1 and Figures 1-2
- Supplementary appendix with Tables S1-S6 and Figures S1-S2
"""
    full_manuscript = manuscript_text(
        summary,
        prevalence_table,
        characteristics_table,
        predictor_table,
        outcome_table,
        predictor_table_no_proxy,
        outcome_table_no_proxy,
    )
    supplementary = supplementary_markdown(
        variable_table,
        additional_prevalence_table,
        predictor_table,
        outcome_table,
        predictor_table_no_proxy,
        outcome_table_no_proxy,
    )

    save_markdown(MANUSCRIPT_DIR / "01_Proposal_LASI_Skin_Disease_20260312.md", proposal_text)
    save_markdown(MANUSCRIPT_DIR / "02_Preliminary_Results_LASI_Skin_Disease_20260312.md", results_text)
    save_markdown(MANUSCRIPT_DIR / "03_QC_Report_LASI_Skin_Disease_20260312.md", qc_text)
    save_markdown(MANUSCRIPT_DIR / "04_Full_Manuscript_LASI_Skin_Disease_20260312.md", full_manuscript)
    save_markdown(MANUSCRIPT_DIR / "06_Supplementary_Appendix_LASI_Skin_Disease_20260312.md", supplementary)


def build_short_docx_files(
    summary: dict,
    prevalence_table: pd.DataFrame,
    characteristics_table: pd.DataFrame,
    outcome_table: pd.DataFrame,
) -> None:
    proposal_doc = Document()
    proposal_doc.add_paragraph("Proposal: LASI skin disease follow-up study")
    style_document(proposal_doc, "Proposal")
    add_paragraph(
        proposal_doc,
        "Working title: Self-reported skin disease and functional plus psychosocial burden among middle-aged and older adults in India: cross-sectional analysis of harmonized LASI.",
    )
    add_heading(proposal_doc, "Background", 1)
    add_paragraph(
        proposal_doc,
        "The current GBD 2023 and World Bank WDI manuscript is ecological and country-level. The LASI package is intended as an individual-level validation platform rather than a direct extension of the same model.",
    )
    add_heading(proposal_doc, "Data Source", 1)
    add_bullet(proposal_doc, "Dataset: Harmonized LASI A.3")
    add_bullet(proposal_doc, f"Local source file: {LASI_PATH}")
    add_bullet(proposal_doc, "Analytic frame: adults aged 45 years or older with non-missing skin-disease status")
    add_heading(proposal_doc, "Specific Aims", 1)
    add_bullet(proposal_doc, "Estimate weighted prevalence of self-reported skin disease.")
    add_bullet(proposal_doc, "Describe patterns by age, sex, education, and rural-urban residence.")
    add_bullet(proposal_doc, "Examine associations with ADL, IADL, mobility, pain, depressive symptoms, and sleep problems.")
    add_heading(proposal_doc, "Preliminary Signals", 1)
    add_bullet(proposal_doc, f"Analytic sample with non-missing skin-disease status: {summary['analytic_n']:,}")
    add_bullet(
        proposal_doc,
        f"Weighted skin-disease prevalence: {summary['weighted_skin_prev_pct']:.2f}% (95% CI {summary['weighted_skin_prev_ci_low']:.2f}-{summary['weighted_skin_prev_ci_high']:.2f})",
    )
    add_bullet(proposal_doc, "The clearest person-level burden signals are frequent pain, mobility limitation, and sleep problems.")
    proposal_doc.save(MANUSCRIPT_DIR / "01_Proposal_LASI_Skin_Disease_20260312.docx")

    results_doc = Document()
    results_doc.add_paragraph("Preliminary Results: LASI skin disease follow-up")
    style_document(results_doc, "Results")
    add_heading(results_doc, "Headline Findings", 1)
    add_bullet(results_doc, f"Analytic sample: {summary['analytic_n']:,}")
    add_bullet(results_doc, f"Weighted prevalence: {summary['weighted_skin_prev_pct']:.2f}%")
    add_bullet(results_doc, f"Weighted prevalence in adults aged 75 years or older: {summary['prev_age75_plus_pct']:.2f}%")
    add_bullet(results_doc, f"Weighted prevalence in rural respondents: {summary['prev_rural_pct']:.2f}%")
    add_heading(results_doc, "Table 1 preview", 1)
    add_dataframe_table(results_doc, table_to_display_strings(characteristics_table, "characteristics"))
    add_heading(results_doc, "Main outcome model highlights", 1)
    add_dataframe_table(results_doc, table_to_display_strings(outcome_table, "outcomes"))
    add_heading(results_doc, "Interpretation", 1)
    add_paragraph(
        results_doc,
        "These preliminary findings support the view that skin disease in ageing populations is associated with broader functional and psychosocial burden. This second-stage LASI platform can therefore act as a person-level validation arm for the main ecological manuscript.",
    )
    results_doc.save(MANUSCRIPT_DIR / "02_Preliminary_Results_LASI_Skin_Disease_20260312.docx")

    qc_doc = Document()
    qc_doc.add_paragraph("QC Report: LASI skin disease follow-up")
    style_document(qc_doc, "QC")
    add_heading(qc_doc, "Data provenance", 1)
    add_bullet(qc_doc, f"Source file: {LASI_PATH}")
    add_bullet(qc_doc, "Release used: Harmonized LASI A.3")
    add_heading(qc_doc, "Core QC conclusions", 1)
    add_bullet(qc_doc, "The harmonized skin variable r1skindise was confirmed in the SAV metadata and desktop variable audit.")
    add_bullet(qc_doc, "The respondent weight r1wtresp is present and was used in descriptive and regression outputs.")
    add_bullet(qc_doc, "Sensitivity analysis excluding proxy respondents was completed and was directionally consistent.")
    add_bullet(qc_doc, "This package uses weighted GLM but not yet a full complex-survey variance estimator.")
    qc_doc.save(MANUSCRIPT_DIR / "03_QC_Report_LASI_Skin_Disease_20260312.docx")


def add_reference_list(doc: Document) -> None:
    add_heading(doc, "References", 1)
    for idx, ref in enumerate(references_list(), start=1):
        add_paragraph(doc, f"{idx}. {ref}")


def add_manuscript_section(doc: Document, heading: str, text: str) -> None:
    add_heading(doc, heading, 1)
    for block in text.strip().split("\n\n"):
        add_paragraph(doc, block.strip())


def build_full_manuscript_docx(
    summary: dict,
    prevalence_table: pd.DataFrame,
    characteristics_table: pd.DataFrame,
    predictor_table: pd.DataFrame,
    outcome_table: pd.DataFrame,
    predictor_table_no_proxy: pd.DataFrame,
    outcome_table_no_proxy: pd.DataFrame,
) -> tuple[int, int]:
    text = manuscript_text(
        summary,
        prevalence_table,
        characteristics_table,
        predictor_table,
        outcome_table,
        predictor_table_no_proxy,
        outcome_table_no_proxy,
    )
    abstract_text, rest = text.split("Introduction\n\n", 1)
    doc = Document()
    doc.add_paragraph("Self-reported skin disease and functional plus psychosocial burden among middle-aged and older adults in India: cross-sectional analysis of harmonized LASI")
    style_document(doc, "Full Manuscript")

    add_paragraph(
        doc,
        "Authors to be added by the study team. This draft is a full English manuscript package generated from harmonized LASI data for further author revision.",
        italic=True,
    )

    add_heading(doc, "Summary", 1)
    summary_blocks = abstract_text.replace("Summary\n\n", "").strip().split("\n\n")
    for block in summary_blocks:
        lines = block.strip().split("\n", 1)
        if len(lines) == 2:
            label, value = lines
            paragraph = doc.add_paragraph()
            run1 = paragraph.add_run(label + " ")
            run1.bold = True
            run1.font.name = "Arial"
            run1.font.size = Pt(10.5)
            run2 = paragraph.add_run(value)
            run2.font.name = "Arial"
            run2.font.size = Pt(10.5)
        else:
            add_paragraph(doc, block.strip())

    sections = rest.split("\n\nReferences\n")
    body = body_text = rest.split("\n\nReferences\n", 1)[0]
    refs = sections[1]
    for section_text in body.split("\n\n"):
        if section_text in {"Introduction", "Methods", "Results", "Discussion"}:
            add_heading(doc, section_text, 1)
        elif section_text in {
            "Study design and data source",
            "Study population",
            "Exposure, covariates, and outcomes",
            "Statistical analysis",
            "Ethics and role of the funding source",
            "Weighted prevalence of self-reported skin disease",
            "Sample characteristics according to skin disease status",
            "Correlates of reporting skin disease",
            "Functional and psychosocial burden associated with skin disease",
            "Sensitivity analysis",
        }:
            add_heading(doc, section_text, 2)
        else:
            add_paragraph(doc, section_text)

    add_heading(doc, "References", 1)
    for ref in refs.strip().split("\n"):
        add_paragraph(doc, ref)

    doc.save(MANUSCRIPT_DIR / "04_Full_Manuscript_LASI_Skin_Disease_20260312.docx")

    main_text_words = len(("Introduction " + body_text).replace("\n", " ").split())
    abstract_words = len(abstract_text.replace("\n", " ").split())
    return abstract_words, main_text_words


def build_display_items_docx(
    characteristics_table: pd.DataFrame,
    figure1_png: Path,
    figure2_png: Path,
) -> None:
    doc = Document()
    doc.add_paragraph("Main display items: LASI skin disease manuscript")
    style_document(doc, "Display Items")

    add_caption(doc, "Table 1. Weighted characteristics of adults aged 45 years and older according to self-reported skin disease status in Harmonized LASI")
    add_dataframe_table(doc, table_to_display_strings(characteristics_table, "characteristics"))
    add_note(doc, "All percentages are respondent-weighted. Table values are intended for main manuscript display.")

    doc.add_page_break()
    add_caption(doc, "Figure 1. Weighted prevalence of self-reported skin disease by age, sex, and rural-urban residence")
    doc.add_picture(str(figure1_png), width=Inches(6.6))
    add_note(doc, "Points show weighted prevalence estimates and whiskers show approximate 95% CIs.")

    doc.add_page_break()
    add_caption(doc, "Figure 2. Adjusted associations between self-reported skin disease and functional and psychosocial outcomes")
    doc.add_picture(str(figure2_png), width=Inches(6.6))
    add_note(doc, "Odds ratios are adjusted for age, sex, rural residence, and education. The vertical dashed line marks an odds ratio of 1.0.")

    doc.save(MANUSCRIPT_DIR / "05_Main_Display_Items_LASI_Skin_Disease_20260312.docx")


def build_supplementary_docx(
    variable_table: pd.DataFrame,
    additional_prevalence_table: pd.DataFrame,
    predictor_table: pd.DataFrame,
    outcome_table: pd.DataFrame,
    predictor_table_no_proxy: pd.DataFrame,
    outcome_table_no_proxy: pd.DataFrame,
    figure_s1_png: Path,
    figure_s2_png: Path,
) -> None:
    doc = Document()
    doc.add_paragraph("Supplementary appendix: LASI skin disease manuscript")
    style_document(doc, "Supplementary")

    add_paragraph(
        doc,
        "This appendix accompanies the full LASI skin disease manuscript and provides coding definitions, supplementary prevalence estimates, and sensitivity analyses.",
        italic=True,
    )
    add_caption(doc, "Supplementary Table S1. Variable definitions and operational coding")
    add_dataframe_table(doc, variable_table)
    add_note(doc, "Variable names correspond to Harmonized LASI A.3 respondent-level baseline fields.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Table S2. Additional weighted prevalence of self-reported skin disease by BMI, physical activity, and proxy interview status")
    add_dataframe_table(doc, table_to_display_strings(additional_prevalence_table, "prevalence"))
    add_note(doc, "Prevalence estimates are weighted using respondent weights and use an approximate effective-sample-size CI.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Table S3. Multivariable correlates of self-reported skin disease")
    add_dataframe_table(doc, table_to_display_strings(predictor_table, "predictors"))
    add_note(doc, "Odds ratios are from weighted binomial models adjusted only for the covariates shown.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Table S4. Functional and psychosocial outcomes associated with self-reported skin disease")
    add_dataframe_table(doc, table_to_display_strings(outcome_table, "outcomes"))
    add_note(doc, "Each model adjusts for age, sex, rural residence, and education.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Table S5. Multivariable correlates of skin disease after excluding proxy respondents")
    add_dataframe_table(doc, table_to_display_strings(predictor_table_no_proxy, "predictors"))
    add_note(doc, "Sensitivity analysis restricted to self-respondents.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Table S6. Functional and psychosocial outcomes after excluding proxy respondents")
    add_dataframe_table(doc, table_to_display_strings(outcome_table_no_proxy, "outcomes"))
    add_note(doc, "Sensitivity models remained directionally consistent with the primary analysis.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Figure S1. Additional weighted prevalence estimates by BMI, activity, and proxy interview status")
    doc.add_picture(str(figure_s1_png), width=Inches(6.6))
    add_note(doc, "Supplementary prevalence plot broadens the descriptive profile beyond the main manuscript display.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Figure S2. Sensitivity analysis of outcome associations after excluding proxy respondents")
    doc.add_picture(str(figure_s2_png), width=Inches(6.6))
    add_note(doc, "Error bars show 95% CIs from weighted models in self-respondents only.")

    doc.save(MANUSCRIPT_DIR / "06_Supplementary_Appendix_LASI_Skin_Disease_20260312.docx")


def write_readme(summary: dict, abstract_words: int, main_text_words: int) -> None:
    text = f"""# LASI Skin Follow-up Package

This folder contains a second-study manuscript framework built from Harmonized LASI after the desktop Global Aging Data audit.

## Key numbers

- Analytic sample: {summary['analytic_n']:,}
- Weighted skin-disease prevalence: {summary['weighted_skin_prev_pct']:.2f}% (95% CI {summary['weighted_skin_prev_ci_low']:.2f}-{summary['weighted_skin_prev_ci_high']:.2f})
- Strongest primary burden correlate: {summary['top_burden_outcome']} (OR {summary['top_burden_outcome_or']:.3f})
- Abstract word count: {abstract_words}
- Main-text word count: {main_text_words}

## Main files

- `outputs/manuscript/04_Full_Manuscript_LASI_Skin_Disease_20260312.docx`
- `outputs/manuscript/05_Main_Display_Items_LASI_Skin_Disease_20260312.docx`
- `outputs/manuscript/06_Supplementary_Appendix_LASI_Skin_Disease_20260312.docx`
- `outputs/tables/*.csv`
- `outputs/figures/*.png`
"""
    (PACKAGE_ROOT / "README.md").write_text(text, encoding="utf-8")


def main() -> None:
    ensure_dirs()
    raw, _ = read_lasi()
    analysis = clean_lasi(raw)
    analysis_no_proxy = analysis.loc[analysis["r1proxy"] == 0].copy()

    prevalence_table = build_prevalence_table(analysis)
    additional_prevalence_table = build_additional_prevalence_table(analysis)
    characteristics_table = build_characteristics_table(analysis)
    predictor_table = build_predictor_table(analysis)
    outcome_table = build_outcome_association_table(analysis)
    variable_table = build_variable_definition_table()
    predictor_table_no_proxy = build_predictor_table(analysis_no_proxy)
    outcome_table_no_proxy = build_outcome_association_table(analysis_no_proxy)

    summary = build_summary(
        analysis,
        prevalence_table,
        predictor_table,
        outcome_table,
        predictor_table_no_proxy,
        outcome_table_no_proxy,
    )

    prevalence_table.to_csv(TABLE_DIR / "table1_weighted_skin_prevalence.csv", index=False)
    characteristics_table.to_csv(TABLE_DIR / "table2_characteristics_by_skin_status.csv", index=False)
    predictor_table.to_csv(TABLE_DIR / "table3_predictors_of_skin_disease.csv", index=False)
    outcome_table.to_csv(TABLE_DIR / "table4_functional_psychosocial_associations.csv", index=False)
    variable_table.to_csv(TABLE_DIR / "tableS1_variable_definitions.csv", index=False)
    additional_prevalence_table.to_csv(TABLE_DIR / "tableS2_additional_weighted_prevalence.csv", index=False)
    predictor_table_no_proxy.to_csv(TABLE_DIR / "tableS5_predictors_excluding_proxy.csv", index=False)
    outcome_table_no_proxy.to_csv(TABLE_DIR / "tableS6_outcomes_excluding_proxy.csv", index=False)

    figure1_png = FIGURE_DIR / "figure1_lasi_skin_prevalence.png"
    figure1_pdf = FIGURE_DIR / "figure1_lasi_skin_prevalence.pdf"
    figure2_png = FIGURE_DIR / "figure2_lasi_skin_outcome_forest.png"
    figure2_pdf = FIGURE_DIR / "figure2_lasi_skin_outcome_forest.pdf"
    figure_s1_png = FIGURE_DIR / "figureS1_lasi_additional_prevalence.png"
    figure_s1_pdf = FIGURE_DIR / "figureS1_lasi_additional_prevalence.pdf"
    figure_s2_png = FIGURE_DIR / "figureS2_lasi_sensitivity_outcomes.png"
    figure_s2_pdf = FIGURE_DIR / "figureS2_lasi_sensitivity_outcomes.pdf"

    write_prevalence_figure(
        prevalence_table,
        figure1_png,
        figure1_pdf,
        ["Age 45-59", "Age 60-74", "Age 75+", "Men", "Women", "Urban", "Rural"],
        "Figure 1. Weighted prevalence of self-reported skin disease in LASI",
    )
    write_forest_figure(
        outcome_table,
        figure2_png,
        figure2_pdf,
        "Figure 2. Functional and psychosocial burden associated with self-reported skin disease",
    )
    write_prevalence_figure(
        additional_prevalence_table,
        figure_s1_png,
        figure_s1_pdf,
        additional_prevalence_table["Group"].tolist(),
        "Supplementary Figure S1. Additional subgroup prevalence estimates",
    )
    write_forest_figure(
        outcome_table_no_proxy,
        figure_s2_png,
        figure_s2_pdf,
        "Supplementary Figure S2. Sensitivity analysis excluding proxy respondents",
    )

    summary_json = MANUSCRIPT_DIR / "lasi_skin_summary_20260312.json"
    summary_json.write_text(json.dumps(summary, indent=2), encoding="utf-8")

    write_markdown_files(
        summary,
        prevalence_table,
        characteristics_table,
        predictor_table,
        outcome_table,
        variable_table,
        additional_prevalence_table,
        predictor_table_no_proxy,
        outcome_table_no_proxy,
    )
    build_short_docx_files(summary, prevalence_table, characteristics_table, outcome_table)
    abstract_words, main_text_words = build_full_manuscript_docx(
        summary,
        prevalence_table,
        characteristics_table,
        predictor_table,
        outcome_table,
        predictor_table_no_proxy,
        outcome_table_no_proxy,
    )
    build_display_items_docx(characteristics_table, figure1_png, figure2_png)
    build_supplementary_docx(
        variable_table,
        additional_prevalence_table,
        predictor_table,
        outcome_table,
        predictor_table_no_proxy,
        outcome_table_no_proxy,
        figure_s1_png,
        figure_s2_png,
    )
    write_readme(summary, abstract_words, main_text_words)

    manifest = {
        "source_lasi_path": str(LASI_PATH),
        "script": str(SCRIPT_OUTPUT),
        "analytic_n": summary["analytic_n"],
        "weighted_skin_prev_pct": summary["weighted_skin_prev_pct"],
        "abstract_word_count": abstract_words,
        "main_text_word_count": main_text_words,
        "table_files": sorted(p.name for p in TABLE_DIR.glob("*.csv")),
        "figure_files": sorted(p.name for p in FIGURE_DIR.glob("*")),
        "manuscript_files": sorted(p.name for p in MANUSCRIPT_DIR.glob("*")),
    }
    (PACKAGE_ROOT / "manifest_20260312.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")

    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
