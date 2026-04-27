from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PACKAGE_ROOT = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万/lasi_skin_followup_package_20260312"
)
MANUSCRIPT_DIR = PACKAGE_ROOT / "outputs" / "manuscript"
TABLE_DIR = PACKAGE_ROOT / "outputs" / "tables"
FIGURE_DIR = PACKAGE_ROOT / "outputs" / "figures"
SUMMARY_PATH = MANUSCRIPT_DIR / "lasi_skin_survey_summary_20260312.json"

BLUE = "1F4E79"
DARK = "163A63"
LIGHT = "EDF3F8"
GRAY = "F5F7FA"


def load_inputs() -> dict:
    summary = json.loads(SUMMARY_PATH.read_text())
    data = {
        "summary": summary,
        "table1": pd.read_csv(TABLE_DIR / "table1_survey_weighted_characteristics.csv"),
        "table2": pd.read_csv(TABLE_DIR / "table2_survey_prevalence_main.csv"),
        "table3": pd.read_csv(TABLE_DIR / "table3_survey_predictors.csv"),
        "table4": pd.read_csv(TABLE_DIR / "table4_survey_outcomes.csv"),
        "tableS1": pd.read_csv(TABLE_DIR / "tableS1_survey_variable_definitions.csv"),
        "tableS2": pd.read_csv(TABLE_DIR / "tableS2_survey_prevalence_additional.csv"),
        "tableS3": pd.read_csv(TABLE_DIR / "tableS3_survey_outcomes_no_proxy.csv"),
        "tableS4": pd.read_csv(TABLE_DIR / "tableS4_survey_outcomes_extended_adjustment.csv"),
        "tableS5": pd.read_csv(TABLE_DIR / "tableS5_survey_interaction_tests.csv"),
        "tableS6": pd.read_csv(TABLE_DIR / "tableS6_survey_stratified_effects.csv"),
    }
    return data


def style_doc(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    for style_name in ["Normal", "List Bullet"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(10.5)
    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor.from_string(BLUE)


def add_title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(DARK)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, italic: bool = False, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label + " ")
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.italic = italic
    r.font.name = "Arial"
    r.font.size = Pt(10.5)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.name = "Arial"
    r.font.size = Pt(10.5)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_cell(cell, bold: bool = False, color: str | None = None, size: int = 9) -> None:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold
            r.font.name = "Arial"
            r.font.size = Pt(size)
            if color:
                r.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, df: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = str(col)
        shade_cell(hdr[i], BLUE)
        style_cell(hdr[i], bold=True, color="FFFFFF")
    for idx, (_, row) in enumerate(df.iterrows(), start=1):
        fill = GRAY if idx % 2 == 0 else "FFFFFF"
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            cells[i].text = str(value)
            shade_cell(cells[i], fill)
            style_cell(cells[i])


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(BLUE)


def note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.name = "Arial"
    r.font.size = Pt(9)


def prevalence_value(table: pd.DataFrame, label: str) -> float:
    return float(table.loc[table["Group"] == label, "Prevalence_pct"].iloc[0])


def pvalue(value: float) -> str:
    if value < 0.001:
        return "<0.001"
    return f"{value:.3f}"


def display_or_table(df: pd.DataFrame, label_col: str) -> pd.DataFrame:
    out = df.copy(deep=True).reset_index(drop=True)
    out["Adjusted OR (95% CI)"] = out.apply(
        lambda row: f"{row['OR']:.3f} ({row['CI_low']:.3f}-{row['CI_high']:.3f})", axis=1
    )
    out = out.assign(**{"p value": out["p_value"].map(pvalue).astype(str)})
    return out[[label_col, "Adjusted OR (95% CI)", "p value", "Analytic_n"]]


def display_prevalence_table(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    out["Weighted prevalence % (95% CI)"] = out.apply(
        lambda row: f"{row['Prevalence_pct']:.2f} ({row['CI_low_pct']:.2f}-{row['CI_high_pct']:.2f})", axis=1
    )
    return out[["Group", "Unweighted_n", "Weighted prevalence % (95% CI)"]]


def display_characteristics_table(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for col in ["Total", "No_skin_disease", "Skin_disease"]:
        out[col] = out.apply(
            lambda row: f"{int(row[col]):,}" if row["Characteristic"] == "Participants, n" else f"{row[col]:.2f}",
            axis=1,
        )
    return out.rename(columns={"No_skin_disease": "No skin disease", "Skin_disease": "Skin disease"})


def references() -> list[str]:
    return [
        "Longitudinal Ageing Study in India. Study design. https://www.iipsindia.ac.in/content/lasi-study-design. Accessed March 12, 2026.",
        "International Institute for Population Sciences, NPHCE, Ministry of Health and Family Welfare, Harvard T H Chan School of Public Health, University of Southern California. Longitudinal Ageing Study in India (LASI) Wave 1, 2017-18, India Report. Mumbai: IIPS; 2020.",
        "Lee J, Banerjee J, Khobragade PY, et al. Gateway to Global Aging Data: resources for studies of aging populations worldwide. J Gerontol B Psychol Sci Soc Sci 2021; 76(Suppl 1): S5-S16.",
        "Lee J, Cantu PA, Ghilarducci T, et al. Cross-country comparisons of disability and morbidity: evidence from the Gateway to Global Aging Data. J Gerontol B Psychol Sci Soc Sci 2024; 79(Suppl 1): S35-S43.",
        "Li D, Lyu J, Xu T, et al. Global burden of skin and subcutaneous diseases: an analysis of the Global Burden of Disease Study 2021. J Eur Acad Dermatol Venereol 2026; 40: e686-e689.",
        "Yakupu A, Wang Y, Pan KH, et al. The global, regional, and national burden of skin and subcutaneous diseases: findings from the Global Burden of Disease Study 2019. Front Public Health 2023; 11: 1089259.",
        "Hahnel E, Lichterfeld-Kottner A, Bohnke N, et al. Epidemiology of skin conditions in the aged: a systematic review. J Tissue Viability 2017; 26: 20-28.",
        "Shah M, Coates M. An assessment of the quality of life in older patients with skin disease. Br J Dermatol 2006; 154: 150-153.",
        "Mostaghimi L. The prevalence of mood and sleep problems in chronic skin diseases. Arch Dermatol 2008; 144: 18-24.",
        "Salfi F, Oliviero A, D'Atri A, et al. The role of insomnia in vulnerability to depressive and anxiety symptoms in adults with atopic dermatitis. Clocks Sleep 2023; 5: 628-639.",
        "Pradhan MR, Saikia D. Prevalence and predictors of insomnia among older adults in India: an analysis based on LASI. Aging Ment Health 2025; 29: 1137-1147.",
        "Muhammad T, Rashid M. Prevalence and correlates of pain and associated depression among community-dwelling older adults in India. BMC Geriatr 2022; 22: 708.",
        "Sharma P, Maurya P, Muhammad T. Number of chronic conditions and associated functional limitations among older adults: cross-sectional findings from the Longitudinal Ageing Study in India. BMC Geriatr 2021; 21: 673.",
        "Banerjee S, Boro B. The role of sleep quality, functional limitation and depressive symptoms in predicting life satisfaction among older adults in India. Sci Rep 2022; 12: 10339.",
        "Lumley T. Analysis of complex survey samples. J Stat Softw 2004; 9: 1-19.",
        "von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement: guidelines for reporting observational studies. Lancet 2007; 370: 1453-1457.",
    ]


def build_texts(data: dict) -> dict:
    s = data["summary"]
    t1 = data["table1"]
    t2 = data["table2"]
    t3 = data["table3"]
    t4 = data["table4"]
    tS3 = data["tableS3"]
    tS4 = data["tableS4"]
    tS5 = data["tableS5"]
    tS6 = data["tableS6"]

    overall = t2.loc[t2["Group"] == "Overall"].iloc[0]
    age75 = t2.loc[t2["Group"] == "Age 75+"].iloc[0]
    men = t2.loc[t2["Group"] == "Men"].iloc[0]
    women = t2.loc[t2["Group"] == "Women"].iloc[0]
    urban = t2.loc[t2["Group"] == "Urban"].iloc[0]
    rural = t2.loc[t2["Group"] == "Rural"].iloc[0]

    age_or = t3.loc[t3["Predictor"] == "Age (per year)"].iloc[0]
    women_or = t3.loc[t3["Predictor"] == "Women vs men"].iloc[0]
    rural_or = t3.loc[t3["Predictor"] == "Rural vs urban"].iloc[0]
    edu_or = t3.loc[t3["Predictor"] == "Secondary education or above vs less"].iloc[0]

    pain = t4.loc[t4["Outcome"] == "Frequent pain"].iloc[0]
    mobility = t4.loc[t4["Outcome"] == "Any mobility limitation"].iloc[0]
    sleep = t4.loc[t4["Outcome"] == "Sleep problems"].iloc[0]
    dep = t4.loc[t4["Outcome"] == "Depressive symptoms"].iloc[0]
    adl = t4.loc[t4["Outcome"] == "Any ADL limitation"].iloc[0]
    iadl = t4.loc[t4["Outcome"] == "Any IADL limitation"].iloc[0]

    sleep_int = tS5.loc[tS5["Outcome"] == "Sleep problems"].iloc[0]
    mobility_int = tS5.loc[tS5["Outcome"] == "Any mobility limitation"].iloc[0]
    sleep_men = tS6.loc[(tS6["Outcome"] == "Sleep problems") & (tS6["Stratum"] == "Men")].iloc[0]
    sleep_women = tS6.loc[(tS6["Outcome"] == "Sleep problems") & (tS6["Stratum"] == "Women")].iloc[0]
    mobility_urban = tS6.loc[(tS6["Outcome"] == "Any mobility limitation") & (tS6["Stratum"] == "Urban")].iloc[0]
    mobility_rural = tS6.loc[(tS6["Outcome"] == "Any mobility limitation") & (tS6["Stratum"] == "Rural")].iloc[0]

    no_proxy_pain = tS3.loc[tS3["Outcome"] == "Frequent pain"].iloc[0]
    ext_pain = tS4.loc[tS4["Outcome"] == "Frequent pain"].iloc[0]

    lancet_md = f"""# Self-reported skin disease and functional plus psychosocial burden among middle-aged and older adults in India: a survey-weighted analysis of harmonized LASI

## Summary

### Background
Skin and subcutaneous diseases are increasingly recognised as part of healthy-ageing agendas, but ecological burden estimates do not show whether older adults who report skin disease also experience higher person-level functional and psychosocial burden. We assessed the prevalence of self-reported skin disease and its association with pain, mobility, daily functioning, depressive symptoms, and sleep problems in a nationally representative ageing dataset from India.

### Methods
We did a cross-sectional analysis of Harmonized LASI A.3. Adults aged 45 years or older with non-missing self-reported skin disease status and respondent weights were included. Because the public harmonized LASI file releases respondent weights but not PSU or design strata identifiers, primary analyses used a conservative survey design with household clustering and state stratification proxy. Survey-weighted logistic models assessed correlates of skin disease and associations with ADL limitation, IADL limitation, mobility limitation, frequent pain, depressive symptoms, and sleep problems. Interaction analyses examined effect modification by sex, rural residence, and age 75 years or older. Sensitivity analyses excluded proxy respondents and additionally adjusted for BMI and physical inactivity.

### Findings
{s['analytic_n']:,} respondents from {s['states']} states and union territories and {s['households']:,} households were included. Survey-weighted prevalence of self-reported skin disease was {overall['Prevalence_pct']:.2f}% (95% CI {overall['CI_low_pct']:.2f}-{overall['CI_high_pct']:.2f}). Prevalence was higher among adults aged 75 years or older ({age75['Prevalence_pct']:.2f}%), men ({men['Prevalence_pct']:.2f}%), and rural respondents ({rural['Prevalence_pct']:.2f}%). In survey-weighted models, skin disease was more common with older age (OR {age_or['OR']:.3f}, 95% CI {age_or['CI_low']:.3f}-{age_or['CI_high']:.3f}) and rural residence (OR {rural_or['OR']:.3f}, {rural_or['CI_low']:.3f}-{rural_or['CI_high']:.3f}), but less common among women (OR {women_or['OR']:.3f}, {women_or['CI_low']:.3f}-{women_or['CI_high']:.3f}). Skin disease was associated with frequent pain (OR {pain['OR']:.3f}, 95% CI {pain['CI_low']:.3f}-{pain['CI_high']:.3f}), mobility limitation ({mobility['OR']:.3f}, {mobility['CI_low']:.3f}-{mobility['CI_high']:.3f}), sleep problems ({sleep['OR']:.3f}, {sleep['CI_low']:.3f}-{sleep['CI_high']:.3f}), ADL limitation ({adl['OR']:.3f}, {adl['CI_low']:.3f}-{adl['CI_high']:.3f}), IADL limitation ({iadl['OR']:.3f}, {iadl['CI_low']:.3f}-{iadl['CI_high']:.3f}), and depressive symptoms ({dep['OR']:.3f}, {dep['CI_low']:.3f}-{dep['CI_high']:.3f}). Sex modified the association with sleep problems (p_interaction={sleep_int['Sex_interaction_p']:.4f}), and rural residence modified the association with mobility limitation (p_interaction={mobility_int['Rural_interaction_p']:.4f}). Findings were directionally stable after excluding proxy respondents and after extended adjustment.

### Interpretation
In LASI, self-reported skin disease marked a wider burden phenotype rather than a narrow dermatologic complaint. The findings strengthen the argument that skin health in ageing populations should be integrated into pain, function, sleep, and mental-health aware care models, and they provide person-level validation for the ageing narrative emerging from global burden analyses.

### Funding
No specific funding was used for this secondary analysis package.

## Research in context

### Evidence before this study
On March 12, 2026, we searched PubMed for combinations of “skin disease”, “dermatologic”, “older adults”, “ageing”, “India”, “LASI”, and “Longitudinal Ageing Study in India”. We found studies on skin-disease quality of life, mood, sleep, and geriatric skin conditions, and multiple ecological papers on global skin burden, but we did not identify a nationally representative survey-weighted LASI analysis that treated broad self-reported skin disease as an exposure and linked it to function, pain, mood, and sleep in one framework.

### Added value of this study
This study moves the project from ecological burden interpretation to person-level validation. It uses a nationally representative ageing dataset, implements survey-weighted regression with household clustering and state stratification proxy, tests pre-specified interaction terms, and adds sensitivity analyses excluding proxy respondents and additionally adjusting for BMI and physical inactivity. It also shows that some signals seen in simpler weighted models change after formal survey variance estimation, including attenuation of the education association.

### Implications of all the available evidence
Skin disease in older adults should be interpreted as part of the broader healthy-ageing and long-term-care agenda rather than as a narrowly cosmetic or outpatient specialty issue. Future work should extend this framework to longitudinal LASI waves and, where design variables are available, to other ageing cohorts in the Gateway to Global Aging Data platform.

## Introduction

Skin and subcutaneous diseases contribute substantial disability worldwide, yet their relevance to ageing policy is often underframed.[5,6] In older adults, skin disease can affect more than diagnosis counts alone: it can shape pain, sleep, mobility, social visibility, mood, and the ability to manage daily life.[7-10] These issues matter because healthy-ageing frameworks increasingly prioritise function, wellbeing, and care needs rather than isolated disease labels.[1-4]

This gap is especially visible in global burden research. Ecological studies describe national or regional burdens well, but they do not show whether older adults who report skin disease also carry worse person-level functional or psychosocial burden. That limitation is relevant to the current skin-burden project because our global manuscript uses GBD and World Bank data to frame ageing at the country level. A complementary person-level validation platform is therefore necessary.

The Longitudinal Ageing Study in India provides such a platform. LASI is large, nationally representative, and harmonized for ageing research.[1-4] Its respondent-level skin-disease item does not provide diagnosis-specific detail, but it does allow a defensible first-stage test of whether self-reported skin disease clusters with pain, mobility problems, sleep disturbance, depressive symptoms, and daily functional limitation.

We therefore aimed to estimate the survey-weighted prevalence of self-reported skin disease in middle-aged and older adults in India and to examine its association with multiple dimensions of burden. We also aimed to test whether these associations differed by sex, rural residence, or older age, and whether the main findings were robust to more conservative sensitivity analyses.

## Methods

### Study design and data source

We did a cross-sectional secondary analysis of Harmonized LASI A.3, derived from the nationally representative LASI baseline survey.[1-4] The public harmonized file provides respondent-level weights and harmonized health measures designed for international ageing comparisons. The local analytic source was `/Users/apple/Desktop/所有数据/global aging data数据/LASI_印度/Harmonized LASI A.3_SPSS/H_LASI_a3.sav`.

### Participants

We included adults aged 45 years or older with non-missing respondent weights and non-missing self-reported skin disease status. This yielded {s['analytic_n']:,} respondents from {s['households']:,} households across {s['states']} states and union territories. Weighted sample characteristics are shown in Table 1.

### Exposure, covariates, and outcomes

The exposure was self-reported skin disease (`r1skindise`). Covariates were age in years, sex, rural residence, and education. The primary outcomes were any ADL limitation, any IADL limitation, any mobility limitation, frequent pain, depressive symptoms, and sleep problems. For sensitivity analysis we additionally used BMI, vigorous physical inactivity, moderate physical inactivity, and proxy interview status.

### Statistical analysis

Primary analyses used survey-weighted methods with household clustering and state stratification proxy, because the public harmonized LASI file releases respondent weights but not the original PSU or design-strata variables. Weighted prevalence was estimated overall and by key demographic groups. Survey-weighted logistic regression assessed correlates of reporting skin disease and associations between skin disease and burden outcomes. Pre-specified interaction terms evaluated modification by sex, rural residence, and age 75 years or older. Sensitivity analyses excluded proxy respondents and, separately, additionally adjusted outcome models for BMI and physical inactivity. Odds ratios are reported with 95% CIs.

## Results

### Survey-weighted prevalence of self-reported skin disease

The survey-weighted prevalence of self-reported skin disease was {overall['Prevalence_pct']:.2f}% (95% CI {overall['CI_low_pct']:.2f}-{overall['CI_high_pct']:.2f}). Prevalence increased with age and was highest among respondents aged 75 years or older ({age75['Prevalence_pct']:.2f}%). Men had higher prevalence than women ({men['Prevalence_pct']:.2f}% vs {women['Prevalence_pct']:.2f}%), and rural respondents had higher prevalence than urban respondents ({rural['Prevalence_pct']:.2f}% vs {urban['Prevalence_pct']:.2f}; Figure 1, Supplementary Table S2).

### Characteristics and correlates

Respondents reporting skin disease were slightly older, more rural, and more burdened across every outcome domain than those without skin disease (Table 1). In survey-weighted regression, older age remained associated with reporting skin disease (OR {age_or['OR']:.3f}, 95% CI {age_or['CI_low']:.3f}-{age_or['CI_high']:.3f}), women were less likely than men to report skin disease (OR {women_or['OR']:.3f}, {women_or['CI_low']:.3f}-{women_or['CI_high']:.3f}), and rural residence was positively associated (OR {rural_or['OR']:.3f}, {rural_or['CI_low']:.3f}-{rural_or['CI_high']:.3f}). By contrast, the adjusted education association was attenuated and no longer conventionally significant after survey variance estimation (OR {edu_or['OR']:.3f}, {edu_or['CI_low']:.3f}-{edu_or['CI_high']:.3f}; p={pvalue(edu_or['p_value'])}; Supplementary Table S3).

### Functional and psychosocial burden

Skin disease was associated with higher odds of every assessed burden outcome (Figure 2). The strongest association was for frequent pain (OR {pain['OR']:.3f}, 95% CI {pain['CI_low']:.3f}-{pain['CI_high']:.3f}), followed by mobility limitation (OR {mobility['OR']:.3f}, {mobility['CI_low']:.3f}-{mobility['CI_high']:.3f}) and sleep problems (OR {sleep['OR']:.3f}, {sleep['CI_low']:.3f}-{sleep['CI_high']:.3f}). More modest but still adverse associations were seen for ADL limitation, IADL limitation, and depressive symptoms.

### Interaction and sensitivity analyses

Most interaction tests were not significant, but two prespecified signals were retained. The association between skin disease and sleep problems was stronger in women (OR {sleep_women['OR']:.3f}, 95% CI {sleep_women['CI_low']:.3f}-{sleep_women['CI_high']:.3f}) than in men (OR {sleep_men['OR']:.3f}, {sleep_men['CI_low']:.3f}-{sleep_men['CI_high']:.3f}; p_interaction={sleep_int['Sex_interaction_p']:.4f}). The association between skin disease and mobility limitation was stronger in urban respondents (OR {mobility_urban['OR']:.3f}, 95% CI {mobility_urban['CI_low']:.3f}-{mobility_urban['CI_high']:.3f}) than in rural respondents (OR {mobility_rural['OR']:.3f}, {mobility_rural['CI_low']:.3f}-{mobility_rural['CI_high']:.3f}; p_interaction={mobility_int['Rural_interaction_p']:.4f}). In sensitivity analyses, the frequent-pain association remained strong after excluding proxy respondents (OR {no_proxy_pain['OR']:.3f}, 95% CI {no_proxy_pain['CI_low']:.3f}-{no_proxy_pain['CI_high']:.3f}) and after extended adjustment (OR {ext_pain['OR']:.3f}, {ext_pain['CI_low']:.3f}-{ext_pain['CI_high']:.3f}; Supplementary Tables S4-S7).

## Discussion

This survey-weighted LASI analysis indicates that self-reported skin disease in older adults is associated with a broad burden phenotype rather than an isolated dermatologic label. Three findings carry most weight. First, prevalence was non-trivial and increased in the oldest age group. Second, the reporting pattern remained higher in men and in rural respondents. Third, and most importantly, skin disease tracked with pain, mobility limitation, sleep problems, depressive symptoms, and impaired daily function even after survey-weighted adjustment.

These results materially strengthen the ongoing global project. Ecological analyses can show where burden is concentrated, but they cannot determine whether the ageing narrative corresponds to lived person-level burden. The LASI findings support that translation. They suggest that older adults reporting skin disease are also more likely to experience symptoms and functional problems that matter to long-term care, rehabilitation, and geriatric assessment.[7-14]

The interaction findings add nuance. Sleep burden was more strongly associated with skin disease in women, which may reflect differences in symptom appraisal, sleep vulnerability, or the psychosocial salience of visible disease. Mobility burden, meanwhile, showed a stronger association in urban respondents than in rural respondents, despite higher prevalence in rural areas. One interpretation is that the same skin condition may interfere differently with participation and movement depending on environment, occupation, care access, and baseline mobility expectations. These findings are exploratory rather than definitive, but they help identify where longitudinal follow-up could be most informative.

The shift in the education estimate also matters methodologically. In earlier simplified weighted models, higher education showed a weak positive association with skin disease. After survey-weighted variance estimation, that signal attenuated and was no longer conventionally significant. This demonstrates why the current stage needed to move beyond simple weighted GLM and into a more defensible design-based framework.

This analysis has several strengths. It uses a large, nationally representative ageing cohort, applies formal survey-weighted regression, pre-specifies effect-modification tests, and demonstrates directional robustness in two sensitivity analyses. It also serves the specific strategic role we intended for desktop Global Aging Data assets: not as a forced extension of an ecological GBD model, but as an individual-level validation platform.

Several limitations remain. The harmonized public LASI file does not release the original PSU and design-strata variables, so we used household clustering and state stratification proxy as a conservative survey specification. The skin-disease variable is self-reported and broad, with no diagnostic subtype, severity, or chronicity detail. The analysis is cross-sectional, so temporality cannot be established. Residual confounding from multimorbidity, health-care access, and occupation is likely. Finally, the present results come from India and should not be overgeneralised.

Taken together, these findings support placing skin health more centrally within healthy-ageing policy. In older adults, skin disease appears linked to pain, function, sleep, and mood, which means it belongs in integrated geriatric care rather than at the margins of health planning. Future work should extend these analyses to longitudinal LASI data and to comparable cohorts in other countries.

## References

"""

    lancet_md += "\n".join([f"{i}. {ref}" for i, ref in enumerate(references(), start=1)]) + "\n"

    bmc_md = f"""# Self-reported skin disease and functional plus psychosocial burden among middle-aged and older adults in India: a survey-weighted analysis of harmonized LASI

## Abstract

### Background
Skin and subcutaneous diseases are common in older adults, but nationally representative evidence linking self-reported skin disease to functional and psychosocial burden in ageing populations is limited. We assessed the prevalence of self-reported skin disease and its associations with pain, mobility, sleep, mood, and daily function among middle-aged and older adults in India.

### Methods
We performed a cross-sectional analysis of Harmonized LASI A.3 and included adults aged 45 years or older with non-missing skin-disease status and respondent weights. Because the public harmonized LASI file does not release PSU or design-strata variables, primary analyses used survey-weighted models with household clustering and state stratification proxy. Survey-weighted logistic regression estimated associations between skin disease and six burden outcomes. Interaction analyses assessed heterogeneity by sex, rural residence, and age 75 years or older. Sensitivity analyses excluded proxy respondents and additionally adjusted for BMI and physical inactivity.

### Results
The analysis included {s['analytic_n']:,} respondents. Survey-weighted prevalence of self-reported skin disease was {overall['Prevalence_pct']:.2f}% (95% CI {overall['CI_low_pct']:.2f}-{overall['CI_high_pct']:.2f}). Skin disease was associated with frequent pain (OR {pain['OR']:.3f}, 95% CI {pain['CI_low']:.3f}-{pain['CI_high']:.3f}), mobility limitation ({mobility['OR']:.3f}, {mobility['CI_low']:.3f}-{mobility['CI_high']:.3f}), sleep problems ({sleep['OR']:.3f}, {sleep['CI_low']:.3f}-{sleep['CI_high']:.3f}), depressive symptoms ({dep['OR']:.3f}, {dep['CI_low']:.3f}-{dep['CI_high']:.3f}), ADL limitation ({adl['OR']:.3f}, {adl['CI_low']:.3f}-{adl['CI_high']:.3f}), and IADL limitation ({iadl['OR']:.3f}, {iadl['CI_low']:.3f}-{iadl['CI_high']:.3f}). Sleep problems showed sex interaction (p={sleep_int['Sex_interaction_p']:.4f}) and mobility limitation showed rural interaction (p={mobility_int['Rural_interaction_p']:.4f}). Associations were stable in sensitivity analyses.

### Conclusions
Self-reported skin disease in older adults in India was associated with broad functional and psychosocial burden. These findings support integrating skin health into healthy-ageing, pain, sleep, and function-oriented care strategies.

## Keywords

Skin disease; ageing; LASI; India; survey-weighted analysis; pain; mobility; sleep

## Background

Skin disorders are often treated as localized clinical problems, but in older adults they can also influence comfort, mobility, sleep, mood, and daily participation.[5-10] This broader burden is important for healthy-ageing policy, especially in low-income and middle-income settings where health systems are simultaneously managing chronic disease, disability, and long-term care demand.[1-4]

Despite this relevance, much of the available literature falls into one of two categories: disease-specific clinical reports or ecological burden analyses. The former can be rich clinically but are rarely nationally representative. The latter provide strong population overviews but cannot establish whether older adults who report skin disease also experience broader individual burden. A nationally representative cohort analysis is therefore a useful next step.

The Longitudinal Ageing Study in India offers a practical setting for this question. LASI captures a wide range of demographic, health, and function variables in a harmonized structure, making it suitable for person-level validation of the ageing narrative emerging from global skin-burden research.[1-4] In this study, we used Harmonized LASI to examine how self-reported skin disease relates to multiple dimensions of burden in adults aged 45 years or older.

## Methods

### Study design and data source

We analysed Harmonized LASI A.3, a public harmonized version of LASI baseline data. The local analytic source was `/Users/apple/Desktop/所有数据/global aging data数据/LASI_印度/Harmonized LASI A.3_SPSS/H_LASI_a3.sav`.

### Participants and variables

Participants were included if they were aged 45 years or older, had non-missing respondent weights, and had non-missing self-reported skin disease status. The exposure was `r1skindise`. Covariates were age, sex, rural residence, and education. Outcomes were ADL limitation, IADL limitation, mobility limitation, frequent pain, depressive symptoms, and sleep problems. Additional variables for sensitivity analysis included BMI, vigorous inactivity, moderate inactivity, and proxy interview status.

### Statistical analysis

Primary analyses used survey-weighted methods with household clustering and state stratification proxy because the public harmonized file provides respondent weights but not the original PSU or design-strata variables. Weighted prevalence estimates were generated overall and by subgroup. Survey-weighted logistic regression models estimated correlates of skin disease and associations between skin disease and each burden outcome. We tested interaction terms for sex, rural residence, and age 75 years or older. Sensitivity analyses excluded proxy respondents and additionally adjusted for BMI and physical inactivity. This report follows the logic of STROBE for observational analyses.[16]

## Results

### Prevalence and participant characteristics

Survey-weighted prevalence of self-reported skin disease was {overall['Prevalence_pct']:.2f}% (95% CI {overall['CI_low_pct']:.2f}-{overall['CI_high_pct']:.2f}). Prevalence was highest in adults aged 75 years or older ({age75['Prevalence_pct']:.2f}%), and was higher in men than women ({men['Prevalence_pct']:.2f}% vs {women['Prevalence_pct']:.2f}%) and in rural than urban respondents ({rural['Prevalence_pct']:.2f}% vs {urban['Prevalence_pct']:.2f}%).

Participants reporting skin disease had higher weighted prevalence of every burden outcome than those without skin disease, especially frequent pain, mobility limitation, and sleep problems (Table 1). Mean age was also slightly higher in the skin-disease group.

### Correlates and burden associations

In survey-weighted models, older age and rural residence were positively associated with skin disease, whereas women were less likely than men to report skin disease. The education association was attenuated and not conventionally significant after survey variance estimation (Supplementary Table S3).

Skin disease was associated with higher odds of all six burden outcomes. The largest effect size was observed for frequent pain (OR {pain['OR']:.3f}, 95% CI {pain['CI_low']:.3f}-{pain['CI_high']:.3f}), followed by mobility limitation (OR {mobility['OR']:.3f}, {mobility['CI_low']:.3f}-{mobility['CI_high']:.3f}) and sleep problems (OR {sleep['OR']:.3f}, {sleep['CI_low']:.3f}-{sleep['CI_high']:.3f}).

### Interaction and sensitivity analyses

The sex interaction was significant for sleep problems, with a stronger skin-disease association in women (OR {sleep_women['OR']:.3f}, 95% CI {sleep_women['CI_low']:.3f}-{sleep_women['CI_high']:.3f}) than in men (OR {sleep_men['OR']:.3f}, {sleep_men['CI_low']:.3f}-{sleep_men['CI_high']:.3f}). The rural interaction was significant for mobility limitation, with stronger effects in urban respondents (OR {mobility_urban['OR']:.3f}, 95% CI {mobility_urban['CI_low']:.3f}-{mobility_urban['CI_high']:.3f}) than in rural respondents (OR {mobility_rural['OR']:.3f}, {mobility_rural['CI_low']:.3f}-{mobility_rural['CI_high']:.3f}).

Sensitivity analyses were reassuring. Excluding proxy respondents preserved the magnitude and direction of the outcome associations, including frequent pain (OR {no_proxy_pain['OR']:.3f}, 95% CI {no_proxy_pain['CI_low']:.3f}-{no_proxy_pain['CI_high']:.3f}). Extended adjustment for BMI and physical inactivity also produced closely similar estimates (frequent pain OR {ext_pain['OR']:.3f}, 95% CI {ext_pain['CI_low']:.3f}-{ext_pain['CI_high']:.3f}).

## Discussion

This study shows that self-reported skin disease in older adults in India is associated with broad multidomain burden. The results are consistent across pain, mobility, daily functioning, sleep, and depressive symptoms, suggesting that skin health should be considered part of integrated geriatric care rather than an isolated specialty concern.

The findings are relevant in two ways. Clinically, they support more routine attention to skin symptoms in older adults presenting with pain, mobility decline, or sleep problems. Epidemiologically, they provide person-level validation for the ageing perspective adopted in the parallel global skin-burden manuscript. They also show the analytic importance of survey-weighted variance estimation, because some signals from simpler weighted models were attenuated when design-based uncertainty was incorporated.

There are limitations. The skin variable is self-reported, broad, and not diagnosis-specific. The harmonized public file does not contain the original PSU and design-strata variables, so household clustering and state stratification proxy were used as a conservative alternative. The study is cross-sectional and residual confounding is likely. Nonetheless, the national scope, survey-weighted design, interaction testing, and sensitivity analyses strengthen the credibility of the findings.

## Conclusions

In Harmonized LASI, self-reported skin disease was associated with substantial functional and psychosocial burden among middle-aged and older adults in India. These findings support integrating skin health into healthy-ageing research and service planning.

## Declarations

### Ethics approval and consent to participate
This study used de-identified secondary data from Harmonized LASI. Ethical approval and participant consent were handled by the original LASI investigators. Any additional local determination should be completed by the submitting author team.

### Consent for publication
Not applicable.

### Availability of data and materials
The underlying data are available through the LASI and Gateway to Global Aging Data access framework. Reproducible local analysis scripts are available in `/Users/apple/Desktop/lancet-research-platform/analysis/r/55_lasi_skin_survey_analysis.R` and `/Users/apple/Desktop/lancet-research-platform/analysis/python/56_build_lasi_skin_journal_variants.py`.

### Competing interests
To be completed by the author team.

### Funding
No specific funding was used for this secondary analysis package. Any broader project funding statements should be completed by the author team.

### Authors' contributions
To be completed by the author team.

### Acknowledgements
We acknowledge the investigators and participants of LASI and the Gateway to Global Aging Data harmonization effort.

## References

"""
    bmc_md += "\n".join([f"{i}. {ref}" for i, ref in enumerate(references(), start=1)]) + "\n"
    return {"lancet_md": lancet_md, "bmc_md": bmc_md}


def save_md(path: Path, text: str) -> int:
    path.write_text(text, encoding="utf-8")
    return len(text.replace("\n", " ").split())


def build_lancet_doc(data: dict, text: str) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Self-reported skin disease and functional plus psychosocial burden among middle-aged and older adults in India: a survey-weighted analysis of harmonized LASI")
    add_para(doc, "Drafted in Lancet-style structure for author revision and journal targeting.", italic=True)

    sections = text.split("\n## ")
    first = sections[0].replace("# ", "")
    if "## Summary" not in text:
        raise RuntimeError("Lancet text missing Summary")

    # Summary
    summary_part = text.split("## Summary\n\n", 1)[1].split("\n## Research in context", 1)[0]
    add_heading(doc, "Summary", 1)
    for block in summary_part.strip().split("\n\n"):
        if block.startswith("### "):
            continue
        if block.startswith("Background\n"):
            label, body = block.split("\n", 1)
            add_para(doc, body, bold_label=label)
    for label in ["Background", "Methods", "Findings", "Interpretation", "Funding"]:
        marker = f"### {label}\n"
        body = summary_part.split(marker, 1)[1].split("\n### ", 1)[0].strip()
        add_para(doc, body, bold_label=label)

    ric = text.split("## Research in context\n\n", 1)[1].split("\n## Introduction", 1)[0]
    add_heading(doc, "Research in context", 1)
    for label in ["Evidence before this study", "Added value of this study", "Implications of all the available evidence"]:
        marker = f"### {label}\n"
        body = ric.split(marker, 1)[1].split("\n### ", 1)[0].strip()
        add_para(doc, body, bold_label=label)

    body_text = text.split("## Introduction\n\n", 1)[1].split("\n## References", 1)[0]
    body_sections = ["Introduction", "Methods", "Results", "Discussion"]
    current = "Introduction"
    add_heading(doc, current, 1)
    for chunk in body_text.split("\n\n"):
        if chunk in body_sections[1:]:
            current = chunk
            add_heading(doc, current, 1)
        elif chunk.startswith("### "):
            add_heading(doc, chunk.replace("### ", ""), 2)
        else:
            add_para(doc, chunk.strip())

    add_heading(doc, "References", 1)
    refs = text.split("## References\n\n", 1)[1].strip().split("\n")
    for ref in refs:
        add_para(doc, ref)

    doc.save(MANUSCRIPT_DIR / "07_Lancet_Style_Manuscript_LASI_Skin_Disease_20260312.docx")


def build_bmc_doc(text: str) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Self-reported skin disease and functional plus psychosocial burden among middle-aged and older adults in India: a survey-weighted analysis of harmonized LASI")
    add_para(doc, "Drafted in BMC Geriatrics style for author revision and journal targeting.", italic=True)

    main = text.split("## References\n\n", 1)[0]
    refs = text.split("## References\n\n", 1)[1].strip().split("\n")
    for chunk in main.split("\n\n"):
        if chunk.startswith("# "):
            continue
        if chunk.startswith("## "):
            add_heading(doc, chunk.replace("## ", ""), 1)
        elif chunk.startswith("### "):
            add_heading(doc, chunk.replace("### ", ""), 2)
        else:
            add_para(doc, chunk.strip())
    add_heading(doc, "References", 1)
    for ref in refs:
        add_para(doc, ref)
    doc.save(MANUSCRIPT_DIR / "08_BMC_Geriatrics_Style_Manuscript_LASI_Skin_Disease_20260312.docx")


def build_display_items_doc(data: dict) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Survey-weighted main display items for LASI skin disease manuscript")
    add_caption(doc, "Table 1. Weighted characteristics of adults aged 45 years and older according to self-reported skin disease status in Harmonized LASI")
    add_table(doc, display_characteristics_table(data["table1"]))
    note(doc, "Values are survey-weighted except participant counts, which are unweighted analytic counts.")

    doc.add_page_break()
    add_caption(doc, "Figure 1. Survey-weighted prevalence of self-reported skin disease by age, sex, and rural residence")
    doc.add_picture(str(FIGURE_DIR / "figure1_survey_weighted_prevalence.png"), width=Inches(6.6))
    note(doc, "Points show weighted prevalence estimates and whiskers show 95% CIs.")

    doc.add_page_break()
    add_caption(doc, "Figure 2. Survey-weighted associations between skin disease and functional and psychosocial outcomes")
    doc.add_picture(str(FIGURE_DIR / "figure2_survey_weighted_outcome_forest.png"), width=Inches(6.6))
    note(doc, "Odds ratios are adjusted for age, sex, rural residence, and education.")

    doc.save(MANUSCRIPT_DIR / "09_Journal_Display_Items_LASI_Skin_Disease_20260312.docx")


def build_supplement_doc(data: dict) -> None:
    doc = Document()
    style_doc(doc)
    add_title(doc, "Survey-weighted supplementary appendix for LASI skin disease manuscript")
    add_para(doc, data["summary"]["design_note"], italic=True)

    tables = [
        ("Supplementary Table S1. Variable definitions and design choices", data["tableS1"]),
        ("Supplementary Table S2. Additional survey-weighted prevalence estimates", display_prevalence_table(data["tableS2"])),
        ("Supplementary Table S3. Survey-weighted correlates of reporting skin disease", display_or_table(data["table3"], "Predictor")),
        ("Supplementary Table S4. Outcome associations after excluding proxy respondents", display_or_table(data["tableS3"], "Outcome")),
        ("Supplementary Table S5. Outcome associations after extended adjustment for BMI and physical inactivity", display_or_table(data["tableS4"], "Outcome")),
        ("Supplementary Table S6. Interaction tests for sex, rural residence, and age 75 years or older", data["tableS5"].assign(
            Sex_interaction_p=data["tableS5"]["Sex_interaction_p"].map(pvalue),
            Rural_interaction_p=data["tableS5"]["Rural_interaction_p"].map(pvalue),
            Age75plus_interaction_p=data["tableS5"]["Age75plus_interaction_p"].map(pvalue),
        )),
        ("Supplementary Table S7. Stratified estimates for interaction-positive findings", display_or_table(data["tableS6"].rename(columns={"Stratum": "Stratum"}), "Stratum")),
    ]
    for idx, (caption, table_df) in enumerate(tables, start=1):
        add_caption(doc, caption)
        add_table(doc, table_df)
        if idx != len(tables):
            doc.add_page_break()

    doc.add_page_break()
    add_caption(doc, "Supplementary Figure S1. Stratified effects for interaction-positive outcomes")
    doc.add_picture(str(FIGURE_DIR / "figureS1_survey_stratified_interactions.png"), width=Inches(6.6))
    note(doc, "The sex interaction was retained for sleep problems and the rural interaction was retained for mobility limitation.")

    doc.add_page_break()
    add_caption(doc, "Supplementary Figure S2. Sensitivity analyses for the association between skin disease and burden outcomes")
    doc.add_picture(str(FIGURE_DIR / "figureS2_survey_sensitivity_comparison.png"), width=Inches(6.8))
    note(doc, "Primary survey estimates are compared against the proxy-exclusion and extended-adjustment sensitivity analyses.")

    doc.save(MANUSCRIPT_DIR / "10_Journal_Supplementary_Appendix_LASI_Skin_Disease_20260312.docx")


def build_qc_doc(lancet_words: int, bmc_words: int, data: dict) -> None:
    s = data["summary"]
    doc = Document()
    style_doc(doc)
    add_title(doc, "QC report for survey-weighted and journal-style LASI package")
    add_heading(doc, "Analysis QC", 1)
    add_bullet(doc, f"Primary survey design used {s['households']:,} households and {s['states']} state strata proxies with respondent weights.")
    add_bullet(doc, "Public harmonized LASI does not release PSU and design-strata variables; the design was therefore approximated conservatively with household clustering and state stratification proxy.")
    add_bullet(doc, f"Survey-weighted prevalence of self-reported skin disease was {s['weighted_prevalence_pct']:.2f}% (95% CI {s['weighted_prevalence_ci'][0]:.2f}-{s['weighted_prevalence_ci'][1]:.2f}).")
    add_bullet(doc, f"Interaction testing identified sleep problems by sex (p={s['significant_interactions']['sex_sleep_problem_p']:.4f}) and mobility limitation by rural residence (p={s['significant_interactions']['rural_mobility_p']:.4f}).")
    add_bullet(doc, "Sensitivity analyses excluding proxy respondents and additionally adjusting for BMI and physical inactivity were directionally consistent.")

    add_heading(doc, "Journal-style outputs", 1)
    add_bullet(doc, f"Lancet-style manuscript word count: {lancet_words}")
    add_bullet(doc, f"BMC Geriatrics-style manuscript word count: {bmc_words}")
    add_bullet(doc, "Main display items document includes Table 1 and Figures 1-2.")
    add_bullet(doc, "Supplementary appendix includes Tables S1-S7 and Figures S1-S2.")

    add_heading(doc, "Style notes", 1)
    add_bullet(doc, "Lancet-style draft was organised with Summary, Research in context, Introduction, Methods, Results, Discussion, and References.")
    add_bullet(doc, "BMC Geriatrics-style draft was organised with structured abstract, keywords, Background, Methods, Results, Discussion, Conclusions, and Declarations.")
    add_bullet(doc, "JEADV was reviewed through the official author hub; the BMC-style clinical IMRAD variant is the closer general-purpose dermatology-journal baseline in this package.")

    doc.save(MANUSCRIPT_DIR / "11_Survey_Journal_QC_Report_LASI_Skin_Disease_20260312.docx")


def build_notes() -> None:
    notes = """# Journal targeting notes

- Lancet structure was aligned to the official submission guidance and formatting portal: https://www.editorialmanager.com/lancet/default2.aspx
- BMC Geriatrics structure was aligned to the official manuscript-formatting page: https://bmcgeriatr.biomedcentral.com/submission-guidelines/preparing-your-manuscript/research-article
- JEADV author guidance was checked through the official EADV author hub: https://eadv.org/our-journals/author-hub/
- These variants are style-targeted writing packages rather than final publisher-formatted submission files.
"""
    (MANUSCRIPT_DIR / "12_Journal_Targeting_Notes_20260312.md").write_text(notes, encoding="utf-8")


def update_readme(lancet_words: int, bmc_words: int) -> None:
    text = f"""# LASI Skin Follow-up Package

This folder contains a second-study manuscript framework built from Harmonized LASI after the desktop Global Aging Data audit.

## Survey-weighted extension

- Analytic sample: 66,412
- Main survey-weighted prevalence: 5.07% (95% CI 4.86-5.29)
- Significant interaction findings: sleep problems by sex, mobility limitation by rural residence
- Lancet-style manuscript word count: {lancet_words}
- BMC Geriatrics-style manuscript word count: {bmc_words}

## Key files

- `outputs/manuscript/07_Lancet_Style_Manuscript_LASI_Skin_Disease_20260312.docx`
- `outputs/manuscript/08_BMC_Geriatrics_Style_Manuscript_LASI_Skin_Disease_20260312.docx`
- `outputs/manuscript/09_Journal_Display_Items_LASI_Skin_Disease_20260312.docx`
- `outputs/manuscript/10_Journal_Supplementary_Appendix_LASI_Skin_Disease_20260312.docx`
- `outputs/manuscript/11_Survey_Journal_QC_Report_LASI_Skin_Disease_20260312.docx`
"""
    (PACKAGE_ROOT / "README.md").write_text(text, encoding="utf-8")


def main() -> None:
    data = load_inputs()
    texts = build_texts(data)
    lancet_words = save_md(MANUSCRIPT_DIR / "07_Lancet_Style_Manuscript_LASI_Skin_Disease_20260312.md", texts["lancet_md"])
    bmc_words = save_md(MANUSCRIPT_DIR / "08_BMC_Geriatrics_Style_Manuscript_LASI_Skin_Disease_20260312.md", texts["bmc_md"])
    build_lancet_doc(data, texts["lancet_md"])
    build_bmc_doc(texts["bmc_md"])
    build_display_items_doc(data)
    build_supplement_doc(data)
    build_qc_doc(lancet_words, bmc_words, data)
    build_notes()
    update_readme(lancet_words, bmc_words)

    manifest = {
        "lancet_words": lancet_words,
        "bmc_words": bmc_words,
        "files": sorted([p.name for p in MANUSCRIPT_DIR.glob("07_*")] + [p.name for p in MANUSCRIPT_DIR.glob("08_*")] + [p.name for p in MANUSCRIPT_DIR.glob("09_*")] + [p.name for p in MANUSCRIPT_DIR.glob("10_*")] + [p.name for p in MANUSCRIPT_DIR.glob("11_*")] + [p.name for p in MANUSCRIPT_DIR.glob("12_*")]),
    }
    (MANUSCRIPT_DIR / "journal_variant_manifest_20260312.json").write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    print(json.dumps(manifest, indent=2))


if __name__ == "__main__":
    main()
