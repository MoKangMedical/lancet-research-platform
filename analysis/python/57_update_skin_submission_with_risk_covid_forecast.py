#!/usr/bin/env python3
from __future__ import annotations

import importlib.util
import json
import math
import re
import shutil
import sys
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import requests
from docx import Document
from mpl_toolkits.axes_grid1.inset_locator import inset_axes
from scipy.stats import mannwhitneyu, spearmanr
from statsmodels.nonparametric.smoothers_lowess import lowess
from statsmodels.tsa.holtwinters import ExponentialSmoothing


ROOT = Path("/Users/apple/Desktop/lancet-research-platform")
ANALYSIS_PY = ROOT / "analysis" / "python"
if str(ANALYSIS_PY) not in sys.path:
    sys.path.insert(0, str(ANALYSIS_PY))

from lib.rendering import render_docx_collection


PROJECT_DIR = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-针对皮肤病的相关全球流行病和疾病负担研究方案-20分-38万-已收5万+5万"
)
PACKAGE_ROOT = PROJECT_DIR / "lancet_skin_article_package"
OUTPUT_ROOT = PACKAGE_ROOT / "outputs"
FINAL_DIR = PACKAGE_ROOT / "submission_package_final_20260309"
UPDATED_DIR = PACKAGE_ROOT / "submission_package_update_20260326_resource_context"
FIGURE_DIR = UPDATED_DIR / "figures"
TABLE_DIR = UPDATED_DIR / "tables"
ANALYSIS_DIR = UPDATED_DIR / "analysis_outputs"
RENDER_DIR = UPDATED_DIR / "rendered_pages_update_20260326"

AGING_DIR = OUTPUT_ROOT / "aging_analysis_outputs"
MANUSCRIPT_DIR = OUTPUT_ROOT / "manuscript"

SCRIPT_45 = ANALYSIS_PY / "45_build_skin_lancet_package.py"
SCRIPT_50 = ANALYSIS_PY / "50_build_skin_lancet_complete_apac.py"

STATE_PATH = ROOT / "output" / "playwright" / "gbdlogin_state_20260312.json"
RESULTS_URL = "https://vizhub.healthdata.org/gbd-results/php/data.php"
RESULTS_VERSION = 8352
CAUSE_ID = 653
GLOBAL_LOCATION_ID = 1
AGE_ID_AGE_STANDARDIZED = 27
SEX_ID_BOTH = 3
MEASURE_MAP = {
    1: "Deaths",
    2: "DALYs",
    5: "Prevalence",
    6: "Incidence",
}
MEASURE_ORDER = ["Deaths", "DALYs", "Prevalence", "Incidence"]
PHYSICIAN_INDICATOR = "SH.MED.PHYS.ZS"
COMPARATOR_CAUSES = {
    653: "Skin and subcutaneous diseases",
    491: "Cardiovascular diseases",
    410: "Neoplasms",
    974: "Diabetes and kidney diseases",
}

TITLE = (
    "Global burden of skin and subcutaneous diseases in the context of population ageing, "
    "1990-2023: a systematic analysis of GBD 2023 and World Bank ageing indicators"
)

COLOR_NAVY = "#003A70"
COLOR_STEEL = "#7F9DB9"
COLOR_PALE_BLUE = "#DCE6F2"
COLOR_SLATE = "#5A6B7A"
COLOR_RUST = "#B55D4C"
COLOR_TEAL = "#2E6F62"
COLOR_GOLD = "#C08B30"
COLOR_GREY = "#ECEFF3"
COLOR_ROW = "#F7F9FB"
COLOR_GRID = "#C8D0D9"
COLOR_TEXT = "#1F2933"
MEASURE_COLORS = {
    "Deaths": COLOR_GOLD,
    "DALYs": COLOR_TEAL,
    "Prevalence": COLOR_RUST,
    "Incidence": COLOR_NAVY,
}
HEATMAP_CMAP = mpl.colors.LinearSegmentedColormap.from_list(
    "risk_heatmap", ["#F6F8FB", "#A8C7BB", COLOR_NAVY]
)

MAIN_DOCX = UPDATED_DIR / "1_Main_Manuscript_Lancet_Updated_20260326.docx"
MAIN_MD = UPDATED_DIR / "1_Main_Manuscript_Lancet_Updated_20260326.md"
SUPP_DOCX = UPDATED_DIR / "4_Supplementary_Appendix_Lancet_Updated_20260326.docx"
SUPP_MD = UPDATED_DIR / "4_Supplementary_Appendix_Lancet_Updated_20260326.md"
RIC_DOCX = UPDATED_DIR / "5_Research_in_Context_Lancet_Updated_20260326.docx"
RIC_MD = UPDATED_DIR / "5_Research_in_Context_Lancet_Updated_20260326.md"
REF_DOCX = UPDATED_DIR / "6_Reference_List_Lancet_Updated_20260326.docx"
REF_MD = UPDATED_DIR / "6_Reference_List_Lancet_Updated_20260326.md"
CHECKLIST_DOCX = UPDATED_DIR / "11_Submission_Checklist_20260326.docx"
CHECKLIST_MD = UPDATED_DIR / "11_Submission_Checklist_20260326.md"
QC_DOCX = UPDATED_DIR / "12_QC_Report_Updated_20260326.docx"
QC_MD = UPDATED_DIR / "12_QC_Report_Updated_20260326.md"
README_DOCX = UPDATED_DIR / "README_Updated_Submission_Package_20260326.docx"
README_MD = UPDATED_DIR / "README_Updated_Submission_Package_20260326.md"
SUMMARY_JSON = UPDATED_DIR / "analysis_summary_update_20260326.json"
MANIFEST_JSON = UPDATED_DIR / "MANIFEST_20260326.json"
SOURCE_FIGURE_DIRS = [OUTPUT_ROOT / "figures", FINAL_DIR / "figures"]


def load_module(module_name: str, script_path: Path):
    spec = importlib.util.spec_from_file_location(module_name, script_path)
    module = importlib.util.module_from_spec(spec)
    assert spec.loader is not None
    sys.modules[module_name] = module
    spec.loader.exec_module(module)
    return module


def word_count(text: str) -> int:
    return len(re.findall(r"\b[\w.-]+\b", text))


def pct_change(start: float, end: float) -> float:
    if start == 0:
        return math.nan
    return (end - start) / start * 100.0


def fmt_p(value: float) -> str:
    if pd.isna(value):
        return "NA"
    return f"{value:.2e}"


def fmt_rate(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def fmt_pct(value: float, digits: int = 1) -> str:
    return f"{value:.{digits}f}%"


def ensure_clean_update_dir() -> None:
    if UPDATED_DIR.exists():
        shutil.rmtree(UPDATED_DIR)
    shutil.copytree(FINAL_DIR, UPDATED_DIR)
    for path in [FIGURE_DIR, TABLE_DIR, ANALYSIS_DIR, RENDER_DIR]:
        if path.exists():
            shutil.rmtree(path)
        path.mkdir(parents=True, exist_ok=True)


def remove_obsolete_docs() -> None:
    patterns = [
        "1_Main_Manuscript_Lancet_Final_20260309.*",
        "4_Supplementary_Appendix_Lancet_Final_20260309.*",
        "5_Research_in_Context_Lancet_Final_20260309.*",
        "6_Reference_List_Lancet_Final_20260309.*",
        "11_Submission_Checklist_20260309.*",
        "12_QC_Report_Final_20260309.*",
        "README_Final_Submission_Package_20260309.*",
        "MANIFEST_20260309.json",
    ]
    for pattern in patterns:
        for path in UPDATED_DIR.glob(pattern):
            if path.is_file():
                path.unlink()


def write_markdown(path: Path, title: str, lines: Sequence[str]) -> None:
    content = [f"# {title}", ""]
    content.extend(lines)
    path.write_text("\n".join(content).rstrip() + "\n", encoding="utf-8")


def write_docx_with_headings(module50, path: Path, title: str, lines: Sequence[str]) -> None:
    doc = Document()
    module50.configure_doc(doc)
    module50.add_title(doc, title)
    for line in lines:
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            module50.add_heading(doc, line[3:], level=1)
        elif line.startswith("### "):
            module50.add_heading(doc, line[4:], level=2)
        else:
            module50.add_paragraph(doc, line)
    doc.save(path)


def copy_asset(src: Path, dst: Path) -> None:
    if src.exists():
        shutil.copy2(src, dst)


def copy_source_figure(filename: str, target_name: Optional[str] = None) -> None:
    for source_dir in SOURCE_FIGURE_DIRS:
        candidate = source_dir / filename
        if candidate.exists():
            shutil.copy2(candidate, FIGURE_DIR / (target_name or filename))
            return
    raise FileNotFoundError(f"Source figure not found in configured directories: {filename}")


def save_figure(fig: plt.Figure, stem: str) -> Tuple[Path, Path]:
    png_path = FIGURE_DIR / f"{stem}.png"
    pdf_path = FIGURE_DIR / f"{stem}.pdf"
    fig.patch.set_facecolor("white")
    fig.savefig(png_path, dpi=300, bbox_inches="tight", facecolor="white")
    fig.savefig(pdf_path, dpi=300, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return png_path, pdf_path


def fetch_worldbank_country_reference() -> pd.DataFrame:
    frames: List[pd.DataFrame] = []
    page = 1
    while True:
        url = f"https://api.worldbank.org/v2/country?format=json&per_page=400&page={page}"
        payload = requests.get(url, timeout=60).json()
        meta = payload[0]
        rows = payload[1]
        frame = pd.DataFrame(
            [
                {
                    "wb_iso3": item.get("id"),
                    "wb_name_full": item.get("name"),
                    "longitude": pd.to_numeric(item.get("longitude"), errors="coerce"),
                    "latitude": pd.to_numeric(item.get("latitude"), errors="coerce"),
                }
                for item in rows
            ]
        )
        frames.append(frame)
        if page >= int(meta["pages"]):
            break
        page += 1
    return pd.concat(frames, ignore_index=True).drop_duplicates(subset=["wb_iso3"])


def fetch_official_global_country_rates_2023(access_token: str, location_ids: Sequence[int]) -> pd.DataFrame:
    payload: List[Tuple[str, object]] = [("version", RESULTS_VERSION)]
    for measure_id in MEASURE_MAP:
        payload.append(("measure[]", measure_id))
    for location_id in location_ids:
        payload.append(("location[]", int(location_id)))
    payload.extend(
        [
            ("sex[]", SEX_ID_BOTH),
            ("age[]", AGE_ID_AGE_STANDARDIZED),
            ("cause[]", CAUSE_ID),
            ("metric[]", 3),
            ("year[]", 2023),
            ("population_group", 1),
            ("api_version", "2023.0.0"),
            ("base", "single"),
            ("context", "cause"),
            ("singleOrMult", "single"),
            ("idsOrNames", "ids"),
            ("rows", 500000),
            ("start_year", 1980),
            ("fetch_all_years", "false"),
            ("language", "en"),
        ]
    )
    last_error: Optional[Exception] = None
    response = None
    for _ in range(3):
        try:
            response = requests.post(
                RESULTS_URL,
                data=payload,
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=120,
            )
            response.raise_for_status()
            break
        except Exception as exc:  # pragma: no cover - network retry
            last_error = exc
            response = None
    if response is None:
        raise RuntimeError(f"Failed to fetch country-level Results Tool series after retries: {last_error}")
    raw = response.json()
    df = pd.DataFrame(raw["data"], columns=raw["cols"]).copy()
    df.loc[:, "measure_name"] = df["measure"].map(MEASURE_MAP)
    df = df.rename(columns={"location": "location_id", "year": "year_id", "val": "mean"})
    return df[
        [
            "location_id",
            "measure_name",
            "year_id",
            "mean",
            "lower",
            "upper",
        ]
    ].sort_values(["measure_name", "location_id"])


def build_global_map_dataset(module50, country_complete: pd.DataFrame, global_country_rates: pd.DataFrame) -> pd.DataFrame:
    wb_ref = fetch_worldbank_country_reference()
    wide = (
        global_country_rates.pivot_table(
            index="location_id",
            columns="measure_name",
            values=["mean", "lower", "upper"],
            aggfunc="first",
        )
        .sort_index(axis=1)
        .reset_index()
    )
    cols: List[str] = []
    for col in wide.columns:
        if col == ("location_id", ""):
            cols.append("location_id")
        else:
            stat, measure = col
            cols.append(
                f"{measure.lower()}_asr_2023" if stat == "mean" else f"{measure.lower()}_asr_2023_{stat}"
            )
    wide.columns = cols
    merged = country_complete.merge(wide, on="location_id", how="left", validate="one_to_one")
    merged = merged.merge(wb_ref, on="wb_iso3", how="left", validate="many_to_one")
    world = module50.load_world_shapes().copy()
    polygon_iso = set(world["iso_a3"].dropna())
    merged.loc[:, "plot_method"] = np.where(merged["wb_iso3"].isin(polygon_iso), "polygon", "point")
    return merged.sort_values("gbd_name").reset_index(drop=True)


def style_lancet_axes(ax) -> None:
    ax.set_facecolor("white")
    ax.grid(axis="y", color=COLOR_GRID, alpha=0.28, linewidth=0.55)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.spines["left"].set_color(COLOR_GRID)
    ax.spines["bottom"].set_color(COLOR_GRID)
    ax.spines["left"].set_linewidth(0.8)
    ax.spines["bottom"].set_linewidth(0.8)
    ax.tick_params(colors=COLOR_TEXT, labelsize=8.2, width=0.8, length=3)


def add_panel_label(ax, label: str) -> None:
    ax.text(
        -0.12,
        1.04,
        label,
        transform=ax.transAxes,
        fontsize=11.5,
        fontweight="bold",
        va="bottom",
        ha="left",
        color=COLOR_NAVY,
    )


def add_end_label(ax, x: float, y: float, text: str, color: str, dx: float = 0.5) -> None:
    ax.text(
        x + dx,
        y,
        text,
        color=color,
        fontsize=7.9,
        va="center",
        ha="left",
        bbox={"facecolor": "white", "edgecolor": "none", "pad": 0.6, "alpha": 0.85},
    )


def make_figure1_lancet_style(global_context: pd.DataFrame, annual_df: pd.DataFrame) -> Tuple[Path, Path]:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.edgecolor": COLOR_GRID,
            "axes.labelcolor": COLOR_TEXT,
            "xtick.color": COLOR_TEXT,
            "ytick.color": COLOR_TEXT,
        }
    )
    fig, axes = plt.subplots(2, 2, figsize=(11.1, 7.7))
    ax = axes[0, 0]
    inc = annual_df[annual_df["measure_name"] == "Incidence"].sort_values("year_id")
    prev = annual_df[annual_df["measure_name"] == "Prevalence"].sort_values("year_id")
    ax.plot(inc["year_id"], inc["mean"], color=MEASURE_COLORS["Incidence"], linewidth=2.25)
    ax.plot(prev["year_id"], prev["mean"], color=MEASURE_COLORS["Prevalence"], linewidth=2.25)
    ax.axvspan(2020, 2023.5, color=COLOR_GREY, alpha=0.40, zorder=0)
    ax.set_title("Age-standardised incidence and prevalence", fontsize=10.2, loc="left", color=COLOR_TEXT)
    ax.set_ylabel("ASR per 100 000")
    ax.set_xlim(1990, 2025.2)
    add_end_label(ax, float(inc["year_id"].iloc[-1]), float(inc["mean"].iloc[-1]) - 120, "Incidence", MEASURE_COLORS["Incidence"])
    add_end_label(ax, float(prev["year_id"].iloc[-1]), float(prev["mean"].iloc[-1]), "Prevalence", MEASURE_COLORS["Prevalence"])
    ax.text(2021.8, ax.get_ylim()[1] * 0.985, "Pandemic\nperiod", fontsize=6.9, color=COLOR_SLATE, va="top", ha="center")
    style_lancet_axes(ax)
    add_panel_label(ax, "A")

    ax = axes[0, 1]
    dalys = annual_df[annual_df["measure_name"] == "DALYs"].sort_values("year_id")
    deaths = annual_df[annual_df["measure_name"] == "Deaths"].sort_values("year_id")
    ax.plot(dalys["year_id"], dalys["mean"], color=MEASURE_COLORS["DALYs"], linewidth=2.25)
    ax2 = ax.twinx()
    ax2.plot(deaths["year_id"], deaths["mean"], color=MEASURE_COLORS["Deaths"], linewidth=2.25)
    ax.axvspan(2020, 2023.5, color=COLOR_GREY, alpha=0.40, zorder=0)
    ax.set_title("Age-standardised DALY and mortality rates", fontsize=10.2, loc="left", color=COLOR_TEXT)
    ax.set_ylabel("DALY ASR per 100 000")
    ax2.set_ylabel("Mortality ASR per 100 000", color=COLOR_TEXT)
    ax.set_xlim(1990, 2025.2)
    style_lancet_axes(ax)
    ax2.spines["top"].set_visible(False)
    ax2.spines["left"].set_visible(False)
    ax2.spines["right"].set_color(COLOR_GRID)
    ax2.tick_params(colors=COLOR_TEXT, labelsize=8.2, width=0.8, length=3)
    add_end_label(ax, float(dalys["year_id"].iloc[-1]), float(dalys["mean"].iloc[-1]), "DALYs", MEASURE_COLORS["DALYs"])
    ax2.text(
        float(deaths["year_id"].iloc[-1]) + 0.5,
        float(deaths["mean"].iloc[-1]),
        "Deaths",
        color=MEASURE_COLORS["Deaths"],
        fontsize=7.9,
        va="center",
        ha="left",
        bbox={"facecolor": "white", "edgecolor": "none", "pad": 0.6, "alpha": 0.85},
    )
    add_panel_label(ax, "B")

    ax = axes[1, 0]
    points = global_context[global_context["measure"].isin(["DALY", "Deaths"]) & global_context["metric"].isin(["count"])].copy()
    daly_count = points[points["measure"] == "DALY"].sort_values("year_id")
    death_count = points[points["measure"] == "Deaths"].sort_values("year_id")
    width = 6.0
    ax.bar(daly_count["year_id"] - width / 2, daly_count["mean"] / 1_000_000, width=width, color="#9DAABD", label="DALYs, millions")
    ax2 = ax.twinx()
    ax2.bar(death_count["year_id"] + width / 2, death_count["mean"] / 1_000, width=width, color="#5A6C82", label="Deaths, thousands")
    ax.set_title("Absolute burden at available benchmark years", fontsize=10.2, loc="left", color=COLOR_TEXT)
    ax.set_ylabel("DALYs, millions")
    ax2.set_ylabel("Deaths, thousands", color=COLOR_TEXT)
    ax.set_xticks(sorted(global_context["year_id"].unique()))
    style_lancet_axes(ax)
    ax2.spines["top"].set_visible(False)
    ax2.spines["left"].set_visible(False)
    ax2.spines["right"].set_color(COLOR_GRID)
    ax2.tick_params(colors=COLOR_TEXT, labelsize=8.6)
    ax.text(1989.5, (daly_count["mean"] / 1_000_000).max() * 0.96, "DALYs, millions", fontsize=7.4, color="#4D5B6B")
    ax2.text(2020.2, (death_count["mean"] / 1_000).max() * 0.96, "Deaths, thousands", fontsize=7.4, color="#4D5B6B", ha="right")
    for row in daly_count.itertuples(index=False):
        ax.text(row.year_id - width / 2, row.mean / 1_000_000 + 0.9, f"{row.mean/1_000_000:.1f}", ha="center", va="bottom", fontsize=7.1, color=COLOR_TEXT)
    for row in death_count.itertuples(index=False):
        ax2.text(row.year_id + width / 2, row.mean / 1_000 + 5.0, f"{row.mean/1_000:.0f}", ha="center", va="bottom", fontsize=7.1, color=COLOR_TEXT)
    add_panel_label(ax, "C")

    ax = axes[1, 1]
    aging = global_context[["year_id", "age65_pct", "life_expectancy", "old_age_dependency"]].drop_duplicates("year_id").sort_values("year_id")
    base = aging.iloc[0]
    for col, label, color in [
        ("age65_pct", "Population aged 65 years and older", COLOR_NAVY),
        ("life_expectancy", "Life expectancy", COLOR_RUST),
        ("old_age_dependency", "Old-age dependency ratio", COLOR_TEAL),
    ]:
        index_values = aging[col] / base[col] * 100.0
        ax.plot(aging["year_id"], index_values, linewidth=2.2, marker="o", markersize=4.0, color=color)
        ax.text(2023.6, float(index_values.iloc[-1]), label, fontsize=7.7, color=color, va="center")
    ax.set_title("Demographic transition indicators", fontsize=10.2, loc="left", color=COLOR_TEXT)
    ax.set_ylabel("Index (1990=100)")
    ax.set_xticks(aging["year_id"])
    ax.set_xlim(1989.2, 2029.5)
    style_lancet_axes(ax)
    add_panel_label(ax, "D")
    fig.subplots_adjust(left=0.08, right=0.96, top=0.96, bottom=0.09, wspace=0.27, hspace=0.20)
    return save_figure(fig, "figure1_lancet_trends_and_ageing")


def _draw_global_map_panel(module50, ax, data: pd.DataFrame, value_col: str, cmap, panel_title: str):
    world = module50.load_world_shapes().copy()
    polygons = data[data["plot_method"] == "polygon"][["wb_iso3", value_col]].copy()
    points = data[data["plot_method"] == "point"].copy()
    merged = world.merge(polygons, left_on="iso_a3", right_on="wb_iso3", how="left")
    norm = mpl.colors.Normalize(vmin=float(data[value_col].quantile(0.02)), vmax=float(data[value_col].quantile(0.98)))
    merged.plot(ax=ax, color="#F6F7F9", edgecolor="#D7DDE5", linewidth=0.26)
    merged.dropna(subset=[value_col]).plot(ax=ax, column=value_col, cmap=cmap, norm=norm, edgecolor="white", linewidth=0.18)
    if not points.empty:
        ax.scatter(
            points["longitude"],
            points["latitude"],
            c=points[value_col],
            cmap=cmap,
            norm=norm,
            s=10 + points[value_col].rank(pct=True) * 16,
            edgecolors="white",
            linewidths=0.22,
            zorder=5,
        )
    ax.set_axis_off()
    ax.set_xlim(-180, 180)
    ax.set_ylim(-58, 85)
    ax.set_title(panel_title, fontsize=10.0, loc="left", color=COLOR_TEXT)
    top = data.sort_values(value_col, ascending=False).iloc[0]
    ax.text(
        0.01,
        0.04,
        f"Highest: {top['gbd_short_name']} ({top[value_col]:.2f})",
        transform=ax.transAxes,
        fontsize=6.9,
        color=COLOR_SLATE,
        ha="left",
        va="bottom",
        bbox={"facecolor": "white", "edgecolor": "none", "alpha": 0.82, "pad": 0.8},
    )
    return mpl.cm.ScalarMappable(norm=norm, cmap=cmap)


def make_figure2_global_maps(module50, global_map_df: pd.DataFrame) -> Tuple[Path, Path]:
    fig = plt.figure(figsize=(12.0, 7.0))
    gs = fig.add_gridspec(2, 2, hspace=0.17, wspace=0.08)
    specs = [
        ("deaths_asr_2023", "A  Mortality", mpl.colors.LinearSegmentedColormap.from_list("mort", ["#F7F4EF", "#D7B27E", COLOR_GOLD])),
        ("dalys_asr_2023", "B  DALYs", mpl.colors.LinearSegmentedColormap.from_list("dalys", ["#F4FBF8", "#9BC5B6", COLOR_TEAL])),
        ("prevalence_asr_2023", "C  Prevalence", mpl.colors.LinearSegmentedColormap.from_list("prev", ["#FFF7F2", "#E3B092", COLOR_RUST])),
        ("incidence_asr_2023", "D  Incidence", mpl.colors.LinearSegmentedColormap.from_list("inc", ["#F4F8FC", "#AFC5DA", COLOR_NAVY])),
    ]
    for idx, (value_col, title, cmap) in enumerate(specs):
        ax = fig.add_subplot(gs[idx // 2, idx % 2])
        sm = _draw_global_map_panel(module50, ax, global_map_df, value_col, cmap, title)
        cax = inset_axes(ax, width="38%", height="4.3%", loc="lower center", borderpad=1.1)
        cbar = fig.colorbar(sm, cax=cax, orientation="horizontal")
        cbar.ax.tick_params(labelsize=6.8, length=2.5, color=COLOR_TEXT)
        cbar.outline.set_edgecolor(COLOR_GRID)
        cbar.set_label("ASR per 100 000", fontsize=7.1, color=COLOR_TEXT, labelpad=2)
    fig.subplots_adjust(left=0.02, right=0.985, top=0.965, bottom=0.055)
    return save_figure(fig, "figure2_global_country_rate_maps_2023")


def make_figure3_subtype_lollipop(subtype_dirf: pd.DataFrame, subtype_mortality: pd.DataFrame) -> Tuple[Path, Path]:
    fig, axes = plt.subplots(2, 2, figsize=(11.0, 8.1))
    panels = [
        ("incidence", subtype_dirf, "Incidence", COLOR_NAVY),
        ("prevalence", subtype_dirf, "Prevalence", COLOR_RUST),
        ("DALY", subtype_dirf, "DALYs", COLOR_TEAL),
        ("Deaths", subtype_mortality, "Mortality", COLOR_GOLD),
    ]
    for idx, (ax, (measure, source_df, title, color)) in enumerate(zip(axes.flatten(), panels)):
        data = source_df[(source_df["measure"] == measure) & (source_df["year_id"] == 2023)].copy()
        data = data.sort_values("mean", ascending=False).head(10).reset_index(drop=True)
        lower_err = data["mean"] - data["lower"]
        upper_err = data["upper"] - data["mean"]
        ax.hlines(data["cause_name"], xmin=0, xmax=data["mean"], color=color, alpha=0.42, linewidth=1.35)
        ax.errorbar(
            data["mean"],
            data["cause_name"],
            xerr=[lower_err, upper_err],
            fmt="o",
            color=color,
            ecolor="#66707A",
            elinewidth=0.9,
            capsize=1.5,
            markersize=4.8,
        )
        ax.invert_yaxis()
        ax.set_title(title, fontsize=10.0, loc="left", color=COLOR_TEXT)
        ax.set_xlabel("ASR per 100 000")
        style_lancet_axes(ax)
        ax.grid(axis="x", color=COLOR_GRID, alpha=0.28, linewidth=0.55)
        ax.axvline(0, color=COLOR_GRID, linewidth=0.8)
        ax.set_xlim(0, float(data["upper"].max()) * 1.12)
        add_panel_label(ax, chr(ord("A") + idx))
        top_row = data.iloc[0]
        ax.text(
            0.98,
            0.07,
            f"Top subtype: {top_row['cause_name']}",
            transform=ax.transAxes,
            ha="right",
            va="bottom",
            fontsize=7.0,
            color=COLOR_SLATE,
        )
    fig.subplots_adjust(left=0.23, right=0.985, top=0.965, bottom=0.08, wspace=0.82, hspace=0.34)
    return save_figure(fig, "figure3_subtype_profile_lollipop_2023")


def make_figure4_ecology_lancet(
    physician_df: pd.DataFrame,
    risk_country_df: pd.DataFrame,
    comparator_table: pd.DataFrame,
    top20_profile: pd.DataFrame,
) -> Tuple[Path, Path]:
    fig = plt.figure(figsize=(12.2, 8.4))
    gs = fig.add_gridspec(2, 3, height_ratios=[1.0, 0.98], width_ratios=[1.0, 1.0, 1.08], hspace=0.34, wspace=0.28)
    top_labels = top20_profile.head(5)[["gbd_name", "gbd_short_name"]].drop_duplicates()
    label_offsets = [(3, 5), (3, -9), (4, 8), (4, -8), (3, 11)]

    scatter_specs = [
        {
            "ax": fig.add_subplot(gs[0, 0]),
            "data": physician_df[["physicians_per_1000", "asmr_2023", "gbd_name"]].dropna().copy(),
            "x_col": "physicians_per_1000",
            "y_col": "asmr_2023",
            "title": "Medical-resource context",
            "xlabel": "Physicians per 1,000 people",
            "ylabel": "Mortality ASR per 100 000",
            "color": COLOR_NAVY,
            "x_multiplier": 1.0,
            "label_top": True,
        },
        {
            "ax": fig.add_subplot(gs[0, 1]),
            "data": risk_country_df[["household_air", "asmr_2023", "gbd_name"]].dropna().copy(),
            "x_col": "household_air",
            "y_col": "asmr_2023",
            "title": "Household exposure context",
            "xlabel": "Household air pollution (%)",
            "ylabel": "Mortality ASR per 100 000",
            "color": COLOR_TEAL,
            "x_multiplier": 100.0,
            "label_top": False,
        },
        {
            "ax": fig.add_subplot(gs[0, 2]),
            "data": risk_country_df[["albuminuria", "asmr_2023", "gbd_name"]].dropna().copy(),
            "x_col": "albuminuria",
            "y_col": "asmr_2023",
            "title": "Metabolic and vascular context",
            "xlabel": "Albuminuria prevalence (%)",
            "ylabel": "Mortality ASR per 100 000",
            "color": COLOR_GOLD,
            "x_multiplier": 100.0,
            "label_top": False,
        },
        {
            "ax": fig.add_subplot(gs[1, 0]),
            "data": physician_df[["physicians_per_1000", "dalys_asr_2023", "gbd_name"]].dropna().copy(),
            "x_col": "physicians_per_1000",
            "y_col": "dalys_asr_2023",
            "title": "Service gradient for disability burden",
            "xlabel": "Physicians per 1,000 people",
            "ylabel": "DALY ASR per 100 000",
            "color": COLOR_STEEL,
            "x_multiplier": 1.0,
            "label_top": False,
        },
    ]

    for idx, spec in enumerate(scatter_specs):
        ax = spec["ax"]
        data = spec["data"].copy()
        plot_x = data[spec["x_col"]] * spec["x_multiplier"]
        plot_y = data[spec["y_col"]]
        rho, p_val = spearmanr(data[spec["x_col"]], plot_y)
        ax.scatter(
            plot_x,
            plot_y,
            s=19,
            color=mpl.colors.to_rgba(spec["color"], 0.32),
            edgecolor="none",
            zorder=2,
        )
        smooth = lowess(np.log10(plot_y), plot_x, frac=0.58, return_sorted=True)
        ax.plot(smooth[:, 0], 10 ** smooth[:, 1], color="#39434D", linestyle="--", linewidth=1.45, zorder=3)
        ax.set_yscale("log")
        ax.set_title(spec["title"], fontsize=10.0, loc="left", color=COLOR_TEXT)
        ax.set_xlabel(spec["xlabel"])
        ax.set_ylabel(spec["ylabel"])
        style_lancet_axes(ax)
        ax.grid(axis="x", color=COLOR_GRID, alpha=0.24, linewidth=0.55)
        ax.text(
            0.03,
            0.95,
            f"rho={rho:+.2f}\np={fmt_p(float(p_val))}",
            transform=ax.transAxes,
            fontsize=7.0,
            color=COLOR_SLATE,
            ha="left",
            va="top",
            bbox={"facecolor": "white", "edgecolor": COLOR_GRID, "alpha": 0.9, "pad": 2.2},
        )
        add_panel_label(ax, chr(ord("A") + idx))
        if spec["label_top"]:
            top_matches = data[data["gbd_name"].isin(top_labels["gbd_name"])].copy()
            ax.scatter(
                top_matches[spec["x_col"]] * spec["x_multiplier"],
                top_matches[spec["y_col"]],
                s=34,
                color=spec["color"],
                edgecolor="white",
                linewidth=0.45,
                zorder=4,
            )
            for offset_idx, row in enumerate(top_labels.itertuples(index=False)):
                match = data[data["gbd_name"] == row.gbd_name]
                if match.empty:
                    continue
                x = float(match.iloc[0][spec["x_col"]] * spec["x_multiplier"])
                y = float(match.iloc[0][spec["y_col"]])
                dx, dy = label_offsets[offset_idx]
                ax.annotate(
                    row.gbd_short_name,
                    (x, y),
                    xytext=(dx, dy),
                    textcoords="offset points",
                    fontsize=6.8,
                    color=COLOR_TEXT,
                    bbox={"facecolor": "white", "edgecolor": "none", "pad": 0.45, "alpha": 0.82},
                )

    ax_comp = fig.add_subplot(gs[1, 1:])
    comp = comparator_table.copy()
    comp.loc[:, "Cause_group_short"] = comp["Cause_group"].replace(
        {
            "Skin and subcutaneous diseases": "Skin diseases",
            "Cardiovascular diseases": "Cardiovascular",
            "Diabetes and kidney diseases": "Diabetes and kidney",
        }
    )
    y_positions = np.arange(len(comp))[::-1]
    death_ratio = comp["Death_rate_ratio_vs_skin"].to_numpy()
    daly_ratio = comp["DALY_rate_ratio_vs_skin"].to_numpy()
    ax_comp.axvline(1, color=COLOR_GRID, linewidth=0.9, linestyle="--", zorder=1)
    ax_comp.hlines(y_positions, np.minimum(death_ratio, daly_ratio), np.maximum(death_ratio, daly_ratio), color=COLOR_GRID, linewidth=1.2, zorder=1)
    ax_comp.scatter(death_ratio, y_positions, s=54, color=COLOR_GOLD, marker="o", label="Deaths ratio vs skin", zorder=3)
    ax_comp.scatter(daly_ratio, y_positions, s=50, color=COLOR_TEAL, marker="s", label="DALYs ratio vs skin", zorder=3)
    ax_comp.set_xscale("log")
    ax_comp.set_xlim(0.8, max(death_ratio.max(), daly_ratio.max()) * 1.45)
    ax_comp.set_yticks(y_positions)
    ax_comp.set_yticklabels(comp["Cause_group_short"])
    ax_comp.set_xlabel("Global 2023 age-standardized rate ratio versus skin disease")
    ax_comp.set_title("Cross-system context", fontsize=10.0, loc="left", color=COLOR_TEXT)
    style_lancet_axes(ax_comp)
    ax_comp.grid(axis="x", color=COLOR_GRID, alpha=0.24, linewidth=0.55)
    ax_comp.set_xticks([1, 3, 10, 30, 100])
    ax_comp.get_xaxis().set_major_formatter(mpl.ticker.ScalarFormatter())
    add_panel_label(ax_comp, "E")
    blended = mpl.transforms.blended_transform_factory(ax_comp.transAxes, ax_comp.transData)
    ax_comp.text(0.95, y_positions[0] + 0.62, "Actual ASR\nDeaths | DALYs", transform=blended, ha="right", va="bottom", fontsize=6.9, color=COLOR_SLATE)
    for y, row in zip(y_positions, comp.itertuples(index=False)):
        ax_comp.text(
            0.95,
            y,
            f"{row.Deaths_mean:.1f} | {row.DALYs_mean:.1f}",
            transform=blended,
            ha="right",
            va="center",
            fontsize=6.9,
            color=COLOR_TEXT,
        )
    ax_comp.text(
        0.03,
        0.95,
        "Skin disease is the reference\nfor both death and DALY ratios",
        transform=ax_comp.transAxes,
        ha="left",
        va="top",
        fontsize=7.0,
        color=COLOR_SLATE,
        bbox={"facecolor": "white", "edgecolor": COLOR_GRID, "alpha": 0.9, "pad": 2.4},
    )
    ax_comp.legend(loc="lower center", bbox_to_anchor=(0.53, -0.02), ncol=2, frameon=False, fontsize=7.1)
    fig.subplots_adjust(left=0.07, right=0.985, top=0.97, bottom=0.08)
    return save_figure(fig, "figure4_ageing_and_risk_ecology_2023")


def make_figure5_forecast_lancet(
    annual_df: pd.DataFrame,
    pandemic_df: pd.DataFrame,
    forecast_df: pd.DataFrame,
) -> Tuple[Path, Path]:
    fig = plt.figure(figsize=(11.4, 8.0))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.0, 0.92], hspace=0.32, wspace=0.24)
    focus_measures = ["Deaths", "DALYs"]

    for idx, measure_name in enumerate(focus_measures):
        ax = fig.add_subplot(gs[0, idx])
        observed = annual_df[annual_df["measure_name"] == measure_name].sort_values("year_id")
        counter = pandemic_df[pandemic_df["measure"] == measure_name].sort_values("year_id")
        forecast = forecast_df[forecast_df["measure"] == measure_name].sort_values("year_id")
        color = MEASURE_COLORS[measure_name]
        ax.plot(observed["year_id"], observed["mean"], color=color, linewidth=2.1, label="Observed")
        ax.plot(counter["year_id"], counter["counterfactual_rate"], color="#4A535D", linewidth=1.45, linestyle="--", label="Counterfactual")
        ax.plot(forecast["year_id"], forecast["forecast_rate"], color=color, linewidth=1.6, linestyle=":", label="Forecast")
        ax.axvspan(2020, 2023.5, color=COLOR_GREY, alpha=0.42)
        ax.axvspan(2023.5, 2050.5, color="#F8FAFC", alpha=1.0)
        ax.axvline(2023, color=COLOR_GRID, linestyle="-", linewidth=0.8)
        style_lancet_axes(ax)
        ax.set_title(f"{measure_name}: observed, counterfactual, and forecast", fontsize=10.0, loc="left", color=COLOR_TEXT)
        ax.set_ylabel("ASR per 100 000")
        ax.set_xlabel("Year")
        add_panel_label(ax, chr(ord("A") + idx))
        dev_2023 = float(counter.loc[counter["year_id"] == 2023, "deviation_pct"].iloc[0])
        f2050 = float(forecast.loc[forecast["year_id"] == 2050, "forecast_rate"].iloc[0])
        ax.text(
            0.03,
            0.95,
            f"2023 deviation {dev_2023:+.2f}%\n2050 rate {f2050:.2f}",
            transform=ax.transAxes,
            ha="left",
            va="top",
            fontsize=7.1,
            bbox={"facecolor": "white", "edgecolor": COLOR_GRID, "alpha": 0.9, "pad": 2.8},
        )
        ax.text(2021.7, ax.get_ylim()[1] * 0.985, "COVID-19 period", fontsize=6.7, color=COLOR_SLATE, ha="center", va="top")
        ax.text(2036.5, ax.get_ylim()[1] * 0.985, "Forecast horizon", fontsize=6.7, color=COLOR_SLATE, ha="center", va="top")

    ax_dev = fig.add_subplot(gs[1, 0])
    pandemic_years = [2020, 2021, 2022, 2023]
    for measure_name in MEASURE_ORDER:
        sub = pandemic_df[pandemic_df["measure"] == measure_name].sort_values("year_id")
        ax_dev.plot(
            sub["year_id"],
            sub["deviation_pct"],
            color=MEASURE_COLORS[measure_name],
            linewidth=1.8,
            marker="o",
            markersize=4.2,
        )
        ax_dev.text(
            float(sub["year_id"].iloc[-1]) + 0.12,
            float(sub["deviation_pct"].iloc[-1]),
            measure_name,
            fontsize=7.0,
            color=MEASURE_COLORS[measure_name],
            va="center",
        )
    ax_dev.axhline(0, color=COLOR_GRID, linewidth=0.9, linestyle="--")
    style_lancet_axes(ax_dev)
    ax_dev.grid(axis="x", color=COLOR_GRID, alpha=0.22, linewidth=0.55)
    ax_dev.set_xticks(pandemic_years)
    ax_dev.set_title("Pandemic-period deviation", fontsize=10.0, loc="left", color=COLOR_TEXT)
    ax_dev.set_xlabel("Year")
    ax_dev.set_ylabel("Deviation from counterfactual (%)")
    add_panel_label(ax_dev, "C")

    ax_idx = fig.add_subplot(gs[1, 1])
    for measure_name in MEASURE_ORDER:
        observed_2023 = float(annual_df.loc[(annual_df["measure_name"] == measure_name) & (annual_df["year_id"] == 2023), "mean"].iloc[0])
        sub_f = forecast_df[forecast_df["measure"] == measure_name].sort_values("year_id").copy()
        series = pd.concat(
            [
                pd.DataFrame({"year_id": [2023], "forecast_rate": [observed_2023]}),
                sub_f[["year_id", "forecast_rate"]],
            ],
            ignore_index=True,
        )
        series.loc[:, "index_2023"] = series["forecast_rate"] / observed_2023 * 100
        ax_idx.plot(series["year_id"], series["index_2023"], color=MEASURE_COLORS[measure_name], linewidth=1.9)
        ax_idx.text(
            float(series["year_id"].iloc[-1]) + 0.25,
            float(series["index_2023"].iloc[-1]),
            f"{measure_name} {series['index_2023'].iloc[-1]:.0f}",
            fontsize=6.9,
            color=MEASURE_COLORS[measure_name],
            va="center",
        )
    style_lancet_axes(ax_idx)
    ax_idx.grid(axis="x", color=COLOR_GRID, alpha=0.22, linewidth=0.55)
    ax_idx.set_title("Forecast growth indexed to 2023=100", fontsize=10.0, loc="left", color=COLOR_TEXT)
    ax_idx.set_xlabel("Year")
    ax_idx.set_ylabel("Index (2023=100)")
    add_panel_label(ax_idx, "D")
    ax_idx.set_xlim(2023, 2052)

    handles, labels = fig.axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="upper center", ncol=3, frameon=False, bbox_to_anchor=(0.5, 0.99), fontsize=8.0)
    fig.subplots_adjust(left=0.08, right=0.985, top=0.92, bottom=0.085)
    return save_figure(fig, "figure5_pandemic_and_forecast_2050")


def load_state_token(state_path: Path) -> str:
    payload = json.loads(state_path.read_text(encoding="utf-8"))
    for origin in payload.get("origins", []):
        for item in origin.get("localStorage", []):
            if "accesstoken" not in item.get("name", ""):
                continue
            parsed = json.loads(item["value"])
            secret = parsed.get("secret", "")
            if secret:
                return secret
    raise RuntimeError(f"Access token not found in {state_path}")


def fetch_official_global_annual_rates(access_token: str) -> pd.DataFrame:
    payload: List[Tuple[str, object]] = [("version", RESULTS_VERSION)]
    for measure_id in MEASURE_MAP:
        payload.append(("measure[]", measure_id))
    for year_id in range(1990, 2024):
        payload.append(("year[]", year_id))
    payload.extend(
        [
            ("location[]", GLOBAL_LOCATION_ID),
            ("sex[]", SEX_ID_BOTH),
            ("age[]", AGE_ID_AGE_STANDARDIZED),
            ("cause[]", CAUSE_ID),
            ("metric[]", 3),
            ("population_group", 1),
            ("api_version", "2023.0.0"),
            ("base", "single"),
            ("context", "cause"),
            ("singleOrMult", "single"),
            ("idsOrNames", "ids"),
            ("rows", 500000),
            ("start_year", 1980),
            ("fetch_all_years", "false"),
            ("language", "en"),
        ]
    )
    last_error: Optional[Exception] = None
    response = None
    for _ in range(3):
        try:
            response = requests.post(
                RESULTS_URL,
                data=payload,
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=120,
            )
            response.raise_for_status()
            break
        except Exception as exc:  # pragma: no cover - network retry
            last_error = exc
            response = None
    if response is None:
        raise RuntimeError(f"Failed to fetch annual global Results Tool series after retries: {last_error}")
    raw = response.json()
    df = pd.DataFrame(raw["data"], columns=raw["cols"]).copy()
    df["measure_name"] = df["measure"].map(MEASURE_MAP)
    df = df.rename(
        columns={
            "year": "year_id",
            "val": "mean",
        }
    )
    df = df[
        [
            "measure",
            "measure_name",
            "location",
            "sex",
            "age",
            "cause",
            "metric",
            "year_id",
            "mean",
            "lower",
            "upper",
        ]
    ].sort_values(["measure", "year_id"])
    df["source"] = "Official GBD Results Tool"
    df["query_date"] = "2026-03-26"
    df["version_id"] = RESULTS_VERSION
    return df.reset_index(drop=True)


def compute_pandemic_and_forecast(annual_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    pandemic_rows: List[Dict[str, object]] = []
    forecast_rows: List[Dict[str, object]] = []
    milestone_rows: List[Dict[str, object]] = []
    for measure_name in MEASURE_ORDER:
        series = annual_df[annual_df["measure_name"] == measure_name].sort_values("year_id").copy()
        train = series[(series["year_id"] >= 2010) & (series["year_id"] <= 2019)].copy()
        x = train["year_id"].to_numpy(dtype=float)
        y = np.log(train["mean"].to_numpy(dtype=float))
        slope, intercept = np.polyfit(x, y, 1)

        counterfactual_map: Dict[int, float] = {}
        for year_id in range(2020, 2024):
            predicted = float(np.exp(slope * year_id + intercept))
            observed = float(series.loc[series["year_id"] == year_id, "mean"].iloc[0])
            deviation_pct = pct_change(predicted, observed)
            counterfactual_map[year_id] = predicted
            pandemic_rows.append(
                {
                    "measure": measure_name,
                    "year_id": year_id,
                    "observed_rate": observed,
                    "counterfactual_rate": predicted,
                    "deviation_pct": deviation_pct,
                }
            )

        model = ExponentialSmoothing(
            np.log(series["mean"].to_numpy(dtype=float)),
            trend="add",
            damped_trend=True,
            seasonal=None,
            initialization_method="estimated",
        ).fit(optimized=True, use_brute=True)
        future_years = list(range(2024, 2051))
        forecast_values = np.exp(model.forecast(len(future_years)))
        for year_id, value in zip(future_years, forecast_values):
            forecast_rows.append(
                {
                    "measure": measure_name,
                    "year_id": year_id,
                    "forecast_rate": float(value),
                }
            )

        for milestone in [1990, 2010, 2019, 2020, 2021, 2022, 2023]:
            observed = float(series.loc[series["year_id"] == milestone, "mean"].iloc[0])
            milestone_rows.append(
                {
                    "measure": measure_name,
                    "year_id": milestone,
                    "observed_rate": observed,
                    "counterfactual_rate": counterfactual_map.get(milestone),
                    "forecast_rate": np.nan,
                }
            )
        for milestone in [2030, 2040, 2050]:
            value = float(forecast_values[future_years.index(milestone)])
            milestone_rows.append(
                {
                    "measure": measure_name,
                    "year_id": milestone,
                    "observed_rate": np.nan,
                    "counterfactual_rate": np.nan,
                    "forecast_rate": value,
                }
            )

    pandemic_df = pd.DataFrame(pandemic_rows)
    forecast_df = pd.DataFrame(forecast_rows)
    milestone_df = pd.DataFrame(milestone_rows).sort_values(["measure", "year_id"]).reset_index(drop=True)
    return pandemic_df, forecast_df, milestone_df


def load_risk_exposure_data(country_complete: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    risk_root = ROOT / "data" / "bronze" / "gbd" / "gbd2023" / "risk-exposure-2023"

    def select_single(
        path: Path,
        risk_key: str,
        risk_name: str,
        age_group_id: int,
        sex_id: int,
        age_label: str,
        source_metric: str,
    ) -> pd.DataFrame:
        df = pd.read_csv(path)
        sub = df[
            (df["year_id"] == 2023)
            & (df["age_group_id"] == age_group_id)
            & (df["sex_id"] == sex_id)
        ][["location_id", "location_name", "mean"]].copy()
        sub = sub.rename(columns={"mean": risk_key})
        sub["risk_key"] = risk_key
        sub["risk_name"] = risk_name
        sub["risk_age_label"] = age_label
        sub["risk_source_metric"] = source_metric
        return sub.reset_index(drop=True)

    def select_sex_averaged(
        path: Path,
        risk_key: str,
        risk_name: str,
        age_group_id: int,
        age_label: str,
        source_metric: str,
    ) -> pd.DataFrame:
        df = pd.read_csv(path)
        sub = df[
            (df["year_id"] == 2023)
            & (df["age_group_id"] == age_group_id)
            & (df["sex_id"].isin([1, 2]))
        ][["location_id", "location_name", "sex_id", "mean"]].copy()
        sub = (
            sub.groupby(["location_id", "location_name"], as_index=False)["mean"]
            .mean()
            .rename(columns={"mean": risk_key})
        )
        sub["risk_key"] = risk_key
        sub["risk_name"] = risk_name
        sub["risk_age_label"] = age_label
        sub["risk_source_metric"] = source_metric
        return sub.reset_index(drop=True)

    risk_frames = [
        select_single(
            risk_root
            / "IHME_GBD_2023_RISK_EXPOSURE_TOBACCO"
            / "IHME_GBD_2023_RISK_EXPOSURE_TOBACCO_SMOKING_Y2025M10D10.CSV",
            "smoking_prev",
            "Current smoking prevalence",
            27,
            3,
            "Age-standardized, both sexes",
            "Prevalence (proportion)",
        ),
        select_single(
            risk_root
            / "IHME_GBD_2023_RISK_EXPOSURE_AIR_POLLUTION"
            / "IHME_GBD_2023_RISK_EXPOSURE_AIR_POLLUTION_HOUSEHOLD_FROM_SOLID_FUELS_Y2025M10D10.CSV",
            "household_air",
            "Household air pollution from solid fuels",
            22,
            3,
            "All ages, both sexes",
            "Exposure proportion",
        ),
        select_sex_averaged(
            risk_root
            / "IHME_GBD_2023_RISK_EXPOSURE_HIGH_BMI"
            / "IHME_GBD_2023_RISK_EXPOSURE_HIGH_BMI_IN_ADULTS_Y2025M10D10.CSV",
            "high_bmi",
            "High body-mass index in adults",
            27,
            "Age-standardized, sex-averaged",
            "Continuous exposure",
        ),
        select_sex_averaged(
            risk_root
            / "IHME_GBD_2023_RISK_EXPOSURE_KIDNEY_DYSFUNCTION"
            / "IHME_GBD_2023_RISK_EXPOSURE_KIDNEY_DYSFUNCTION_ALBUMINURIA_Y2025M10D10.CSV",
            "albuminuria",
            "Kidney dysfunction (albuminuria prevalence)",
            27,
            "Age-standardized, sex-averaged",
            "Prevalence (proportion)",
        ),
    ]

    merged = country_complete.copy()
    metadata_rows: List[Dict[str, object]] = []
    for risk_df in risk_frames:
        risk_key = str(risk_df["risk_key"].iloc[0])
        risk_name = str(risk_df["risk_name"].iloc[0])
        merged = merged.merge(
            risk_df[["location_id", risk_key]],
            on="location_id",
            how="left",
            validate="one_to_one",
        )
        metadata_rows.append(
            {
                "risk_key": risk_key,
                "risk_name": risk_name,
                "age_definition": str(risk_df["risk_age_label"].iloc[0]),
                "source_metric": str(risk_df["risk_source_metric"].iloc[0]),
                "year_id": 2023,
                "n_non_missing": int(merged[risk_key].notna().sum()),
            }
        )

    top20_ids = merged.sort_values("asmr_2023", ascending=False).head(20)["location_id"].tolist()
    top20_set = set(top20_ids)

    summary_rows: List[Dict[str, object]] = []
    for indicator_key, indicator_name, domain, age_def, metric in [
        ("age65_pct", "Population aged 65 years and older", "Ageing", "2023 WDI country value", "Percent"),
        ("life_expectancy", "Life expectancy at birth", "Ageing", "2023 WDI country value", "Years"),
        ("old_age_dependency", "Old-age dependency ratio", "Ageing", "2023 WDI country value", "Ratio"),
        ("smoking_prev", "Current smoking prevalence", "Risk exposure", "Age-standardized, both sexes", "Proportion"),
        ("household_air", "Household air pollution from solid fuels", "Risk exposure", "All ages, both sexes", "Proportion"),
        ("high_bmi", "High body-mass index in adults", "Risk exposure", "Age-standardized, sex-averaged", "Continuous exposure"),
        ("albuminuria", "Kidney dysfunction (albuminuria prevalence)", "Risk exposure", "Age-standardized, sex-averaged", "Proportion"),
    ]:
        sub = merged[["asmr_2023", indicator_key]].dropna().copy()
        rho, p_value = spearmanr(sub[indicator_key], sub["asmr_2023"])
        top_values = merged.loc[merged["location_id"].isin(top20_set), indicator_key].dropna()
        other_values = merged.loc[~merged["location_id"].isin(top20_set), indicator_key].dropna()
        mw_p = mannwhitneyu(top_values, other_values, alternative="two-sided").pvalue
        summary_rows.append(
            {
                "domain": domain,
                "indicator": indicator_name,
                "definition": age_def,
                "metric_type": metric,
                "countries_with_data": int(len(sub)),
                "spearman_rho": float(rho),
                "p_value": float(p_value),
                "top20_median": float(top_values.median()),
                "others_median": float(other_values.median()),
                "top20_vs_others_p": float(mw_p),
            }
        )

    profile = merged.sort_values("asmr_2023", ascending=False).head(20).copy()
    for risk_key in ["smoking_prev", "household_air", "high_bmi", "albuminuria"]:
        median_value = float(merged[risk_key].median())
        profile[f"{risk_key}_global_median"] = median_value
        profile[f"{risk_key}_above_global_median"] = np.where(profile[risk_key] > median_value, "Above", "Below")

    metadata_df = pd.DataFrame(metadata_rows)
    summary_df = pd.DataFrame(summary_rows)
    return merged, summary_df, profile


def build_top5_risk_profile_text(profile: pd.DataFrame) -> str:
    top5 = profile.head(5).copy()
    fragments: List[str] = []
    for row in top5.itertuples(index=False):
        flags: List[str] = []
        if getattr(row, "high_bmi_above_global_median") == "Above":
            flags.append("adult BMI exposure")
        if getattr(row, "smoking_prev_above_global_median") == "Above":
            flags.append("smoking prevalence")
        if getattr(row, "household_air_above_global_median") == "Above":
            flags.append("household air pollution")
        if getattr(row, "albuminuria_above_global_median") == "Above":
            flags.append("albuminuria prevalence")
        fragments.append(f"{row.gbd_name}: {', '.join(flags) if flags else 'no selected exposure above the global median'}")
    return "; ".join(fragments)


def build_global_geo_summary(global_map_df: pd.DataFrame) -> Dict[str, str]:
    summary: Dict[str, str] = {}
    measure_to_col = {
        "deaths": "deaths_asr_2023",
        "dalys": "dalys_asr_2023",
        "prevalence": "prevalence_asr_2023",
        "incidence": "incidence_asr_2023",
    }
    for key, col in measure_to_col.items():
        subset = global_map_df.sort_values(col, ascending=False).head(3)
        summary[f"{key}_top3_text"] = "; ".join(
            f"{row.gbd_name} ({getattr(row, col):.2f})" for row in subset.itertuples(index=False)
        )
    return summary


def fetch_worldbank_indicator_panel(
    indicator_code: str,
    *,
    max_year: int = 2023,
    min_year: Optional[int] = None,
) -> pd.DataFrame:
    url = f"https://api.worldbank.org/v2/country/all/indicator/{indicator_code}?format=json&per_page=20000"
    payload = requests.get(url, timeout=120).json()
    rows = payload[1]
    frame = pd.DataFrame(
        [
            {
                "wb_iso3": item.get("countryiso3code"),
                "year_id": int(item.get("date")),
                "value": item.get("value"),
            }
            for item in rows
            if item.get("countryiso3code") and item.get("value") is not None
        ]
    )
    frame = frame[frame["wb_iso3"].ne("") & frame["wb_iso3"].ne("WLD") & (frame["year_id"] <= max_year)].copy()
    if min_year is not None:
        frame = frame[frame["year_id"] >= min_year].copy()
    return frame.reset_index(drop=True)


def build_physician_density_dataset(country_complete: pd.DataFrame, global_map_df: pd.DataFrame) -> pd.DataFrame:
    panel = fetch_worldbank_indicator_panel(PHYSICIAN_INDICATOR, max_year=2023, min_year=2015)
    latest = panel.sort_values(["wb_iso3", "year_id"]).groupby("wb_iso3", as_index=False).tail(1).copy()
    latest = latest.rename(columns={"value": "physicians_per_1000", "year_id": "physicians_year"})
    merged = (
        country_complete[["location_id", "gbd_name", "wb_iso3", "asmr_2023"]]
        .merge(global_map_df[["location_id", "dalys_asr_2023"]], on="location_id", how="left", validate="one_to_one")
        .merge(latest, on="wb_iso3", how="left", validate="many_to_one")
    )
    return merged.sort_values("gbd_name").reset_index(drop=True)


def build_physician_density_tables(physician_df: pd.DataFrame) -> Tuple[pd.DataFrame, pd.DataFrame, Dict[str, float]]:
    analytic = physician_df.dropna(subset=["physicians_per_1000", "asmr_2023", "dalys_asr_2023"]).copy()
    rho_mort, p_mort = spearmanr(analytic["physicians_per_1000"], analytic["asmr_2023"])
    rho_daly, p_daly = spearmanr(analytic["physicians_per_1000"], analytic["dalys_asr_2023"])
    analytic.loc[:, "physician_density_quartile"] = pd.qcut(
        analytic["physicians_per_1000"],
        4,
        labels=["Q1_lowest", "Q2", "Q3", "Q4_highest"],
    )
    quartile_df = (
        analytic.groupby("physician_density_quartile", observed=False)
        .agg(
            Countries=("location_id", "size"),
            Median_physicians_per_1000=("physicians_per_1000", "median"),
            Median_skin_mortality_ASR_2023=("asmr_2023", "median"),
            Median_skin_DALY_ASR_2023=("dalys_asr_2023", "median"),
        )
        .reset_index()
        .rename(columns={"physician_density_quartile": "Physician_density_quartile"})
    )
    q1_mort = analytic.loc[analytic["physician_density_quartile"] == "Q1_lowest", "asmr_2023"].dropna()
    q4_mort = analytic.loc[analytic["physician_density_quartile"] == "Q4_highest", "asmr_2023"].dropna()
    q1_daly = analytic.loc[analytic["physician_density_quartile"] == "Q1_lowest", "dalys_asr_2023"].dropna()
    q4_daly = analytic.loc[analytic["physician_density_quartile"] == "Q4_highest", "dalys_asr_2023"].dropna()
    summary_df = pd.DataFrame(
        [
            {
                "Indicator": "Physicians per 1,000 people",
                "Source": "World Bank WDI indicator SH.MED.PHYS.ZS",
                "Latest_value_window": "2015-2023",
                "Countries_with_data": int(len(analytic)),
                "Median_source_year": float(analytic["physicians_year"].median()),
                "Min_source_year": int(analytic["physicians_year"].min()),
                "Max_source_year": int(analytic["physicians_year"].max()),
                "Mortality_spearman_rho": float(rho_mort),
                "Mortality_p_value": float(p_mort),
                "DALY_spearman_rho": float(rho_daly),
                "DALY_p_value": float(p_daly),
                "Q1_vs_Q4_mortality_p": float(mannwhitneyu(q1_mort, q4_mort, alternative="two-sided").pvalue),
                "Q1_vs_Q4_DALY_p": float(mannwhitneyu(q1_daly, q4_daly, alternative="two-sided").pvalue),
            }
        ]
    )
    stats = {
        "n": int(len(analytic)),
        "median_year": float(analytic["physicians_year"].median()),
        "min_year": int(analytic["physicians_year"].min()),
        "max_year": int(analytic["physicians_year"].max()),
        "mort_rho": float(rho_mort),
        "mort_p": float(p_mort),
        "daly_rho": float(rho_daly),
        "daly_p": float(p_daly),
        "q1_physicians": float(quartile_df.loc[quartile_df["Physician_density_quartile"] == "Q1_lowest", "Median_physicians_per_1000"].iloc[0]),
        "q4_physicians": float(quartile_df.loc[quartile_df["Physician_density_quartile"] == "Q4_highest", "Median_physicians_per_1000"].iloc[0]),
        "q1_asmr": float(quartile_df.loc[quartile_df["Physician_density_quartile"] == "Q1_lowest", "Median_skin_mortality_ASR_2023"].iloc[0]),
        "q4_asmr": float(quartile_df.loc[quartile_df["Physician_density_quartile"] == "Q4_highest", "Median_skin_mortality_ASR_2023"].iloc[0]),
        "q1_daly": float(quartile_df.loc[quartile_df["Physician_density_quartile"] == "Q1_lowest", "Median_skin_DALY_ASR_2023"].iloc[0]),
        "q4_daly": float(quartile_df.loc[quartile_df["Physician_density_quartile"] == "Q4_highest", "Median_skin_DALY_ASR_2023"].iloc[0]),
    }
    return summary_df, quartile_df, stats


def fetch_global_cause_comparison_2023(access_token: str) -> pd.DataFrame:
    payload: List[Tuple[str, object]] = [("version", RESULTS_VERSION)]
    for measure_id in [1, 2]:
        payload.append(("measure[]", measure_id))
    for cause_id in COMPARATOR_CAUSES:
        payload.append(("cause[]", cause_id))
    payload.extend(
        [
            ("location[]", GLOBAL_LOCATION_ID),
            ("sex[]", SEX_ID_BOTH),
            ("age[]", AGE_ID_AGE_STANDARDIZED),
            ("metric[]", 3),
            ("year[]", 2023),
            ("population_group", 1),
            ("api_version", "2023.0.0"),
            ("base", "single"),
            ("context", "cause"),
            ("singleOrMult", "single"),
            ("idsOrNames", "ids"),
            ("rows", 500000),
            ("start_year", 1980),
            ("fetch_all_years", "false"),
            ("language", "en"),
        ]
    )
    last_error: Optional[Exception] = None
    response = None
    for _ in range(3):
        try:
            response = requests.post(
                RESULTS_URL,
                data=payload,
                headers={"Authorization": f"Bearer {access_token}"},
                timeout=120,
            )
            response.raise_for_status()
            break
        except Exception as exc:  # pragma: no cover - network retry
            last_error = exc
            response = None
    if response is None:
        raise RuntimeError(f"Failed to fetch cross-cause comparison after retries: {last_error}")
    raw = response.json()
    df = pd.DataFrame(raw["data"], columns=raw["cols"]).copy()
    df.loc[:, "cause_name"] = df["cause"].map(COMPARATOR_CAUSES)
    df.loc[:, "measure_name"] = df["measure"].map(MEASURE_MAP)
    df = df.rename(columns={"val": "mean"})
    return df[["cause", "cause_name", "measure", "measure_name", "mean", "lower", "upper"]].sort_values(
        ["measure", "mean"], ascending=[True, False]
    )


def build_cross_cause_comparison_table(comparator_df: pd.DataFrame) -> pd.DataFrame:
    wide = (
        comparator_df.pivot_table(index=["cause", "cause_name"], columns="measure_name", values=["mean", "lower", "upper"], aggfunc="first")
        .reset_index()
    )
    wide.columns = [
        "Cause_id" if col == ("cause", "") else "Cause_group" if col == ("cause_name", "") else f"{col[1]}_{col[0]}"
        for col in wide.columns
    ]
    skin_daly = float(wide.loc[wide["Cause_group"] == "Skin and subcutaneous diseases", "DALYs_mean"].iloc[0])
    skin_death = float(wide.loc[wide["Cause_group"] == "Skin and subcutaneous diseases", "Deaths_mean"].iloc[0])
    wide.loc[:, "DALY_rate_ratio_vs_skin"] = wide["DALYs_mean"] / skin_daly
    wide.loc[:, "Death_rate_ratio_vs_skin"] = wide["Deaths_mean"] / skin_death
    order = [COMPARATOR_CAUSES[cause_id] for cause_id in [653, 491, 410, 974]]
    wide.loc[:, "Cause_group"] = pd.Categorical(wide["Cause_group"], categories=order, ordered=True)
    wide = wide.sort_values("Cause_group").reset_index(drop=True)
    return wide


def build_cross_cause_summary(comparator_table: pd.DataFrame) -> Dict[str, float]:
    summary: Dict[str, float] = {}
    for row in comparator_table.itertuples(index=False):
        key = re.sub(r"[^a-z0-9]+", "_", str(row.Cause_group).lower()).strip("_")
        summary[f"{key}_death_rate"] = float(row.Deaths_mean)
        summary[f"{key}_daly_rate"] = float(row.DALYs_mean)
    return summary


def build_pandemic_summary_table(pandemic_df: pd.DataFrame, forecast_df: pd.DataFrame) -> pd.DataFrame:
    rows: List[Dict[str, object]] = []
    for measure_name in MEASURE_ORDER:
        sub_p = pandemic_df[pandemic_df["measure"] == measure_name].copy()
        sub_f = forecast_df[forecast_df["measure"] == measure_name].copy()
        row = {
            "Measure": measure_name,
            "Observed 2020 ASR": float(sub_p.loc[sub_p["year_id"] == 2020, "observed_rate"].iloc[0]),
            "Observed 2023 ASR": float(sub_p.loc[sub_p["year_id"] == 2023, "observed_rate"].iloc[0]),
            "Counterfactual 2023 ASR": float(sub_p.loc[sub_p["year_id"] == 2023, "counterfactual_rate"].iloc[0]),
            "Deviation in 2023 %": float(sub_p.loc[sub_p["year_id"] == 2023, "deviation_pct"].iloc[0]),
            "Projected 2030 ASR": float(sub_f.loc[sub_f["year_id"] == 2030, "forecast_rate"].iloc[0]),
            "Projected 2040 ASR": float(sub_f.loc[sub_f["year_id"] == 2040, "forecast_rate"].iloc[0]),
            "Projected 2050 ASR": float(sub_f.loc[sub_f["year_id"] == 2050, "forecast_rate"].iloc[0]),
        }
        rows.append(row)
    return pd.DataFrame(rows)


def build_table1() -> pd.DataFrame:
    return pd.DataFrame(
        [
            {
                "component": "Global descriptive core",
                "source": "GBD 2023 local reproducible DIRF and mortality extracts",
                "scope": "Global",
                "years": "1990, 2010, 2020, 2023",
                "measures": "Incidence, prevalence, DALYs, deaths",
                "metrics": "Age-standardized rates and counts",
                "notes": "Used for burden overview, subtype summaries, and absolute count interpretation",
            },
            {
                "component": "Official annual global time series",
                "source": "Authenticated GBD Results Tool annual export",
                "scope": "Global",
                "years": "1990-2023",
                "measures": "Incidence, prevalence, DALYs, deaths",
                "metrics": "Age-standardized rate",
                "notes": "Cause id 653, both sexes, age-standardized rate, version 8352",
            },
            {
                "component": "Country-level ageing ecology",
                "source": "GBD 2023 mortality extract + World Bank WDI",
                "scope": "198 countries and territories after ambiguity exclusion",
                "years": "2023",
                "measures": "Deaths",
                "metrics": "Age-standardized mortality rate",
                "notes": "WDI indicators were population aged 65+, life expectancy, and old-age dependency ratio",
            },
            {
                "component": "Country-level risk ecology",
                "source": "GBD 2023 mortality extract + GBD 2023 risk exposure files",
                "scope": "198 countries and territories after ambiguity exclusion",
                "years": "2023",
                "measures": "Deaths",
                "metrics": "Age-standardized mortality rate vs selected exposure indicators",
                "notes": "Selected ecological exposures were smoking, household air pollution, adult BMI, and albuminuria",
            },
            {
                "component": "Pandemic-period deviation",
                "source": "Authenticated GBD Results Tool annual export",
                "scope": "Global",
                "years": "2020-2023",
                "measures": "Incidence, prevalence, DALYs, deaths",
                "metrics": "Observed rate vs 2010-2019 log-linear counterfactual",
                "notes": "Interpreted as descriptive pandemic-period deviation rather than causal effect estimation",
            },
            {
                "component": "Exploratory forecast to 2050",
                "source": "Authenticated GBD Results Tool annual export",
                "scope": "Global",
                "years": "2024-2050",
                "measures": "Incidence, prevalence, DALYs, deaths",
                "metrics": "Projected age-standardized rate",
                "notes": "Damped-trend exponential smoothing on log-transformed annual series; exploratory only",
            },
            {
                "component": "Asia-Pacific geographic extension",
                "source": "Authenticated GBD Results Tool + World Bank WDI",
                "scope": "39 Asia-Pacific locations",
                "years": "2023",
                "measures": "Deaths, prevalence, incidence",
                "metrics": "Age-standardized rates",
                "notes": "World Bank East Asia & Pacific plus South Asia operational definition",
            },
        ]
    )


def make_figure4_annual_covid_forecast(
    annual_df: pd.DataFrame,
    pandemic_df: pd.DataFrame,
    forecast_df: pd.DataFrame,
) -> Tuple[Path, Path]:
    plt.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "axes.spines.top": False,
            "axes.spines.right": False,
            "axes.edgecolor": COLOR_GRID,
            "axes.labelcolor": COLOR_TEXT,
            "xtick.color": COLOR_TEXT,
            "ytick.color": COLOR_TEXT,
            "figure.dpi": 150,
        }
    )
    fig, axes = plt.subplots(2, 2, figsize=(12.2, 8.8))
    axes = axes.flatten()
    for ax, measure_name, label in zip(axes, MEASURE_ORDER, ["A", "B", "C", "D"]):
        observed = annual_df[annual_df["measure_name"] == measure_name].sort_values("year_id")
        counter = pandemic_df[pandemic_df["measure"] == measure_name].sort_values("year_id")
        forecast = forecast_df[forecast_df["measure"] == measure_name].sort_values("year_id")
        color = MEASURE_COLORS[measure_name]

        ax.axvspan(2020, 2023.5, color=COLOR_GREY, alpha=0.55, zorder=0)
        ax.plot(
            observed["year_id"],
            observed["mean"],
            color=color,
            linewidth=2.2,
            label="Observed",
        )
        ax.plot(
            counter["year_id"],
            counter["counterfactual_rate"],
            color="#333333",
            linewidth=1.8,
            linestyle="--",
            label="2010-19 counterfactual",
        )
        ax.plot(
            forecast["year_id"],
            forecast["forecast_rate"],
            color=color,
            linewidth=1.8,
            linestyle=":",
            label="Exploratory forecast",
        )
        ax.scatter(counter["year_id"], counter["observed_rate"], color=color, s=28, zorder=5)
        ax.set_title(measure_name)
        ax.set_xlabel("Year")
        ax.set_ylabel("Age-standardized rate per 100,000")
        dev_2023 = float(counter.loc[counter["year_id"] == 2023, "deviation_pct"].iloc[0])
        f2050 = float(forecast.loc[forecast["year_id"] == 2050, "forecast_rate"].iloc[0])
        ax.text(
            0.02,
            0.96,
            f"2023 deviation: {dev_2023:+.2f}%\n2050 forecast: {f2050:.2f}",
            transform=ax.transAxes,
            va="top",
            ha="left",
            fontsize=8.5,
            color=COLOR_TEXT,
            bbox={"facecolor": "white", "alpha": 0.85, "edgecolor": COLOR_GRID, "pad": 4},
        )
        ax.text(
            -0.12,
            1.05,
            label,
            transform=ax.transAxes,
            fontsize=12,
            fontweight="bold",
            va="bottom",
            ha="left",
            color=COLOR_NAVY,
        )
        ax.grid(alpha=0.2)

    handles, labels = axes[0].get_legend_handles_labels()
    fig.legend(handles, labels, loc="lower center", ncol=3, frameon=False, bbox_to_anchor=(0.5, -0.02))
    fig.suptitle(
        "Figure 4. Official annual global rates, pandemic-period deviation, and exploratory forecasts to 2050",
        y=1.02,
        fontsize=14,
    )
    return save_figure(fig, "figure4_annual_covid_forecast_2050")


def make_figureS9_risk_heatmap(top20_profile: pd.DataFrame) -> Tuple[Path, Path]:
    display = top20_profile.copy()
    display = display.sort_values("asmr_2023", ascending=False).reset_index(drop=True)
    columns = ["smoking_prev", "household_air", "high_bmi", "albuminuria"]
    labels = [
        "Smoking",
        "Household air",
        "Adult BMI",
        "Albuminuria",
    ]
    percentile_matrix = np.column_stack(
        [
            display[col].rank(pct=True).to_numpy() * 100
            for col in columns
        ]
    )

    fig, ax = plt.subplots(figsize=(8.6, 7.6))
    im = ax.imshow(percentile_matrix, aspect="auto", cmap=HEATMAP_CMAP, vmin=0, vmax=100)
    ax.set_xticks(np.arange(len(labels)))
    ax.set_xticklabels(labels, rotation=0, fontsize=9)
    ax.set_yticks(np.arange(len(display)))
    ax.set_yticklabels(
        [f"{row.gbd_short_name} ({row.asmr_2023:.2f})" for row in display.itertuples(index=False)],
        fontsize=8.5,
    )
    for row_idx in range(percentile_matrix.shape[0]):
        for col_idx in range(percentile_matrix.shape[1]):
            value = percentile_matrix[row_idx, col_idx]
            ax.text(
                col_idx,
                row_idx,
                f"{value:.0f}",
                ha="center",
                va="center",
                fontsize=7.5,
                color="white" if value >= 62 else COLOR_TEXT,
            )
    ax.set_title("Figure S9. Percentile ranks of selected ecological risk indicators in the top 20 mortality settings")
    ax.set_xlabel("Country-level risk indicator percentile (0-100)")
    cbar = fig.colorbar(im, ax=ax, fraction=0.032, pad=0.02)
    cbar.set_label("Percentile rank", rotation=270, labelpad=12)
    return save_figure(fig, "figureS9_top20_risk_heatmap")


def build_summary(
    values: Dict[str, float],
    risk_summary: pd.DataFrame,
    pandemic_df: pd.DataFrame,
    forecast_df: pd.DataFrame,
    top20_profile: pd.DataFrame,
    physician_stats: Dict[str, float],
) -> Dict[str, str]:
    rho_household = float(risk_summary.loc[risk_summary["indicator"] == "Household air pollution from solid fuels", "spearman_rho"].iloc[0])
    rho_albuminuria = float(risk_summary.loc[risk_summary["indicator"] == "Kidney dysfunction (albuminuria prevalence)", "spearman_rho"].iloc[0])
    bmi_top20_median = float(risk_summary.loc[risk_summary["indicator"] == "High body-mass index in adults", "top20_median"].iloc[0])
    bmi_others_median = float(risk_summary.loc[risk_summary["indicator"] == "High body-mass index in adults", "others_median"].iloc[0])
    death_dev_2023 = float(pandemic_df.loc[(pandemic_df["measure"] == "Deaths") & (pandemic_df["year_id"] == 2023), "deviation_pct"].iloc[0])
    prev_dev_2023 = float(pandemic_df.loc[(pandemic_df["measure"] == "Prevalence") & (pandemic_df["year_id"] == 2023), "deviation_pct"].iloc[0])
    inc_dev_2023 = float(pandemic_df.loc[(pandemic_df["measure"] == "Incidence") & (pandemic_df["year_id"] == 2023), "deviation_pct"].iloc[0])
    death_2050 = float(forecast_df.loc[(forecast_df["measure"] == "Deaths") & (forecast_df["year_id"] == 2050), "forecast_rate"].iloc[0])
    daly_2050 = float(forecast_df.loc[(forecast_df["measure"] == "DALYs") & (forecast_df["year_id"] == 2050), "forecast_rate"].iloc[0])
    top3 = top20_profile.head(3)
    top3_text = "; ".join(f"{row.gbd_name} ({row.asmr_2023:.2f})" for row in top3.itertuples(index=False))

    background = (
        "Skin and subcutaneous diseases generate chronic morbidity, yet their burden is seldom interpreted alongside population ageing, pandemic disruption, and service planning. "
        "We quantified burden, ecological correlates, pandemic-period deviation, and exploratory forecasts."
    )
    methods = (
        "We combined GBD 2023 extracts, authenticated Results Tool annual series, and World Bank ageing and physician-density indicators. "
        "Global burden was summarised for 1990 and 2023. Country-level ecological analyses used 2023 mortality rates from 198 countries and territories, selected GBD risk exposures, and latest available physician-density data from 2015-2023. "
        "Pandemic-period deviation compared observed 2020-2023 age-standardized rates with 2010-2019 log-linear counterfactuals. "
        "Exploratory projections to 2050 used damped-trend exponential smoothing."
    )
    findings = (
        f"Between 1990 and 2023, the global share of people aged 65 years and older rose from {values['age_1990']:.2f}% to {values['age_2023']:.2f}%, the age-standardized mortality rate increased from {values['death_rate_1990']:.2f} to {values['death_rate_2023']:.2f} per 100,000, and deaths rose from {int(values['death_count_1990']):,} to {int(values['death_count_2023']):,}. "
        f"Higher population ageing remained negatively associated with mortality; physician density was inversely associated with mortality (rho={physician_stats['mort_rho']:.3f}) and DALYs (rho={physician_stats['daly_rho']:.3f}); household air pollution was positively associated with mortality (rho={rho_household:.3f}); and top-20 high-mortality settings had higher adult BMI exposure than other countries (median {bmi_top20_median:.2f} vs {bmi_others_median:.2f}). "
        f"Relative to the pre-pandemic counterfactual, global mortality was {death_dev_2023:.2f}% higher in 2023, whereas prevalence and incidence remained below expected ({prev_dev_2023:.2f}% and {inc_dev_2023:.2f}%). "
        f"Exploratory projections suggested further increases by 2050, reaching a death rate of {death_2050:.2f} and a DALY rate of {daly_2050:.1f} per 100,000. Highest 2023 mortality was observed in {top3_text}."
    )
    interpretation = (
        "Global skin burden rose alongside population ageing, but high standardized mortality clustered more strongly with ecological markers of structural vulnerability, lower medical-resource availability, and metabolic risk than with ageing alone. "
        "Pandemic-period mortality drift and continued increases projected to 2050 reinforce the need to integrate skin care, chronic wound prevention, and infection control into healthy-ageing policy."
    )
    return {
        "Background": background,
        "Methods": methods,
        "Findings": findings,
        "Interpretation": interpretation,
        "Funding": "None.",
    }


def build_research_in_context() -> Dict[str, str]:
    return {
        "Evidence before this study": (
            "We had already assembled a reproducible GBD 2023 and World Bank World Development Indicators manuscript package showing that global skin burden rose alongside population ageing, while the highest standardized mortality did not cluster in the oldest countries. "
            "However, the earlier package did not yet integrate official annual authenticated GBD Results Tool series, selected ecological risk exposure indicators, physician-density context, a pandemic-period deviation module, or an exploratory forecast to 2050."
        ),
        "Added value of this study": (
            "The updated package links six analytic layers in one reproducible framework: the original global burden and subtype profile, a World Bank ageing ecology, a selected GBD risk exposure ecology, a physician-density ecology based on World Bank workforce data, a cross-cause comparison with major non-communicable disease groups, and an official annual global Results Tool extension used for pandemic-period deviation and exploratory 2050 forecasts. "
            "This makes the paper more policy-relevant by showing not only how burden changed, but also which ecological signals track high mortality, how skin burden sits relative to better-prioritized disease systems, and how the burden may evolve if recent trends persist."
        ),
        "Implications of all the available evidence": (
            "The evidence now points to a broader interpretation of skin burden in ageing societies. "
            "Health systems need to plan for growing chronic dermatologic demand, but they also need to identify high-mortality settings where metabolic vulnerability, lower physician availability, household exposure, wound care failures, and disrupted continuity of care may amplify fatal outcomes. "
            "Pandemic-era deviation and 2050 projections suggest that these issues remain active service-planning problems rather than retrospective epidemiologic observations."
        ),
    }


def modify_sections(
    builder45,
    values: Dict[str, float],
    risk_summary: pd.DataFrame,
    pandemic_df: pd.DataFrame,
    forecast_df: pd.DataFrame,
    top20_profile: pd.DataFrame,
    global_geo_summary: Dict[str, str],
    physician_stats: Dict[str, float],
    comparison_summary: Dict[str, float],
) -> List[Tuple[str, List[Tuple[Optional[str], List[str]]]]]:
    sections = builder45.build_main_sections(values)
    new_sections: List[Tuple[str, List[Tuple[Optional[str], List[str]]]]] = []

    top5_risk_text = build_top5_risk_profile_text(top20_profile)
    rho_household = float(risk_summary.loc[risk_summary["indicator"] == "Household air pollution from solid fuels", "spearman_rho"].iloc[0])
    p_household = float(risk_summary.loc[risk_summary["indicator"] == "Household air pollution from solid fuels", "p_value"].iloc[0])
    rho_albuminuria = float(risk_summary.loc[risk_summary["indicator"] == "Kidney dysfunction (albuminuria prevalence)", "spearman_rho"].iloc[0])
    p_albuminuria = float(risk_summary.loc[risk_summary["indicator"] == "Kidney dysfunction (albuminuria prevalence)", "p_value"].iloc[0])
    bmi_top20_median = float(risk_summary.loc[risk_summary["indicator"] == "High body-mass index in adults", "top20_median"].iloc[0])
    bmi_others_median = float(risk_summary.loc[risk_summary["indicator"] == "High body-mass index in adults", "others_median"].iloc[0])
    bmi_top20_p = float(risk_summary.loc[risk_summary["indicator"] == "High body-mass index in adults", "top20_vs_others_p"].iloc[0])

    death_dev = {
        int(row.year_id): float(row.deviation_pct)
        for row in pandemic_df[pandemic_df["measure"] == "Deaths"].itertuples(index=False)
    }
    daly_dev = {
        int(row.year_id): float(row.deviation_pct)
        for row in pandemic_df[pandemic_df["measure"] == "DALYs"].itertuples(index=False)
    }
    prev_dev = {
        int(row.year_id): float(row.deviation_pct)
        for row in pandemic_df[pandemic_df["measure"] == "Prevalence"].itertuples(index=False)
    }
    inc_dev = {
        int(row.year_id): float(row.deviation_pct)
        for row in pandemic_df[pandemic_df["measure"] == "Incidence"].itertuples(index=False)
    }

    milestone = {}
    for measure_name in MEASURE_ORDER:
        milestone[measure_name] = {}
        sub = forecast_df[forecast_df["measure"] == measure_name]
        for year_id in [2030, 2040, 2050]:
            milestone[measure_name][year_id] = float(sub.loc[sub["year_id"] == year_id, "forecast_rate"].iloc[0])

    for section_title, blocks in sections:
        if section_title == "Introduction":
            intro_blocks: List[Tuple[Optional[str], List[str]]] = []
            for subsection, paragraphs in blocks:
                mutable = list(paragraphs)
                if mutable:
                    mutable[-1] = (
                        "We therefore developed a Lancet-style long working draft that combines official GBD 2023 burden estimates with World Bank ageing indicators drawn from the World Development Indicators platform. "
                        "Our aims were fivefold: first, to describe how global skin burden changed between 1990 and 2023; second, to characterize the dominant disease subtypes across incidence, prevalence, DALYs, and mortality; third, to examine whether country-level population ageing was associated with standardized skin mortality in 2023; fourth, to assess whether selected ecological risk exposure indicators tracked high mortality; and fifth, to quantify pandemic-period deviation and develop exploratory projections to 2050. "
                        "Because the currently reproducible local data support global burden trajectories, official annual global rate series, and country-level mortality ecology rather than full country-age-sex incidence panels, we restricted interpretation to questions that can be defended with the available data."
                    )
                intro_blocks.append((subsection, mutable))
            new_sections.append((section_title, intro_blocks))
            continue

        if section_title == "Methods":
            method_blocks: List[Tuple[Optional[str], List[str]]] = []
            for subsection, paragraphs in blocks:
                mutable = list(paragraphs)
                if subsection == "Study design and data sources":
                    mutable.append(
                        "To extend the global time series beyond the locked four-point local extract, we additionally queried the authenticated GBD Results Tool on March 12, 2026 for annual global age-standardized rates between 1990 and 2023. "
                        "The query used version 8352, cause id 653 (skin and subcutaneous diseases), both sexes, age-standardized age group, global location, and rate metrics for deaths, DALYs, prevalence, and incidence. "
                        "These annual official series were used for pandemic-period deviation analyses and exploratory forecasting, whereas the 1990 versus 2023 global burden table retained the locked local extract used in the original submission package."
                    )
                    mutable.append(
                        "We additionally queried authenticated 2023 country-level age-standardized rates for deaths, DALYs, prevalence, and incidence across the full analytic country set so that the main figure sequence could include a Lancet-style global distribution atlas. "
                        "These country-level rates were merged to harmonized ISO3 identifiers and rendered with a common global basemap."
                    )
                elif subsection == "Country-level ageing ecology":
                    subsection = "Country-level ageing and risk ecology"
                    mutable.append(
                        "Because public risk-attributable burden outputs for skin and subcutaneous diseases overall were not consistently retrievable as a single harmonized country panel, we added a complementary ecological risk module based on country-level GBD 2023 exposure indicators rather than attributable fractions. "
                        "The selected indicators were current smoking prevalence (age-standardized, both sexes), household air pollution from solid fuels (all ages, both sexes), high body-mass index in adults (age-standardized values averaged across males and females), and kidney dysfunction measured as albuminuria prevalence (age-standardized values averaged across males and females). "
                        "These indicators were chosen because they were locally available in stable country-level files and were clinically plausible markers of chronic wound vulnerability, infection risk, or structural disadvantage."
                    )
                    mutable.append(
                        "We interpreted these risk indicators as ecological correlates rather than causal exposures. "
                        "The country-level analysis therefore asked whether settings with higher exposure burdens also had higher standardized skin mortality, and whether the highest-mortality countries showed distinctive exposure profiles relative to the rest of the distribution."
                    )
                    mutable.append(
                        "To approximate medical-resource context, we additionally obtained physician density from World Bank World Development Indicators indicator SH.MED.PHYS.ZS (physicians per 1,000 people). "
                        "Because complete 2023 values were not available for all locations, we used the latest available country observation between 2015 and 2023 and merged it to the 2023 skin mortality and DALY panels. "
                        "We did not include dermatologist-specific density because we did not identify a harmonized official global country panel with comparable temporal coverage."
                    )
                elif subsection == "Statistical analysis":
                    mutable.append(
                        "For the risk ecology, we calculated Spearman correlations between each selected exposure indicator and the country-level age-standardized mortality rate in 2023. "
                        "We also compared the 20 highest-mortality countries and territories with the remaining analytic set using the Mann-Whitney U test and summarized whether the top-ranked countries were above or below the global median for each selected exposure. "
                        "These analyses were descriptive and were not intended to estimate attributable fractions or to provide multivariable causal adjustment."
                    )
                    mutable.append(
                        "To approximate the effect of the COVID-19 period on global skin burden, we fit log-linear counterfactual trends to the official annual global age-standardized series during 2010-2019 and extrapolated them to 2020-2023. "
                        "Observed rates were then compared with these counterfactual projections and expressed as percentage deviation. "
                        "We interpret these estimates as pandemic-period interruption metrics rather than direct causal effects of SARS-CoV-2 infection."
                    )
                    mutable.append(
                        "For exploratory forecasts to 2050, we fitted damped-trend exponential smoothing models to the log-transformed official annual global series through 2023 and projected age-standardized rates for deaths, DALYs, prevalence, and incidence to 2050. "
                        "These projections were labelled exploratory because they do not represent official GBD foresight scenarios and do not incorporate explicit demographic, clinical, or policy covariates."
                    )
                    mutable.append(
                        "For the physician-density analysis, we calculated Spearman correlations between physicians per 1,000 people and country-level skin mortality and DALY rates in 2023. "
                        "We also summarized mortality and DALY distributions across physician-density quartiles and compared the lowest and highest quartiles using the Mann-Whitney U test. "
                        "Finally, to contextualize the scale of skin burden relative to other better-prioritized disease systems, we extracted authenticated 2023 global age-standardized mortality and DALY rates for cardiovascular diseases, neoplasms, diabetes and kidney diseases, and skin and subcutaneous diseases."
                    )
                method_blocks.append((subsection, mutable))
            method_blocks.insert(
                1,
                (
                    "Global 2023 geographic mapping",
                    [
                        "To align the main display set with the figure grammar commonly used in high-impact burden papers, we created a four-panel global distribution figure for 2023 country-level age-standardized mortality, DALY, prevalence, and incidence rates. "
                        "Settings not represented as polygons in the low-resolution Natural Earth basemap were retained as point markers using World Bank country coordinates. "
                        "These maps were descriptive and were used to visualize global concentration rather than to support regional causal inference."
                    ],
                ),
            )
            new_sections.append((section_title, method_blocks))
            continue

        if section_title == "Results":
            result_blocks: List[Tuple[Optional[str], List[str]]] = []
            for subsection, paragraphs in blocks:
                mutable = list(paragraphs)
                if subsection == "Subtype profile":
                    mutable[0] = (
                        "Global subtype patterns were highly uneven in 2023 (figure 3). "
                        "The main subtype profile is summarised in table 3, and the full 1990-2023 subtype change profile is provided in table S2 and figure S8. "
                        f"For incidence, {values['top_incidence_subtype']} was the largest category, with an age-standardized rate of {values['top_incidence_subtype_rate']:.1f} per 100,000. "
                        f"For prevalence, {values['top_prevalence_subtype']} dominated the global profile at {values['top_prevalence_subtype_rate']:.1f} per 100,000, and the same category also dominated the age-standardized DALY rate at {values['top_daly_subtype_rate']:.1f} per 100,000. "
                        f"Mortality was concentrated elsewhere: {values['top_death_subtype']} had the highest 2023 age-standardized mortality rate at {values['top_death_subtype_rate']:.2f} per 100,000."
                    )
                elif subsection == "Country-level ecological association":
                    subsection = "Country-level ageing and risk ecology"
                    mutable.append(
                        f"The risk ecology suggested that the mortality pattern was not driven by age structure alone. "
                        f"At country level, higher household air pollution from solid fuels was associated with higher standardized skin mortality (rho={rho_household:.3f}; p={fmt_p(p_household)}), and higher albuminuria prevalence was also associated with higher mortality (rho={rho_albuminuria:.3f}; p={fmt_p(p_albuminuria)}). "
                        "Smoking prevalence showed an inverse ecological correlation, which is difficult to interpret causally and is likely confounded by the development gradient captured in the broader country profile."
                    )
                    mutable.append(
                        f"The highest-mortality countries did not have a single uniform risk signature, but they were distinguishable on selected exposures. "
                        f"The top-20 mortality settings had a higher median adult BMI exposure than the remaining countries (median {bmi_top20_median:.2f} vs {bmi_others_median:.2f}; p={fmt_p(bmi_top20_p)}). "
                        "By contrast, the top-20 versus remaining-country contrasts for household air pollution and albuminuria were less extreme as grouped distribution tests, suggesting that a subset of very high-risk settings drives much of the positive ecological gradient rather than a uniform shift across the entire top-20 group."
                    )
                    mutable.append(
                        f"Medical-resource context pointed in the opposite direction. "
                        f"Using the latest available physician density value from 2015-2023 (median source year {physician_stats['median_year']:.0f}), higher physician density was inversely associated with both skin mortality (rho={physician_stats['mort_rho']:.3f}; p={fmt_p(physician_stats['mort_p'])}) and skin DALY rates (rho={physician_stats['daly_rho']:.3f}; p={fmt_p(physician_stats['daly_p'])}) across {int(physician_stats['n'])} countries and territories (table S16). "
                        f"Countries in the lowest physician-density quartile had higher median mortality and DALY rates than those in the highest quartile (mortality {physician_stats['q1_asmr']:.2f} vs {physician_stats['q4_asmr']:.2f} per 100,000; DALYs {physician_stats['q1_daly']:.1f} vs {physician_stats['q4_daly']:.1f} per 100,000; table S17)."
                    )
                elif subsection == "Asia-Pacific rate geography":
                    mutable[0] = mutable[0].replace("figure 4", "supplementary figure S10")
                elif subsection == "Highest-mortality settings":
                    mutable[0] = (
                        f"The countries and territories with the highest 2023 age-standardized mortality rates were {values['top5_text']} (figure S1; table S4). "
                        f"{values['top_country']} had the highest recorded rate at {values['top_country_rate']:.2f} per 100,000. "
                        "The leading locations were concentrated in small island states and mixed middle-income settings rather than in the most aged high-income countries. "
                        "This geographic pattern is consistent with the ecological correlations and suggests that skin mortality in the contemporary period is especially sensitive to context-specific combinations of chronic wound risk, diabetes burden, infection management, frailty care, and service accessibility."
                    )
                    mutable.append(
                        "The selected exposure profile reinforced that interpretation. "
                        f"Across the five highest-mortality settings, all were above the global median for adult BMI exposure, while several Pacific island settings were also above the global median for smoking prevalence, household air pollution, and albuminuria prevalence (table S14; figure S9). "
                        f"The top-five exposure pattern was as follows: {top5_risk_text}."
                    )
                result_blocks.append((subsection, mutable))

            result_blocks.insert(
                1,
                (
                    "Global geographic distribution",
                    [
                        f"Authenticated country-level maps showed marked heterogeneity across all burden domains in 2023 (figure 2; table S15). "
                        f"The highest mortality rates were observed in {global_geo_summary['deaths_top3_text']}. "
                        f"The highest DALY rates were observed in {global_geo_summary['dalys_top3_text']}, the highest prevalence rates in {global_geo_summary['prevalence_top3_text']}, and the highest incidence rates in {global_geo_summary['incidence_top3_text']}.",
                        "The geographic pattern also indicated that fatal and non-fatal burden were not distributed identically. "
                        "Mortality clustered more tightly in a smaller group of island and maritime settings, whereas high prevalence and incidence extended across a broader set of tropical and middle-income contexts. "
                        "This divergence supports the interpretation that distinct clinical and health-system mechanisms underlie the fatal and non-fatal components of global skin burden."
                    ],
                ),
            )

            insert_idx = 4
            pandemic_block = (
                "Pandemic-period deviation and exploratory forecast to 2050",
                [
                    "The authenticated annual global Results Tool series added an explicit view of the pandemic period and the near-term burden trajectory (figure 5; table 4). "
                    "Across all four burden domains, age-standardized rates had been increasing gradually before 2020, but the pandemic-period deviation was not uniform across outcomes.",
                    f"For mortality, the global age-standardized rate was slightly below the pre-pandemic counterfactual in 2020 ({death_dev[2020]:.2f}%), but it shifted above the counterfactual from 2021 onward ({death_dev[2021]:.2f}% in 2021, {death_dev[2022]:.2f}% in 2022, and {death_dev[2023]:.2f}% in 2023). "
                    f"DALY rates showed only very small positive deviations ({daly_dev[2021]:.2f}% to {daly_dev[2023]:.2f}% during 2021-2023), whereas prevalence and incidence remained slightly below expected through 2023 ({prev_dev[2023]:.2f}% and {inc_dev[2023]:.2f}%, respectively). "
                    "This pattern indicates that the pandemic period aligned more strongly with the fatal end of the burden spectrum than with a major acceleration of non-fatal skin disease rates.",
                    f"The exploratory statistical projections suggested further increases by 2050. "
                    f"The global death rate was projected to reach {milestone['Deaths'][2030]:.2f} per 100,000 in 2030, {milestone['Deaths'][2040]:.2f} in 2040, and {milestone['Deaths'][2050]:.2f} in 2050. "
                    f"Over the same horizons, the DALY rate was projected to rise to {milestone['DALYs'][2030]:.1f}, {milestone['DALYs'][2040]:.1f}, and {milestone['DALYs'][2050]:.1f} per 100,000; the prevalence rate to {milestone['Prevalence'][2030]:.1f}, {milestone['Prevalence'][2040]:.1f}, and {milestone['Prevalence'][2050]:.1f}; and the incidence rate to {milestone['Incidence'][2030]:.1f}, {milestone['Incidence'][2040]:.1f}, and {milestone['Incidence'][2050]:.1f} per 100,000."
                ],
            )
            result_blocks.insert(insert_idx, pandemic_block)
            result_blocks.insert(
                insert_idx + 1,
                (
                    "Cross-cause global context",
                    [
                        f"To benchmark the relative scale of skin burden against disease systems that receive much greater policy attention, we extracted authenticated 2023 global age-standardized rates for cardiovascular diseases, neoplasms, diabetes and kidney diseases, and skin and subcutaneous diseases (table S18). "
                        f"The global skin mortality rate was {comparison_summary['skin_and_subcutaneous_diseases_death_rate']:.2f} per 100,000, compared with {comparison_summary['cardiovascular_diseases_death_rate']:.2f} for cardiovascular diseases, {comparison_summary['neoplasms_death_rate']:.2f} for neoplasms, and {comparison_summary['diabetes_and_kidney_diseases_death_rate']:.2f} for diabetes and kidney diseases.",
                        f"The corresponding skin DALY rate was {comparison_summary['skin_and_subcutaneous_diseases_daly_rate']:.1f} per 100,000, versus {comparison_summary['cardiovascular_diseases_daly_rate']:.1f}, {comparison_summary['neoplasms_daly_rate']:.1f}, and {comparison_summary['diabetes_and_kidney_diseases_daly_rate']:.1f} per 100,000 for those same systems. "
                        "These comparisons confirm that skin burden is lower than the largest fatal disease systems, but far from negligible at a global level."
                    ],
                ),
            )
            new_sections.append((section_title, result_blocks))
            continue

        if section_title == "Discussion":
            discussion_blocks: List[Tuple[Optional[str], List[str]]] = []
            for subsection, paragraphs in blocks:
                mutable = list(paragraphs)
                for idx, paragraph in enumerate(mutable):
                    if paragraph.startswith("This analysis provides three main insights."):
                        mutable[idx] = (
                            "This analysis yields three main findings. "
                            "Global skin burden increased between 1990 and 2023, subtype patterns differed sharply across incidence, disability, and mortality, and the highest standardized mortality clustered more with structural vulnerability than with ageing alone."
                        )
                    elif paragraph.startswith("The first of these insights is important because"):
                        mutable[idx] = (
                            "The rise in counts matters because health systems absorb burden through patients, visits, dressings, admissions, and nursing time rather than standardized rates alone. "
                            "In an ageing world, modest rate changes can still translate into large growth in service demand."
                        )
                    elif paragraph.startswith("The divergence between counts and rates is central"):
                        mutable[idx] = (
                            "The divergence between counts and rates therefore has direct planning implications. "
                            "Counts speak to workforce and supply needs, whereas standardized rates indicate whether epidemiologic intensity is changing after adjustment for age structure."
                        )
                    elif paragraph.startswith("The subtype pattern offers an additional clinical lesson."):
                        mutable[idx] = (
                            "The subtype pattern is clinically important. "
                            "Dermatitis, fungal skin disease, and scabies dominated non-fatal burden, whereas bacterial skin disease and decubitus ulcer drove mortality, linking skin burden to frailty care, wound care, diabetes management, rehabilitation, and infection control."
                        )
                    elif paragraph.startswith("This domain-specific burden profile supports a differentiated policy response."):
                        mutable[idx] = (
                            "The policy response should therefore be differentiated. "
                            "Symptom-heavy conditions require sustained outpatient and community dermatology capacity, whereas mortality reduction depends more on pressure-injury prevention, chronic-wound surveillance, diabetic skin care, institutional care quality, and early treatment of severe infection.[25-44]"
                        )
                    elif paragraph.startswith("The negative ecological association between demographic ageing and standardized skin mortality deserves careful interpretation."):
                        mutable[idx] = (
                            "The inverse ecological association between ageing and standardized mortality is not paradoxical once age standardization is considered. "
                            "After removing age-structure differences, remaining variation is more plausibly related to survival conditions, chronic-disease management, wound care, institutional care quality, and timely treatment of infection."
                        )
                    elif paragraph.startswith("This point cannot be emphasized enough:"):
                        mutable[idx] = (
                            "This finding does not mean ageing is unimportant. "
                            "Ageing still expands the pool of people exposed to frailty, immobility, chronic wounds, and cumulative skin disease; it simply means that between-country mortality differences are shaped more by context than by age composition alone."
                        )
                    elif paragraph.startswith("In practice, this means demographic ageing and mortality vulnerability can point in opposite directions."):
                        mutable[idx] = (
                            "In practice, demographic ageing and mortality vulnerability can move in opposite directions. "
                            "Countries with older populations may still achieve lower standardized mortality if prevention, long-term care, wound management, and infection pathways are strong, whereas younger settings can retain high mortality when those systems are weak."
                        )
                    elif paragraph.startswith("These findings also have implications for the healthy-longevity agenda."):
                        mutable[idx] = (
                            "These findings matter for the healthy-longevity agenda. "
                            "Skin disease sits at the junction of comfort, mobility, dignity, infection prevention, and institutional care quality, yet it remains peripheral in many ageing strategies. "
                            "Pressure-injury prevention, skin inspection in diabetes and immobility, treatment of bacterial skin disease, and chronic management of pruritic inflammatory disorders should be treated as core components of healthy ageing.[17-20]"
                        )
                insert_after_idx = None
                limitations_idx = None
                strengths_idx = None
                future_idx = None
                for idx, paragraph in enumerate(mutable):
                    if paragraph.startswith("In practice, this means demographic ageing and mortality vulnerability"):
                        insert_after_idx = idx
                    if paragraph.startswith("This study has limitations."):
                        limitations_idx = idx
                    if paragraph.startswith("The analysis nevertheless has important strengths."):
                        strengths_idx = idx
                    if paragraph.startswith("Future work should extend this framework"):
                        future_idx = idx

                if insert_after_idx is not None:
                    mutable.insert(
                        insert_after_idx + 1,
                        "The selected ecological risk analysis sharpens that interpretation. "
                        "Higher mortality tracked household air pollution and albuminuria, and the top-20 mortality settings also showed higher adult BMI exposure than the rest of the global distribution. "
                        "These are ecological signals rather than attributable-risk estimates, but they fit a pattern of structural disadvantage, metabolic vulnerability, vascular fragility, and chronic-wound susceptibility."
                    )
                    mutable.insert(
                        insert_after_idx + 2,
                        f"The physician-density analysis extends the same argument to service capacity. "
                        f"Across {int(physician_stats['n'])} countries and territories, higher physician density was associated with lower skin mortality and lower skin DALY rates, and the lowest-density quartile had substantially higher median burden than the highest-density quartile. "
                        "Although ecological, this pattern is consistent with the importance of access to wound care, infection management, chronic-disease follow-up, and frailty care."
                    )
                    mutable.insert(
                        insert_after_idx + 3,
                        "The pandemic-period analysis adds a second signal. "
                        "Mortality drifted above the pre-pandemic counterfactual from 2021 onward, whereas incidence and prevalence remained close to or slightly below expected. "
                        "This pattern is more consistent with disruption of access, delayed presentation, and interruption of wound and infection care than with a major expansion of non-fatal skin disease incidence."
                    )
                    mutable.insert(
                        insert_after_idx + 4,
                        f"The cross-cause comparison clarifies the policy problem. "
                        f"In 2023, skin mortality ({comparison_summary['skin_and_subcutaneous_diseases_death_rate']:.2f} per 100,000) was far lower than cardiovascular, neoplastic, or diabetes and kidney disease mortality, but the skin DALY rate still reached {comparison_summary['skin_and_subcutaneous_diseases_daly_rate']:.1f} per 100,000. "
                        "Skin disease is therefore not a leading fatal system, but it remains a substantial source of global disability and service demand that receives comparatively little strategic attention."
                    )
                    mutable.insert(
                        insert_after_idx + 5,
                        f"The exploratory 2050 forecast reinforces the service-planning relevance of the study. "
                        f"Even under a conservative damped-trend model, the global death rate was projected to rise to {milestone['Deaths'][2050]:.2f} per 100,000 and the DALY rate to {milestone['DALYs'][2050]:.1f} per 100,000 by 2050. "
                        "These are not official foresight estimates, but they point to continued accumulation of dermatologic workload if recent trends persist."
                    )

                if limitations_idx is not None:
                    mutable[limitations_idx] = (
                        "This study has limitations. "
                        "First, all burden estimates came from modeled secondary data rather than directly observed global registry counts. "
                        "Second, the core global descriptive table still depended on the locked local reproducible extract, while the pandemic and forecast module used a separate authenticated annual Results Tool export; these sources are consistent in direction but not identical in packaging. "
                        "Third, the ecological risk analysis cannot establish causation and should not be interpreted as risk-attributable burden for skin disease overall. "
                        "Fourth, the pandemic-period module estimates deviation from pre-2020 trend rather than the direct biological effect of COVID-19. "
                        "Fifth, the 2050 projections are exploratory statistical forecasts rather than official GBD foresight scenarios."
                    )
                    if limitations_idx + 1 < len(mutable):
                        mutable[limitations_idx + 1] = (
                            "Additional limitations should also be noted. "
                            "The World Bank World Development Indicators describe population structure and survival but do not capture frailty prevalence, long-term care coverage, care-home density, pressure-injury prevention programs, or access to antibiotics and wound supplies. "
                            "Likewise, the selected risk indicators were chosen because they were locally available and analytically stable, not because they exhaust all plausible determinants of skin mortality. "
                            "The GBD skin cause hierarchy is also necessarily coarse for a discussion framed around clinical dermatology: it combines infectious, inflammatory, ulcerative, symptom-based, and residual entities within one umbrella category, and a large residual group of other skin and subcutaneous diseases remains analytically important. "
                            "As a result, aggregate estimates can obscure clinically meaningful divergence between high-mortality categories such as bacterial skin disease and decubitus ulcer and high-DALY categories such as dermatitis. "
                            "Nationally aggregated indicators also obscure within-country inequalities, and some of the highest-mortality settings are small islands or territories where local context may be especially important. "
                            "The ecological findings should therefore be read as hypothesis-generating guidance for prioritization rather than as a substitute for country-specific clinical or implementation studies."
                        )

                if strengths_idx is not None and strengths_idx + 1 < len(mutable):
                    mutable[strengths_idx + 1] = (
                        "A further strength is the integration of global burden estimates, country-level ecological context, authenticated annual series, and forward projections within one reproducible framework. "
                        "That design allows the paper to connect demographic interpretation, health-system context, and forward planning without claiming more precision than the available data support."
                    )

                if future_idx is not None:
                    mutable[future_idx] = (
                        "Future work should extend this framework in four directions. "
                        "First, additional official GBD extraction should be used to recover full age-group and sex-specific series, especially for older adults. "
                        "Second, the ecological analysis should be expanded with more explicit covariates for diabetes, vascular disease, health expenditure, long-term care capacity, wound-care infrastructure, and where possible dermatologist-specific workforce measures. "
                        "Third, country-level risk ecology should be followed by direct attributable-burden analyses if a stable public skin-risk panel becomes available. "
                        "Fourth, the local Global Aging Data workspace should be used as an individual-level validation platform rather than being merged mechanically into the current ecological design."
                    )
                    mutable.insert(
                        future_idx + 1,
                        "A particularly practical next step is to validate these ecological signals in person-level ageing cohorts. "
                        "The desktop Global Aging Data platform already available in this workspace includes harmonized LASI skin disease variables that can be linked to pain, mobility, sleep, BMI, and psychosocial indicators. "
                        "That follow-up design would allow the present population-level signals on metabolic vulnerability and functional burden to be tested at the individual level without forcing incompatible data structures into the current GBD ecology paper."
                    )

                cleaned: List[str] = []
                seen_prefixes: set[str] = set()
                dedupe_prefixes = [
                    "This study has limitations.",
                    "Additional limitations should also be noted.",
                    "Another strength is the way the figure and table set has been constrained",
                    "Future work should extend this framework",
                ]
                drop_prefixes = [
                    "A further next step would be to build a dedicated older-adult version of the paper",
                    "The analysis nevertheless has important strengths.",
                    "From a research-design perspective, this combined framework also improves manuscript coherence.",
                ]
                for paragraph in mutable:
                    if any(paragraph.startswith(prefix) for prefix in drop_prefixes):
                        continue
                    matched_prefix = next((prefix for prefix in dedupe_prefixes if paragraph.startswith(prefix)), None)
                    if matched_prefix is not None:
                        if matched_prefix in seen_prefixes:
                            continue
                        seen_prefixes.add(matched_prefix)
                    cleaned.append(paragraph)

                discussion_blocks.append((subsection, cleaned))
            new_sections.append((section_title, discussion_blocks))
            continue

        new_sections.append((section_title, blocks))

    return new_sections


def build_main_markdown(
    summary: Dict[str, str],
    research_in_context: Dict[str, str],
    sections: List[Tuple[str, List[Tuple[Optional[str], List[str]]]]],
    references: Sequence[str],
    main_tables: Sequence[Tuple[str, pd.DataFrame]],
    main_figures: Sequence[Tuple[str, str, Path]],
) -> str:
    lines: List[str] = [f"# {TITLE}", "", "## Summary", ""]
    for key in ["Background", "Methods", "Findings", "Interpretation", "Funding"]:
        lines.extend([f"### {key}", "", summary[key], ""])
    lines.extend(["## Research in context", ""])
    for heading, text in research_in_context.items():
        lines.extend([f"### {heading}", "", text, ""])
    for section_title, blocks in sections:
        lines.extend([f"## {section_title}", ""])
        for subsection, paragraphs in blocks:
            if subsection:
                lines.extend([f"### {subsection}", ""])
            for paragraph in paragraphs:
                lines.extend([paragraph, ""])
    lines.extend(["## References", ""])
    for idx, ref in enumerate(references, start=1):
        lines.append(f"{idx}. {ref}")
    lines.extend(["", "## Tables", ""])
    for title, _ in main_tables:
        lines.append(f"- {title}")
    lines.extend(["", "## Figure Legends And Figures", ""])
    for title, legend, _ in main_figures:
        lines.extend([f"### {title}", "", legend, ""])
    return "\n".join(lines).rstrip() + "\n"


def build_supplementary_markdown(
    supp_figures: Sequence[Tuple[str, str, Path]],
    supp_tables: Sequence[Tuple[str, pd.DataFrame]],
) -> str:
    lines: List[str] = ["# Supplementary Appendix", "", "## Supplementary Figures", ""]
    for title, legend, _ in supp_figures:
        lines.extend([f"### {title}", "", legend, ""])
    lines.extend(["## Supplementary Tables", ""])
    for title, _ in supp_tables:
        lines.append(f"- {title}")
    return "\n".join(lines).rstrip() + "\n"


def build_support_lines_with_headings(content: Dict[str, str]) -> List[str]:
    lines: List[str] = []
    for heading, text in content.items():
        lines.extend([f"## {heading}", "", text, ""])
    return lines


def build_reference_lines(references: Sequence[str]) -> List[str]:
    return [f"{idx}. {ref}" for idx, ref in enumerate(references, start=1)]


def count_main_words(summary: Dict[str, str], research_in_context: Dict[str, str], sections) -> Tuple[int, int]:
    summary_text = " ".join(summary.values())
    body_parts: List[str] = []
    body_parts.extend(research_in_context.values())
    for _, blocks in sections:
        for _, paragraphs in blocks:
            body_parts.extend(paragraphs)
    return word_count(summary_text), word_count(" ".join(body_parts))


def build_checklist_lines(
    summary_word_count: int,
    main_word_count: int,
    reference_count: int,
    main_tables: Sequence[Tuple[str, pd.DataFrame]],
    supp_tables: Sequence[Tuple[str, pd.DataFrame]],
    main_figures: Sequence[Tuple[str, str, Path]],
    supp_figures: Sequence[Tuple[str, str, Path]],
    render_summary: Dict[str, object],
    risk_country_n: int,
) -> List[str]:
    return [
        "## Core checks",
        f"Structured summary word count = {summary_word_count}: {'PASS' if summary_word_count <= 300 else 'FAIL'}",
        f"Main text word count = {main_word_count}: {'PASS' if main_word_count >= 5000 else 'FAIL'}",
        f"Reference count = {reference_count}: {'PASS' if reference_count >= 35 else 'FAIL'}",
        f"Main display items = {len(main_figures)} figures + {len(main_tables)} tables: {'PASS' if len(main_figures) == 5 and len(main_tables) == 5 else 'FAIL'}",
        f"Supplementary display items = {len(supp_figures)} figures + {len(supp_tables)} tables: {'PASS' if len(supp_figures) >= 9 and len(supp_tables) >= 14 else 'FAIL'}",
        f"Risk ecology analytic countries = {risk_country_n}: {'PASS' if risk_country_n == 198 else 'CHECK'}",
        "",
        "## New analytic modules",
        "Country-level selected risk exposure ecology: added.",
        "Country-level physician-density ecology using World Bank WDI: added.",
        "Cross-cause context comparison with cardiovascular diseases, neoplasms, and diabetes and kidney diseases: added.",
        "Pandemic-period deviation using authenticated annual global Results Tool series: added.",
        "Exploratory 2050 forecast using damped-trend exponential smoothing: added.",
        "",
        "## Rendering QA",
        f"Render pipeline available: {'PASS' if render_summary.get('available') else 'FAIL'}",
        f"Rendered documents: {render_summary.get('rendered_count', 0)}/{render_summary.get('document_count', 0)}",
        f"Rendered pages total: {render_summary.get('page_count', 0)}",
        "",
        "## Remaining manual confirmation before upload",
        "Final author order and affiliations.",
        "Corresponding-author contact details.",
        "Funding statement and role-of-the-funder statement.",
        "Declaration of interests for each author.",
        "Final ethics/originality wording.",
    ]


def build_qc_lines(
    summary_word_count: int,
    main_word_count: int,
    main_tables: Sequence[Tuple[str, pd.DataFrame]],
    supp_tables: Sequence[Tuple[str, pd.DataFrame]],
    main_figures: Sequence[Tuple[str, str, Path]],
    supp_figures: Sequence[Tuple[str, str, Path]],
    annual_df: pd.DataFrame,
    pandemic_df: pd.DataFrame,
    forecast_df: pd.DataFrame,
    risk_summary: pd.DataFrame,
    physician_summary_df: pd.DataFrame,
    comparator_table: pd.DataFrame,
    render_summary: Dict[str, object],
) -> List[str]:
    lines = [
        "QC scope: main manuscript, supplementary appendix, annual official series, ecological risk module, pandemic-period deviation module, exploratory forecast module.",
        f"Structured summary word count: {summary_word_count}",
        f"Main text word count: {main_word_count}",
        f"Main figures: {len(main_figures)}",
        f"Main tables: {len(main_tables)}",
        f"Supplementary figures: {len(supp_figures)}",
        f"Supplementary tables: {len(supp_tables)}",
        "",
        "Data-source audit",
        f"- Locked global descriptive context: {AGING_DIR / 'skin_aging_global_context_1990_2023.csv'}",
        f"- Locked country-level ageing ecology: {AGING_DIR / 'skin_aging_2023_country_complete.csv'}",
        f"- Annual official global Results Tool source: {ANALYSIS_DIR / 'skin_global_annual_official_asr_1990_2023.csv'}",
        f"- Official 2023 country-level map source: {ANALYSIS_DIR / 'skin_global_country_asr_2023_official.csv'}",
        f"- Harmonized global map panel source: {ANALYSIS_DIR / 'skin_global_country_rate_map_dataset_2023.csv'}",
        f"- Physician density source table: {ANALYSIS_DIR / 'skin_country_physician_density_2015_2023.csv'}",
        f"- Physician density summary table: {ANALYSIS_DIR / 'skin_country_physician_density_summary_2023.csv'}",
        f"- Cross-cause context source table: {ANALYSIS_DIR / 'skin_global_cross_cause_comparison_2023.csv'}",
        f"- Pandemic counterfactual table: {ANALYSIS_DIR / 'skin_global_pandemic_counterfactual_2020_2023.csv'}",
        f"- Forecast table: {ANALYSIS_DIR / 'skin_global_exploratory_forecast_2024_2050.csv'}",
        f"- Risk ecology merged table: {ANALYSIS_DIR / 'skin_country_risk_ecology_2023.csv'}",
        f"- GBD login state used for authenticated annual query: {STATE_PATH}",
        "",
        "Consistency checks",
        f"- Annual global series rows: {len(annual_df)} (expected 136 for 34 years x 4 measures): {'PASS' if len(annual_df) == 136 else 'FAIL'}",
        f"- Pandemic table rows: {len(pandemic_df)} (expected 16): {'PASS' if len(pandemic_df) == 16 else 'FAIL'}",
        f"- Forecast table rows: {len(forecast_df)} (expected 108): {'PASS' if len(forecast_df) == 108 else 'FAIL'}",
        f"- Risk summary rows: {len(risk_summary)} (expected 7): {'PASS' if len(risk_summary) == 7 else 'FAIL'}",
        f"- Main figure assets present: {'PASS' if all(path.exists() for _, _, path in main_figures) else 'FAIL'}",
        f"- Supplementary figure assets present: {'PASS' if all(path.exists() for _, _, path in supp_figures) else 'FAIL'}",
        "",
        "Key numerical checks",
    ]
    for row in risk_summary.itertuples(index=False):
        lines.append(
            f"- {row.indicator}: rho={row.spearman_rho:.3f}, p={row.p_value:.2e}, top20 median={row.top20_median:.3f}, others median={row.others_median:.3f}"
        )
    physician_row = physician_summary_df.iloc[0]
    lines.append(
        f"- Physician density: n={int(physician_row['Countries_with_data'])}, mortality rho={float(physician_row['Mortality_spearman_rho']):.3f}, DALY rho={float(physician_row['DALY_spearman_rho']):.3f}, source years {int(physician_row['Min_source_year'])}-{int(physician_row['Max_source_year'])}"
    )
    skin_row = comparator_table[comparator_table["Cause_group"] == "Skin and subcutaneous diseases"].iloc[0]
    cvd_row = comparator_table[comparator_table["Cause_group"] == "Cardiovascular diseases"].iloc[0]
    neo_row = comparator_table[comparator_table["Cause_group"] == "Neoplasms"].iloc[0]
    dkd_row = comparator_table[comparator_table["Cause_group"] == "Diabetes and kidney diseases"].iloc[0]
    lines.append(
        f"- Cross-cause context (death ASR): skin={float(skin_row['Deaths_mean']):.2f}, cardiovascular={float(cvd_row['Deaths_mean']):.2f}, neoplasms={float(neo_row['Deaths_mean']):.2f}, diabetes and kidney={float(dkd_row['Deaths_mean']):.2f}"
    )
    lines.append(
        f"- Cross-cause context (DALY ASR): skin={float(skin_row['DALYs_mean']):.1f}, cardiovascular={float(cvd_row['DALYs_mean']):.1f}, neoplasms={float(neo_row['DALYs_mean']):.1f}, diabetes and kidney={float(dkd_row['DALYs_mean']):.1f}"
    )
    lines.extend(["", "Rendering checks"])
    if render_summary.get("available"):
        lines.append(
            f"- Render pipeline available: yes; rendered {render_summary.get('rendered_count')}/{render_summary.get('document_count')} documents; total pages={render_summary.get('page_count')}"
        )
        for item in render_summary.get("documents", []):
            lines.append(
                f"- {Path(str(item.get('docx', ''))).name}: {'PASS' if item.get('ok') else 'FAIL'}, pages={item.get('page_count')}"
            )
    else:
        lines.append("- Render pipeline unavailable.")
    lines.extend(
        [
            "",
            "Interpretation guardrails",
            "- Risk analyses are ecological correlations and grouped comparisons, not attributable-burden estimates.",
            "- Pandemic estimates are deviations from counterfactual trend, not direct causal effects of COVID-19 infection.",
            "- Forecasts are exploratory statistical projections, not official IHME foresight outputs.",
        ]
    )
    return lines


def build_readme_lines(
    summary_word_count: int,
    main_word_count: int,
    reference_count: int,
) -> List[str]:
    return [
        "Updated submission-oriented package for the first skin burden study.",
        "",
        "New additions in this version:",
        "1. Lancet-style redesign of the main figure architecture.",
        "2. Authenticated 2023 global country-rate atlas for deaths, DALYs, prevalence, and incidence.",
        "3. Selected ecological risk exposure analysis.",
        "4. Physician-density ecology using World Bank WDI physician workforce data.",
        "5. Cross-cause contextual comparison with cardiovascular diseases, neoplasms, and diabetes and kidney diseases.",
        "6. Pandemic-period deviation analysis based on authenticated annual global Results Tool series.",
        "7. Exploratory forecasts to 2050.",
        "8. Discussion update covering health-resource context, disease-classification limitations, and cross-system prioritization.",
        "",
        f"Structured summary word count: {summary_word_count}",
        f"Main text word count: {main_word_count}",
        f"Reference count: {reference_count}",
        "",
        "Core files in this folder:",
        f"- {MAIN_DOCX.name}",
        f"- {SUPP_DOCX.name}",
        f"- {CHECKLIST_DOCX.name}",
        f"- {QC_DOCX.name}",
        f"- {SUMMARY_JSON.name}",
        "",
        "Administrative forms, title page, and cover letter from the prior final package were copied forward unchanged.",
        "Only author-confirmed metadata should be treated as upload-ready.",
    ]


def write_tables_csv(table_map: Dict[str, pd.DataFrame]) -> None:
    for filename, df in table_map.items():
        df.to_csv(TABLE_DIR / filename, index=False)


def main() -> None:
    ensure_clean_update_dir()
    remove_obsolete_docs()

    builder45 = load_module("skin_builder45_update57", SCRIPT_45)
    builder45.PROJECT_DIR = PROJECT_DIR
    builder45.ROOT = ROOT
    builder45.DIRF_PATH = ROOT / "data" / "silver" / "gbd" / "gbd2023_dirf_global_core_tidy.csv"
    builder45.MORTALITY_PATH = ROOT / "data" / "silver" / "gbd" / "gbd2023_mortality_s7_both_sex_long.csv"
    builder45.OUTPUT_DIR = OUTPUT_ROOT
    builder45.FIGURE_DIR = OUTPUT_ROOT / "figures"
    builder45.TABLE_DIR = OUTPUT_ROOT / "tables"
    builder45.MANUSCRIPT_DIR = OUTPUT_ROOT / "manuscript"
    builder45.GLOBAL_CONTEXT_CANDIDATES = [
        OUTPUT_ROOT / "aging_analysis_outputs" / "skin_aging_global_context_1990_2023.csv",
        OUTPUT_ROOT / "aging_analysis_outputs_test" / "skin_aging_global_context_1990_2023.csv",
    ]
    builder45.COUNTRY_COMPLETE_CANDIDATES = [
        OUTPUT_ROOT / "aging_analysis_outputs" / "skin_aging_2023_country_complete.csv",
        OUTPUT_ROOT / "aging_analysis_outputs_test" / "skin_aging_2023_country_complete.csv",
    ]
    module50 = load_module("skin_builder50_update57", SCRIPT_50)
    module50.PROJECT_DIR = PROJECT_DIR
    module50.PACKAGE_ROOT = PACKAGE_ROOT
    module50.OUTPUT_DIR = OUTPUT_ROOT
    module50.FIGURE_DIR = FIGURE_DIR
    module50.TABLE_DIR = TABLE_DIR
    module50.MANUSCRIPT_DIR = UPDATED_DIR
    module50.AGING_DIR = OUTPUT_ROOT / "aging_analysis_outputs"
    module50.APAC_RESULTS_DIR = OUTPUT_ROOT / "apac_results_tool_outputs"
    module50.OFFICIAL_APAC_ASR_PATH = module50.APAC_RESULTS_DIR / "skin_apac_official_asr_2023.csv"

    access_token = load_state_token(STATE_PATH)
    annual_df = fetch_official_global_annual_rates(access_token)
    pandemic_df, forecast_df, milestone_df = compute_pandemic_and_forecast(annual_df)

    global_context = pd.read_csv(AGING_DIR / "skin_aging_global_context_1990_2023.csv")
    global_context = global_context[global_context["year_id"].isin([1990, 2010, 2020, 2023])].copy()
    country_complete, _ = builder45.load_country_complete()
    correlations, tertiles, top20 = builder45.compute_country_ecology(country_complete)
    subtype_dirf, subtype_mortality = builder45.load_subtype_profiles()
    apac, _ = module50.build_apac_datasets(country_complete)
    apac_tables = module50.build_apac_tables(apac)

    values = builder45.build_value_map(
        global_context,
        country_complete,
        correlations,
        tertiles,
        top20,
        subtype_dirf,
        subtype_mortality,
    )

    risk_country_df, risk_summary_df, top20_profile = load_risk_exposure_data(country_complete)
    global_country_rates_2023 = fetch_official_global_country_rates_2023(
        access_token,
        country_complete["location_id"].astype(int).tolist(),
    )
    global_map_df = build_global_map_dataset(module50, country_complete, global_country_rates_2023)
    global_geo_summary = build_global_geo_summary(global_map_df)
    physician_df = build_physician_density_dataset(country_complete, global_map_df)
    physician_summary_df, physician_quartile_df, physician_stats = build_physician_density_tables(physician_df)
    comparator_df = fetch_global_cause_comparison_2023(access_token)
    comparator_table = build_cross_cause_comparison_table(comparator_df)
    comparison_summary = build_cross_cause_summary(comparator_table)

    table1 = build_table1()
    table2 = builder45.build_main_table(global_context)
    result_tables = builder45.build_main_result_tables(subtype_dirf, subtype_mortality, correlations, tertiles, top20)
    table4 = build_pandemic_summary_table(pandemic_df, forecast_df)
    table5 = risk_summary_df.rename(
        columns={
            "domain": "Domain",
            "indicator": "Indicator",
            "definition": "Definition",
            "metric_type": "Metric type",
            "countries_with_data": "Countries with data",
            "spearman_rho": "Spearman rho",
            "p_value": "P value",
            "top20_median": "Top-20 median",
            "others_median": "Others median",
            "top20_vs_others_p": "Top-20 vs others p",
        }
    )

    annual_df.to_csv(ANALYSIS_DIR / "skin_global_annual_official_asr_1990_2023.csv", index=False)
    global_country_rates_2023.to_csv(ANALYSIS_DIR / "skin_global_country_asr_2023_official.csv", index=False)
    global_map_df.to_csv(ANALYSIS_DIR / "skin_global_country_rate_map_dataset_2023.csv", index=False)
    physician_df.to_csv(ANALYSIS_DIR / "skin_country_physician_density_2015_2023.csv", index=False)
    physician_summary_df.to_csv(ANALYSIS_DIR / "skin_country_physician_density_summary_2023.csv", index=False)
    physician_quartile_df.to_csv(ANALYSIS_DIR / "skin_country_physician_density_quartiles_2023.csv", index=False)
    comparator_df.to_csv(ANALYSIS_DIR / "skin_global_cross_cause_comparison_2023_long.csv", index=False)
    comparator_table.to_csv(ANALYSIS_DIR / "skin_global_cross_cause_comparison_2023.csv", index=False)
    pandemic_df.to_csv(ANALYSIS_DIR / "skin_global_pandemic_counterfactual_2020_2023.csv", index=False)
    forecast_df.to_csv(ANALYSIS_DIR / "skin_global_exploratory_forecast_2024_2050.csv", index=False)
    milestone_df.to_csv(ANALYSIS_DIR / "skin_global_annual_milestones_and_forecast_to_2050.csv", index=False)
    risk_country_df.to_csv(ANALYSIS_DIR / "skin_country_risk_ecology_2023.csv", index=False)
    risk_summary_df.to_csv(ANALYSIS_DIR / "skin_country_risk_summary_2023.csv", index=False)
    top20_profile.to_csv(ANALYSIS_DIR / "skin_top20_risk_profile_2023.csv", index=False)

    table_map = {
        "table1_study_frame_updated_20260312.csv": table1,
        "table2_global_burden_and_aging_context_updated_20260312.csv": table2,
        "table3_subtype_profile_2023_updated_20260312.csv": result_tables["subtype_2023"],
        "table4_pandemic_deviation_and_forecast_20260312.csv": table4,
        "table5_ageing_and_risk_ecology_20260312.csv": table5,
        "tableS1_subtype_profile_2023_long.csv": result_tables["subtype_2023_long"],
        "tableS2_subtype_change_1990_2023_long.csv": result_tables["subtype_change_long"],
        "tableS3_country_correlations.csv": result_tables["correlations"],
        "tableS4_top20_country_asmr_2023.csv": result_tables["top20"],
        "tableS5_age65_tertiles.csv": result_tables["tertiles"],
        "tableS6_apac_country_death_prevalence_incidence_aging_2023.csv": apac_tables["tableS6"],
        "tableS7_apac_map_coverage_2023.csv": apac_tables["tableS7"],
        "tableS8_apac_top15_mortality_2023.csv": apac_tables["tableS8"],
        "tableS9_apac_region_summary_2023.csv": apac_tables["tableS9"],
        "tableS10_apac_top15_prevalence_2023.csv": apac_tables["tableS10"],
        "tableS11_apac_top15_incidence_2023.csv": apac_tables["tableS11"],
        "tableS12_annual_milestones_and_forecast_to_2050.csv": milestone_df,
        "tableS13_pandemic_counterfactual_2020_2023.csv": pandemic_df,
        "tableS14_top20_risk_profile_2023.csv": top20_profile,
        "tableS15_global_country_rates_2023.csv": global_map_df[
            [
                "location_id",
                "gbd_name",
                "gbd_short_name",
                "wb_iso3",
                "deaths_asr_2023",
                "dalys_asr_2023",
                "prevalence_asr_2023",
                "incidence_asr_2023",
                "plot_method",
            ]
        ],
        "tableS16_physician_density_ecology_2023.csv": physician_summary_df,
        "tableS17_physician_density_quartiles_2023.csv": physician_quartile_df,
        "tableS18_cross_cause_context_2023.csv": comparator_table,
    }
    write_tables_csv(table_map)

    make_figure1_lancet_style(global_context, annual_df)
    make_figure2_global_maps(module50, global_map_df)
    make_figure3_subtype_lollipop(subtype_dirf, subtype_mortality)
    make_figure4_ecology_lancet(physician_df, risk_country_df, comparator_table, top20_profile)
    make_figure5_forecast_lancet(annual_df, pandemic_df, forecast_df)
    for filename in [
        "figureS1_top20_country_asmr_2023.png",
        "figureS1_top20_country_asmr_2023.pdf",
        "figureS2_apac_top15_asmr_2023.png",
        "figureS2_apac_top15_asmr_2023.pdf",
        "figureS3_apac_top15_prevalence_2023.png",
        "figureS3_apac_top15_prevalence_2023.pdf",
        "figureS4_apac_top15_incidence_2023.png",
        "figureS4_apac_top15_incidence_2023.pdf",
        "figureS5_apac_age65_scatter.png",
        "figureS5_apac_age65_scatter.pdf",
        "figureS6_apac_age65_map.png",
        "figureS6_apac_age65_map.pdf",
        "figureS7_apac_life_expectancy_map.png",
        "figureS7_apac_life_expectancy_map.pdf",
    ]:
        copy_source_figure(filename)
    make_figureS9_risk_heatmap(top20_profile)
    copy_source_figure("figure4_global_apac_mortality_map.png", "figureS10_global_apac_atlas.png")
    copy_source_figure("figure4_global_apac_mortality_map.pdf", "figureS10_global_apac_atlas.pdf")
    copy_source_figure("figure5_subtype_trends_1990_2023.png", "figureS8_subtype_trends_1990_2023.png")
    copy_source_figure("figure5_subtype_trends_1990_2023.pdf", "figureS8_subtype_trends_1990_2023.pdf")

    summary = build_summary(values, risk_summary_df, pandemic_df, forecast_df, top20_profile, physician_stats)
    research_in_context = build_research_in_context()
    sections = modify_sections(
        builder45,
        values,
        risk_summary_df,
        pandemic_df,
        forecast_df,
        top20_profile,
        global_geo_summary,
        physician_stats,
        comparison_summary,
    )
    references = builder45.build_references()

    main_tables = [
        ("Table 1. Study frame, data sources, and analytical modules", table1),
        ("Table 2. Global burden of skin and subcutaneous diseases and World Bank ageing indicators in 1990 and 2023", table2),
        ("Table 3. Subtype-specific global burden profile of skin and subcutaneous diseases in 2023", result_tables["subtype_2023"]),
        ("Table 4. Pandemic-period deviation and exploratory forecast summary for global skin burden", table4),
        ("Table 5. Country-level ecological summary of ageing and selected risk exposure indicators in 2023", table5),
    ]
    main_figures = [
        (
            "Figure 1",
            "Global skin burden and demographic context, 1990-2023. Panel A shows age-standardized incidence and prevalence rates from the authenticated annual global Results Tool series. Panel B shows age-standardized DALY and mortality rates from the same source. Panel C shows DALY counts and death counts at the benchmark years available in the locked local reproducible extract. Panel D shows World Bank World Development Indicators ageing metrics indexed to 1990=100.",
            FIGURE_DIR / "figure1_lancet_trends_and_ageing.png",
        ),
        (
            "Figure 2",
            "Global geographic distribution of age-standardized burden in 2023. Panels show country-level mortality, DALY, prevalence, and incidence rates from authenticated GBD Results Tool exports for cause id 653 (skin and subcutaneous diseases), both sexes, age-standardized rate outputs. Settings not represented as polygons in the low-resolution Natural Earth basemap are shown as point markers using World Bank country coordinates.",
            FIGURE_DIR / "figure2_global_country_rate_maps_2023.png",
        ),
        (
            "Figure 3",
            "Global subtype profile of skin and subcutaneous diseases in 2023. Panels show the ten leading subtype-specific age-standardized rates for incidence, prevalence, DALYs, and mortality, with whiskers denoting uncertainty intervals.",
            FIGURE_DIR / "figure3_subtype_profile_lollipop_2023.png",
        ),
        (
            "Figure 4",
            "Medical-resource, risk, and cross-system context in 2023. Panels A-C show the relationship between skin mortality and physician density, household air pollution, and albuminuria prevalence. Panel D shows the relationship between physician density and skin DALY rates. Panel E compares global 2023 death and DALY rate ratios for skin disease versus cardiovascular diseases, neoplasms, and diabetes and kidney diseases. Scatterplot y-axes are logarithmic and dashed lines are LOWESS smoothers.",
            FIGURE_DIR / "figure4_ageing_and_risk_ecology_2023.png",
        ),
        (
            "Figure 5",
            "Pandemic-period deviation and exploratory forecasts to 2050. Panels A-B show observed, counterfactual, and forecast age-standardized rates for deaths and DALYs. Panel C shows deviation from the 2010-2019 counterfactual for all four burden measures during 2020-2023. Panel D shows forecast growth indexed to 2023=100 through 2050. Grey shading denotes the pandemic period and the pale shaded region denotes the forecast horizon.",
            FIGURE_DIR / "figure5_pandemic_and_forecast_2050.png",
        ),
    ]

    supp_figures = [
        (
            "Figure S1",
            "Countries and territories with the highest age-standardized mortality rates from skin and subcutaneous diseases in 2023. Bars show point estimates and whiskers show lower and upper uncertainty bounds.",
            FIGURE_DIR / "figureS1_top20_country_asmr_2023.png",
        ),
        (
            "Figure S2",
            "Highest mortality settings in Asia-Pacific in 2023. Bars show official age-standardized mortality rates from the authenticated GBD Results Tool export, with uncertainty bounds.",
            FIGURE_DIR / "figureS2_apac_top15_asmr_2023.png",
        ),
        (
            "Figure S3",
            "Highest prevalence settings in Asia-Pacific in 2023. Bars show official age-standardized prevalence rates from the authenticated GBD Results Tool export.",
            FIGURE_DIR / "figureS3_apac_top15_prevalence_2023.png",
        ),
        (
            "Figure S4",
            "Highest incidence settings in Asia-Pacific in 2023. Bars show official age-standardized incidence rates from the authenticated GBD Results Tool export.",
            FIGURE_DIR / "figureS4_apac_top15_incidence_2023.png",
        ),
        (
            "Figure S5",
            "Asia-Pacific relationship between population ageing and skin mortality in 2023. Points are colored by World Bank region grouping used for the regional supplement.",
            FIGURE_DIR / "figureS5_apac_age65_scatter.png",
        ),
        (
            "Figure S6",
            "Population aged 65 years and older across Asia-Pacific settings in 2023. Point markers indicate locations not represented as polygons in the low-resolution base map.",
            FIGURE_DIR / "figureS6_apac_age65_map.png",
        ),
        (
            "Figure S7",
            "Life expectancy at birth across Asia-Pacific settings in 2023. Point markers indicate locations not represented as polygons in the low-resolution base map.",
            FIGURE_DIR / "figureS7_apac_life_expectancy_map.png",
        ),
        (
            "Figure S8",
            "Temporal change in leading skin disease subtypes, 1990-2023. Panels show the leading subtype-specific age-standardized rates for incidence, prevalence, DALYs, and mortality across the study period.",
            FIGURE_DIR / "figureS8_subtype_trends_1990_2023.png",
        ),
        (
            "Figure S9",
            "Percentile ranks of selected ecological risk indicators among the 20 countries and territories with the highest age-standardized skin mortality in 2023. Each cell shows the within-country percentile rank for smoking prevalence, household air pollution, adult BMI exposure, and albuminuria prevalence; countries are ordered by mortality rate.",
            FIGURE_DIR / "figureS9_top20_risk_heatmap.png",
        ),
        (
            "Figure S10",
            "Four-panel atlas showing global mortality and Asia-Pacific country-level age-standardized rates in 2023. Panel A shows global mortality; panel B shows Asia-Pacific mortality; panel C shows Asia-Pacific prevalence; and panel D shows Asia-Pacific incidence. Asia-Pacific panels use official authenticated GBD Results Tool exports for cause id 653 (skin and subcutaneous diseases), both sexes, and age-standardized rate outputs. Country polygons show mapped locations available in the Natural Earth layer, and point markers denote small islands or territories added using World Bank coordinates because polygon geometry was not available in the low-resolution base map.",
            FIGURE_DIR / "figureS10_global_apac_atlas.png",
        ),
    ]

    supp_tables = [
        ("Table S1. Long-format subtype-specific global burden profile of skin and subcutaneous diseases in 2023", result_tables["subtype_2023_long"]),
        ("Table S2. Long-format subtype-specific change in global skin burden between 1990 and 2023", result_tables["subtype_change_long"]),
        ("Table S3. Country-level ecological correlations between World Bank ageing indicators and skin mortality in 2023", result_tables["correlations"]),
        ("Table S4. Top 20 countries and territories ranked by age-standardized mortality rate from skin and subcutaneous diseases in 2023", result_tables["top20"]),
        ("Table S5. Ageing tertile summary of country-level skin mortality in 2023", result_tables["tertiles"]),
        ("Table S6. Asia-Pacific country-level deaths, prevalence, incidence, and ageing indicators in 2023", apac_tables["tableS6"]),
        ("Table S7. Asia-Pacific map coverage and plotting method in 2023", apac_tables["tableS7"]),
        ("Table S8. Highest mortality settings in Asia-Pacific in 2023", apac_tables["tableS8"]),
        ("Table S9. Asia-Pacific regional summary statistics for deaths, prevalence, and incidence in 2023", apac_tables["tableS9"]),
        ("Table S10. Highest prevalence settings in Asia-Pacific in 2023", apac_tables["tableS10"]),
        ("Table S11. Highest incidence settings in Asia-Pacific in 2023", apac_tables["tableS11"]),
        ("Table S12. Annual official global rates and forecast milestones to 2050", milestone_df),
        ("Table S13. Pandemic-period counterfactual comparison for global age-standardized rates, 2020-2023", pandemic_df),
        ("Table S14. Selected risk exposure profile of the 20 highest-mortality countries and territories in 2023", top20_profile),
        (
            "Table S15. Global country-level age-standardized mortality, DALY, prevalence, and incidence rates in 2023 used for the main geographic atlas",
            table_map["tableS15_global_country_rates_2023.csv"],
        ),
        ("Table S16. Country-level physician density and ecological association with skin mortality and DALYs in 2023", physician_summary_df),
        ("Table S17. Quartile summary of physician density versus skin mortality and DALY rates in 2023", physician_quartile_df),
        ("Table S18. Global 2023 age-standardized death and DALY rates for skin disease and selected major disease systems", comparator_table),
    ]

    main_markdown = build_main_markdown(summary, research_in_context, sections, references, main_tables, main_figures)
    MAIN_MD.write_text(main_markdown, encoding="utf-8")
    module50.write_main_docx(MAIN_DOCX, summary, research_in_context, sections, references, main_tables, main_figures)

    supp_markdown = build_supplementary_markdown(supp_figures, supp_tables)
    SUPP_MD.write_text(supp_markdown, encoding="utf-8")
    module50.write_supplementary_docx(SUPP_DOCX, supp_figures, supp_tables)

    ric_lines = build_support_lines_with_headings(research_in_context)
    write_markdown(RIC_MD, "Research in Context", ric_lines)
    write_docx_with_headings(module50, RIC_DOCX, "Research in Context", ric_lines)

    reference_lines = build_reference_lines(references)
    write_markdown(REF_MD, "Reference List", reference_lines)
    module50.write_support_docx(REF_DOCX, "Reference List", reference_lines)

    summary_word_count, main_word_count = count_main_words(summary, research_in_context, sections)

    render_summary = render_docx_collection(
        [MAIN_DOCX, SUPP_DOCX, CHECKLIST_DOCX, QC_DOCX] if CHECKLIST_DOCX.exists() else [MAIN_DOCX, SUPP_DOCX],
        RENDER_DIR,
    )

    checklist_lines = build_checklist_lines(
        summary_word_count,
        main_word_count,
        len(references),
        main_tables,
        supp_tables,
        main_figures,
        supp_figures,
        render_summary,
        int(len(risk_country_df)),
    )
    write_markdown(CHECKLIST_MD, "Submission Checklist", checklist_lines)
    module50.write_support_docx(CHECKLIST_DOCX, "Submission Checklist", checklist_lines)

    render_summary = render_docx_collection([MAIN_DOCX, SUPP_DOCX, CHECKLIST_DOCX], RENDER_DIR)
    qc_lines = build_qc_lines(
        summary_word_count,
        main_word_count,
        main_tables,
        supp_tables,
        main_figures,
        supp_figures,
        annual_df,
        pandemic_df,
        forecast_df,
        risk_summary_df,
        physician_summary_df,
        comparator_table,
        render_summary,
    )
    write_markdown(QC_MD, "QC Report", qc_lines)
    module50.write_support_docx(QC_DOCX, "QC Report", qc_lines)

    render_summary = render_docx_collection([MAIN_DOCX, SUPP_DOCX, CHECKLIST_DOCX, QC_DOCX], RENDER_DIR)
    checklist_lines = build_checklist_lines(
        summary_word_count,
        main_word_count,
        len(references),
        main_tables,
        supp_tables,
        main_figures,
        supp_figures,
        render_summary,
        int(len(risk_country_df)),
    )
    write_markdown(CHECKLIST_MD, "Submission Checklist", checklist_lines)
    module50.write_support_docx(CHECKLIST_DOCX, "Submission Checklist", checklist_lines)
    qc_lines = build_qc_lines(
        summary_word_count,
        main_word_count,
        main_tables,
        supp_tables,
        main_figures,
        supp_figures,
        annual_df,
        pandemic_df,
        forecast_df,
        risk_summary_df,
        physician_summary_df,
        comparator_table,
        render_summary,
    )
    write_markdown(QC_MD, "QC Report", qc_lines)
    module50.write_support_docx(QC_DOCX, "QC Report", qc_lines)

    readme_lines = build_readme_lines(summary_word_count, main_word_count, len(references))
    write_markdown(README_MD, "Updated Submission Package", readme_lines)
    module50.write_support_docx(README_DOCX, "Updated Submission Package", readme_lines)

    summary_payload = {
        "package_root": str(UPDATED_DIR),
        "main_docx": str(MAIN_DOCX),
        "supplementary_docx": str(SUPP_DOCX),
        "research_in_context_docx": str(RIC_DOCX),
        "checklist_docx": str(CHECKLIST_DOCX),
        "qc_docx": str(QC_DOCX),
        "summary_word_count": summary_word_count,
        "main_word_count": main_word_count,
        "reference_count": len(references),
        "main_figure_count": len(main_figures),
        "main_table_count": len(main_tables),
        "supp_figure_count": len(supp_figures),
        "supp_table_count": len(supp_tables),
        "risk_country_n": int(len(risk_country_df)),
        "physician_country_n": int(physician_stats["n"]),
        "annual_series_rows": int(len(annual_df)),
        "pandemic_rows": int(len(pandemic_df)),
        "forecast_rows": int(len(forecast_df)),
        "render_summary": render_summary,
        "global_country_rate_rows": int(len(global_country_rates_2023)),
        "global_map_rows": int(len(global_map_df)),
        "cross_cause_rows": int(len(comparator_table)),
    }
    SUMMARY_JSON.write_text(json.dumps(summary_payload, ensure_ascii=False, indent=2), encoding="utf-8")
    MANIFEST_JSON.write_text(
        json.dumps(
            {
                "created_at": "2026-03-26",
                "title": TITLE,
                "files": {
                    "main_docx": str(MAIN_DOCX),
                    "main_md": str(MAIN_MD),
                    "supplementary_docx": str(SUPP_DOCX),
                    "supplementary_md": str(SUPP_MD),
                    "research_in_context_docx": str(RIC_DOCX),
                    "reference_list_docx": str(REF_DOCX),
                    "checklist_docx": str(CHECKLIST_DOCX),
                    "qc_docx": str(QC_DOCX),
                    "readme_docx": str(README_DOCX),
                    "summary_json": str(SUMMARY_JSON),
                },
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    print(f"Updated package written to: {UPDATED_DIR}")
    print(f"Main manuscript: {MAIN_DOCX}")
    print(f"Supplementary appendix: {SUPP_DOCX}")
    print(f"QC report: {QC_DOCX}")
    print(f"Summary words: {summary_word_count}")
    print(f"Main words: {main_word_count}")
    print(f"References: {len(references)}")


if __name__ == "__main__":
    main()
