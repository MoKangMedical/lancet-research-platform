#!/usr/bin/env python3
"""Build a submission-facing Lancet-style DOCX draft for the DR-T2D project."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_DIR = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-基于全球疾病负担(GBD)和英国生物数据银行（UK Biobank）的糖尿病视网膜病变与2型糖尿病（T2D）交互作用-10分以上/"
    "Codex研究产出_2026-03-07/04_多期刊投稿包/GBD2023_UKB_DR_T2D_2026-03-07/journal_01_eClinicalMedicine"
)
MANUSCRIPT_PATH = PROJECT_DIR / "13_MAIN_MANUSCRIPT_Lancet_main_v3.md"
FIGURE_LEGENDS_PATH = PROJECT_DIR / "07_LANCET_MAIN_FIGURE_LEGENDS.md"
TABLE_DIR = PROJECT_DIR / "lancet_main_tables"
OUTPUT_PATH = PROJECT_DIR / "12_Lancet_main_submission_draft.docx"

TABLE_SPECS = [
    {
        "title": (
            "Table 1. Baseline profile and crude event counts of the incident UK Biobank "
            "cohort by joint type 2 diabetes-retinal disease status"
        ),
        "path": TABLE_DIR / "Table_1_UKB_Cohort_Profile_Lancet.csv",
        "note": (
            "Values are mean or percentage unless otherwise stated. Event counts are shown "
            "before complete-case restriction."
        ),
    },
    {
        "title": (
            "Table 2. Global burden summary for type 2 diabetes and type 2 diabetes-related "
            "vision loss in 1990 and 2023"
        ),
        "path": TABLE_DIR / "Table_2_GBD_Global_Burden_Lancet.csv",
        "note": (
            "The retinal burden construct refers to the public GBD 2023 proxy of blindness "
            "and vision loss attributable to type 2 diabetes."
        ),
    },
    {
        "title": (
            "Table 3. Adjusted associations of joint type 2 diabetes-retinal disease status "
            "with all-cause mortality and major adverse cardiovascular events in UK Biobank"
        ),
        "path": TABLE_DIR / "Table_3_UKB_Adjusted_Associations_Lancet_main.csv",
        "note": (
            "Reference group is no type 2 diabetes and no retinopathy. Adjusted models "
            "included age, sex, white ethnicity, Townsend deprivation, body-mass index, "
            "systolic blood pressure, total cholesterol, HbA1c, former smoking, and current smoking."
        ),
    },
]

APPENDIX_TABLE_SPEC = {
    "title": (
        "Appendix Table 3. Expanded joint, interaction, and additive interaction estimates "
        "for all-cause mortality and major adverse cardiovascular events"
    ),
    "path": TABLE_DIR / "Appendix_Table_3_UKB_Interaction_and_Full_Models_Lancet.csv",
    "note": (
        "This appendix-facing version preserves the full interaction model output for reviewer "
        "transparency and should remain secondary to the concise main-text Table 3."
    ),
}


def set_font(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.styles["Title"].font.size = Pt(16)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(13)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(11.5)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(11)
    doc.styles["Heading 3"].font.bold = True


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_section_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_font(run, size=13 if level == 1 else 11.5, bold=True)
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)


def add_body_paragraph(doc: Document, text: str, *, size: float = 11, align: int = WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=size)


def add_inline_subhead_paragraph(doc: Document, label: str, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(f"{label} ")
    set_font(run, size=11, bold=True)
    body = paragraph.add_run(text)
    set_font(body, size=11)


def add_centered_meta_line(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: float = 10.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)


def get_text(path: Path) -> str:
    return path.read_text(encoding="utf-8").strip()


def extract_title(markdown_text: str) -> str:
    first_line = markdown_text.splitlines()[0]
    return first_line.removeprefix("# ").strip()


def extract_block(markdown_text: str, heading: str, next_heading: str | None) -> str:
    start = markdown_text.index(heading)
    if next_heading:
        end = markdown_text.index(next_heading, start + len(heading))
        return markdown_text[start:end].strip()
    return markdown_text[start:].strip()


def iter_subsections(block_text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_heading: str | None = None
    buffer: list[str] = []
    for line in block_text.splitlines():
        if line.startswith("### "):
            if current_heading is not None:
                sections.append((current_heading, "\n".join(buffer).strip()))
            current_heading = line[4:].strip()
            buffer = []
            continue
        if line.startswith("## "):
            continue
        buffer.append(line)
    if current_heading is not None:
        sections.append((current_heading, "\n".join(buffer).strip()))
    return sections


def iter_level_two_sections(block_text: str) -> list[tuple[str, str]]:
    sections: list[tuple[str, str]] = []
    current_heading: str | None = None
    buffer: list[str] = []
    for line in block_text.splitlines():
        if line.startswith("## "):
            if current_heading is not None:
                sections.append((current_heading, "\n".join(buffer).strip()))
            current_heading = line[3:].strip()
            buffer = []
            continue
        if line.startswith("# "):
            continue
        buffer.append(line)
    if current_heading is not None:
        sections.append((current_heading, "\n".join(buffer).strip()))
    return sections


def clean_paragraphs(text: str) -> list[str]:
    paragraphs = []
    for part in re.split(r"\n\s*\n", text.strip()):
        part = re.sub(r"\s+", " ", part).strip()
        if part:
            paragraphs.append(part)
    return paragraphs


def add_summary_page(doc: Document, manuscript_text: str) -> None:
    block = extract_block(manuscript_text, "## Summary", "## Research in context")
    add_section_heading(doc, "Summary", level=1)
    for label, text in iter_subsections(block):
        add_inline_subhead_paragraph(doc, f"{label}:", " ".join(clean_paragraphs(text)))


def add_research_in_context_page(doc: Document, manuscript_text: str) -> None:
    block = extract_block(manuscript_text, "## Research in context", "## Introduction")
    add_section_heading(doc, "Research in context", level=1)
    for label, text in iter_subsections(block):
        add_inline_subhead_paragraph(doc, f"{label}:", " ".join(clean_paragraphs(text)))


def add_markdown_section(doc: Document, heading: str, body_text: str) -> None:
    add_section_heading(doc, heading, level=1)
    current_subheading: str | None = None
    buffer: list[str] = []
    for line in body_text.splitlines():
        if line.startswith("### "):
            if buffer:
                for paragraph in clean_paragraphs("\n".join(buffer)):
                    add_body_paragraph(doc, paragraph)
                buffer = []
            current_subheading = line[4:].strip()
            add_section_heading(doc, current_subheading, level=2)
            continue
        if line.strip():
            buffer.append(line)
        else:
            if buffer:
                for paragraph in clean_paragraphs("\n".join(buffer)):
                    add_body_paragraph(doc, paragraph)
                buffer = []
    if buffer:
        for paragraph in clean_paragraphs("\n".join(buffer)):
            add_body_paragraph(doc, paragraph)


def add_references(doc: Document, block_text: str) -> None:
    add_section_heading(doc, "References", level=1)
    for line in block_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if match:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.1
            run = paragraph.add_run(match.group(2))
            set_font(run, size=10.5)
        else:
            add_body_paragraph(doc, stripped, size=10.5)


def add_title_page(doc: Document, title: str, submission_word_count: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(title)
    set_font(run, size=17, bold=True)

    add_centered_meta_line(doc, "Target journal: The Lancet", bold=True)
    add_centered_meta_line(doc, "Article type: Original research", italic=True)
    add_centered_meta_line(doc, "Running title: Retinal disease, T2D, and cardiovascular risk")
    add_centered_meta_line(
        doc,
        f"Draft package word count (including Summary, Research in Context, legends, and table titles): {submission_word_count}",
    )
    add_centered_meta_line(doc, "Figures: 4 main figures")
    add_centered_meta_line(doc, "Tables: 3 main tables plus 1 appendix table")

    doc.add_paragraph()
    add_centered_meta_line(doc, "Authors: To be finalised by the study team", bold=True, size=11)
    add_centered_meta_line(doc, "Affiliations: To be finalised by the study team", size=11)
    add_centered_meta_line(doc, "Corresponding author: To be finalised by the study team", size=11)
    add_centered_meta_line(doc, "Funding statement: No dedicated external funding supported this analytical workflow.", size=10.5)
    add_centered_meta_line(doc, "This file is a submission-facing internal review draft generated from the quality-controlled analysis package.", size=10)


def add_table_from_dataframe(doc: Document, dataframe: pd.DataFrame, *, font_size: float = 9.5) -> None:
    table = doc.add_table(rows=dataframe.shape[0] + 1, cols=dataframe.shape[1])
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for col_idx, column in enumerate(dataframe.columns):
        cell = table.cell(0, col_idx)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(column))
        set_font(run, size=font_size, bold=True)

    for row_idx, (_, row) in enumerate(dataframe.iterrows(), start=1):
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx else WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(value))
            set_font(run, size=font_size)

    doc.add_paragraph()


def add_figure_legends(doc: Document, figure_legends_text: str) -> None:
    add_section_heading(doc, "Figure legends", level=1)
    for label, text in iter_level_two_sections(figure_legends_text):
        add_section_heading(doc, label, level=2)
        for paragraph in clean_paragraphs(text):
            add_body_paragraph(doc, paragraph, size=10.5)


def set_section_landscape(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def add_table_section(doc: Document) -> None:
    section = doc.add_section(WD_SECTION_START.NEW_PAGE)
    set_section_landscape(section)
    add_section_heading(doc, "Table titles and review tables", level=1)
    for index, spec in enumerate(TABLE_SPECS, start=1):
        if index == 3:
            add_page_break(doc)
        add_section_heading(doc, spec["title"], level=2)
        add_body_paragraph(doc, spec["note"], size=9.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        dataframe = pd.read_csv(spec["path"]).fillna("")
        add_table_from_dataframe(doc, dataframe, font_size=9)

    add_page_break(doc)
    add_section_heading(doc, APPENDIX_TABLE_SPEC["title"], level=2)
    add_body_paragraph(doc, APPENDIX_TABLE_SPEC["note"], size=9.3, align=WD_ALIGN_PARAGRAPH.LEFT)
    appendix_df = pd.read_csv(APPENDIX_TABLE_SPEC["path"]).fillna("")
    add_table_from_dataframe(doc, appendix_df, font_size=8.5)


def count_submission_words(manuscript_text: str, figure_legends_text: str) -> int:
    words = re.findall(r"\b[\w.-]+\b", manuscript_text + "\n" + figure_legends_text)
    table_words = []
    for spec in TABLE_SPECS + [APPENDIX_TABLE_SPEC]:
        dataframe = pd.read_csv(spec["path"]).fillna("")
        table_words.extend(re.findall(r"\b[\w.-]+\b", spec["title"]))
        for value in dataframe.astype(str).to_numpy().flatten():
            table_words.extend(re.findall(r"\b[\w.-]+\b", value))
    return len(words) + len(table_words)


def main() -> None:
    manuscript_text = get_text(MANUSCRIPT_PATH)
    figure_legends_text = get_text(FIGURE_LEGENDS_PATH)
    title = extract_title(manuscript_text)
    submission_word_count = count_submission_words(manuscript_text, figure_legends_text)

    doc = Document()
    configure_document(doc)
    add_title_page(doc, title, submission_word_count)
    add_page_break(doc)
    add_summary_page(doc, manuscript_text)
    add_page_break(doc)
    add_research_in_context_page(doc, manuscript_text)
    add_page_break(doc)

    sections = [
        ("Introduction", extract_block(manuscript_text, "## Introduction", "## Methods").replace("## Introduction", "", 1).strip()),
        ("Methods", extract_block(manuscript_text, "## Methods", "## Results").replace("## Methods", "", 1).strip()),
        ("Results", extract_block(manuscript_text, "## Results", "## Discussion").replace("## Results", "", 1).strip()),
        ("Discussion", extract_block(manuscript_text, "## Discussion", "## Contributors").replace("## Discussion", "", 1).strip()),
        ("Contributors", extract_block(manuscript_text, "## Contributors", "## Declaration of interests").replace("## Contributors", "", 1).strip()),
        (
            "Declaration of interests",
            extract_block(manuscript_text, "## Declaration of interests", "## Data sharing").replace("## Declaration of interests", "", 1).strip(),
        ),
        ("Data sharing", extract_block(manuscript_text, "## Data sharing", "## Acknowledgments").replace("## Data sharing", "", 1).strip()),
        ("Acknowledgments", extract_block(manuscript_text, "## Acknowledgments", "## References").replace("## Acknowledgments", "", 1).strip()),
    ]
    for heading, body in sections:
        add_markdown_section(doc, heading, body)

    add_references(doc, extract_block(manuscript_text, "## References", None).replace("## References", "", 1).strip())
    add_page_break(doc)
    add_figure_legends(doc, figure_legends_text)
    add_table_section(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved DOCX draft to {OUTPUT_PATH}")
    print(f"Submission package word count: {submission_word_count}")


if __name__ == "__main__":
    main()
