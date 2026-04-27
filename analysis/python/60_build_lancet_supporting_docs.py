#!/usr/bin/env python3
"""Build supporting submission documents for the Lancet DR-T2D package."""

from __future__ import annotations

import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_DIR = Path(
    "/Users/apple/Desktop/研究方案-赵老师项目/0 研究方案-基于全球疾病负担(GBD)和英国生物数据银行（UK Biobank）的糖尿病视网膜病变与2型糖尿病（T2D）交互作用-10分以上/"
    "Codex研究产出_2026-03-07/04_多期刊投稿包/GBD2023_UKB_DR_T2D_2026-03-07/journal_01_eClinicalMedicine"
)
MANUSCRIPT_MD = PROJECT_DIR / "13_MAIN_MANUSCRIPT_Lancet_main_v3.md"
SUPP_DIR = PROJECT_DIR / "supplementary_tables"

TITLE_PAGE_MD = PROJECT_DIR / "14_TITLE_PAGE_TheLancet.md"
TITLE_PAGE_DOCX = PROJECT_DIR / "14_TITLE_PAGE_TheLancet.docx"
COVER_LETTER_MD = PROJECT_DIR / "15_COVER_LETTER_TheLancet.md"
COVER_LETTER_DOCX = PROJECT_DIR / "15_COVER_LETTER_TheLancet.docx"
CHECKLIST_MD = PROJECT_DIR / "16_REPORTING_CHECKLIST_STROBE_GATHER_ALIGNED.md"
CHECKLIST_DOCX = PROJECT_DIR / "16_REPORTING_CHECKLIST_STROBE_GATHER_ALIGNED.docx"
APPENDIX_MD = PROJECT_DIR / "17_SUPPLEMENTARY_APPENDIX_TheLancet.md"
APPENDIX_DOCX = PROJECT_DIR / "17_SUPPLEMENTARY_APPENDIX_TheLancet.docx"

SUPPLEMENTARY_TABLE_TITLES = {
    "Supplementary_Table_S1_UKB_Missingness_Incident_Cohort.csv": "Supplementary Table S1. Missingness across key UK Biobank variables in the incident cohort",
    "Supplementary_Table_S2_UKB_Event_Rates_Per1000PY.csv": "Supplementary Table S2. Event counts, person-years, and event rates by joint exposure group",
    "Supplementary_Table_S3_UKB_Model_Hierarchy.csv": "Supplementary Table S3. Hierarchical Cox models for all-cause mortality, MACE, and heart failure",
    "Supplementary_Table_S4_UKB_T2D_Definition_Sensitivity.csv": "Supplementary Table S4. Sensitivity analysis using the raw cohort T2D definition",
    "Supplementary_Table_S5_UKB_Sex_Stratified_Full_Models.csv": "Supplementary Table S5. Sex-stratified full-model associations",
    "Supplementary_Table_S6_UKB_Heart_Failure_Associations.csv": "Supplementary Table S6. Heart-failure associations across model structures",
    "Supplementary_Table_S7_GBD_EAPC.csv": "Supplementary Table S7. Estimated annual percentage change for global GBD burden series",
    "Supplementary_Table_S8_GBD_Regional_Prevalence_ASR_2023.csv": "Supplementary Table S8. Regional age-standardised prevalence rates in 2023",
    "Supplementary_Table_S9_GBD_Country_Top10_VisionLoss_ASR_2023.csv": "Supplementary Table S9. Top 10 countries or territories for age-standardised T2D-related vision-loss prevalence in 2023",
    "Supplementary_Table_S10_GBD_Severity_Split_2023.csv": "Supplementary Table S10. Global severity structure of T2D-related vision loss in 2023",
    "Supplementary_Table_S11_GBD_Global_Sex_ASR_2023.csv": "Supplementary Table S11. Global sex-specific age-standardised prevalence rates in 2023",
}


def set_font(run, *, size: float | None = None, bold: bool | None = None, italic: bool | None = None, superscript: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if superscript is not None:
        run.font.superscript = superscript


def configure_document(doc: Document, *, landscape: bool = False) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9 if not landscape else 0.6)
    section.right_margin = Inches(0.9 if not landscape else 0.6)
    if landscape:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
    else:
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    doc.styles["Title"].font.size = Pt(16)
    doc.styles["Title"].font.bold = True
    doc.styles["Heading 1"].font.size = Pt(13)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(11.5)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(11)
    doc.styles["Heading 3"].font.bold = True


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(level=level)
    run = paragraph.add_run(text)
    set_font(run, size=13 if level == 1 else 11.5, bold=True)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, *, align: int = WD_ALIGN_PARAGRAPH.LEFT, size: float = 11, bold: bool = False, italic: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def manuscript_text() -> str:
    return MANUSCRIPT_MD.read_text(encoding="utf-8")


def manuscript_title(text: str) -> str:
    return text.splitlines()[0].removeprefix("# ").strip()


def block(text: str, start_heading: str, end_heading: str | None) -> str:
    start = text.index(start_heading)
    if end_heading is None:
        return text[start:].strip()
    end = text.index(end_heading, start + len(start_heading))
    return text[start:end].strip()


def words_in(text: str) -> int:
    return len(re.findall(r"\b[\w.-]+\b", text))


def title_page_markdown(text: str) -> str:
    title = manuscript_title(text)
    summary_words = words_in(block(text, "## Summary", "## Research in context"))
    research_context_words = words_in(block(text, "## Research in context", "## Introduction"))
    main_text_words = words_in(block(text, "## Introduction", "## Contributors"))
    total_words = words_in(text)
    lines = [
        "# Title Page",
        "",
        f"**Full title:** {title}",
        "**Short running title:** Retinal disease, T2D, and cardiovascular risk",
        "",
        "**Authors**",
        "- [Author 1], [highest degree], [affiliation superscripts]",
        "- [Author 2], [highest degree], [affiliation superscripts]",
        "- [Author 3], [highest degree], [affiliation superscripts]",
        "",
        "**Affiliations**",
        "1. [Department], [Institution], [City], [Country]",
        "2. [Department], [Institution], [City], [Country]",
        "3. [Department], [Institution], [City], [Country]",
        "",
        "**Corresponding author**",
        "- [Name], [Department], [Institution], [Full postal address]",
        "- Email: [email]",
        "- Telephone: [telephone]",
        "",
        "**Manuscript metrics**",
        f"- Summary words: {summary_words}",
        f"- Research in Context words: {research_context_words}",
        f"- Main-text words (Introduction through Discussion): {main_text_words}",
        f"- Total manuscript words in the current markdown draft: {total_words}",
        "- Figures: 4 main figures",
        "- Tables: 3 main tables",
        "- Supplementary tables: 11",
        "",
        "**Funding statement**",
        "- No dedicated external funding supported this analytical workflow.",
        "",
        "**Declaration of interests**",
        "- To be completed by each author.",
        "",
        "**Key abbreviations**",
        "- ASR, age-standardised rate",
        "- DALY, disability-adjusted life-year",
        "- HF, heart failure",
        "- MACE, major adverse cardiovascular events",
        "- T2D, type 2 diabetes",
        "- YLD, years lived with disability",
    ]
    return "\n".join(lines) + "\n"


def build_title_page_docx(text: str) -> None:
    doc = Document()
    configure_document(doc)
    title = manuscript_title(text)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    set_font(run, size=16, bold=True)
    paragraph.paragraph_format.space_after = Pt(14)

    add_paragraph(doc, "Running title: Retinal disease, T2D, and cardiovascular risk", bold=True)
    add_paragraph(doc, "Authors", bold=True)
    for author_line in [
        "[Author 1], [highest degree]",
        "[Author 2], [highest degree]",
        "[Author 3], [highest degree]",
    ]:
        add_paragraph(doc, author_line)

    add_paragraph(doc, "Affiliations", bold=True)
    for idx in range(1, 4):
        paragraph = doc.add_paragraph()
        run_num = paragraph.add_run(f"{idx}. ")
        set_font(run_num, size=11, bold=True)
        run_text = paragraph.add_run("[Department], [Institution], [City], [Country]")
        set_font(run_text, size=11)
        paragraph.paragraph_format.space_after = Pt(3)

    add_paragraph(doc, "Corresponding author", bold=True)
    add_paragraph(doc, "[Name], [Department], [Institution], [Full postal address]")
    add_paragraph(doc, "Email: [email]")
    add_paragraph(doc, "Telephone: [telephone]")

    text_md = title_page_markdown(text)
    for line in text_md.splitlines():
        if line.startswith("- Summary words:") or line.startswith("- Research in Context") or line.startswith("- Main-text words") or line.startswith("- Total manuscript words") or line.startswith("- Figures:") or line.startswith("- Tables:") or line.startswith("- Supplementary tables:"):
            add_paragraph(doc, line[2:])

    add_paragraph(doc, "Funding statement", bold=True)
    add_paragraph(doc, "No dedicated external funding supported this analytical workflow.")
    add_paragraph(doc, "Declaration of interests", bold=True)
    add_paragraph(doc, "To be completed by each author.")
    add_paragraph(doc, "Key abbreviations", bold=True)
    for item in [
        "ASR, age-standardised rate",
        "DALY, disability-adjusted life-year",
        "HF, heart failure",
        "MACE, major adverse cardiovascular events",
        "T2D, type 2 diabetes",
        "YLD, years lived with disability",
    ]:
        add_paragraph(doc, item)

    doc.save(TITLE_PAGE_DOCX)


def cover_letter_markdown(text: str) -> str:
    title = manuscript_title(text)
    return (
        "# Cover Letter\n\n"
        "Date: 2026-03-26\n\n"
        "To the Editors of The Lancet,\n\n"
        f"We submit our manuscript, \"{title}\", for consideration as an original research Article in *The Lancet*.\n\n"
        "This study addresses a clinically and globally relevant problem by integrating two complementary scales of evidence. "
        "Using GBD 2023, we show that the public burden proxy for type 2 diabetes-related vision loss has risen faster than the broader burden of type 2 diabetes itself. "
        "Using UK Biobank, we show that adults with both type 2 diabetes and retinal disease phenotype have the highest subsequent risk of major adverse cardiovascular events, elevated all-cause mortality, and a strong supplementary heart-failure signal.\n\n"
        "We believe the manuscript fits *The Lancet* because it links a growing global chronic-disease burden to a practical clinical stratification opportunity: retinal findings that are already collected in diabetes care may identify a subgroup with substantially higher systemic vascular risk. "
        "The paper is not framed as a narrow ophthalmology analysis. Instead, it argues for tighter integration of diabetes, retinal, and cardiovascular prevention pathways, which aligns with the journal's broad clinical and population-health readership.\n\n"
        "The analysis also takes a conservative methodological stance. We distinguish the stronger message of joint prognostic enrichment from the weaker claim of formal biological interaction, and we provide robustness analyses across hierarchical covariate adjustment, an alternative diabetes definition, sex-stratified models, and a supplementary heart-failure outcome.\n\n"
        "This manuscript is not under consideration elsewhere. All authors will need to confirm approval of the submitted version before formal submission. Any required declarations, authorship forms, funding statements, and competing-interest disclosures will be provided separately through the journal's submission system as requested.\n\n"
        "Sincerely,\n\n"
        "[Corresponding author name]\n"
        "[Institution]\n"
        "[Email]\n"
    )


def build_cover_letter_docx(text: str) -> None:
    doc = Document()
    configure_document(doc)
    add_paragraph(doc, "Date: 2026-03-26")
    add_paragraph(doc, "To the Editors of The Lancet,")
    for paragraph in [p for p in cover_letter_markdown(text).split("\n\n") if p and not p.startswith("# Cover Letter") and not p.startswith("Date: 2026-03-26") and not p.startswith("To the Editors")]:
        add_paragraph(doc, paragraph)
    doc.save(COVER_LETTER_DOCX)


def checklist_sections() -> tuple[list[dict[str, str]], list[dict[str, str]]]:
    strobe_rows = [
        {"Item": "1", "Checklist item": "Title and abstract indicate the study design and provide an informative balanced summary.", "Location in draft": "Summary; title; Methods"},
        {"Item": "2", "Checklist item": "Background and rationale.", "Location in draft": "Introduction paragraphs 1-6"},
        {"Item": "3", "Checklist item": "Objectives and prespecified aims.", "Location in draft": "Introduction final paragraph"},
        {"Item": "4", "Checklist item": "Key elements of study design.", "Location in draft": "Methods, Study design and data sources"},
        {"Item": "5", "Checklist item": "Setting, locations, and relevant dates.", "Location in draft": "Methods, Study design and data sources; UK Biobank cohort construction"},
        {"Item": "6", "Checklist item": "Participants, eligibility, and cohort construction.", "Location in draft": "Methods, UK Biobank cohort construction"},
        {"Item": "7", "Checklist item": "Definitions of exposures, outcomes, predictors, confounders, and effect modifiers.", "Location in draft": "Methods, Exposure, outcomes, and covariates"},
        {"Item": "8", "Checklist item": "Data sources and measurement methods.", "Location in draft": "Methods, GBD 2023 query specification; Exposure, outcomes, and covariates"},
        {"Item": "9", "Checklist item": "Efforts to address bias.", "Location in draft": "Methods, cohort QC description; Discussion, Strengths and limitations"},
        {"Item": "10", "Checklist item": "Study size.", "Location in draft": "Results, Study profile; Table 1; Supplementary Tables S1-S3"},
        {"Item": "11", "Checklist item": "Handling of quantitative variables.", "Location in draft": "Methods, Exposure, outcomes, and covariates; Statistical analysis"},
        {"Item": "12", "Checklist item": "Statistical methods including confounder adjustment and subgroup/sensitivity analyses.", "Location in draft": "Methods, Statistical analysis; Supplementary Tables S1-S6"},
        {"Item": "13", "Checklist item": "Participant flow and numbers at each stage.", "Location in draft": "Results, Study profile; Figure 1; Table 1"},
        {"Item": "14", "Checklist item": "Descriptive data and missingness.", "Location in draft": "Results, UK Biobank cohort profile; Table 1; Supplementary Table S1"},
        {"Item": "15", "Checklist item": "Outcome data.", "Location in draft": "Results, UK Biobank cohort profile; Supplementary Table S2"},
        {"Item": "16", "Checklist item": "Main results with estimates and precision.", "Location in draft": "Results, Joint associations; Table 3; Figure 4"},
        {"Item": "17", "Checklist item": "Other analyses including subgroups and sensitivity analyses.", "Location in draft": "Results, Interaction analyses; Sensitivity analyses; Supplementary Tables S3-S6"},
        {"Item": "18", "Checklist item": "Key results in relation to objectives.", "Location in draft": "Discussion, Principal findings"},
        {"Item": "19", "Checklist item": "Limitations.", "Location in draft": "Discussion, Strengths and limitations"},
        {"Item": "20", "Checklist item": "Interpretation with caution.", "Location in draft": "Discussion, Interpretation and relation to previous evidence; Conclusion"},
        {"Item": "21", "Checklist item": "Generalisability.", "Location in draft": "Discussion, Strengths and limitations"},
        {"Item": "22", "Checklist item": "Funding and role of funders.", "Location in draft": "Summary, Funding; Methods; Acknowledgments"},
    ]
    gather_rows = [
        {"Item": "1", "Checklist item": "Define the indicator(s), populations, and time period.", "Location in draft": "Methods, GBD 2023 query specification"},
        {"Item": "2", "Checklist item": "List funding sources.", "Location in draft": "Summary, Funding; Acknowledgments"},
        {"Item": "3", "Checklist item": "Describe data inputs and data-identification methods.", "Location in draft": "Methods, GBD 2023 query specification"},
        {"Item": "4", "Checklist item": "Specify inclusion and exclusion criteria for data inputs.", "Location in draft": "Methods, GBD 2023 query specification; Discussion, scope note on public proxy"},
        {"Item": "5", "Checklist item": "Provide data-input metadata.", "Location in draft": "Methods, GBD 2023 query specification; Supplementary Tables S7-S11"},
        {"Item": "6", "Checklist item": "Document data inputs unavailable for sharing.", "Location in draft": "Data sharing; Methods, secondary use of public IHME outputs"},
        {"Item": "7", "Checklist item": "Describe analytical or statistical methods.", "Location in draft": "Methods, Statistical analysis"},
        {"Item": "8", "Checklist item": "Describe model selection and evaluation where relevant.", "Location in draft": "Methods, Statistical analysis; note that GBD estimates were extracted, not re-estimated"},
        {"Item": "9", "Checklist item": "Describe performance metrics and sensitivity analyses where relevant.", "Location in draft": "Methods, Statistical analysis; Supplementary Tables S7-S11"},
        {"Item": "10", "Checklist item": "Provide access to analytic or statistical source code if possible.", "Location in draft": "Data sharing; local analysis scripts in project workspace"},
        {"Item": "11", "Checklist item": "Present published estimates in a file format from which data can be extracted.", "Location in draft": "Supplementary Tables S7-S11; project CSV outputs"},
        {"Item": "12", "Checklist item": "Report uncertainty for estimates.", "Location in draft": "Results, burden trends and heterogeneity; Table 2; Supplementary Tables S7-S11"},
        {"Item": "13", "Checklist item": "State how to access additional results.", "Location in draft": "Data sharing; Supplementary appendix"},
        {"Item": "14", "Checklist item": "Discuss limitations of input data.", "Location in draft": "Discussion, Strengths and limitations"},
        {"Item": "15", "Checklist item": "Interpret results in light of existing evidence.", "Location in draft": "Discussion, Interpretation and relation to previous evidence"},
        {"Item": "16", "Checklist item": "Discuss implications for users of the estimates.", "Location in draft": "Discussion, Clinical implications"},
        {"Item": "17", "Checklist item": "Explain differences from previously published estimates when relevant.", "Location in draft": "Discussion, Interpretation and relation to previous evidence"},
        {"Item": "18", "Checklist item": "Clarify that the GBD component is secondary use of public estimates rather than de novo IHME model generation.", "Location in draft": "Methods, GBD 2023 query specification; Discussion, Strengths and limitations"},
    ]
    return strobe_rows, gather_rows


def checklist_markdown() -> str:
    strobe_rows, gather_rows = checklist_sections()
    lines = [
        "# Reporting Checklist",
        "",
        "This draft reporting map combines a STROBE cohort checklist for the UK Biobank component with a GATHER-aligned checklist for the secondary analysis of public GBD 2023 estimates.",
        "",
        "## STROBE Cohort Checklist",
        "",
        "| Item | Checklist item | Location in draft |",
        "| --- | --- | --- |",
    ]
    for row in strobe_rows:
        lines.append(f"| {row['Item']} | {row['Checklist item']} | {row['Location in draft']} |")
    lines.extend(["", "## GATHER-Aligned Reporting Map", "", "| Item | Checklist item | Location in draft |", "| --- | --- | --- |"])
    for row in gather_rows:
        lines.append(f"| {row['Item']} | {row['Checklist item']} | {row['Location in draft']} |")
    lines.append("")
    return "\n".join(lines)


def add_table(doc: Document, rows: list[dict[str, str]], headers: list[str], *, font_size: float = 9.3) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, header in enumerate(headers):
        cell = table.cell(0, col_idx)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_font(run, size=font_size, bold=True)
    for row_idx, row in enumerate(rows, start=1):
        for col_idx, header in enumerate(headers):
            cell = table.cell(row_idx, col_idx)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(str(row[header]))
            set_font(run, size=font_size)
    doc.add_paragraph()


def build_checklist_docx() -> None:
    doc = Document()
    configure_document(doc)
    add_heading(doc, "Reporting checklist", level=1)
    add_paragraph(
        doc,
        "This draft reporting map combines a STROBE cohort checklist for the UK Biobank component with a GATHER-aligned checklist for the secondary analysis of public GBD 2023 estimates.",
    )
    strobe_rows, gather_rows = checklist_sections()
    add_heading(doc, "STROBE cohort checklist", level=2)
    add_table(doc, strobe_rows, ["Item", "Checklist item", "Location in draft"])
    add_heading(doc, "GATHER-aligned reporting map", level=2)
    add_table(doc, gather_rows, ["Item", "Checklist item", "Location in draft"])
    doc.save(CHECKLIST_DOCX)


def appendix_markdown(text: str) -> str:
    title = manuscript_title(text)
    lines = [
        "# Supplementary Appendix",
        "",
        f"**Manuscript title:** {title}",
        "**Version date:** 2026-03-26",
        "",
        "## Scope note",
        "",
        "This appendix accompanies the Lancet-main draft and contains supplementary cohort diagnostics, sensitivity analyses, and expanded GBD descriptive tables. The UK Biobank supplement should be interpreted as a robustness package rather than a new primary-analysis layer.",
        "",
        "## Contents",
        "",
    ]
    for filename, title_line in SUPPLEMENTARY_TABLE_TITLES.items():
        lines.append(f"- {title_line} (`{filename}`)")
    lines.append("")
    return "\n".join(lines)


def set_landscape_section(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def add_dataframe_table(doc: Document, dataframe: pd.DataFrame, *, font_size: float = 8.7) -> None:
    table = doc.add_table(rows=dataframe.shape[0] + 1, cols=dataframe.shape[1])
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col_idx, column in enumerate(dataframe.columns):
        paragraph = table.cell(0, col_idx).paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(str(column))
        set_font(run, size=font_size, bold=True)
    for row_idx, (_, row) in enumerate(dataframe.iterrows(), start=1):
        for col_idx, value in enumerate(row):
            paragraph = table.cell(row_idx, col_idx).paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx < 2 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(str(value))
            set_font(run, size=font_size)
    doc.add_paragraph()


def build_appendix_docx(text: str) -> None:
    doc = Document()
    configure_document(doc)
    title = manuscript_title(text)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Supplementary Appendix")
    set_font(run, size=16, bold=True)
    paragraph.paragraph_format.space_after = Pt(12)
    add_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_paragraph(doc, "Version date: 2026-03-26", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(
        doc,
        "This appendix contains supplementary cohort diagnostics, sensitivity analyses, and expanded GBD descriptive tables for the dual-source manuscript.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(doc, "Contents", bold=True)
    for title_line in SUPPLEMENTARY_TABLE_TITLES.values():
        add_paragraph(doc, title_line)

    first_table = True
    for filename, title_line in SUPPLEMENTARY_TABLE_TITLES.items():
        if first_table:
            section = doc.add_section(WD_SECTION_START.NEW_PAGE)
            set_landscape_section(section)
            first_table = False
        else:
            add_page_break(doc)
        add_heading(doc, title_line, level=1)
        dataframe = pd.read_csv(SUPP_DIR / filename).fillna("")
        add_dataframe_table(doc, dataframe)

    doc.save(APPENDIX_DOCX)


def main() -> None:
    text = manuscript_text()

    TITLE_PAGE_MD.write_text(title_page_markdown(text), encoding="utf-8")
    COVER_LETTER_MD.write_text(cover_letter_markdown(text), encoding="utf-8")
    CHECKLIST_MD.write_text(checklist_markdown(), encoding="utf-8")
    APPENDIX_MD.write_text(appendix_markdown(text), encoding="utf-8")

    build_title_page_docx(text)
    build_cover_letter_docx(text)
    build_checklist_docx()
    build_appendix_docx(text)

    print(f"Saved supporting docs under {PROJECT_DIR}")


if __name__ == "__main__":
    main()
